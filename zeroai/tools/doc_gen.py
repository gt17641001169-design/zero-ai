"""文档生成工具（Word / Excel / PDF）

迁移来源：tui_agent.py 行 3959-3991（_add_formatted_text）、
5084-5114（render_latex_in_text）、5117-6184（Word/Excel/PDF 生成函数）

提供以下纯函数：
- _add_formatted_text：解析 Markdown 粗体/斜体/行内代码/删除线并添加到段落
- render_latex_in_text：检测文本中的 LaTeX 公式并渲染为 Unicode
- generate_word：生成 Word 文档（.docx），支持 8 种模板（含学术论文模板）
- generate_excel：生成 Excel 文档（.xlsx），支持图表、公式、多工作表
- generate_pdf：生成 PDF 文档，支持 Markdown 标记和多种模板
- _latex_to_image：将 LaTeX 公式渲染为 PNG 图片（PDF 学术模板用）
- _parse_color / _set_cell_shading / _add_page_number_field 等辅助函数

依赖：
- 标准库：re, os, pathlib
- zeroai.core.paths：_resolve_save_path（文档默认保存到桌面）
- .file_manager：_load_svg_icon（返回结果中的图标标签）
- .academic：_latex_to_unicode（LaTeX 公式渲染为 Unicode，用于学术模板）
- 可选第三方库：python-docx（Word）、openpyxl（Excel）、reportlab（PDF）、matplotlib（PDF 公式图片）

注意：本模块的所有文档生成函数均支持「path 为空或只含文件名时自动保存到桌面」的便捷特性。
"""
import re
from pathlib import Path

from zeroai.core.paths import _resolve_save_path
from .file_manager import _load_svg_icon
from .academic import _latex_to_unicode


# ====== Markdown 格式化文本解析 ======
# 迁移来源：tui_agent.py 行 3959-3991

def _add_formatted_text(paragraph, text: str):
    """解析 Markdown 粗体/斜体/行内代码/删除线并添加到段落"""
    # 匹配 **粗体** *斜体* `代码` ~~删除线~~
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|~~(.+?)~~)'
    last_end = 0

    for match in re.finditer(pattern, text):
        # 添加前面的普通文本
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # **粗体**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *斜体*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `行内代码`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            from docx.shared import Pt as _Pt, RGBColor as _RC
            run.font.size = _Pt(10)
            run.font.color.rgb = _RC(0xC0, 0x39, 0x2B)
        elif match.group(5):  # ~~删除线~~
            run = paragraph.add_run(match.group(5))
            run.font.strike = True

        last_end = match.end()

    # 添加剩余文本
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


# ====== LaTeX 公式渲染（终端 Unicode 渲染预处理）======
# 迁移来源：tui_agent.py 行 5084-5114

def render_latex_in_text(text: str) -> str:
    """检测文本中的 LaTeX 公式并渲染为 Unicode

    支持：
    - 行内公式：$E=mc^2$
    - 块级公式：$$\\int_0^1 f(x)dx$$
    - \\( \\) 行内公式
    - \\[ \\] 块级公式

    用于终端 Markdown 渲染前的预处理
    """
    if not text or "$" not in text and "\\(" not in text and "\\[" not in text:
        return text

    # 1. 块级公式 $$...$$
    def _block_formula(m):
        rendered = _latex_to_unicode(m.group(1))
        return f"\n   {rendered}\n"
    text = re.sub(r"\$\$([^$]+)\$\$", _block_formula, text)

    # 2. 块级公式 \[...\]
    text = re.sub(r"\\\[([^\]]+)\]", lambda m: f"\n   {_latex_to_unicode(m.group(1))}\n", text)

    # 3. 行内公式 \(...\)
    text = re.sub(r"\\\(([^)]+)\)", lambda m: _latex_to_unicode(m.group(1)), text)

    # 4. 行内公式 $...$（最后处理，避免误伤 $$）
    # 使用非贪婪匹配，且内容不含 $
    text = re.sub(r"(?<!\$)\$([^$\n]+)\$(?!\$)", lambda m: _latex_to_unicode(m.group(1)), text)

    return text


# ====== Word 文档格式模板 ======
# 迁移来源：tui_agent.py 行 5117-5162
WORD_TEMPLATES = {
    "default": {
        "font": "Microsoft YaHei", "font_size": 11,
        "heading_color": "1A1A1A", "margin": [2.54, 2.54, 2.54, 2.54],
        "header": "", "footer": "",
    },
    "report": {  # 正式报告
        "font": "SimSun", "font_size": 12,
        "heading_color": "1F4E79", "margin": [3.0, 2.5, 2.5, 2.5],
        "header": "", "footer": "第 {page} 页",
    },
    "contract": {  # 合同
        "font": "SimSun", "font_size": 12,
        "heading_color": "000000", "margin": [2.54, 3.0, 2.54, 3.0],
        "header": "", "footer": "— {page} —",
    },
    "resume": {  # 简历
        "font": "Microsoft YaHei", "font_size": 10,
        "heading_color": "2E75B6", "margin": [1.5, 2.0, 1.5, 2.0],
        "header": "", "footer": "",
    },
    "thesis": {  # 学术论文
        "font": "Times New Roman", "font_size": 12,
        "heading_color": "000000", "margin": [2.54, 3.17, 2.54, 3.17],
        "header": "", "footer": "{page}",
    },
    "letter": {  # 信函
        "font": "KaiTi", "font_size": 14,
        "heading_color": "333333", "margin": [2.54, 2.54, 2.54, 2.54],
        "header": "", "footer": "",
    },
    "technical": {  # 技术文档
        "font": "Microsoft YaHei", "font_size": 10,
        "heading_color": "0070C0", "margin": [2.0, 2.0, 2.0, 2.0],
        "header": "技术文档", "footer": "第 {page} 页 / 共 {numpages} 页",
    },
    "academic": {  # 学术论文（严谨版：双倍行距、页码、Times New Roman）
        "font": "Times New Roman", "font_size": 12,
        "heading_color": "000000", "margin": [2.54, 2.54, 2.54, 2.54],
        "header": "", "footer": "{page}",
        "line_spacing": 2.0,  # 双倍行距（学术规范）
        "abstract_label": "摘要", "keywords_label": "关键词",
        "references_label": "参考文献", "doi_support": True,
    },
}

# 颜色名映射
# 迁移来源：tui_agent.py 行 5164-5171
COLOR_MAP = {
    "black": "000000", "white": "FFFFFF", "red": "FF0000", "green": "008000",
    "blue": "0000FF", "yellow": "FFFF00", "gray": "808080", "grey": "808080",
    "darkgray": "A9A9A9", "darkblue": "00008B", "navy": "1F4E79",
    "orange": "FFA500", "purple": "800080", "brown": "A52A2A",
    "cyan": "00FFFF", "magenta": "FF00FF", "lime": "00FF00",
}


def _parse_color(color_str: str):
    """解析颜色字符串，返回 RGBColor

    迁移来源：tui_agent.py 行 5174-5187
    """
    from docx.shared import RGBColor
    color_str = color_str.strip().lstrip("#")
    # 颜色名
    if color_str.lower() in COLOR_MAP:
        color_str = COLOR_MAP[color_str.lower()]
    # 6位十六进制
    if len(color_str) == 6 and all(c in "0123456789abcdefABCDEF" for c in color_str):
        r = int(color_str[0:2], 16)
        g = int(color_str[2:4], 16)
        b = int(color_str[4:6], 16)
        return RGBColor(r, g, b)
    return None


def _set_cell_shading(cell, color_hex: str):
    """设置表格单元格背景色

    迁移来源：tui_agent.py 行 5190-5200
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex.lstrip("#"))
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    except Exception:
        pass


def _add_page_number_field(paragraph, field: str = "PAGE"):
    """在段落中添加页码域代码

    迁移来源：tui_agent.py 行 5203-5220
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f" {field} "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    except Exception:
        pass


def _apply_footer_template(section, template: dict):
    """应用页脚模板

    迁移来源：tui_agent.py 行 5223-5245
    """
    from docx.shared import Pt
    footer_text = template.get("footer", "")
    if not footer_text:
        return
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = 1  # 居中

    # 解析 {page} {numpages} 占位符
    parts = re.split(r'(\{page\}|\{numpages\})', footer_text)
    for part in parts:
        if part == "{page}":
            _add_page_number_field(p, "PAGE")
        elif part == "{numpages}":
            _add_page_number_field(p, "NUMPAGES")
        elif part:
            run = p.add_run(part)
            run.font.size = Pt(9)
            run.font.color.rgb = _parse_color("888888")


def _apply_header_template(section, template: dict):
    """应用页眉模板

    迁移来源：tui_agent.py 行 5248-5261
    """
    from docx.shared import Pt
    header_text = template.get("header", "")
    if not header_text:
        return
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = 1  # 居中
    run = p.add_run(header_text)
    run.font.size = Pt(9)
    run.font.color.rgb = _parse_color("888888")


def _parse_table(lines: list, start_idx: int):
    """解析 Markdown 表格语法，返回 (rows, end_idx)

    迁移来源：tui_agent.py 行 5264-5285

    表格格式：
    | 列1 | 列2 | 列3 |
    |----|----|----|
    | a  | b  | c  |
    """
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # 跳过分隔行 |---|---|
        if re.match(r'^\|[\s\-:]+\|', line):
            i += 1
            continue
        # 解析单元格
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def generate_word(path: str, content: str, title: str = "", template: str = "default",
                  font: str = "", font_size: int = 0, margin: list = None,
                  heading_color: str = "", align: str = "", header: str = "",
                  footer: str = "") -> str:
    """生成 Word 文档（.docx）- 支持指定格式

    迁移来源：tui_agent.py 行 5288-5645

    Args:
        path: 保存路径（为空或只含文件名时自动保存到桌面）
        content: 文档内容（支持 Markdown 风格标记）
        title: 文档标题（可选）
        template: 格式模板（default/report/contract/resume/thesis/letter/technical/academic）
                    academic=学术论文（双倍行距/摘要/关键词/参考文献自动编号/LaTeX公式渲染）
        font: 覆盖模板字体（如 "SimSun"/"Microsoft YaHei"/"KaiTi"）
        font_size: 覆盖模板字号
        margin: 页边距 [上, 右, 下, 左]（厘米）
        heading_color: 标题颜色（如 "1F4E79" 或 "navy"）
        align: 全文对齐（left/center/right/justify）
        header: 页眉文字
        footer: 页脚文字（支持 {page} {numpages} 占位符）

    内容支持的标记：
    #/##/### 标题 | -/1. 列表 | > 引用 | ```代码``` | **粗体** *斜体* `代码` ~~删除线~~
    | 表格语法 | ---（分隔线） | [居中]行首标记 [右对齐] | {color:红色}文字{/color}
    """
    # 默认保存到桌面（path 为空或只含文件名时自动拼到桌面）
    path = _resolve_save_path(path, "未命名.docx")
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        # 获取模板配置
        tpl = WORD_TEMPLATES.get(template, WORD_TEMPLATES["default"]).copy()

        # 应用参数覆盖
        if font:
            tpl["font"] = font
        if font_size:
            tpl["font_size"] = font_size
        if margin:
            tpl["margin"] = margin
        if heading_color:
            tpl["heading_color"] = heading_color
        if header:
            tpl["header"] = header
        if footer:
            tpl["footer"] = footer

        doc = Document()

        # 设置页边距
        section = doc.sections[0]
        margins = tpl["margin"]
        if len(margins) == 4:
            section.top_margin = Cm(margins[0])
            section.right_margin = Cm(margins[1])
            section.bottom_margin = Cm(margins[2])
            section.left_margin = Cm(margins[3])

        # 设置页眉页脚
        _apply_header_template(section, tpl)
        _apply_footer_template(section, tpl)

        # 设置默认字体
        style = doc.styles["Normal"]
        font_obj = style.font
        font_obj.name = tpl["font"]
        font_obj.size = Pt(tpl["font_size"])
        style.element.rPr.rFonts.set(qn("w:eastAsia"), tpl["font"])

        # 学术模板：设置双倍行距（APA/学术规范）
        if tpl.get("line_spacing"):
            from docx.shared import Pt as _Pt2
            from docx.enum.text import WD_LINE_SPACING
            style.paragraph_format.line_spacing = tpl["line_spacing"]
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        # 全文对齐
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        default_align = align_map.get(align.lower(), None)

        # 解析标题
        doc_title = title if title else ""
        if not doc_title and content:
            first_line = content.split("\n")[0].strip()
            if first_line.startswith("# "):
                doc_title = first_line[2:].strip()
                content = "\n".join(content.split("\n")[1:])

        if doc_title:
            heading = doc.add_heading(doc_title, level=0)
            h_color = _parse_color(tpl["heading_color"])
            if h_color:
                for run in heading.runs:
                    run.font.color.rgb = h_color
                    run.font.name = tpl["font"]
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), tpl["font"])

        # 解析内容
        lines = content.split("\n")
        i = 0
        in_code_block = False
        code_lines = []
        # 学术论文状态变量
        in_references = False
        ref_counter = 0

        while i < len(lines):
            line = lines[i]

            # 代码块
            if line.strip().startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.0)
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # 空行
            if not line.strip():
                i += 1
                continue

            # 表格检测
            if line.strip().startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
                rows, end_idx = _parse_table(lines, i)
                if rows:
                    # 创建表格
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(table.rows[r_idx].cells):
                                cell = table.rows[r_idx].cells[c_idx]
                                cell.text = ""
                                p = cell.paragraphs[0]
                                _add_formatted_text(p, cell_text)
                                # 表头加粗+背景色
                                if r_idx == 0:
                                    for run in p.runs:
                                        run.bold = True
                                    _set_cell_shading(cell, "D9E2F3")
                    i = end_idx
                    continue

            # 分隔线
            if line.strip() in ("---", "***", "___"):
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("─" * 40)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                i += 1
                continue

            # 行首对齐标记：[居中] [右对齐] [左对齐]
            line_align = default_align
            align_match = re.match(r'^\[(居中|右对齐|左对齐|两端对齐)\]\s*(.*)', line)
            if align_match:
                align_text_map = {"居中": "center", "右对齐": "right", "左对齐": "left", "两端对齐": "justify"}
                line_align = align_map.get(align_text_map.get(align_match.group(1), ""), default_align)
                line = align_match.group(2)

            # 颜色标记：{color:红色}文字{/color}
            color_pattern = r'\{color:([^}]+)\}(.*?)\{/color\}'

            # ===== 学术论文段落特殊处理（academic/thesis 模板）=====
            is_academic = template in ("academic", "thesis")
            stripped_line = line.strip()

            if is_academic:
                # 检测学术段落标记：摘要/Abstract/关键词/Keywords/参考文献/References
                academic_match = re.match(
                    r'^(摘要|Abstract|关键词|Keywords|参考文献|References|引言|Introduction|结论|Conclusion|致谢|Acknowledgments?)\s*[:：]?\s*(.*)',
                    stripped_line, re.IGNORECASE
                )

                if academic_match:
                    section_name = academic_match.group(1)
                    section_content = academic_match.group(2)

                    # 摘要/Abstract：小字号、缩进、两端对齐
                    if section_name.lower() in ("摘要", "abstract"):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1.0)
                        p.paragraph_format.right_indent = Cm(1.0)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run = p.add_run(f"【{section_name}】")
                        run.bold = True
                        run.font.size = Pt(tpl["font_size"] - 1)
                        if section_content:
                            # 渲染行内公式
                            rendered = render_latex_in_text(section_content)
                            run2 = p.add_run(rendered)
                            run2.font.size = Pt(tpl["font_size"] - 1)
                            run2.italic = True
                        i += 1
                        continue

                    # 关键词/Keywords：加粗标签、分号分隔
                    elif section_name.lower() in ("关键词", "keywords"):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1.0)
                        run = p.add_run(f"【{section_name}】 ")
                        run.bold = True
                        if section_content:
                            run2 = p.add_run(section_content)
                            run2.font.size = Pt(tpl["font_size"] - 1)
                        i += 1
                        continue

                    # 参考文献/References：标题 + 后续条目自动编号
                    elif section_name.lower() in ("参考文献", "references"):
                        h = doc.add_heading(f"{section_name}", level=2)
                        hc = _parse_color(tpl["heading_color"])
                        if hc:
                            for run in h.runs:
                                run.font.color.rgb = hc
                        # 标记进入参考文献区域
                        in_references = True
                        ref_counter = 0
                        i += 1
                        continue

                    # 引言/结论等：作为一级标题
                    elif section_name.lower() in ("引言", "introduction", "结论", "conclusion"):
                        h = doc.add_heading(f"{section_name}", level=1)
                        hc = _parse_color(tpl["heading_color"])
                        if hc:
                            for run in h.runs:
                                run.font.color.rgb = hc
                        i += 1
                        continue

            # 参考文献条目自动编号（[1] [2] [3]...）
            if is_academic and in_references and stripped_line:
                # 跳过空行和已有编号的条目
                if not re.match(r'^\[\d+\]', stripped_line) and not stripped_line.startswith("#"):
                    ref_counter += 1
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.5)  # 悬挂缩进
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(f"[{ref_counter}] ")
                    run.bold = True
                    _add_formatted_text(p, stripped_line)
                    i += 1
                    continue

            # 标题（支持 # ~ #### 四级，符合 GB/T 7713.1-2025 论文标题层级规范）
            if line.startswith("#### "):
                h = doc.add_heading(line[5:].strip(), level=4)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            elif line.startswith("### "):
                h = doc.add_heading(line[4:].strip(), level=3)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            elif line.startswith("## "):
                h = doc.add_heading(line[3:].strip(), level=2)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            elif line.startswith("# "):
                h = doc.add_heading(line[2:].strip(), level=1)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            # 引用
            elif line.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                run = p.add_run(line[2:].strip())
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True
            # 无序列表
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_text(p, line[2:].strip())
            # 有序列表
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^(\d+)\.\s(.*)', line)
                if match:
                    p = doc.add_paragraph(style="List Number")
                    _add_formatted_text(p, match.group(2).strip())
            # 普通段落
            else:
                p = doc.add_paragraph()
                if line_align:
                    p.paragraph_format.alignment = line_align
                # 学术模板：两端对齐（默认）
                if is_academic and not line_align:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # 渲染行内 LaTeX 公式（$...$ → Unicode）
                processed_line = render_latex_in_text(line.strip()) if is_academic else line.strip()
                # 处理颜色标记
                if re.search(color_pattern, processed_line):
                    _add_colored_text(p, processed_line, color_pattern)
                else:
                    _add_formatted_text(p, processed_line)

            # 设置中文字体
            for p in doc.paragraphs[-1:]:
                for run in p.runs:
                    if not run.font.name:
                        run.font.name = tpl["font"]
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), tpl["font"])

            i += 1

        # 未关闭的代码块
        if in_code_block and code_lines:
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        # 保存
        full = Path(path).resolve()
        full.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(full))

        info = f"{_load_svg_icon('document')} Word 文档已生成：{full}\n"
        info += f"  标题：{doc_title or '（无标题）'}\n"
        info += f"  模板：{template}（字体:{tpl['font']} 字号:{tpl['font_size']}）\n"
        info += f"  段落：{len(doc.paragraphs)} | 表格：{len(doc.tables)}\n"
        info += f"  大小：{full.stat().st_size} 字节"
        return info
    except ImportError:
        return "错误：未安装 python-docx，请运行 pip install python-docx"
    except Exception as e:
        return f"错误：{e}"


def _add_colored_text(paragraph, text: str, color_pattern: str):
    """处理 {color:xxx}文字{/color} 标记

    迁移来源：tui_agent.py 行 5648-5663
    """
    last_end = 0
    for match in re.finditer(color_pattern, text):
        # 前面的普通文本
        if match.start() > last_end:
            _add_formatted_text(paragraph, text[last_end:match.start()])
        color_str = match.group(1)
        colored_text = match.group(2)
        run = paragraph.add_run(colored_text)
        color = _parse_color(color_str)
        if color:
            run.font.color.rgb = color
        last_end = match.end()
    if last_end < len(text):
        _add_formatted_text(paragraph, text[last_end:])


# ====== Excel 文档生成 ======
def generate_excel(path: str, sheets: list, template: str = "default",
                   charts: list = None, formulas: list = None) -> str:
    """生成 Excel 文档（.xlsx）- 增强版：支持图表、公式

    迁移来源：tui_agent.py 行 5667-5854

    Args:
        path: 保存路径（为空或只含文件名时自动保存到桌面）
        sheets: 工作表列表，每个工作表是 dict：
            {
                "name": "Sheet1",          # 工作表名（可选，默认 Sheet1）
                "data": [                   # 数据（二维数组）
                    ["姓名", "年龄", "成绩"],
                    ["张三", 18, 95.5],
                    ["李四", 19, 88],
                ],
                "header": true,             # 是否有表头（默认 true）
            }
        template: 格式模板（default/report/data/financial）
        charts: 图表列表（可选），每个图表是 dict：
            {
                "type": "bar",              # bar(柱状图) / line(折线图) / pie(饼图)
                "title": "成绩对比",         # 图表标题
                "sheet": "Sheet1",          # 数据所在工作表名
                "data_range": "A1:C4",      # 数据范围（含表头）
                "categories_col": "A",      # 分类轴列（如姓名列）
                "values_cols": ["B", "C"],  # 值轴列（可多列）
                "position": "E2",           # 图表放置位置（单元格）
            }
        formulas: 公式列表（可选），每个公式是 dict：
            {
                "sheet": "Sheet1",          # 工作表名
                "cell": "D2",              # 写入单元格
                "formula": "=AVERAGE(C2:C4)", # 公式
            }
    """
    # 默认保存到桌面（path 为空或只含文件名时自动拼到桌面）
    path = _resolve_save_path(path, "未命名.xlsx")
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "错误：缺少 openpyxl 库，请运行 pip install openpyxl"

    # 模板配置
    templates = {
        "default": {"header_bg": "4472C4", "header_fg": "FFFFFF", "alt_row": "F2F2F2", "border": True},
        "report": {"header_bg": "1F4E79", "header_fg": "FFFFFF", "alt_row": "D6E4F0", "border": True},
        "data": {"header_bg": "70AD47", "header_fg": "FFFFFF", "alt_row": "E2EFDA", "border": True},
        "financial": {"header_bg": "5B5B5B", "header_fg": "FFFFFF", "alt_row": "FFF2CC", "border": True},
    }
    tpl = templates.get(template, templates["default"])

    wb = Workbook()
    wb.remove(wb.active)

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    ) if tpl["border"] else None

    # 记录工作表引用，供图表/公式使用
    ws_map = {}

    for i, sheet_def in enumerate(sheets):
        sheet_name = sheet_def.get("name", f"Sheet{i+1}")
        ws = wb.create_sheet(title=sheet_name)
        ws_map[sheet_name] = ws
        data = sheet_def.get("data", [])
        has_header = sheet_def.get("header", True)

        for row_idx, row_data in enumerate(data, 1):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                if has_header and row_idx == 1:
                    cell.font = Font(bold=True, color=tpl["header_fg"], name="Microsoft YaHei", size=11)
                    cell.fill = PatternFill(start_color=tpl["header_bg"], end_color=tpl["header_bg"], fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                else:
                    cell.font = Font(name="Microsoft YaHei", size=10)
                    if tpl["alt_row"] and row_idx > 1 and (row_idx % 2 == 0):
                        cell.fill = PatternFill(start_color=tpl["alt_row"], end_color=tpl["alt_row"], fill_type="solid")
                    cell.alignment = Alignment(vertical="center")
                if thin_border:
                    cell.border = thin_border

        # 自动列宽
        for col_idx in range(1, ws.max_column + 1):
            max_len = 0
            for row_idx in range(1, ws.max_row + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                val = str(cell.value) if cell.value else ""
                length = sum(2 if ord(c) > 127 else 1 for c in val)
                if length > max_len:
                    max_len = length
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 50)

    # 写入公式
    formula_count = 0
    if formulas:
        for f in formulas:
            ws = ws_map.get(f.get("sheet", sheets[0].get("name", "Sheet1")))
            if ws:
                cell_ref = f.get("cell", "")
                formula = f.get("formula", "")
                if cell_ref and formula:
                    ws[cell_ref] = formula
                    # 公式单元格样式
                    ws[cell_ref].font = Font(name="Microsoft YaHei", size=10, bold=True, color=tpl["header_bg"])
                    if thin_border:
                        ws[cell_ref].border = thin_border
                    formula_count += 1

    # 生成图表
    chart_count = 0
    if charts:
        try:
            from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        except ImportError:
            pass
        else:
            for ch in charts:
                ws = ws_map.get(ch.get("sheet", sheets[0].get("name", "Sheet1")))
                if not ws:
                    continue
                chart_type = ch.get("type", "bar").lower()
                chart_title = ch.get("title", "")
                categories_col = ch.get("categories_col", "A")
                values_cols = ch.get("values_cols", ["B"])
                position = ch.get("position", "E2")

                # 解析列号为数字
                def col_to_num(col_str):
                    """A→1, B→2, ..."""
                    num = 0
                    for c in col_str.upper():
                        num = num * 26 + (ord(c) - ord('A') + 1)
                    return num

                cat_col_num = col_to_num(categories_col)
                val_col_nums = [col_to_num(c) for c in values_cols]

                # 数据范围：从第1行（表头）到最后一行
                max_row = ws.max_row
                min_row = 2 if ws.cell(1, 1).value else 1  # 跳过表头

                if chart_type == "bar":
                    chart = BarChart()
                    chart.type = "col"
                    chart.style = 10
                elif chart_type == "line":
                    chart = LineChart()
                    chart.style = 12
                elif chart_type == "pie":
                    chart = PieChart()
                    chart.style = 10
                else:
                    continue

                chart.title = chart_title

                # 添加数据（值列）
                for vcn in val_col_nums:
                    data_ref = Reference(ws, min_col=vcn, min_row=1, max_row=max_row)
                    chart.add_data(data_ref, titles_from_data=True)

                # 设置分类轴（X轴标签）
                if chart_type != "pie":
                    cat_ref = Reference(ws, min_col=cat_col_num, min_row=min_row, max_row=max_row)
                    chart.set_categories(cat_ref)
                else:
                    # 饼图分类
                    cat_ref = Reference(ws, min_col=cat_col_num, min_row=min_row, max_row=max_row)
                    chart.set_categories(cat_ref)

                # 图表尺寸
                chart.width = 18
                chart.height = 12

                # 放置图表
                ws.add_chart(chart, position)
                chart_count += 1

    wb.save(path)
    result = f"{_load_svg_icon('document')} Excel 已生成：{path}（{len(sheets)} 个工作表"
    if chart_count:
        result += f"，{chart_count} 个图表"
    if formula_count:
        result += f"，{formula_count} 个公式"
    result += "）"
    return result


# ====== PDF 文档生成 ======
def _latex_to_image(latex: str, font_size: int = 16, dpi: int = 200) -> object:
    """将 LaTeX 公式渲染为 PNG 图片（reportlab Image 对象）

    迁移来源：tui_agent.py 行 5858-5890

    用于 PDF 学术模板的独立公式 $$...$$ 渲染，比纯 Unicode 更标准。
    失败时返回 None，调用方降级为 Unicode 文本。

    Args:
        latex: LaTeX 公式（不含 $$ 包裹）
        font_size: 字号
        dpi: 分辨率
    """
    try:
        import matplotlib
        matplotlib.use("Agg")  # 非交互后端
        import matplotlib.pyplot as plt
        import io, os, tempfile

        # 去除首尾 $ 符号
        formula = latex.strip().strip("$").strip()

        # 创建图片：用 mathtext 渲染公式
        fig = plt.figure(figsize=(0.01, 0.01), dpi=dpi)
        fig.text(0, 0, f"${formula}$", fontsize=font_size, color="black")
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                    pad_inches=0.05, transparent=False, facecolor="white")
        plt.close(fig)
        buf.seek(0)

        # 保存为临时文件（reportlab Image 需要文件路径或 BytesIO）
        return buf
    except Exception:
        return None


def generate_pdf(path: str, content: str, title: str = "", template: str = "default") -> str:
    """生成 PDF 文档

    迁移来源：tui_agent.py 行 5893-6184

    Args:
        path: 保存路径（为空或只含文件名时自动保存到桌面）
        content: 文档内容（支持 Markdown 风格标记）
        title: 文档标题（可选）
        template: 格式模板（default/report/contract/resume/letter/technical/academic）
                    academic=学术论文（1.5倍行距/摘要/关键词/参考文献自动编号/LaTeX公式渲染）

    内容支持：
    # 一级标题 / ## 二级标题 / ### 三级标题
    - 无序列表 / 1. 有序列表
    > 引用 / ```代码块``` / **粗体** *斜体*
    ---（分隔线）/ 普通段落 / |表格语法|
    """
    # 默认保存到桌面（path 为空或只含文件名时自动拼到桌面）
    path = _resolve_save_path(path, "未命名.pdf")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm, mm
        from reportlab.lib.colors import HexColor, black, white, grey
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, HRFlowable, ListFlowable, ListItem, Preformatted
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return "错误：缺少 reportlab 库，请运行 pip install reportlab"

    # 注册中文字体
    font_name = "Helvetica"
    for font_path, font_id in [
        (r"C:\Windows\Fonts\msyh.ttc", "MSYH"),
        (r"C:\Windows\Fonts\msyhbd.ttc", "MSYHBD"),
        (r"C:\Windows\Fonts\simsun.ttc", "SimSun"),
        (r"C:\Windows\Fonts\simhei.ttf", "SimHei"),
    ]:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont(font_id, font_path))
                font_name = font_id
                break
            except Exception:
                continue

    # 模板配置
    templates = {
        "default": {"title_size": 20, "h1_size": 16, "h2_size": 14, "h3_size": 12, "body_size": 10, "color": "#1F4E79", "margin": [2.54, 2.54, 2.54, 2.54]},
        "report": {"title_size": 22, "h1_size": 16, "h2_size": 14, "h3_size": 12, "body_size": 10, "color": "#1F4E79", "margin": [3, 2.5, 3, 2.5]},
        "contract": {"title_size": 18, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 10, "color": "#333333", "margin": [2.54, 2.54, 2.54, 2.54]},
        "resume": {"title_size": 24, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 10, "color": "#2E5C8A", "margin": [2, 2.5, 2, 2.5]},
        "letter": {"title_size": 18, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 11, "color": "#333333", "margin": [3, 3, 3, 3]},
        "technical": {"title_size": 20, "h1_size": 16, "h2_size": 14, "h3_size": 12, "body_size": 9, "color": "#0B5394", "margin": [2, 2, 2, 2]},
        "academic": {"title_size": 18, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 10, "color": "#000000", "margin": [2.54, 2.54, 2.54, 2.54], "line_spacing": 1.5},
    }
    tpl = templates.get(template, templates["default"])
    heading_color = HexColor(tpl["color"])
    m = tpl["margin"]
    is_academic = template == "academic"
    line_spacing_factor = tpl.get("line_spacing", 1.5)

    doc = SimpleDocTemplate(path, pagesize=A4,
                           topMargin=m[0]*cm, rightMargin=m[1]*cm,
                           bottomMargin=m[2]*cm, leftMargin=m[3]*cm)

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                 fontName=font_name, fontSize=tpl["title_size"],
                                 textColor=heading_color, alignment=TA_CENTER, spaceAfter=20)
    style_h1 = ParagraphStyle('CustomH1', parent=styles['Heading1'],
                              fontName=font_name, fontSize=tpl["h1_size"],
                              textColor=heading_color, spaceBefore=16, spaceAfter=8)
    style_h2 = ParagraphStyle('CustomH2', parent=styles['Heading2'],
                              fontName=font_name, fontSize=tpl["h2_size"],
                              textColor=heading_color, spaceBefore=12, spaceAfter=6)
    style_h3 = ParagraphStyle('CustomH3', parent=styles['Heading3'],
                              fontName=font_name, fontSize=tpl["h3_size"],
                              textColor=heading_color, spaceBefore=10, spaceAfter=4)
    # 四级标题（GB/T 7713.1-2025 允许四级，h4_size 缺省时用 h3_size-1）
    style_h4 = ParagraphStyle('CustomH4', parent=styles['Heading4'],
                              fontName=font_name, fontSize=tpl.get("h4_size", tpl["h3_size"] - 1),
                              textColor=heading_color, spaceBefore=8, spaceAfter=3)
    style_body = ParagraphStyle('CustomBody', parent=styles['Normal'],
                                fontName=font_name, fontSize=tpl["body_size"],
                                leading=tpl["body_size"]*line_spacing_factor, spaceAfter=6, alignment=TA_JUSTIFY)
    style_quote = ParagraphStyle('CustomQuote', parent=style_body,
                                 leftIndent=20, textColor=grey, fontSize=tpl["body_size"]-1)
    style_code = ParagraphStyle('CustomCode', parent=styles['Code'],
                                fontName='Courier', fontSize=tpl["body_size"]-1,
                                backColor=HexColor("#F5F5F5"), leftIndent=10, rightIndent=10,
                                spaceBefore=4, spaceAfter=4)
    # 学术论文专用样式
    style_abstract = ParagraphStyle('AcademicAbstract', parent=style_body,
                                    fontSize=tpl["body_size"]-1, leading=(tpl["body_size"]-1)*1.5,
                                    leftIndent=15, rightIndent=15, textColor=HexColor("#333333"),
                                    spaceBefore=8, spaceAfter=8, alignment=TA_JUSTIFY)
    style_keywords = ParagraphStyle('AcademicKeywords', parent=style_body,
                                    fontSize=tpl["body_size"]-1, leftIndent=15,
                                    spaceBefore=4, spaceAfter=8)
    style_reference = ParagraphStyle('AcademicReference', parent=style_body,
                                     fontSize=tpl["body_size"]-1, leading=(tpl["body_size"]-1)*1.3,
                                     leftIndent=20, firstLineIndent=-20,
                                     spaceBefore=2, spaceAfter=2, alignment=TA_JUSTIFY)

    story = []

    # 标题
    if title:
        story.append(Paragraph(title, style_title))
        story.append(Spacer(1, 10))

    # 解析内容
    lines = content.split('\n')
    in_code = False
    code_lines = []
    in_table = False
    table_rows = []
    # 学术论文状态变量
    in_references = False
    ref_counter = 0

    for line in lines:
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            if in_code:
                story.append(Preformatted('\n'.join(code_lines), style_code))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        # 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # 跳过分隔行 |---|---|
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # 表格结束
            if table_rows:
                t = Table(table_rows)
                t.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, -1), font_name, tpl["body_size"]),
                    ('BACKGROUND', (0, 0), (-1, 0), heading_color),
                    ('TEXTCOLOR', (0, 0), (-1, 0), white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('GRID', (0, 0), (-1, -1), 0.5, grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, HexColor("#F5F5F5")]),
                ]))
                story.append(t)
                story.append(Spacer(1, 8))
                table_rows = []
            in_table = False

        # 分隔线
        if stripped == '---' or stripped == '***':
            story.append(HRFlowable(width="100%", thickness=1, color=heading_color))
            story.append(Spacer(1, 6))
            continue

        # ===== 学术论文段落特殊处理（academic 模板）=====
        if is_academic:
            academic_match = re.match(
                r'^(摘要|Abstract|关键词|Keywords|参考文献|References|引言|Introduction|结论|Conclusion|致谢|Acknowledgments?)\s*[:：]?\s*(.*)',
                stripped, re.IGNORECASE
            )
            if academic_match:
                section_name = academic_match.group(1)
                section_content = academic_match.group(2)

                # 摘要/Abstract
                if section_name.lower() in ("摘要", "abstract"):
                    abstract_text = f"<b>【{section_name}】</b> "
                    if section_content:
                        rendered = render_latex_in_text(section_content)
                        abstract_text += rendered
                    story.append(Paragraph(abstract_text, style_abstract))
                    continue

                # 关键词/Keywords
                elif section_name.lower() in ("关键词", "keywords"):
                    kw_text = f"<b>【{section_name}】</b> {section_content}"
                    story.append(Paragraph(kw_text, style_keywords))
                    continue

                # 参考文献/References
                elif section_name.lower() in ("参考文献", "references"):
                    story.append(Paragraph(f"<b>{section_name}</b>", style_h2))
                    in_references = True
                    ref_counter = 0
                    continue

                # 引言/结论
                elif section_name.lower() in ("引言", "introduction", "结论", "conclusion"):
                    story.append(Paragraph(f"<b>{section_name}</b>", style_h1))
                    in_references = False  # 退出参考文献区域
                    continue

        # 参考文献条目自动编号
        if is_academic and in_references and stripped and not stripped.startswith('#'):
            if not re.match(r'^\[\d+\]', stripped):
                ref_counter += 1
                ref_text = f"[{ref_counter}] {stripped}"
                # 处理粗体斜体
                ref_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', ref_text)
                ref_text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', ref_text)
                story.append(Paragraph(ref_text, style_reference))
                continue

        # 标题（支持 # ~ #### 四级，符合 GB/T 7713.1-2025 论文标题层级规范）
        if stripped.startswith('#### '):
            story.append(Paragraph(stripped[5:], style_h4))
        elif stripped.startswith('### '):
            story.append(Paragraph(stripped[4:], style_h3))
        elif stripped.startswith('## '):
            story.append(Paragraph(stripped[3:], style_h2))
        elif stripped.startswith('# '):
            story.append(Paragraph(stripped[2:], style_h1))
        elif is_academic and stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            # 学术模板：独立公式 $$...$$ 用 matplotlib 渲染为图片（比 Unicode 更标准）
            formula = stripped[2:-2].strip()
            img_buf = _latex_to_image(formula)
            if img_buf is not None:
                try:
                    from reportlab.lib.utils import ImageReader
                    from reportlab.platypus import Image as RLImage
                    img_reader = ImageReader(img_buf)
                    iw, ih = img_reader.getSize()
                    # 按宽度缩放（最大宽度 400pt）
                    max_w = 400
                    if iw > max_w:
                        ratio = max_w / iw
                        iw, ih = max_w, ih * ratio
                    img_buf.seek(0)
                    story.append(RLImage(img_buf, width=iw, height=ih))
                    story.append(Spacer(1, 6))
                except Exception:
                    # 图片插入失败 → 降级为 Unicode
                    story.append(Paragraph(_latex_to_unicode(formula), style_body))
            else:
                # matplotlib 渲染失败 → 降级为 Unicode
                story.append(Paragraph(_latex_to_unicode(formula), style_body))
        elif stripped.startswith('> '):
            story.append(Paragraph(stripped[2:], style_quote))
        elif stripped.startswith('- '):
            item = Paragraph(stripped[2:], style_body)
            story.append(ListFlowable([ListItem(item)], bulletType='bullet'))
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            item = Paragraph(text, style_body)
            story.append(ListFlowable([ListItem(item)], bulletType='1'))
        elif stripped:
            # 处理粗体斜体（**粗体** *斜体* → <b>粗体</b> <i>斜体</i>）
            text = stripped
            # 学术模板：先渲染 LaTeX 公式为 Unicode
            if is_academic:
                text = render_latex_in_text(text)
            # 先处理 **粗体**
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            # 再处理 *斜体*
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            # 转义 XML 特殊字符（但保留已生成的标签）
            text = text.replace('&', '&amp;').replace('<b>', '<b>').replace('</b>', '</b>')
            story.append(Paragraph(text, style_body))
        else:
            story.append(Spacer(1, 6))

    # 处理未关闭的表格
    if table_rows:
        t = Table(table_rows)
        t.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), font_name, tpl["body_size"]),
            ('BACKGROUND', (0, 0), (-1, 0), heading_color),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, grey),
        ]))
        story.append(t)

    doc.build(story)
    return f"{_load_svg_icon('document')} PDF 已生成：{path}（模板：{template}）"
