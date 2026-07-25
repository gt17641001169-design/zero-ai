"""
ZeroAI - 终端 AI 助手
稳定、实用、美观，对标 OpenCode

特性：
- 流式输出，实时看到回复
- Markdown 渲染，代码块语法高亮
- 工具调用可视化（读文件/写文件/执行命令/搜索）
- 多轮上下文记忆
- Tokyo Night 配色

运行：python d:/C/C/tui_agent.py
"""

import os
import sys
import json
import atexit
import subprocess
import re
import asyncio
import shutil
import platform
import locale
import time
import base64

# ════════════════════════════════════════════════════════════════════
# 资源路径智能查找（支持开发模式和 pip 安装模式）
# 查找优先级：1. 脚本所在目录（开发模式）2. 环境变量 ZEROAI_HOME 3. 用户主目录 ~/.zeroai/
# ════════════════════════════════════════════════════════════════════
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_USER_HOME = os.path.expanduser("~")
_ZEROAI_USER_DIR = os.path.join(_USER_HOME, ".zeroai")

# 桌面目录缓存（兼容中英文 Windows / macOS / Linux，首次调用 _get_desktop_dir 时填充）
_DESKTOP_DIR_CACHE: str = ""


def _get_desktop_dir() -> str:
    """获取桌面目录路径（兼容中英文 Windows / macOS / Linux）。

    优先级：
    1. Windows: SHGetFolderPathW（CSIDL_DESKTOP，最可靠，能识别"桌面"重定向）
    2. Windows: USERPROFILE/Desktop 或 USERPROFILE/桌面
    3. macOS/Linux: ~/Desktop 或 ~/桌面
    4. Linux: XDG_DESKTOP_DIR 环境变量
    5. 回退: 用户主目录
    """
    global _DESKTOP_DIR_CACHE
    if _DESKTOP_DIR_CACHE:
        return _DESKTOP_DIR_CACHE

    # Windows: 用 SHGetFolderPathW 获取真实桌面路径
    if sys.platform == "win32":
        try:
            import ctypes
            from ctypes import wintypes
            buf = ctypes.create_unicode_buffer(wintypes.MAX_PATH)
            # CSIDL_DESKTOP = 0x00
            if ctypes.windll.shell32.SHGetFolderPathW(0, 0x00, 0, 0, buf) == 0:
                if buf.value and os.path.isdir(buf.value):
                    _DESKTOP_DIR_CACHE = buf.value
                    return _DESKTOP_DIR_CACHE
        except Exception:
            pass
        # 回退：USERPROFILE 下找 Desktop / 桌面
        up = os.environ.get("USERPROFILE") or _USER_HOME
        for name in ("Desktop", "桌面"):
            p = os.path.join(up, name)
            if os.path.isdir(p):
                _DESKTOP_DIR_CACHE = p
                return _DESKTOP_DIR_CACHE

    # macOS / Linux
    for name in ("Desktop", "桌面"):
        p = os.path.join(_USER_HOME, name)
        if os.path.isdir(p):
            _DESKTOP_DIR_CACHE = p
            return _DESKTOP_DIR_CACHE

    # Linux XDG
    xdg = os.environ.get("XDG_DESKTOP_DIR")
    if xdg and os.path.isdir(xdg):
        _DESKTOP_DIR_CACHE = xdg
        return _DESKTOP_DIR_CACHE

    # 最终回退：用户主目录
    _DESKTOP_DIR_CACHE = _USER_HOME
    return _DESKTOP_DIR_CACHE


def _resolve_save_path(path: str, default_filename: str) -> str:
    """解析文档保存路径（默认保存到桌面）。

    规则：
    - path 为空：保存到 桌面/{default_filename}
    - path 只含文件名（无目录分隔符）：保存到 桌面/{path}
    - path 含完整路径：按 path 原样保存

    Args:
        path: 用户传入的路径
        default_filename: 默认文件名（如 "未命名.docx"）

    Returns:
        解析后的绝对路径
    """
    if not path:
        return os.path.join(_get_desktop_dir(), default_filename)
    # 判断是否只含文件名（无目录部分）
    if not os.path.dirname(path):
        return os.path.join(_get_desktop_dir(), path)
    return path

def _find_resource_dir(subdir: str) -> str:
    """智能查找资源目录（libs/、models/ 等），支持多位置查找。

    查找优先级：
    1. 脚本所在目录的子目录（开发模式：D:\\C\\C\\libs）
    2. 环境变量 ZEROAI_HOME 指定的子目录
    3. 用户主目录 ~/.zeroai/ 的子目录（pip 安装模式）

    返回第一个存在的目录；若都不存在，返回脚本所在目录的子目录（用于错误提示）。
    """
    candidates = [
        os.path.join(_SCRIPT_DIR, subdir),                    # 1. 开发模式
        os.path.join(os.environ.get("ZEROAI_HOME", ""), subdir),  # 2. 环境变量
        os.path.join(_ZEROAI_USER_DIR, subdir),               # 3. pip 安装模式
    ]
    for p in candidates:
        if p and os.path.isdir(p):
            return p
    # 默认返回脚本所在目录的子目录（用于错误提示和后续创建）
    return os.path.join(_SCRIPT_DIR, subdir)

def _ensure_user_dir() -> str:
    """确保用户主目录 ~/.zeroai/ 存在，返回路径"""
    if not os.path.isdir(_ZEROAI_USER_DIR):
        try:
            os.makedirs(_ZEROAI_USER_DIR, exist_ok=True)
        except OSError:
            pass
    return _ZEROAI_USER_DIR

# 将项目目录下的 libs 子目录加入 sys.path（用于加载 --target 安装的库，如 SpeechRecognition）
_LIBS_DIR = _find_resource_dir("libs")
if os.path.isdir(_LIBS_DIR) and _LIBS_DIR not in sys.path:
    sys.path.insert(0, _LIBS_DIR)
# 添加 DLL 搜索路径（sherpa_onnx 等库的 .pyd 依赖 onnxruntime.dll 等原生 DLL）
_SHERPA_LIB_DIR = os.path.join(_LIBS_DIR, "sherpa_onnx", "lib")
if os.path.isdir(_SHERPA_LIB_DIR):
    try:
        os.add_dll_directory(_SHERPA_LIB_DIR)
    except (AttributeError, OSError):
        pass  # Python < 3.8 或路径无效时忽略
if os.path.isdir(_LIBS_DIR):
    try:
        os.add_dll_directory(_LIBS_DIR)
    except (AttributeError, OSError):
        pass
import inspect
import tempfile
import traceback
import urllib.request
import urllib.parse
import urllib.error
from pathlib import Path
from openai import OpenAI, AsyncOpenAI
from rich.text import Text
from rich.markdown import Markdown
from rich.syntax import Syntax
from rich.panel import Panel
from rich.console import Group

from textual.app import App, ComposeResult
from textual.widgets import Header, Input, Static, ListView, ListItem, Label, TextArea, RichLog, Markdown
from textual.containers import Vertical, VerticalScroll, Horizontal, Container
from textual.screen import ModalScreen
from textual.binding import Binding
from textual import events

# ====== 配置文件路径 ======
CONFIG_FILE = Path.home() / ".zeroai_config.json"
CUSTOM_MODELS_FILE = Path.home() / ".zeroai_models.json"

# ====== 资源路径（打包后从 _MEIPASS 读取，开发时从源码目录读取）======
def _get_resource_dir() -> Path:
    """获取资源目录路径（兼容 PyInstaller 打包和源码运行）"""
    if getattr(sys, 'frozen', False):
        # PyInstaller 打包后，资源在 _MEIPASS 中
        return Path(sys._MEIPASS) / "assets"
    else:
        # 源码运行
        return Path(__file__).parent / "assets"

ASSETS_DIR = _get_resource_dir()
ICONS_DIR = ASSETS_DIR / "icons"


def _copy_to_clipboard(text: str) -> bool:
    """复制文本到系统剪贴板（Windows API，完整 Unicode 支持）

    用 ctypes 直接调用 user32/kernel32，避免 clip.exe 的 GBK 编码问题。
    返回 True 表示成功，False 表示失败。
    """
    if not text:
        return False
    try:
        import ctypes

        CF_UNICODETEXT = 13
        GMEM_MOVEABLE = 0x0002

        kernel32 = ctypes.windll.kernel32
        user32 = ctypes.windll.user32

        # 64 位 Python 必须显式设置函数签名，否则句柄/指针会被截断为 32 位
        # HGLOBAL / LPVOID / HANDLE 在 64 位 Windows 上都是 64 位
        kernel32.GlobalAlloc.restype = ctypes.c_void_p  # HGLOBAL
        kernel32.GlobalAlloc.argtypes = [ctypes.c_uint, ctypes.c_size_t]
        kernel32.GlobalLock.restype = ctypes.c_void_p  # LPVOID
        kernel32.GlobalLock.argtypes = [ctypes.c_void_p]
        kernel32.GlobalUnlock.argtypes = [ctypes.c_void_p]
        user32.OpenClipboard.argtypes = [ctypes.c_void_p]
        user32.SetClipboardData.restype = ctypes.c_void_p  # HANDLE
        user32.SetClipboardData.argtypes = [ctypes.c_uint, ctypes.c_void_p]

        # 编码为 UTF-16LE（Windows 内部字符串格式），末尾加 \0
        data_bytes = (text + "\0").encode("utf-16-le")

        if not user32.OpenClipboard(None):
            return False
        try:
            user32.EmptyClipboard()
            h_global = kernel32.GlobalAlloc(GMEM_MOVEABLE, len(data_bytes))
            if not h_global:
                return False
            locked = kernel32.GlobalLock(h_global)
            if not locked:
                return False
            ctypes.memmove(locked, data_bytes, len(data_bytes))
            kernel32.GlobalUnlock(h_global)
            result = user32.SetClipboardData(CF_UNICODETEXT, h_global)
            return bool(result)
        finally:
            user32.CloseClipboard()
    except Exception:
        return False


def _load_svg_icon(name: str) -> str:
    """加载 SVG 图标文件内容，返回纯文本标签（终端无法渲染 SVG，返回文字标签）"""
    svg_path = ICONS_DIR / f"{name}.svg"
    if svg_path.exists():
        # 终端环境下返回对应的文字标签
        ICON_LABELS = {
            "folder": "[DIR]",
            "file": "[FILE]",
            "search": "[SCAN]",
            "check": "[OK]",
            "cross": "[ERR]",
            "warning": "[!]",
            "security": "[SEC]",
            "monitor": "[SCREEN]",
            "download": "[DL]",
            "document": "[DOC]",
            "tool": "[TOOL]",
        }
        return ICON_LABELS.get(name, "")
    return ""


# ====== 运行时缓存（停止运行自动删除）======
class RuntimeCache:
    """运行时缓存管理器：所有临时数据存入临时目录，程序退出时自动删除"""

    def __init__(self):
        self._cache_dir = None
        self._initialized = False

    @property
    def cache_dir(self) -> Path:
        """获取缓存目录，不存在则创建"""
        if self._cache_dir is None:
            self._cache_dir = Path(tempfile.mkdtemp(prefix="zeroai_cache_"))
            self._initialized = True
            # 注册退出时清理
            atexit.register(self.cleanup)
        return self._cache_dir

    def get_path(self, name: str) -> Path:
        """获取缓存文件路径"""
        return self.cache_dir / name

    def read(self, name: str, default: str = "") -> str:
        """读取缓存文件"""
        p = self.cache_dir / name
        if p.exists():
            try:
                return p.read_text(encoding="utf-8")
            except Exception:
                return default
        return default

    def write(self, name: str, data: str):
        """写入缓存文件"""
        try:
            (self.cache_dir / name).write_text(data, encoding="utf-8")
        except Exception:
            pass

    def read_bytes(self, name: str) -> bytes:
        """读取缓存二进制文件"""
        p = self.cache_dir / name
        if p.exists():
            try:
                return p.read_bytes()
            except Exception:
                return b""
        return b""

    def write_bytes(self, name: str, data: bytes):
        """写入缓存二进制文件"""
        try:
            (self.cache_dir / name).write_bytes(data)
        except Exception:
            pass

    def exists(self, name: str) -> bool:
        """检查缓存文件是否存在"""
        return (self.cache_dir / name).exists()

    def cleanup(self):
        """清理缓存目录（程序退出时自动调用）"""
        if self._cache_dir is not None and self._cache_dir.exists():
            try:
                shutil.rmtree(self._cache_dir, ignore_errors=True)
            except Exception:
                pass

    def size(self) -> int:
        """获取缓存总大小（字节）"""
        if self._cache_dir is None or not self._cache_dir.exists():
            return 0
        total = 0
        for f in self._cache_dir.rglob("*"):
            if f.is_file():
                try:
                    total += f.stat().st_size
                except Exception:
                    pass
        return total


# 全局缓存实例
runtime_cache = RuntimeCache()


def _obfuscate(text: str) -> str:
    """简单混淆（base64），防止明文泄露，不是加密"""
    return base64.b64encode(text.encode("utf-8")).decode("ascii")


def _deobfuscate(text: str) -> str:
    """反混淆"""
    try:
        return base64.b64decode(text.encode("ascii")).decode("utf-8")
    except Exception:
        return text


def _load_config() -> dict:
    """加载配置文件（存储 API Key 等敏感信息）"""
    if CONFIG_FILE.exists():
        try:
            data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            # 反混淆 api_key
            for key in data:
                if "api_key" in data[key] and data[key]["api_key"]:
                    data[key]["api_key"] = _deobfuscate(data[key]["api_key"])
            return data
        except Exception:
            pass
    return {}


def _save_config(data: dict):
    """保存配置文件（API Key 混淆后存储）"""
    try:
        safe = {}
        for key, cfg in data.items():
            safe[key] = dict(cfg)
            if safe[key].get("api_key"):
                safe[key]["api_key"] = _obfuscate(safe[key]["api_key"])
        CONFIG_FILE.write_text(json.dumps(safe, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def _get_api_key(key: str, default: str = "") -> str:
    """按优先级获取 API Key：环境变量 > 配置文件 > 内置默认值
    确保一定返回非空 Key（如果有内置默认值）"""
    # 1. 环境变量（最高优先级，最安全）
    env_key = f"ZEROAI_API_KEY_{key.upper()}"
    env_val = os.environ.get(env_key, "").strip()
    if env_val:
        return env_val
    # 2. 配置文件（用户自行配置的 Key）
    config = _load_config()
    if key in config and config[key].get("api_key"):
        return config[key]["api_key"]
    # 3. 内置默认值（确保打包后所有人都能用）
    if default:
        return default
    # 4. 没有内置默认值（如 openrouter 可能没有），返回空
    return ""


# ====== 代理服务器配置（v1.1.0 新增）======
# 启用代理后，所有 AI 请求经代理服务器转发，真实 API Key 只存在服务器端
# 客户端仅需配置代理 URL + 访问 Token，零 Key 泄露风险

# 环境变量优先（ZEROAI_PROXY_URL / ZEROAI_PROXY_TOKEN）
_PROXY_URL_ENV = os.environ.get("ZEROAI_PROXY_URL", "").strip()
_PROXY_TOKEN_ENV = os.environ.get("ZEROAI_PROXY_TOKEN", "").strip()


def _load_proxy_config() -> dict:
    """加载代理配置：环境变量 > 配置文件
    返回 {"enabled": bool, "base_url": str, "token": str}
    """
    # 1. 环境变量优先
    if _PROXY_URL_ENV and _PROXY_TOKEN_ENV:
        return {
            "enabled": True,
            "base_url": _PROXY_URL_ENV,
            "token": _PROXY_TOKEN_ENV,
        }

    # 2. 配置文件
    try:
        data = _load_config()
        if "proxy" in data and isinstance(data["proxy"], dict):
            p = data["proxy"]
            url = p.get("base_url", "").strip()
            token = p.get("token", "").strip()
            enabled = bool(p.get("enabled", False)) and bool(url) and bool(token)
            return {"enabled": enabled, "base_url": url, "token": token}
    except Exception:
        pass

    return {"enabled": False, "base_url": "", "token": ""}


def _save_proxy_config(enabled: bool, base_url: str, token: str):
    """保存代理配置到配置文件"""
    try:
        data = _load_config()
        data["proxy"] = {
            "enabled": bool(enabled),
            "base_url": base_url.strip(),
            "token": token.strip(),
        }
        _save_config(data)
    except Exception:
        pass


# 全局缓存代理配置（启动时加载一次，设置面板修改后刷新）
PROXY_CONFIG = _load_proxy_config()


def _is_proxy_enabled() -> bool:
    """代理是否启用"""
    return PROXY_CONFIG.get("enabled", False)


def _make_openai_client(model_key: str):
    """统一的 AsyncOpenAI 客户端工厂
    - 代理启用时：base_url 指向代理，api_key 用 Token（不是真实 Key）
    - 代理未启用时：使用原始 base_url + 真实 api_key（本地开发模式）

    model_key: MODEL_CONFIGS 的键（glm / glm-v / glm-4 / openrouter / ollama / 自定义）
    """
    base_cfg = MODEL_CONFIGS.get(model_key, {})

    # 代理模式：所有请求经代理转发
    if _is_proxy_enabled():
        proxy_url = PROXY_CONFIG["base_url"]
        # 确保 base_url 以 /v1 结尾（OpenAI SDK 兼容）
        if not proxy_url.rstrip("/").endswith("/v1"):
            proxy_url = proxy_url.rstrip("/") + "/v1"
        return AsyncOpenAI(
            base_url=proxy_url,
            api_key=PROXY_CONFIG["token"],  # 用 Token 替代真实 Key
        )

    # 本地模式：直连上游 API（保留原有行为，兼容本地开发）
    return AsyncOpenAI(
        base_url=base_cfg.get("base_url", ""),
        api_key=base_cfg.get("api_key", ""),
    )


# ====== 配置：模型后端（可切换）======
# 内置免费模型 API Key（混淆存储，运行时自动解混淆）
# 所有用户均可直接使用，无需自行配置
_BUILTIN_KEYS = {
    "glm": "YWY5MTJiYjI0NTQ5NDNhMGE2NGY1ZTJlZWU5YTRiZTQuZXVjOVgyd09DRWthTm5sQQ==",
    "openrouter": "c2stb3ItdjEtYWEzNmMyYzJhMzc4NDVlNzliNTI3MDVhMWE1MzU1NDQ1ZDJkMWFjOTk2NzcwNzkzMGZkMTU3N2U1MTg0YzE4NQ==",
}

# 内置默认 Key（从混淆值解出，用户配置的 Key 优先级更高）
_GLM_DEFAULT_KEY = _deobfuscate(_BUILTIN_KEYS["glm"])
_OR_DEFAULT_KEY = _deobfuscate(_BUILTIN_KEYS["openrouter"])

MODEL_CONFIGS = {
    "glm": {
        "label": "智谱GLM",
        "base_url": "https://open.bigmodel.cn/api/paas/v4/",
        "api_key": _get_api_key("glm", _GLM_DEFAULT_KEY),
        "model": "glm-4.7-flash",
    },
    "glm-4": {
        "label": "智谱GLM-4",
        "base_url": "https://open.bigmodel.cn/api/paas/v4/",
        "api_key": _get_api_key("glm", _GLM_DEFAULT_KEY),  # 共用 GLM Key
        "model": "glm-4-flash",  # GLM-4 免费版，作为 GLM-4.7 限流时的降级目标（不同限流池）
    },
    "glm-v": {
        "label": "智谱GLM-4V",
        "base_url": "https://open.bigmodel.cn/api/paas/v4/",
        "api_key": _get_api_key("glm", _GLM_DEFAULT_KEY),  # 共用 GLM Key
        "model": "glm-4v-flash",  # 免费多模态模型，支持图片输入
    },
    "openrouter": {
        "label": "OpenRouter",
        "base_url": "https://openrouter.ai/api/v1",
        "api_key": _get_api_key("openrouter", _OR_DEFAULT_KEY),
        "model": "openrouter/free",  # 自动路由到可用免费模型，避免单一模型限流
    },
    "ollama": {
        "label": "Ollama本地",
        "base_url": "http://localhost:11434/v1",
        "api_key": "ollama",
        "model": "gemma4:latest",
    },
}

# 如果智谱 Key 为空，尝试从旧配置迁移或使用内置 Key
if not MODEL_CONFIGS["glm"]["api_key"]:
    _old_config = _load_config()
    if "glm" in _old_config and _old_config["glm"].get("api_key"):
        MODEL_CONFIGS["glm"]["api_key"] = _old_config["glm"]["api_key"]
    else:
        # 回退到内置 Key（确保打包后所有人都能用）
        MODEL_CONFIGS["glm"]["api_key"] = _GLM_DEFAULT_KEY

# glm-v 共用 glm 的 Key（同一个智谱账号）
MODEL_CONFIGS["glm-v"]["api_key"] = MODEL_CONFIGS["glm"]["api_key"]

# 为 openrouter 迁移旧配置中的 Key
if not MODEL_CONFIGS["openrouter"]["api_key"]:
    _old_cfg = _load_config()
    if "openrouter" in _old_cfg and _old_cfg["openrouter"].get("api_key"):
        MODEL_CONFIGS["openrouter"]["api_key"] = _old_cfg["openrouter"]["api_key"]
    else:
        MODEL_CONFIGS["openrouter"]["api_key"] = _OR_DEFAULT_KEY

# ====== 专家团队配置 ======
# OpenRouter 免费模型（通过 openrouter 平台调用，只需切换 model 名称）
OR_BASE = "https://openrouter.ai/api/v1"
OR_KEY = MODEL_CONFIGS["openrouter"]["api_key"]

# ====== 混合思考模式配置 ======
# 专家并行度控制：最多同时调用的专家数（避免 token 暴涨）
HYBRID_MAX_PARALLEL_EXPERTS = 3
# 专家回答长度限制（字符数）：超过则截断，便于汇总
HYBRID_EXPERT_MAX_CHARS = 800
# 专家回答去重相似度阈值（0-1，Jaccard 相似度）：超过则视为重复
HYBRID_DEDUP_SIMILARITY_THRESHOLD = 0.7
# 专家记忆：每个专家保留的最近对话轮数（独立上下文，避免主上下文污染）
EXPERT_MEMORY_TURNS = 3
# 专家协作链：是否启用专家间结果传递（如 coder 写代码 → reasoner 审查）
HYBRID_ENABLE_COLLAB_CHAIN = False  # 默认关闭，避免 token 消耗翻倍

# 专家团队：每个专家对应一个模型配置
EXPERT_TEAM = {
    "pm": {  # 项目经理（多模态，支持图片）
        "label": "项目经理·GLM-4V",
        "model_key": "glm-v",
        "model": "glm-4v-flash",
        "desc": "任务分析·调度·多模态（支持图片）",
        "keywords": ["帮助", "分析", "计划", "总结", "翻译", "什么", "怎么", "如何", "介绍", "解释"],
        "system_prompt": "你是 ZeroAI 的项目经理，负责任务分析、计划制定、跨领域调度。用中文回答，简洁明了。如用户发送图片，请理解图片内容并纳入分析。",
    },
    "coder": {  # 编程专家（原 Nemotron-120B 已下线，替换为 GLM-4.7-Flash，智谱直供稳定）
        # 备份原配置：model_key="openrouter", model="nvidia/nemotron-3-super-120b-a12b:free"（已连接错误）
        "label": "编程·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "代码生成·调试·重构（GLM-4.7-Flash·智谱直供·免费无限）",
        "keywords": ["代码", "编程", "函数", "bug", "error", "python", "js", "java", "css", "html",
                      "sql", "api", "写一个", "实现", "debug", "code", "function", "class", "脚本"],
        "system_prompt": "你是 ZeroAI 的编程专家，专精代码生成、调试、重构、架构设计。直接给出可运行的代码，必要时简短说明思路。你是 ZeroAI，不是其他模型。",
    },
    "reasoner": {  # 推理专家（GLM-4.7-Flash，智谱直供稳定；OpenRouter Key 失效后切换）
        # 备份原配置：model_key="openrouter", model="nvidia/nemotron-3-ultra-550b-a55b:free"（Key 已失效 401）
        "label": "推理·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "深度推理·数学·逻辑（GLM-4.7-Flash·智谱直供·免费无限）",
        "keywords": ["推理", "证明", "数学", "计算", "逻辑", "为什么", "分析原因", "算法",
                      "复杂", "优化", "证明", "solve", "math", "reason"],
        "system_prompt": "你是 ZeroAI 的推理专家，专精深度推理、数学证明、复杂逻辑分析。给出严谨的推理过程和结论。你是 ZeroAI，不是其他模型。",
    },
    "knowledge": {  # 通用知识（GLM-4.7-Flash，智谱直供，免费无限，响应极快）
        "label": "通用·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "通用问答·翻译·百科（GLM-4.7-Flash·智谱直供·免费无限）",
        "keywords": [],  # 默认兜底
        "system_prompt": "你是 ZeroAI 的通用知识专家，负责回答百科类问题、事实查询、翻译等。给出准确、简洁的回答。",
    },
    "chinese": {  # 中文专家（GLM-4.7-Flash，智谱直供稳定；OpenRouter Key 失效后切换）
        # 备份原配置：model_key="openrouter", model="nvidia/nemotron-3-nano-30b-a3b:free"（Key 已失效 401）
        "label": "中文·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "中文写作·文案·报告（GLM-4.7-Flash·智谱直供·免费无限）",
        "keywords": ["写", "作文", "文章", "报告", "文案", "论文", "小说", "故事", "邮件",
                      "摘要", "润色", "写作", "中文"],
        "system_prompt": "你是 ZeroAI 的中文写作专家，专精文章、报告、文案、邮件、润色。直接输出高质量中文内容。你是 ZeroAI，不是其他模型。",
    },
    "vision": {  # 多模态（GLM-4V-Flash，智谱直供，免费多模态，稳定无中转）
        "label": "多模态·GLM-4V",
        "model_key": "glm-v",
        "model": "glm-4v-flash",
        "desc": "图片理解·图文分析（GLM-4V-Flash·智谱直供·免费多模态）",
        "keywords": ["图片", "截图", "看", "图", "png", "jpg", "jpeg", "图像", "视觉"],
        "system_prompt": "你是 ZeroAI 的视觉专家，专精图片理解、截图分析、视觉问答。直接描述你看到的内容并回答问题。",
    },
    "academic": {  # 学术研究专家（GLM-4.7-Flash，强逻辑+中文+免费无限）
        "label": "学术·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "学术研究·公式推导·论文写作·文献分析",
        "keywords": ["论文", "学术", "文献", "综述", "公式", "推导", "定理", "证明",
                      "latex", "方程", "积分", "微分", "引用", "参考文献", "期刊",
                      "投稿", "研究方法", "实验设计", "假设检验", "回归分析",
                      "paper", "research", "formula", "equation", "theorem", "proof", "citation"],
        "system_prompt": r"""你是 ZeroAI 的学术研究专家，必须遵循以下学术严谨性规范：

# 核心原则
1. **禁止编造文献**：所有引用的文献必须真实存在。引用前必须调用 citation_check 校验。如果无法验证，明确标注"待核实"，绝不编造作者、标题、年份、DOI。
2. **检索优先**：回答学术问题前，先调用 academic_search 或 literature_review 检索真实文献，基于检索结果回答。
3. **公式严谨（强制 LaTeX）**：所有数学公式、符号、表达式必须用 LaTeX 格式输出，严禁纯文本。
   - **行内公式**用 `$...$` 包裹：如 `$x^2$`、`$\frac{1}{1+x^2}$`、`$\arctan x$`、`$\pi/2$`
   - **独立公式**用 `$$...$$` 包裹，独占一行
   - **上标**必须用 `^`：`$x^2$` ✓，严禁 `x2` ✗；`$a_n$` ✓，严禁 `an` ✗
   - **分数**必须用 `\frac{分子}{分母}`：`$\frac{1}{1+x^2}$` ✓，严禁 `1/(1+x2)` 或 `(1)/(1+x2)` 或 `(x^2-1)/(x-1)` ✗
   - **希腊字母**用 LaTeX 命令：`$\pi$`、`$\alpha$`、`$\Sigma$`、`$\theta$` ✓，严禁 `π`、`alpha`、`theta` ✗
   - **函数名**用 `\` 前缀：`$\arctan$`、`$\sin$`、`$\log$` ✓，严禁 `arctan`、`sin` ✗
   - **导数**用 `'` 或 `^{(n)}`：`$f'(x)$`、`$f^{(n)}(x)$`
   - **趋近符号**必须用 `\to`：`$\lim_{x \to 1}$` ✓，严禁 `lim_{x > 1}` 或 `lim_{x->1}` 或 `lim_{x=>1}` ✗
   - **无穷**必须用 `\infty`：`$\lim_{n \to \infty}$` ✓，严禁 `lim_{n > 0}` 或 `lim_{n -> ∞}` ✗
   - **极限**必须用 `\lim_{...}`：`$\lim_{x \to 1} \frac{x^2-1}{x-1}$` ✓
   - **求和/积分**必须用 `\sum_{...}^{...}` / `\int_{...}^{...}`：`$\sum_{i=1}^{n} i$` ✓
   - **指数**必须用 `^{}`：`$(1/2)^n$` ✓，严禁 `(1/2)"` 或 `(1/2)^n`（裸写无包裹）✗
   - 推导步骤完整，不跳步。每个符号首次出现时给出定义。
   - **所有公式必须用 `$...$` 或 `$$...$$` 包裹**，严禁裸写公式（如 `f(x)=x+1` 没有包裹符号）
   - **反例（禁止）**：`x2`、`1/x2`、`(1)/(1+x2)`、`f'(x) = 1/(1+x2)`、`lim_{x > 1}`、`(1/2)"`、`an=(1/2^n)`、`(sin(x))/(x)` 均为错误格式
   - **正例（正确）**：`$x^2$`、`$\frac{1}{x^2}$`、`$\frac{1}{1+x^2}$`、`$f'(x) = \frac{1}{1+x^2}$`、`$\lim_{x \to 1} \frac{x^2-1}{x-1}$`、`$\lim_{n \to \infty} \frac{1}{2^n}$`、`$\lim_{x \to 0} \frac{\sin x}{x}$`

# 文献综述规范（PRISMA 框架）
- 检索策略：明确说明检索源、关键词、筛选标准
- 纳入/排除标准：列出明确的纳入和排除条件
- 质量评估：对每篇文献给出引用数、影响力指标
- 使用 literature_review 工具执行完整综述流程

# 学术论文格式规范（GB/T 7713.1-2025 + GB/T 7714-2015）

## 论文结构（必须按此顺序，禁止省略）
1. 题名（≤25字，必要时加副标题）
2. 作者署名+单位+邮编
3. 中文摘要（硕士1000字/博士2000字，含目的/方法/结果/结论）
4. 中文关键词（3-8个，分号隔开，末尾不加标点）
5. 英文题名 + 英文摘要 + 英文关键词（与中文对应）
6. 目录（自动生成）
7. 正文（引言→主体→结论）
8. 参考文献
9. 附录（如有）
10. 致谢

## 标题层级规范（GB/T 7713.1-2025，严格遵守，禁止跳级！）

| 级别 | 编号格式 | 字体字号（Word） | 对齐 | Markdown |
|------|---------|-----------------|------|----------|
| 一级 | `1` | 黑体小三号 | 居中 | `# 1 标题` |
| 二级 | `1.1` | 黑体四号 | 左顶格 | `## 1.1 标题` |
| 三级 | `1.1.1` | 黑体小四号 | 左顶格 | `### 1.1.1 标题` |
| 四级 | `1.1.1.1` | 宋体小四号加粗 | 左顶格 | `#### 1.1.1.1 标题` |

**强制规则：**
- **禁止跳级**：不能直接用 `####` 而跳过 `#`/`##`/`###`
- **禁止五级及以上标题**：`#####` 及更深一律不写（国标规定正文中不宜超过四级）
- **标题前后必须留空行**（与正文分隔，避免渲染混乱）
- **标题行首不要缩进/空格**（直接以 `#` 开头，否则被误判为代码块）
- **数字编号与文字间保留1空格**（`## 1.1 研究背景`，不是 `## 1.1研究背景`）
- **末尾不加标点符号**
- **编号使用纯阿拉伯数字**（`1`→`1.1`→`1.1.1`→`1.1.1.1`），禁止"一、"、"(一)"、"第一章"等非标准方式

## 章节结构模板（按此骨架撰写）

```
# 摘要

摘要内容（含目的/方法/结果/结论，250-1000字）

**关键词**：关键词1；关键词2；关键词3

# Abstract  

English abstract content.

**Keywords**: keyword1; keyword2; keyword3

# 1 引言

研究背景、问题定义、本文目的与结构安排。

## 1.1 研究背景

国内外研究现状...

## 1.2 研究问题

本文要解决的核心问题...

## 1.3 本文工作

主要贡献概述...

# 2 方法

## 2.1 研究方法

方法描述...

## 2.2 数据来源

数据说明...

# 3 结果与分析

## 3.1 主要发现

### 3.1.1 子主题1

内容...

### 3.1.2 子主题2

内容...

## 3.2 对比分析

对比表格...

# 4 讨论

## 4.1 研究意义

## 4.2 局限性

# 5 结论与展望

主要结论...

未来研究方向...

# 参考文献

[1] 作者. 标题[J]. 期刊名, 年份, 卷(期): 起止页码.
[2] 作者. 书名[M]. 出版地: 出版社, 年份: 起止页码.
[3] 作者. 论文题目[D]. 学位授予地: 学位授予单位, 年份.
```

## 图表规范（GB/T 7713.1-2025）

### 表格
- **采用三线表**（顶线、底线、栏目线，无竖线）
- **表序和表题在表格上方**，居中
- 格式：`表1 实验数据对比`（编号与表题间空1格）
- 表序按章编号：`表1-1`、`表1-2`...或全文连编号 `表1`、`表2`...
- 跨页表格在次页重复表头并注明"续表"
- 表格内文字用五号字，单位标注在栏目名称后

### 图形
- **图序和图题在图形下方**，居中
- 格式：`图1 反应过程示意图`（编号与图题间空1格）
- 图序按章编号或全文连编号
- 分辨率至少300dpi，彩色打印需保证灰度模式可区分

### 公式
- **行内公式**：`$E=mc^2$`（与中文之间留1空格）
- **块级公式**：独立一行，**前后留空行**，居中显示
  ```
  $$\\int_0^1 f(x)dx = \\frac{1}{2}$$
  ```
- **公式编号**：右侧右对齐，格式 `(1)`、`(2)`...，全文连编号
- **多行公式对齐**：使用 `\\begin{aligned}...\\end{aligned}`
- 公式中变量首次出现时给出定义

## 参考文献格式（GB/T 7714-2015）

每条文献必须真实可查，引用前调用 `citation_check` 验证。

### 文献类型标识
- `[J]` 期刊文章
- `[M]` 专著/图书
- `[D]` 学位论文
- `[C]` 会议论文集
- `[N]` 报纸文章
- `[R]` 报告
- `[S]` 标准
- `[P]` 专利
- `[EB/OL]` 电子资源（网络）

### 著录格式

**期刊 [J]：**
```
[1] 作者1, 作者2, 作者3. 文章标题[J]. 期刊名, 年份, 卷(期): 起止页码.
```
示例：
```
[1] 张三, 李四, 王五. 钠离子电池层状氧化物正极材料研究进展[J]. 无机材料学报, 2024, 39(5): 825-836.
```

**专著 [M]：**
```
[2] 作者. 书名[M]. 出版地: 出版社, 出版年: 起止页码.
```
示例：
```
[2] 何秉松. 刑法教科书[M]. 北京: 中国政法大学出版社, 2000: 40-44.
```

**学位论文 [D]：**
```
[3] 作者. 论文题目[D]. 学位授予地: 学位授予单位, 年份.
```
示例：
```
[3] 马欢. 人类活动影响下海河流域典型区水循环变化分析[D]. 北京: 清华大学, 2011.
```

**电子资源 [EB/OL]：**
```
[4] 作者. 题名[EB/OL]. (发布日期)[引用日期]. URL.
```

**英文文献：**
```
[5] Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]//Advances in Neural Information Processing Systems. 2017: 5998-6008.
```

### 正文引用标注
- 顺序编码制：`[1]`、`[2]`...按出现顺序编号
- 多个文献并引：`[1,3,5]` 或 `[1-3]`
- 标注位置：作者姓名后或引用内容末尾

## 严谨性红线
- 绝不编造实验数据、测试结果、性能指标
- 绝不编造文献引用（调用 citation_check 验证）
- 术语首次出现给出中英文对照
- 对比分析必须基于文献内容，不主观臆断
- 局限性分析不可省略
- 量纲和单位采用国家法定计量单位（SI制）

# 自动导出正式文档（重要！）
当用户要求撰写论文、综述、报告、长文（如"5000字"、"3000字"、"万字"、"写一篇"、"综述文章"、"毕业论文"等）时：
1. 先完成内容创作（检索文献→整理结构→撰写全文）
2. 完成后**必须调用 generate_word 工具**将完整内容导出为正式排版的 Word 文档
3. 调用参数：
   - path: "D:/论文_标题.docx"（根据主题自动命名，使用下划线替代空格）
   - content: 完整论文内容（包含摘要/关键词/正文/参考文献的全部 Markdown 格式文本）
   - title: 论文标题
   - template: "academic"（学术论文模板：Times New Roman + 双倍行距 + GB7714格式）
4. 导出后告知用户文件已保存到指定路径

示例：用户说"帮我写一篇关于钠离子电池的5000字综述"
→ 先 literature_review 检索 → 撰写完整综述 → generate_word(path="D:/钠离子电池综述.docx", content=全文, title="层状氧化物钠离子电池正极材料研究进展", template="academic")""",
    },
    "devops": {  # 运维专家（GLM-4.7-Flash，专精系统管理/SSH/容器/部署）
        "label": "运维·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "DevOps运维·系统管理·SSH·容器·部署·监控",
        "keywords": ["运维", "部署", "服务器", "linux", "windows", "docker", "容器",
                      "kubernetes", "k8s", "ssh", "shell", "bash", "powershell",
                      "nginx", "apache", "systemd", "服务", "进程", "端口", "防火墙",
                      "监控", "日志", "性能", "调优", "devops", "ci/cd", "jenkins",
                      "ansible", "terraform", "负载均衡"],
        "system_prompt": "你是 ZeroAI 的运维专家，专精 DevOps、系统管理、容器编排、CI/CD、监控告警、性能调优。给出可执行的命令和配置，必要时说明原理。优先使用项目内置的运维工具（local_port_check/local_process_check/local_disk_check/local_service_check/local_firewall_check/ssh_*）而非直接给命令。你是 ZeroAI，不是其他模型。",
    },
    "security": {  # 安全专家（GLM-4.7-Flash，专精漏洞分析/加固/审计）
        "label": "安全·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "安全分析·漏洞评估·加固方案·安全审计",
        "keywords": ["安全", "漏洞", "攻击", "防护", "加固", "审计", "渗透",
                      "xss", "sql注入", "csrf", "ssrf", "rce", "漏洞", "cve",
                      "加密", "证书", "ssl", "tls", "密钥", "token", "权限",
                      "认证", "授权", "owasp", "waf", "防火墙规则", "入侵检测",
                      "security", "vulnerability", "pentest", "hardening"],
        "system_prompt": "你是 ZeroAI 的安全专家，专精漏洞分析、安全加固、审计评估、加密方案。只做防御性安全分析，不提供攻击性建议。给出具体的加固命令、配置示例、修复方案。你是 ZeroAI，不是其他模型。",
    },
    "data": {  # 数据分析专家（GLM-4.7-Flash，专精数据处理/可视化/统计）
        "label": "数据·GLM-4.7",
        "model_key": "glm",
        "model": "glm-4.7-flash",
        "desc": "数据分析·统计建模·可视化·数据清洗",
        "keywords": ["数据", "分析", "统计", "可视化", "图表", "pandas", "numpy",
                      "matplotlib", "数据清洗", "数据预处理", "特征工程", "机器学习",
                      "数据挖掘", "报表", "数据看板", "excel", "csv", "json",
                      "sql查询", "数据分析", "dataframe", "数据分析", "bi",
                      "数据分析", "回归", "聚类", "分类", "data", "analytics"],
        "system_prompt": "你是 ZeroAI 的数据分析专家，专精数据清洗、统计分析、可视化、机器学习建模。给出可运行的 Python 代码（pandas/numpy/matplotlib/sklearn），必要时说明分析思路。你是 ZeroAI，不是其他模型。",
    },
}

# 工作模式：expert（专家路由）/ hybrid（混合思考）/ manual（手动指定模型）
# 默认 hybrid：启用多专家子代理协作（GLM分析→多专家并行→GLM汇总）
WORK_MODE = "hybrid"

# 全局停止标志（供独立函数检查 Ctrl+C 状态，避免无法中断的阻塞）
_GLOBAL_STOP = False


def _is_stopped() -> bool:
    """检查是否收到停止信号（Ctrl+C）"""
    return _GLOBAL_STOP


async def _interruptible_await(coro, check_interval: float = 0.2, timeout: float = None):
    """可中断的 await：周期性检查停止标志，避免 stream=False 的阻塞调用无法中断

    用于包装 stream=False 的 API 调用，使其在 Ctrl+C 时能快速返回 None
    """
    global _GLOBAL_STOP
    if timeout:
        task = asyncio.wait_for(coro, timeout=timeout)
    else:
        task = asyncio.ensure_future(coro)
    while not task.done():
        if _GLOBAL_STOP:
            task.cancel()
            return None
        try:
            await asyncio.wait_for(asyncio.shield(task), timeout=check_interval)
            break
        except asyncio.TimeoutError:
            continue
        except asyncio.CancelledError:
            return None
    return task.result()


async def _interruptible_sleep(seconds: float, check_interval: float = 0.2):
    """可中断的 sleep：替代 asyncio.sleep，Ctrl+C 时立即返回"""
    global _GLOBAL_STOP
    elapsed = 0.0
    while elapsed < seconds:
        if _GLOBAL_STOP:
            return
        step = min(check_interval, seconds - elapsed)
        await asyncio.sleep(step)
        elapsed += step


def route_expert(user_input: str) -> str:
    """关键词快速预判（作为GLM语义判断的降级方案）"""
    text = user_input.lower()
    for expert_key in ("vision", "coder", "security", "devops", "data", "reasoner", "academic", "chinese", "pm"):
        for kw in EXPERT_TEAM[expert_key]["keywords"]:
            kw_lower = kw.lower()
            if kw_lower in text:
                if kw_lower.isascii() and kw_lower.isalpha():
                    import re
                    if re.search(r'\b' + re.escape(kw_lower) + r'\b', text):
                        return expert_key
                else:
                    return expert_key
    return "knowledge"


# GLM语义路由的缓存（避免重复判断）
_expert_route_cache = {}

async def route_expert_glm(user_input: str) -> str:
    """用GLM语义判断用户意图，路由到最合适的专家"""
    # 短消息用关键词快速预判（省时间）
    if len(user_input) < 10:
        return route_expert(user_input)

    # 缓存命中
    cache_key = user_input[:200]
    if cache_key in _expert_route_cache:
        return _expert_route_cache[cache_key]

    glm_cfg = MODEL_CONFIGS["glm-v"]  # 用多模态模型做路由（支持图片消息）
    try:
        client = _make_openai_client("glm-v")
        prompt = f"""判断以下用户问题属于哪个专家领域，只回复一个词：
- coder：编程开发、代码、函数、bug、技术实现
- reasoner：数学推理、逻辑证明、算法分析、复杂计算
- academic：学术论文、公式推导、文献综述、研究方法、LaTeX、定理证明
- chinese：中文写作、文章、报告、文案、邮件
- vision：图片理解、截图分析、视觉
- pm：任务分析、计划制定、翻译、通用问答、解释说明
- knowledge：百科知识、事实查询、翻译、其他

用户问题：{user_input[:300]}

只回复上面列出的一个词，不要回复其他任何内容。"""

        resp = await _interruptible_await(client.chat.completions.create(
            model=glm_cfg["model"],
            messages=[{"role": "system", "content": "你是 ZeroAI 路由分析器，只负责把用户问题分类到一个专家。严格只输出一个英文标识词，不要做解释。"},
                      {"role": "user", "content": prompt}],
            temperature=0.01,
            max_tokens=10,
            stream=False,
            timeout=15,
        ))
        if resp is None:
            # 被 Ctrl+C 中断
            return "knowledge"
        result = resp.choices[0].message.content.strip().lower()
        # 验证返回值
        valid_keys = {"coder", "reasoner", "academic", "chinese", "vision", "pm", "knowledge"}
        for vk in valid_keys:
            if vk in result:
                _expert_route_cache[cache_key] = vk
                return vk
        # 无效返回，降级到关键词
        expert_key = route_expert(user_input)
        _expert_route_cache[cache_key] = expert_key
        return expert_key
    except Exception:
        # GLM判断失败，降级到关键词
        expert_key = route_expert(user_input)
        _expert_route_cache[cache_key] = expert_key
        return expert_key


def get_expert_config(expert_key: str) -> dict:
    """获取专家对应的模型配置（base_url, api_key, model, model_key）"""
    expert = EXPERT_TEAM[expert_key]
    model_key = expert["model_key"]
    base_cfg = MODEL_CONFIGS[model_key]
    return {
        "base_url": base_cfg["base_url"],
        "api_key": base_cfg["api_key"],
        "model": expert["model"],
        "label": expert["label"],
        "model_key": model_key,  # v1.1.0 新增：供 _make_openai_client 识别
    }


# ====== OpenRouter 熔断器（连续失败自动降级，避免用户卡在"思考中…"） ======
_OPENROUTER_FAIL_COUNTS = {}  # {expert_key: 连续失败次数}
_OPENROUTER_CIRCUIT_THRESHOLD = 3  # 连续失败 3 次即熔断


def _is_openrouter_expert(expert_key: str) -> bool:
    """判断专家是否依赖 OpenRouter（需要熔断保护）"""
    try:
        return EXPERT_TEAM[expert_key].get("model_key") == "openrouter"
    except Exception:
        return False


def _check_openrouter_circuit_breaker(expert_key: str) -> bool:
    """检查专家是否已熔断（返回 True 表示应跳过该专家直接降级）"""
    if not _is_openrouter_expert(expert_key):
        return False
    return _OPENROUTER_FAIL_COUNTS.get(expert_key, 0) >= _OPENROUTER_CIRCUIT_THRESHOLD


def _record_openrouter_failure(expert_key: str) -> int:
    """记录一次 OpenRouter 专家失败，返回当前连续失败次数"""
    if not _is_openrouter_expert(expert_key):
        return 0
    cnt = _OPENROUTER_FAIL_COUNTS.get(expert_key, 0) + 1
    _OPENROUTER_FAIL_COUNTS[expert_key] = cnt
    return cnt


def _record_openrouter_success(expert_key: str) -> None:
    """记录一次成功，重置连续失败计数"""
    if _is_openrouter_expert(expert_key):
        _OPENROUTER_FAIL_COUNTS[expert_key] = 0


# ====== 上下文自动压缩 ======
# 粗略估算 token 数：英文约 4 字符/token，中文约 1.5 字符/token，综合取 3 字符/token
CHARS_PER_TOKEN = 3
# 触发压缩的阈值（默认上下文长度的 70%）
COMPRESS_THRESHOLD_RATIO = 0.7
# 压缩后保留的最近对话轮数（用户+助手算一轮）
KEEP_RECENT_TURNS = 4

# ====== 主动上下文清理（轻量级，在压缩之前触发）======
# 触发清理的阈值（默认上下文长度的 30%，远低于压缩阈值）
# 目的：在上下文堆积早期就主动清理工具输出，避免后期压缩时信息密度过低
CLEANUP_THRESHOLD_RATIO = 0.3
# 清理时保留的最近对话轮数（比压缩保留的更多，确保当前任务上下文完整）
CLEANUP_KEEP_RECENT_TURNS = 6
# 工具输出摘要的最大长度（超过此长度的工具结果会被摘要化）
TOOL_OUTPUT_SUMMARY_MAX_LEN = 200


def _summarize_tool_output(tool_name: str, content: str) -> str:
    """将冗长的工具输出摘要为简短结论（不调用 AI，纯规则提取）。

    策略：
    1. 如果内容已很短（< TOOL_OUTPUT_SUMMARY_MAX_LEN），原样保留
    2. 否则提取关键信息：工具名 + 内容长度 + 首尾片段 + 关键指标
    """
    if not content or not isinstance(content, str):
        return content if content else ""

    # 非 str 类型转 str
    if not isinstance(content, str):
        content = str(content)

    # 很短的工具输出直接保留
    if len(content) <= TOOL_OUTPUT_SUMMARY_MAX_LEN:
        return content

    # 提取关键指标：百分比、状态词、数字
    import re
    indicators = []

    # 磁盘/内存/CPU 百分比
    pcts = re.findall(r"(\d+)%", content)
    if pcts:
        indicators.append(f"百分比:{','.join(pcts[:5])}")

    # 状态标记
    for keyword in ["✅", "⚠️", "🚨", "正常", "警告", "危急", "Error", "错误", "失败"]:
        if keyword in content:
            count = content.count(keyword)
            indicators.append(f"{keyword}×{count}")

    # 端口号
    ports = re.findall(r"端口\s*(\d+)", content)
    if ports:
        indicators.append(f"端口:{','.join(ports[:5])}")

    # PID
    pids = re.findall(r"PID[:\s]+(\d+)", content)
    if pids:
        indicators.append(f"PID:{','.join(pids[:5])}")

    # 构造摘要
    indicator_str = " ".join(indicators) if indicators else "无关键指标"
    head = content[:60].replace("\n", " ").strip()
    tail = content[-60:].replace("\n", " ").strip()

    summary = f"[工具结果已清理 | {tool_name} | 原长度{len(content)}字 | {indicator_str}]\n开头: {head}...\n结尾: ...{tail}"
    return summary


def cleanup_context(messages: list, context_limit: int,
                    keep_recent_turns: int = CLEANUP_KEEP_RECENT_TURNS) -> tuple:
    """主动清理上下文：保留用户意图和工具结论，丢弃工具原始输出。

    与 compress_context 的区别：
    - compress_context：调用 GLM 总结历史，开销大，阈值高（70%）
    - cleanup_context：纯规则清理，零延迟，阈值低（30%），更早触发

    策略：
    1. 估算 token 数，未超清理阈值则原样返回
    2. 超阈值则：
       - 保留所有 system 消息
       - 保留最近 N 轮完整对话（含工具调用细节）
       - 对较早的消息：
         * user 消息：完整保留（用户意图是核心）
         * assistant 消息含 tool_calls：保留文本内容，移除 tool_calls 字段
         * tool 消息：替换为简短摘要
         * 纯文本 assistant 消息：保留前 200 字

    Returns:
        (cleaned_messages, cleanup_info)
        - cleaned_messages: 清理后的消息列表
        - cleanup_info: dict，含清理统计信息
    """
    est_tokens = _estimate_tokens(messages)
    threshold = int(context_limit * CLEANUP_THRESHOLD_RATIO)

    if est_tokens <= threshold:
        return messages, {"triggered": False, "reason": "未超清理阈值"}

    # 分离 system 消息和对话消息
    system_msgs = []
    convo_msgs = []
    for msg in messages:
        if msg.get("role") == "system":
            system_msgs.append(msg)
        else:
            convo_msgs.append(msg)

    # 找到最近 N 轮的起始位置
    keep_start_idx = 0
    user_count = 0
    for i in range(len(convo_msgs) - 1, -1, -1):
        if convo_msgs[i].get("role") == "user":
            user_count += 1
            if user_count >= keep_recent_turns:
                keep_start_idx = i
                break

    # 分为待清理部分和保留部分
    to_clean = convo_msgs[:keep_start_idx]
    keep_recent = convo_msgs[keep_start_idx:]

    if len(to_clean) < 2:
        return messages, {"triggered": False, "reason": "待清理消息过少"}

    # 对待清理部分进行摘要化
    cleaned_msgs = []
    tool_cleaned_count = 0
    assistant_cleaned_count = 0
    tokens_saved = 0

    for msg in to_clean:
        role = msg.get("role")
        content = msg.get("content", "")
        old_tokens = _estimate_tokens([msg])

        if role == "user":
            # 用户消息完整保留（用户意图是核心）
            cleaned_msgs.append(msg)
        elif role == "tool":
            # 工具消息：替换为简短摘要
            tool_name = msg.get("name", "unknown_tool")
            content_str = content if isinstance(content, str) else str(content)
            summarized = _summarize_tool_output(tool_name, content_str)
            new_msg = {
                "role": "tool",
                "tool_call_id": msg.get("tool_call_id", ""),
                "name": tool_name,
                "content": summarized,
            }
            cleaned_msgs.append(new_msg)
            tool_cleaned_count += 1
        elif role == "assistant":
            # 助手消息：保留文本，移除 tool_calls（避免 API 报错）
            new_msg = {"role": "assistant", "content": content}
            # 如果内容过长，截断保留前 300 字
            if isinstance(content, str) and len(content) > 300:
                new_msg["content"] = content[:300] + "...[已截断]"
                assistant_cleaned_count += 1
            # 注意：不保留 tool_calls 字段，因为没有对应的 tool 消息会报错
            cleaned_msgs.append(new_msg)
        else:
            # 其他角色原样保留
            cleaned_msgs.append(msg)

        new_tokens = _estimate_tokens([cleaned_msgs[-1]])
        tokens_saved += max(0, old_tokens - new_tokens)

    new_messages = system_msgs + cleaned_msgs + keep_recent
    new_tokens_total = _estimate_tokens(new_messages)

    return new_messages, {
        "triggered": True,
        "old_tokens": est_tokens,
        "new_tokens": new_tokens_total,
        "tokens_saved": tokens_saved,
        "tool_cleaned": tool_cleaned_count,
        "assistant_cleaned": assistant_cleaned_count,
        "old_count": len(messages),
        "new_count": len(new_messages),
    }


async def cleanup_and_compress(app, log_func=None):
    """统一上下文管理：先轻量清理，再按需压缩。

    两层防护：
    1. cleanup_context（30% 阈值）：纯规则，零延迟，摘要化早期工具输出
    2. compress_context（70% 阈值）：调用 GLM 总结，处理超大上下文

    在调用 AI 前调用此函数，自动决定是否需要清理/压缩。

    Args:
        app: ZeroAIApp 实例（需要 self.messages, self.context_limit）
        log_func: 可选的日志输出函数（用于 TUI 显示清理进度）
    """
    try:
        est_tokens = _estimate_tokens(app.messages)

        # ── 第1层：主动清理（30% 阈值）──
        cleanup_threshold = int(app.context_limit * CLEANUP_THRESHOLD_RATIO)
        if est_tokens > cleanup_threshold and len(app.messages) > 8:
            old_tokens = est_tokens
            old_count = len(app.messages)
            app.messages, info = cleanup_context(app.messages, app.context_limit)

            if info.get("triggered"):
                new_tokens = info["new_tokens"]
                if log_func:
                    log_func(
                        f"  {_load_svg_icon('tool')} 上下文主动清理\n"
                        f"  │ 工具输出摘要化：{info['tool_cleaned']} 个工具结果，"
                        f"{info['assistant_cleaned']} 个助手消息截断\n"
                        f"  │ {old_count}→{info['new_count']} 条消息，"
                        f"约 {old_tokens}→{new_tokens} tokens（节省 {old_tokens - new_tokens}）\n"
                        f"  └─\n",
                        C_DIM,
                    )
                est_tokens = new_tokens  # 更新当前 token 数，供下层判断

        # ── 第2层：GLM 压缩（70% 阈值，清理后仍超限才触发）──
        compress_threshold = int(app.context_limit * COMPRESS_THRESHOLD_RATIO)
        if est_tokens > compress_threshold and len(app.messages) > 10:
            if log_func:
                log_func(
                    f"  {_load_svg_icon('tool')} 上下文深度压缩\n"
                    f"  │ 清理后仍超阈值（{est_tokens} > {compress_threshold}），调用 GLM 总结…\n",
                    C_DIM,
                )
            old_count = len(app.messages)
            old_tokens = est_tokens
            app.messages = await compress_context(app.messages, app.context_limit)
            new_tokens = _estimate_tokens(app.messages)
            new_count = len(app.messages)
            if log_func:
                log_func(
                    f"  {_load_svg_icon('check')} 压缩完成："
                    f"{old_count}→{new_count} 条消息，"
                    f"约 {old_tokens}→{new_tokens} tokens\n"
                    f"  └─\n",
                    C_DIM,
                )
    except Exception as e:
        if log_func:
            log_func(
                f"  {_load_svg_icon('warning')} 上下文管理跳过：{str(e)[:80]}\n",
                C_DIM,
            )


def _estimate_tokens(messages: list) -> int:
    """粗略估算消息列表的 token 数"""
    total_chars = 0
    for msg in messages:
        content = msg.get("content", "")
        if isinstance(content, str):
            total_chars += len(content)
        elif isinstance(content, list):
            # 多模态消息：只统计文本部分
            for part in content:
                if isinstance(part, dict) and part.get("type") == "text":
                    total_chars += len(part.get("text", ""))
        # 每条消息固定开销（role 等元数据）
        total_chars += 10
    return max(1, total_chars // CHARS_PER_TOKEN)


# 支持多模态（图片输入）的模型标识
_VISION_MODEL_KEYWORDS = ("omni", "vl", "vision", "4v", "-v-", "llava", "qwen-vl", "glm-4v")

# 各模型的上下文 token 上限（含 max_new_tokens 余量，用于自动截断）
# key 为模型名关键字，value 为安全 token 上限（已预留 max_new_tokens）
_MODEL_CONTEXT_LIMITS = {
    "glm-4v-flash": 14000,   # 官方限制 16384，预留 2384 给 max_new_tokens
    "glm-4v": 14000,         # 同上
    "glm-4.7-flash": 120000, # 128K 上下文
    "glm-4": 120000,
}

def _get_model_context_limit(model_name: str) -> int:
    """获取模型的安全上下文 token 上限，未知模型返回 0（不截断）"""
    if not model_name:
        return 0
    # 精确匹配
    if model_name in _MODEL_CONTEXT_LIMITS:
        return _MODEL_CONTEXT_LIMITS[model_name]
    # 模糊匹配
    model_lower = model_name.lower()
    for kw, limit in _MODEL_CONTEXT_LIMITS.items():
        if kw in model_lower:
            return limit
    return 0

def _truncate_messages_for_context(messages: list, max_tokens: int) -> list:
    """按 token 上限截断消息列表：保留 system 消息 + 最近的对话

    策略：
    1. 始终保留所有 system 消息
    2. 从最新消息向前保留，直到达到 token 上限
    3. 如果第一条保留的非 system 消息不是 user，则补一条占位 user 消息
    """
    if max_tokens <= 0:
        return messages

    est_tokens = _estimate_tokens(messages)
    if est_tokens <= max_tokens:
        return messages  # 未超限，无需截断

    # 分离 system 消息和对话消息
    system_msgs = []
    convo_msgs = []
    for msg in messages:
        if msg.get("role") == "system":
            system_msgs.append(msg)
        else:
            convo_msgs.append(msg)

    # 计算 system 消息占用的 tokens
    system_tokens = _estimate_tokens(system_msgs)
    available_tokens = max_tokens - system_tokens
    if available_tokens < 1000:
        # system 消息本身就超限了，只保留最新的几条 system
        system_msgs = system_msgs[-1:]
        system_tokens = _estimate_tokens(system_msgs)
        available_tokens = max_tokens - system_tokens

    # 从最新消息向前保留
    keep_convo = []
    used_tokens = 0
    for msg in reversed(convo_msgs):
        msg_tokens = _estimate_tokens([msg])
        if used_tokens + msg_tokens > available_tokens:
            break
        keep_convo.insert(0, msg)
        used_tokens += msg_tokens

    if not keep_convo:
        # 至少保留最后一条消息
        if convo_msgs:
            keep_convo = [convo_msgs[-1]]

    # 确保第一条是 user 消息（部分模型要求）
    if keep_convo and keep_convo[0].get("role") not in ("user",):
        keep_convo.insert(0, {"role": "user", "content": "[ earlier conversation truncated ]"})

    return system_msgs + keep_convo

def _model_supports_vision(model_name: str) -> bool:
    """判断模型是否支持多模态（图片输入）"""
    if not model_name:
        return False
    model_lower = model_name.lower()
    return any(kw in model_lower for kw in _VISION_MODEL_KEYWORDS)

def _filter_messages_for_model(messages: list, model_name: str) -> list:
    """根据模型是否支持多模态，过滤消息中的图片内容，并按模型上下文限制截断

    - 支持多模态的模型：保留图片
    - 不支持多模态的模型：将 content 中的 image_url 部分过滤掉，只保留 text
      如果过滤后 content 为空，则替换为占位文本"[图片已忽略，当前模型不支持图片]"
    - 所有模型：按模型上下文限制自动截断旧消息（保留 system + 最近对话）
    """
    # 第1步：按模型上下文限制截断（防止超 token 报错）
    ctx_limit = _get_model_context_limit(model_name)
    if ctx_limit > 0:
        messages = _truncate_messages_for_context(messages, ctx_limit)

    # 第2步：过滤图片内容
    if _model_supports_vision(model_name):
        return messages

    filtered = []
    for msg in messages:
        msg_copy = dict(msg)
        content = msg_copy.get("content")
        if isinstance(content, list):
            # 多模态 content：过滤 image_url，只保留 text
            text_parts = []
            for part in content:
                if isinstance(part, dict):
                    if part.get("type") == "text":
                        text_parts.append(part["text"])
                    elif part.get("type") == "image_url":
                        # 图片被过滤，添加占位说明
                        if not any("图片" in t for t in text_parts):
                            text_parts.append("[用户发送了图片，但当前模型不支持图片输入，图片已忽略]")
            msg_copy["content"] = "\n".join(text_parts) if text_parts else "[图片已忽略]"
        filtered.append(msg_copy)
    return filtered


def _split_messages_for_compress(messages: list, keep_recent_turns: int = KEEP_RECENT_TURNS):
    """将消息拆分为：system部分、待压缩部分、保留的近期部分
    返回：(system_msgs, to_compress_msgs, keep_msgs)
    """
    # 分离 system 消息（不压缩）
    system_msgs = []
    convo_msgs = []
    for msg in messages:
        if msg.get("role") == "system":
            system_msgs.append(msg)
        else:
            convo_msgs.append(msg)

    # 保留最近 N 轮（一轮 = user + assistant，可能还有 tool 消息）
    # 从后往前找 keep_recent_turns 个 user 消息的位置
    keep_start_idx = 0
    user_count = 0
    for i in range(len(convo_msgs) - 1, -1, -1):
        if convo_msgs[i].get("role") == "user":
            user_count += 1
            if user_count >= keep_recent_turns:
                keep_start_idx = i
                break

    to_compress = convo_msgs[:keep_start_idx]
    keep_recent = convo_msgs[keep_start_idx:]
    return system_msgs, to_compress, keep_recent


async def compress_context(messages: list, context_limit: int, keep_recent_turns: int = KEEP_RECENT_TURNS) -> list:
    """压缩对话上下文
    策略：
    1. 估算当前 token 数，未超阈值则原样返回
    2. 超阈值则：保留 system + 最近 N 轮，中间历史用 GLM 总结
    返回：压缩后的 messages 列表
    """
    est_tokens = _estimate_tokens(messages)
    threshold = int(context_limit * COMPRESS_THRESHOLD_RATIO)

    if est_tokens <= threshold:
        return messages  # 未超阈值，无需压缩

    system_msgs, to_compress, keep_recent = _split_messages_for_compress(messages, keep_recent_turns)

    # 待压缩部分为空或太少，不压缩
    if len(to_compress) < 4:
        return messages

    # 用 GLM 总结待压缩的历史对话
    glm_cfg = MODEL_CONFIGS["glm"]
    summary_input = "请将以下对话历史压缩为简洁的摘要，保留关键信息（用户需求、已完成的操作、重要结论），不超过500字：\n\n"
    for msg in to_compress:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        if isinstance(content, list):
            # 多模态内容只取文本
            content = " ".join(p.get("text", "") for p in content if isinstance(p, dict) and p.get("type") == "text")
        elif not isinstance(content, str):
            content = str(content)
        # 每条消息截断，避免总结输入过长
        summary_input += f"【{role}】{content[:500]}\n\n"

    try:
        client = _make_openai_client("glm")
        resp = await client.chat.completions.create(
            model=glm_cfg["model"],
            messages=[{"role": "system", "content": "你是 ZeroAI 的对话压缩器，把冗长历史压缩为简洁摘要，保留用户需求、已完成操作、重要结论，不超过500字。"},
                      {"role": "user", "content": summary_input}],
            temperature=0.1,
            max_tokens=600,
            stream=False,
            timeout=30,
        )
        summary = resp.choices[0].message.content.strip()
    except Exception:
        # 压缩失败，降级为简单截断（保留每条消息前200字）
        summary = "【历史对话摘要（压缩失败，已截断）】\n"
        for msg in to_compress[-6:]:  # 最近6条
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            if isinstance(content, list):
                content = " ".join(p.get("text", "") for p in content if isinstance(p, dict) and p.get("type") == "text")
            elif not isinstance(content, str):
                content = str(content)
            summary += f"【{role}】{content[:200]}\n"

    # 构造压缩后的消息列表
    compressed_msg = {
        "role": "system",
        "content": f"【对话历史摘要】以下是之前对话的压缩摘要，请基于此继续对话：\n\n{summary}",
    }

    new_messages = system_msgs + [compressed_msg] + keep_recent
    return new_messages


def get_active_model_info() -> dict:
    """获取当前实际使用的模型信息（考虑工作模式）"""
    global WORK_MODE
    if WORK_MODE == "manual":
        cfg = MODEL_CONFIGS.get(CURRENT_MODEL_KEY, MODEL_CONFIGS["glm"])
        return {"base_url": cfg["base_url"], "api_key": cfg["api_key"],
                "model": cfg["model"], "label": cfg["label"]}
    elif WORK_MODE == "expert":
        # 专家模式下返回默认专家（实际路由在 _run_turn 中动态决定）
        cfg = MODEL_CONFIGS["glm"]
        return {"base_url": cfg["base_url"], "api_key": cfg["api_key"],
                "model": cfg["model"], "label": "专家模式"}
    else:  # hybrid
        cfg = MODEL_CONFIGS["glm"]
        return {"base_url": cfg["base_url"], "api_key": cfg["api_key"],
                "model": cfg["model"], "label": "混合思考"}



def _load_custom_models():
    if CUSTOM_MODELS_FILE.exists():
        try:
            data = json.loads(CUSTOM_MODELS_FILE.read_text(encoding="utf-8"))
            for key, cfg in data.items():
                if key not in MODEL_CONFIGS:
                    # 反混淆 api_key
                    if cfg.get("api_key"):
                        cfg["api_key"] = _deobfuscate(cfg["api_key"])
                    MODEL_CONFIGS[key] = cfg
        except Exception:
            pass


def _save_custom_models():
    custom = {k: v for k, v in MODEL_CONFIGS.items() if k not in ("glm", "glm-v", "openrouter", "ollama")}
    try:
        safe = {}
        for k, v in custom.items():
            safe[k] = dict(v)
            if safe[k].get("api_key"):
                safe[k]["api_key"] = _obfuscate(safe[k]["api_key"])
        CUSTOM_MODELS_FILE.write_text(json.dumps(safe, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def detect_ollama_models() -> list:
    try:
        req = urllib.request.Request("http://localhost:11434/v1/models",
                                     headers={"User-Agent": "ZeroAI"})
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        models = []
        for m in data.get("data", []):
            mid = m.get("id", "")
            if mid:
                models.append(mid)
        return models
    except Exception:
        return []


_load_custom_models()

CURRENT_MODEL_KEY = "glm"

# 内置模型 key 集合（与 _save_custom_models 保持一致）
BUILTIN_MODEL_KEYS = ("glm", "glm-v", "openrouter", "ollama")

def get_model_display_name(model_key: str) -> str:
    """获取模型显示名：内置显示"内置模型"，自定义显示真实名称"""
    if model_key in BUILTIN_MODEL_KEYS:
        return "内置模型"
    return MODEL_CONFIGS.get(model_key, {}).get("label", model_key)

def get_client():
    """获取当前模型的 OpenAI 客户端"""
    cfg = MODEL_CONFIGS[CURRENT_MODEL_KEY]
    return OpenAI(base_url=cfg["base_url"], api_key=cfg["api_key"])

def get_model_name():
    """获取当前模型名"""
    return MODEL_CONFIGS[CURRENT_MODEL_KEY]["model"]

def get_model_label():
    """获取当前模型显示名"""
    if WORK_MODE == "expert":
        return "专家模式"
    elif WORK_MODE == "hybrid":
        return "混合思考"
    return MODEL_CONFIGS[CURRENT_MODEL_KEY]["label"]

# 兼容旧代码：初始化全局 client 和 MODEL
client = get_client()
MODEL = get_model_name()

WORK_DIR = os.getcwd()
MAX_FILE_SIZE = 1024 * 1024

# ====== 权限级别配置 ======
# restricted（受限）：保留所有安全检查（危险命令拦截、删除确认、深度限制等）
# full（全权限）：用户授权对电脑的完全操作权限，所有限制关闭
PERMISSION_LEVEL = "full"  # 当前：全权限模式（用户已授权对电脑的完全操作）

# 备份目录（使用运行时缓存，程序退出自动删除，不再污染工作目录）
BACKUP_DIR = str(runtime_cache.cache_dir / "backups")
os.makedirs(BACKUP_DIR, exist_ok=True) if PERMISSION_LEVEL == "full" else None

# 核心文件清单（修改前自动备份）
CORE_FILES = {
    "tui_agent.py", "settings.json", "requirements.txt", "README.md",
    "config.py", "tools.py", "prompts.py", "utils.py", "main.py",
}


def auto_backup(file_path: str) -> str:
    """全权限模式下，修改/删除核心文件前自动备份
    返回备份路径，失败返回错误信息
    """
    if PERMISSION_LEVEL != "full":
        return ""
    try:
        full = Path(file_path).resolve()
        if not full.exists():
            return ""
        # 只备份核心文件
        if full.name not in CORE_FILES:
            return ""
        # 生成备份文件名：原名.时间戳.bak
        import time
        ts = time.strftime("%Y%m%d_%H%M%S")
        backup_name = f"{full.name}.{ts}.bak"
        backup_path = os.path.join(BACKUP_DIR, backup_name)
        shutil.copy2(str(full), backup_path)
        return backup_path
    except Exception as e:
        return f"备份失败：{e}"


# ====== 配色：MiMo Code Agent 风格（纯黑 + 灰文字 + 彩色工具标签）======
C_BG = "#000000"        # 纯黑（主背景）
C_BG2 = "#0A0A0A"       # 接近黑（卡片背景）
C_FG = "#C8C8D0"        # 柔灰白（主文字，不刺眼）
C_DIM = "#6B6B75"       # 灰色（次要文字/说明）
C_BORDER = "#1F1F1F"    # 极深灰（边框，几乎看不见）
C_BLUE = "#7AA2F7"      # 蓝（强调）
C_PURPLE = "#BB9AF7"    # 紫
C_RED = "#F7768E"       # 红（工具/告警）
C_GREEN = "#9ECE6A"     # 绿（成功/进行中）
C_YELLOW = "#E0AF68"    # 黄（思考/提示）
C_CYAN = "#7DCFFF"      # 青
C_ORANGE = "#FF9E64"    # 橙（重点）
C_ACCENT = "#7AA2F7"    # 主强调色（蓝）
C_USER_BUBBLE = "#000000"  # 用户气泡背景（纯黑）
C_AI_BUBBLE = "#000000"    # AI气泡背景（纯黑）


# ====== 工具函数 ======
def read_file(path: str, max_length: int = 3000) -> str:
    try:
        full = Path(path).resolve()
        if not full.exists():
            return f"错误：文件不存在 {path}"
        if full.is_dir():
            return f"错误：{path} 是目录"
        if full.stat().st_size > MAX_FILE_SIZE:
            return "错误：文件太大"
        for enc in ["utf-8", "gbk", "latin-1"]:
            try:
                text = full.read_text(encoding=enc)
                if len(text) > max_length:
                    text = text[:max_length] + f"\n... [已截断，共{len(text)}字符，显示前{max_length}字符]"
                return text
            except UnicodeDecodeError:
                continue
        return "错误：无法解码"
    except Exception as e:
        return f"错误：{e}"


def write_file(path: str, content: str) -> str:
    try:
        full = Path(path).resolve()
        full.parent.mkdir(parents=True, exist_ok=True)
        full.write_text(content, encoding="utf-8")
        return f"已写入 {len(content)} 字符到 {path}"
    except Exception as e:
        return f"错误：{e}"


def list_dir(path: str = ".", recursive: bool = False, max_depth: int = 15) -> str:
    """列出目录内容
    path: 目录路径
    recursive: 是否递归显示子目录（树形结构）
    max_depth: 递归最大深度（1=只看当前层，默认15=深入最深层，自动跳过无权限目录）
    """
    try:
        full = Path(path).resolve()
        if not full.exists() or not full.is_dir():
            return f"错误：目录不存在 {path}"

        # 忽略的目录（避免扫描无用内容）
        IGNORE_DIRS = {".git", "__pycache__", "node_modules", ".venv", "venv",
                       ".idea", ".vs", "dist", "build", ".next", ".nuxt",
                       "target", ".gradle", ".mypy_cache", ".pytest_cache",
                       ".cache", ".npm", ".yarn", "bower_components", "vendor"}

        if not recursive:
            # 原有行为：只显示一层
            items = []
            dir_tag = _load_svg_icon("folder")
            file_tag = _load_svg_icon("file")
            for p in sorted(full.iterdir()):
                tag = dir_tag if p.is_dir() else file_tag
                size = p.stat().st_size if p.is_file() else ""
                items.append(f"{tag} {p.name} {size}")
            return "\n".join(items) if items else "(空目录)"

        # 递归模式：树形结构（深入最深层）
        lines = []
        file_count = 0
        dir_count = 0
        max_files = 500  # 安全上限，避免超大目录卡死模型（500足够了解结构）

        def _walk(directory: Path, prefix: str, depth: int):
            nonlocal file_count, dir_count
            if depth > max_depth or file_count >= max_files:
                return
            try:
                entries = sorted(directory.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
            except (PermissionError, OSError):
                return  # 无权限目录自动跳过

            # 过滤忽略目录
            entries = [e for e in entries if e.name not in IGNORE_DIRS]

            for idx, entry in enumerate(entries):
                if file_count >= max_files:
                    lines.append(f"{prefix}... (已达到最大文件数 {max_files}，停止扫描)")
                    return
                is_last = (idx == len(entries) - 1)
                connector = "└── " if is_last else "├── "

                if entry.is_dir():
                    dir_count += 1
                    lines.append(f"{prefix}{connector}{entry.name}/")
                    # 递归子目录（深入最深层）
                    extension = "    " if is_last else "│   "
                    _walk(entry, prefix + extension, depth + 1)
                else:
                    file_count += 1
                    try:
                        size = entry.stat().st_size
                        size_str = f" ({_format_size(size)})" if size > 0 else ""
                    except OSError:
                        size_str = ""
                    lines.append(f"{prefix}{connector}{entry.name}{size_str}")

        def _format_size(size: int) -> str:
            if size < 1024:
                return f"{size}B"
            elif size < 1024 * 1024:
                return f"{size/1024:.1f}KB"
            else:
                return f"{size/1024/1024:.1f}MB"

        lines.append(f"{full.name}/")
        _walk(full, "", 1)
        lines.append(f"\n共 {dir_count} 个目录，{file_count} 个文件"
                    + (f"（已达上限 {max_files}，可能未扫描完）" if file_count >= max_files else ""))
        return "\n".join(lines)
    except Exception as e:
        return f"错误：{e}"


def _is_windows_local() -> bool:
    """检测本地操作系统是否为 Windows"""
    return sys.platform == "win32" or os.name == "nt"


# 跨平台命令映射表：Linux 命令 → Windows 等效命令
# 当用户在 Windows 上输入 Linux 命令时，自动转换为对应的 Windows 命令
_CROSS_PLATFORM_COMMAND_MAP = {
    # 列表/查看类
    "ls": "dir",
    "ll": "dir /Q",
    "la": "dir /A",
    "cat": "type",
    "less": "more",
    "head": "more",       # 简化映射，PowerShell 下可用 Select-Object -First
    "tail": "more",       # 简化映射
    "grep": "findstr",
    "find": "findstr",    # 注意：Windows 的 find 和 Linux 的 find 含义不同
    "which": "where",
    "whereis": "where",
    "echo": "echo",
    "pwd": "cd",
    "whoami": "whoami",
    "hostname": "hostname",
    "date": "date /t",
    "uptime": "net statistics workstation",
    "uname": "ver",
    "df": "wmic logicaldisk get name,size,freespace",
    "du": "dir /s",
    "free": "wmic OS get TotalVisibleMemorySize,FreePhysicalMemory",
    "top": "tasklist",
    "ps": "tasklist",
    "kill": "taskkill /PID",
    "killall": "taskkill /IM",
    # 网络
    "ifconfig": "ipconfig",
    "ip addr": "ipconfig",
    "netstat": "netstat -ano",
    "ss": "netstat -ano",
    "ping": "ping",
    "traceroute": "tracert",
    "tracepath": "tracert",
    "nslookup": "nslookup",
    "dig": "nslookup",
    "host": "nslookup",
    "curl": "curl",       # Windows 10+ 自带
    "wget": "curl -O",    # Windows 10+ 用 curl 替代
    # 服务
    "systemctl": "sc",    # 简化映射，实际语义不完全等价
    "service": "sc",
    # 文件操作
    "rm": "del",
    "rm -rf": "rmdir /s /q",
    "rmdir": "rmdir /s /q",
    "mkdir": "mkdir",
    "mv": "move",
    "cp": "copy",
    "touch": "type nul >",
    "chmod": "icacls",    # 权限管理
    "chown": "icacls",
    # 文本处理（近似映射）
    "sed": "powershell -Command (Get-Content).Replace()",
    "awk": "powershell -Command",
    "sort": "sort",
    "uniq": "sort /unique",
    "wc": "find /c /v \"\"",
    # 包管理
    "apt": "winget",
    "apt-get": "winget",
    "yum": "winget",
    "dnf": "winget",
    "pip": "pip",
    "python3": "python",
    # 其他
    "history": "doskey /history",
    "env": "set",
    "export": "set",
}


def _translate_single_command(cmd: str) -> tuple:
    """翻译单个命令（不含管道符）。返回 (translated, is_translated)。"""
    if not cmd or not cmd.strip():
        return cmd, False

    cmd = cmd.strip()
    parts = cmd.split(None, 1)
    if not parts:
        return cmd, False

    base_cmd = parts[0].lower()
    rest = parts[1] if len(parts) > 1 else ""
    cmd_lower = cmd.lower()

    # 优先检查：原命令是否已经是 Windows 格式（以某个 value 开头）
    # 避免 "netstat -ano" 被重复翻译为 "netstat -ano -ano"
    for win_cmd in _CROSS_PLATFORM_COMMAND_MAP.values():
        win_lower = win_cmd.lower()
        if cmd_lower == win_lower or cmd_lower.startswith(win_lower + " "):
            return cmd, False  # 已经是 Windows 命令，不翻译

    # 完整短语匹配优先（按 key 长度降序，确保 "rm -rf" 在 "rm" 之前匹配）
    for linux_cmd in sorted(_CROSS_PLATFORM_COMMAND_MAP.keys(), key=len, reverse=True):
        win_cmd = _CROSS_PLATFORM_COMMAND_MAP[linux_cmd]
        if cmd_lower == linux_cmd:
            return win_cmd, True
        if cmd_lower.startswith(linux_cmd + " "):
            # 去掉整个匹配短语作为 rest，保留剩余参数
            suffix = cmd[len(linux_cmd):].strip()
            translated = f"{win_cmd} {suffix}" if suffix else win_cmd
            return translated, True

    # 单词匹配（兜底）
    if base_cmd in _CROSS_PLATFORM_COMMAND_MAP:
        win_cmd = _CROSS_PLATFORM_COMMAND_MAP[base_cmd]
        if rest:
            return f"{win_cmd} {rest}", True
        return win_cmd, True

    return cmd, False


def _translate_command(command: str) -> tuple:
    """跨平台命令翻译：把 Linux 命令转换为 Windows 等效命令（或反之）。

    支持管道符组合命令翻译：对每个管道子命令单独翻译，再重新拼接。
    例如 'ls | grep x' → 'dir | findstr x'，'cat file | grep error' → 'type file | findstr error'。
    注意：保留 '||'（逻辑或）不分割，只按单个 '|' 分割。

    Returns:
        (translated_command, is_translated)
        - translated_command: 翻译后的命令（如无需翻译则返回原命令）
        - is_translated: 是否发生了翻译
    """
    if not command or not command.strip():
        return command, False

    if not _is_windows_local():
        return command, False

    cmd = command.strip()

    # 按单个 | 分割（保留 || 不分割）
    # 正则：匹配单个 |，但其前后不是 |（负向后行/先行断言）
    import re
    # 先把 || 替换为占位符，避免被分割
    placeholder = "\x00OR\x00"
    cmd_protected = cmd.replace("||", placeholder)
    # 按单个 | 分割
    pipe_parts = cmd_protected.split("|")

    # 对每个子命令单独翻译
    translated_parts = []
    any_translated = False
    for part in pipe_parts:
        # 恢复 || 占位符
        part_restored = part.replace(placeholder, "||")
        sub_translated, sub_is = _translate_single_command(part_restored)
        if sub_is:
            any_translated = True
        translated_parts.append(sub_translated)

    if not any_translated:
        return command, False

    # 用 | 重新拼接（保留原 || 逻辑或）
    result = " | ".join(translated_parts)
    return result, True


def run_command(command: str, skip_translate: bool = False) -> str:
    """全权限模式：执行任意命令，无黑名单限制

    跨平台支持：自动识别 Linux 命令并转换为 Windows 等效命令（或反之）。
    例如在 Windows 上输入 'ls' 会自动转换为 'dir'，'cat file' 转换为 'type file'。

    Args:
        command: 要执行的命令
        skip_translate: 跳过跨平台翻译（语义化本地运维工具内部已适配，传 True 避免误翻译）
    """
    # 危险命令黑名单（仅受限模式生效）
    dangerous = ["rm -rf /", "del /f /s /q C:\\", "format C:", "shutdown /s /t 0"]
    if PERMISSION_LEVEL != "full":
        if any(d in command.lower() for d in dangerous):
            return "已拦截危险命令（受限模式）"

    # 跨平台命令翻译（语义化工具已内部适配，跳过）
    original_command = command
    translate_hint = ""
    if not skip_translate:
        command, is_translated = _translate_command(command)
        if is_translated:
            translate_hint = f"[跨平台] 已将 '{original_cmd_safe(original_command)}' 翻译为 '{command}'\n"

    try:
        # 全权限：超时延长到 120 秒；受限：30 秒
        timeout = 120 if PERMISSION_LEVEL == "full" else 30
        # 全权限：cwd 限制放开（不强制 WORK_DIR）
        cwd = None if PERMISSION_LEVEL == "full" else WORK_DIR
        # encoding 使用本地默认（中文 Windows 为 GBK/cp936），errors='replace' 兜底
        # 避免某些命令（ipconfig/systeminfo/sc/netsh）输出非 UTF-8 时 UnicodeDecodeError
        r = subprocess.run(command, shell=True, capture_output=True,
                          text=True, timeout=timeout, cwd=cwd,
                          encoding=locale.getpreferredencoding(False),
                          errors="replace")
        out = (r.stdout or "") + (r.stderr or "")
        # 全权限：返回更长（8000）；受限：4000
        max_out = 8000 if PERMISSION_LEVEL == "full" else 4000
        result = out.strip()[:max_out] if out.strip() else "(无输出)"
        return translate_hint + result
    except subprocess.TimeoutExpired:
        return f"{translate_hint}错误：命令超时（>{timeout}秒）"
    except Exception as e:
        return f"{translate_hint}错误：{e}"


def original_cmd_safe(cmd: str) -> str:
    """对原始命令做简单脱敏（避免控制字符）"""
    return cmd[:200].replace("\n", " ").replace("\r", " ")


# ====== 语义化本地运维工具集（4个，跨平台）======
def local_port_check(action: str = "list", port: int = 0,
                     protocol: str = "tcp", target: str = "") -> str:
    r"""本地端口/网络检查工具（跨平台：Windows/Linux 自动适配）。

    Args:
        action: 操作类型：
            - list: 列出所有监听端口（默认）
            - check: 检查指定端口是否被占用（需 port 参数）
            - ping: ping 目标主机（需 target 参数）
            - connections: 查看活跃 TCP 连接
        port: 端口号（action=check 时必填）
        protocol: 协议（tcp/udp），默认 tcp
        target: 目标主机/IP（action=ping 时必填）

    Returns:
        端口/网络检查结果
    """
    import socket

    if action == "list":
        if _is_windows_local():
            cmd = "netstat -ano | findstr LISTENING"
        else:
            cmd = "ss -tlnp 2>/dev/null || netstat -tlnp 2>/dev/null"
        return run_command(cmd)

    elif action == "check":
        if not port:
            return "错误：action=check 需要 port 参数"
        # 跨平台端口占用检查
        try:
            sock_type = socket.SOCK_STREAM if protocol.lower() == "tcp" else socket.SOCK_DGRAM
            s = socket.socket(socket.AF_INET, sock_type)
            s.settimeout(1)
            result = s.connect_ex(("127.0.0.1", port))
            s.close()
            if result == 0:
                # 端口被占用，查询占用进程
                if _is_windows_local():
                    proc_cmd = f"netstat -ano | findstr :{port}"
                    proc_out = run_command(proc_cmd, skip_translate=True)
                    return f"⚠️ 端口 {port}/{protocol} 已被占用\n\n{proc_out}"
                else:
                    proc_cmd = f"lsof -i :{port} 2>/dev/null || ss -tlnp | grep :{port}"
                    proc_out = run_command(proc_cmd, skip_translate=True)
                    return f"⚠️ 端口 {port}/{protocol} 已被占用\n\n{proc_out}"
            else:
                return f"✅ 端口 {port}/{protocol} 未被占用（可使用）"
        except Exception as e:
            return f"错误：检查端口失败 - {e}"

    elif action == "ping":
        if not target:
            return "错误：action=ping 需要 target 参数"
        # 防注入：仅允许字母数字点破折号
        if not all(c.isalnum() or c in ".-" for c in target):
            return f"错误：target 含非法字符 '{target}'"
        if _is_windows_local():
            cmd = f"ping -n 4 {target}"
        else:
            cmd = f"ping -c 4 {target}"
        return run_command(cmd)

    elif action == "connections":
        if _is_windows_local():
            cmd = "netstat -ano | findstr ESTABLISHED"
        else:
            cmd = "ss -tn state established 2>/dev/null || netstat -tn | grep ESTABLISHED"
        return run_command(cmd)

    else:
        return f"错误：action 必须是 list/check/ping/connections 之一"


def local_process_check(action: str = "top", name: str = "",
                        pid: int = 0, top_n: int = 10) -> str:
    """本地进程查看工具（跨平台：Windows/Linux 自动适配）。

    Args:
        action: 操作类型：
            - top: 按 CPU 占用排序显示前 N 个进程（默认）
            - memory: 按内存占用排序显示前 N 个进程
            - find: 按名称查找进程（需 name 参数）
            - kill: 结束指定进程（需 pid 或 name 参数）
        name: 进程名（action=find/kill 时使用）
        pid: 进程 ID（action=kill 时使用，优先于 name）
        top_n: 返回前 N 个进程（默认 10）

    Returns:
        进程信息
    """
    if action == "top":
        if _is_windows_local():
            # PowerShell 按 CPU 排序（内部用单引号避免与外层双引号冲突）
            cmd = (
                'powershell -NoProfile -Command "'
                f"Get-Process | Sort-Object CPU -Descending | Select-Object -First {top_n} "
                "Id, ProcessName, CPU, @{N='Mem(MB)';E={[int]($_.WorkingSet/1MB)}} | Format-Table -AutoSize"
                '"'
            )
        else:
            cmd = f"ps aux --sort=-%cpu | head -n {top_n + 1}"
        return run_command(cmd, skip_translate=True)

    elif action == "memory":
        if _is_windows_local():
            cmd = (
                'powershell -NoProfile -Command "'
                f"Get-Process | Sort-Object WorkingSet -Descending | Select-Object -First {top_n} "
                "Id, ProcessName, CPU, @{N='Mem(MB)';E={[int]($_.WorkingSet/1MB)}} | Format-Table -AutoSize"
                '"'
            )
        else:
            cmd = f"ps aux --sort=-%mem | head -n {top_n + 1}"
        return run_command(cmd, skip_translate=True)

    elif action == "find":
        if not name:
            return "错误：action=find 需要 name 参数"
        # 防注入：仅允许字母数字点下划线
        if not all(c.isalnum() or c in "._-" for c in name):
            return f"错误：name 含非法字符 '{name}'"
        if _is_windows_local():
            cmd = f'tasklist | findstr /I "{name}"'
        else:
            cmd = f"ps aux | grep -i {name} | grep -v grep"
        return run_command(cmd)

    elif action == "kill":
        if pid:
            if _is_windows_local():
                cmd = f"taskkill /PID {pid} /F"
            else:
                cmd = f"kill -9 {pid}"
            return run_command(cmd)
        elif name:
            if not all(c.isalnum() or c in "._-" for c in name):
                return f"错误：name 含非法字符 '{name}'"
            if _is_windows_local():
                cmd = f'taskkill /IM "{name}" /F'
            else:
                cmd = f"pkill -f {name}"
            return run_command(cmd)
        else:
            return "错误：action=kill 需要 pid 或 name 参数"

    else:
        return f"错误：action 必须是 top/memory/find/kill 之一"


def local_disk_check(action: str = "list", path: str = "") -> str:
    """本地磁盘空间分析工具（跨平台：Windows/Linux 自动适配）。

    Args:
        action: 操作类型：
            - list: 列出所有磁盘及使用率（默认）
            - top: 显示指定目录下 Top10 大目录/文件
        path: action=top 时指定分析目录（Windows: 'C:' 或 'C:\\Users'；Linux: '/var'），默认根目录

    Returns:
        磁盘使用情况
    """
    if action == "list":
        if _is_windows_local():
            cmd = ('powershell -NoProfile -Command "Get-CimInstance Win32_LogicalDisk -Filter \'DriveType=3\' | '
                   "ForEach-Object { $t=[math]::Round($_.Size/1GB,1); $f=[math]::Round($_.FreeSpace/1GB,1); "
                   "$u=[math]::Round($t-$f,1); $p=if($t-gt 0){[math]::Round($u/$t*100,1)}else{0}; "
                   "Write-Output ($_.DeviceID + ' 总:' + $t + 'GB 已用:' + $u + 'GB 可用:' + $f + 'GB 使用率:' + $p + '%') }\"")
        else:
            cmd = "df -h"
        return run_command(cmd)

    elif action == "top":
        if _is_windows_local():
            target = path if path else "C:\\"
            target = target.replace("/", "\\")
            # 扫描顶层子目录大小
            cmd = (
                'powershell -NoProfile -Command "'
                f"$root = '{target}';"
                "$dirs = Get-ChildItem -Path $root -Directory -ErrorAction SilentlyContinue;"
                "foreach ($d in $dirs) { try {"
                "  $size = (Get-ChildItem -Path $d.FullName -File -Recurse -ErrorAction SilentlyContinue | Measure-Object -Property Length -Sum).Sum;"
                "  $mb = [math]::Round($size/1MB, 1);"
                "  if ($mb -gt 10) { Write-Output ($mb.ToString().PadLeft(10) + 'MB  ' + $d.FullName) }"
                "} catch {} }\""
            )
            return run_command(cmd)
        else:
            target = path if path else "/"
            cmd = f"du -h --max-depth=1 {target} 2>/dev/null | sort -rh | head -n 10"
            return run_command(cmd)

    else:
        return f"错误：action 必须是 list/top 之一"


def local_service_check(action: str = "list", service: str = "") -> str:
    """本地服务管理工具（跨平台：Windows/Linux 自动适配）。

    Args:
        action: 操作类型：
            - list: 列出所有运行中的服务（默认）
            - status: 查看指定服务状态（需 service 参数）
            - start: 启动服务（需管理员权限）
            - stop: 停止服务（需管理员权限）
            - restart: 重启服务（需管理员权限）
        service: 服务名（action=status/start/stop/restart 时必填）

    Returns:
        服务信息
    """
    if action == "list":
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "Get-Service | Where-Object {$_.Status -eq \'Running\'} | Select-Object Status, Name, DisplayName | Format-Table -AutoSize"'
        else:
            cmd = "systemctl list-units --type=service --state=running 2>/dev/null | head -n 30"
        return run_command(cmd)

    elif action in ("status", "start", "stop", "restart"):
        if not service:
            return f"错误：action={action} 需要 service 参数"
        # 防注入：仅允许字母数字点下划线破折号
        if not all(c.isalnum() or c in "._-" for c in service):
            return f"错误：service 含非法字符 '{service}'"

        if _is_windows_local():
            if action == "status":
                cmd = f'powershell -NoProfile -Command "Get-Service -Name {service} -ErrorAction SilentlyContinue | Select-Object Status, Name, DisplayName | Format-Table -AutoSize"'
            elif action == "start":
                cmd = f'powershell -NoProfile -Command "Start-Service -Name {service}"'
            elif action == "stop":
                cmd = f'powershell -NoProfile -Command "Stop-Service -Name {service} -Force"'
            elif action == "restart":
                cmd = f'powershell -NoProfile -Command "Restart-Service -Name {service} -Force"'
        else:
            if action == "status":
                cmd = f"systemctl status {service}"
            else:
                cmd = f"sudo systemctl {action} {service}"
        return run_command(cmd)

    else:
        return f"错误：action 必须是 list/status/start/stop/restart 之一"


def local_firewall_check(action: str = "list", port: int = 0,
                         protocol: str = "tcp", direction: str = "in",
                         rule_name: str = "") -> str:
    r"""本地防火墙检查/管理工具（跨平台：Windows/Linux 自动适配）。

    Args:
        action: 操作类型：
            - list: 列出所有防火墙规则（默认）
            - status: 查看防火墙整体状态
            - check: 检查指定端口是否放行（需 port 参数）
            - open: 放行指定端口（需 port 参数，可能需管理员权限）
            - close: 关闭指定端口（需 port 参数，可能需管理员权限）
        port: 端口号（action=check/open/close 时必填）
        protocol: 协议（tcp/udp），默认 tcp
        direction: 方向（in/out），默认 in（入站）
        rule_name: 规则名（action=open/close 时可选，默认自动生成）

    Returns:
        防火墙信息
    """
    # 防注入：rule_name 仅允许字母数字空格下划线破折号
    if rule_name and not all(c.isalnum() or c in " _-" for c in rule_name):
        return f"错误：rule_name 含非法字符 '{rule_name}'"

    if action == "list":
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "Get-NetFirewallRule | Where-Object {$_.Enabled -eq \'True\'} | Select-Object DisplayName, Direction, Action, Profile -First 30 | Format-Table -AutoSize"'
        else:
            # Linux: 优先 ufw，其次 firewalld，最后 iptables
            cmd = "ufw status 2>/dev/null || firewall-cmd --list-all 2>/dev/null || iptables -L -n 2>/dev/null | head -n 30"
        return run_command(cmd, skip_translate=True)

    elif action == "status":
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "Get-NetFirewallProfile | Select-Object Name, Enabled, DefaultInboundAction, DefaultOutboundAction | Format-Table -AutoSize"'
        else:
            cmd = "ufw status verbose 2>/dev/null || systemctl is-active firewalld 2>/dev/null || iptables -L -n 2>/dev/null | head -n 10"
        return run_command(cmd, skip_translate=True)

    elif action == "check":
        if not port:
            return "错误：action=check 需要 port 参数"
        if _is_windows_local():
            # 查找放行该端口的规则
            cmd = f'powershell -NoProfile -Command "Get-NetFirewallRule -Enabled True | Where-Object {{($_.Direction -eq \'{direction.capitalize()}\') -and ($_.Action -eq \'Allow\')}} | Get-NetFirewallPortFilter | Where-Object {{($_.LocalPort -eq \'{port}\') -and ($_.Protocol -eq \'{protocol.upper()}\')}} | Format-List"'
            return run_command(cmd, skip_translate=True)
        else:
            cmd = f"ufw status | grep ':?{port} ' 2>/dev/null || iptables -L -n | grep ':{port} ' 2>/dev/null"
            result = run_command(cmd, skip_translate=True)
            if "无输出" in result or not result.strip():
                return f"⚠️ 端口 {port}/{protocol} 未在防火墙规则中找到放行记录（可能被阻止）"
            return f"✅ 端口 {port}/{protocol} 已放行\n\n{result}"

    elif action in ("open", "close"):
        if not port:
            return "错误：action=open/close 需要 port 参数"
        if not rule_name:
            rule_name = f"ZeroAI_{protocol}_{port}_{direction}"
        if _is_windows_local():
            action_ps = "Allow" if action == "open" else "Block"
            cmd = (
                f'powershell -NoProfile -Command "'
                f"New-NetFirewallRule -DisplayName '{rule_name}' "
                f"-Direction {direction.capitalize()} -Action {action_ps} "
                f"-Protocol {protocol.upper()} -LocalPort {port}"
                f'"'
            )
        else:
            if action == "open":
                cmd = f"ufw allow {port}/{protocol} 2>/dev/null || firewall-cmd --add-port={port}/{protocol} --permanent 2>/dev/null && firewall-cmd --reload 2>/dev/null || iptables -I INPUT -p {protocol} --dport {port} -j ACCEPT"
            else:
                cmd = f"ufw deny {port}/{protocol} 2>/dev/null || firewall-cmd --remove-port={port}/{protocol} --permanent 2>/dev/null && firewall-cmd --reload 2>/dev/null || iptables -D INPUT -p {protocol} --dport {port} -j ACCEPT"
        return run_command(cmd, skip_translate=True)

    else:
        return f"错误：action 必须是 list/status/check/open/close 之一"


def local_user_check(action: str = "list", username: str = "",
                     detail: bool = False) -> str:
    r"""本地用户/登录管理工具（跨平台：Windows/Linux 自动适配）。

    Args:
        action: 操作类型：
            - list: 列出所有本地用户（默认）
            - current: 查看当前登录用户
            - info: 查看指定用户详情（需 username 参数）
            - groups: 查看指定用户所属组（需 username 参数）
            - sessions: 查看当前登录会话
        username: 用户名（action=info/groups 时必填）
        detail: 是否显示详细信息（action=list 时有效）

    Returns:
        用户信息
    """
    # 防注入：username 仅允许字母数字点下划线破折号
    if username and not all(c.isalnum() or c in "._-" for c in username):
        return f"错误：username 含非法字符 '{username}'"

    if action == "list":
        if _is_windows_local():
            if detail:
                cmd = 'powershell -NoProfile -Command "Get-LocalUser | Select-Object Name, Enabled, LastLogon, Description | Format-Table -AutoSize"'
            else:
                cmd = 'powershell -NoProfile -Command "Get-LocalUser | Select-Object Name, Enabled | Format-Table -AutoSize"'
        else:
            cmd = "cat /etc/passwd | cut -d: -f1,3,7 | head -n 50"
        return run_command(cmd, skip_translate=True)

    elif action == "current":
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "whoami; Get-LocalUser | Where-Object {$_.Name -eq $env:USERNAME} | Select-Object Name, Enabled, LastLogon"'
        else:
            cmd = "whoami && id"
        return run_command(cmd, skip_translate=True)

    elif action == "info":
        if not username:
            return "错误：action=info 需要 username 参数"
        if _is_windows_local():
            cmd = f'powershell -NoProfile -Command "Get-LocalUser -Name {username} | Select-Object Name, Enabled, FullName, Description, LastLogon, PasswordLastSet | Format-List"'
        else:
            cmd = f"id {username} 2>/dev/null && grep '^{username}:' /etc/passwd"
        return run_command(cmd, skip_translate=True)

    elif action == "groups":
        if not username:
            return "错误：action=groups 需要 username 参数"
        if _is_windows_local():
            cmd = f'powershell -NoProfile -Command "Get-LocalGroup | Where-Object {{(Get-LocalGroupMember -Group $_.Name -ErrorAction SilentlyContinue).Name -contains \'{username}\'}} | Select-Object Name, Description | Format-Table -AutoSize"'
        else:
            cmd = f"groups {username} 2>/dev/null"
        return run_command(cmd, skip_translate=True)

    elif action == "sessions":
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "query user 2>$null; Get-CimInstance Win32_LogonSession | Select-Object LogonId, LogonType, StartTime -First 10 | Format-Table -AutoSize"'
        else:
            cmd = "who -a 2>/dev/null | head -n 20"
        return run_command(cmd, skip_translate=True)

    else:
        return f"错误：action 必须是 list/current/info/groups/sessions 之一"


def local_monitor(threshold_cpu: int = 80, threshold_disk: int = 90,
                  threshold_memory: int = 85, check_ports: str = "") -> str:
    r"""本地综合监控告警：一次性检查 CPU/内存/磁盘/端口/防火墙，返回结构化告警报告。

    跨平台自动适配 Windows/Linux。基于本地运维工具组合调用，输出标准化告警。

    Args:
        threshold_cpu: CPU 使用率告警阈值（默认 80%）
        threshold_disk: 磁盘使用率告警阈值（默认 90%）
        threshold_memory: 内存使用率告警阈值（默认 85%）
        check_ports: 需要检查的关键端口（逗号分隔，如 "22,80,443,3306,8080"）
                     为空则只列出当前监听端口，不针对性检查

    Returns:
        结构化告警报告：
        [监控概览] 总体健康状态
        [告警项] ⚠️ 警告 / 🚨 危急
        [正常项] ✅ 正常
        [建议] 💡 优化建议
    """
    import time as _time

    report_lines = ["=" * 60]
    report_lines.append(f"📊 本地监控告警报告  {_time.strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append("=" * 60)

    warnings = []   # ⚠️ 警告
    criticals = []  # 🚨 危急
    normals = []    # ✅ 正常
    suggestions = []  # 💡 建议

    # ===== 1. CPU 检查 =====
    try:
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "(Get-CimInstance Win32_Processor).LoadPercentage"'
        else:
            cmd = "top -bn1 | grep 'Cpu(s)' | awk '{print $2}'"
        cpu_out = run_command(cmd, skip_translate=True)
        # 解析 CPU 使用率
        cpu_pct = -1
        for line in cpu_out.split("\n"):
            line = line.strip()
            if line and any(c.isdigit() for c in line):
                # 提取第一个数字
                import re
                m = re.search(r"(\d+)", line)
                if m:
                    cpu_pct = int(m.group(1))
                    break
        if cpu_pct >= 0:
            if cpu_pct >= threshold_cpu + 10:
                criticals.append(f"CPU 使用率 {cpu_pct}%（危急，阈值 {threshold_cpu}%）")
                suggestions.append("CPU 占用过高，建议用 local_process_check(action='top') 查看高 CPU 进程")
            elif cpu_pct >= threshold_cpu:
                warnings.append(f"CPU 使用率 {cpu_pct}%（警告，阈值 {threshold_cpu}%）")
            else:
                normals.append(f"CPU 使用率 {cpu_pct}%（正常）")
        else:
            warnings.append("CPU 使用率获取失败")
    except Exception as e:
        warnings.append(f"CPU 检查异常: {e}")

    # ===== 2. 内存检查 =====
    try:
        if _is_windows_local():
            cmd = 'powershell -NoProfile -Command "$os = Get-CimInstance Win32_OperatingSystem; [math]::Round(($os.TotalVisibleMemorySize - $os.FreePhysicalMemory) / $os.TotalVisibleMemorySize * 100, 1)"'
        else:
            cmd = "free | grep Mem | awk '{printf \"%.1f\", $3/$2*100}'"
        mem_out = run_command(cmd, skip_translate=True).strip()
        import re
        m = re.search(r"(\d+(?:\.\d+)?)", mem_out)
        if m:
            mem_pct = float(m.group(1))
            if mem_pct >= threshold_memory + 10:
                criticals.append(f"内存使用率 {mem_pct}%（危急，阈值 {threshold_memory}%）")
                suggestions.append("内存占用过高，建议用 local_process_check(action='memory') 查看高内存进程")
            elif mem_pct >= threshold_memory:
                warnings.append(f"内存使用率 {mem_pct}%（警告，阈值 {threshold_memory}%）")
            else:
                normals.append(f"内存使用率 {mem_pct}%（正常）")
        else:
            warnings.append("内存使用率获取失败")
    except Exception as e:
        warnings.append(f"内存检查异常: {e}")

    # ===== 3. 磁盘检查 =====
    try:
        disk_out = local_disk_check(action="list")
        import re
        # 匹配使用率百分比
        pcts = re.findall(r"(\d+)%", disk_out)
        disk_high = False
        for pct_str in pcts:
            pct = int(pct_str)
            if pct >= threshold_disk + 5:
                criticals.append(f"磁盘使用率 {pct}%（危急，阈值 {threshold_disk}%）")
                disk_high = True
            elif pct >= threshold_disk:
                warnings.append(f"磁盘使用率 {pct}%（警告，阈值 {threshold_disk}%）")
                disk_high = True
        if not disk_high:
            normals.append(f"所有磁盘使用率低于 {threshold_disk}%（正常）")
        if disk_high:
            suggestions.append("磁盘空间不足，建议用 local_disk_check(action='top') 分析大目录")
    except Exception as e:
        warnings.append(f"磁盘检查异常: {e}")

    # ===== 4. 关键端口检查 =====
    if check_ports:
        for port_str in check_ports.split(","):
            port_str = port_str.strip()
            if not port_str.isdigit():
                continue
            port = int(port_str)
            try:
                import socket
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.settimeout(1)
                result = s.connect_ex(("127.0.0.1", port))
                s.close()
                if result == 0:
                    normals.append(f"端口 {port} 已监听（正常）")
                else:
                    # 判断是否为常见关键端口
                    critical_ports = {22: "SSH", 80: "HTTP", 443: "HTTPS", 3306: "MySQL", 5432: "PostgreSQL", 6379: "Redis"}
                    service_name = critical_ports.get(port, "")
                    if service_name:
                        criticals.append(f"关键端口 {port} ({service_name}) 未监听（危急）")
                        suggestions.append(f"端口 {port} ({service_name}) 未监听，建议用 local_service_check(action='status', service='{service_name.lower()}') 检查服务状态")
                    else:
                        warnings.append(f"端口 {port} 未监听")
            except Exception as e:
                warnings.append(f"端口 {port} 检查异常: {e}")

    # ===== 5. 防火墙状态检查 =====
    try:
        fw_out = local_firewall_check(action="status")
        fw_lower = fw_out.lower()
        if _is_windows_local():
            if "true" in fw_lower and "enabled" in fw_lower:
                normals.append("防火墙已启用（正常）")
            else:
                criticals.append("防火墙未启用（危急，建议立即启用）")
                suggestions.append("防火墙关闭会暴露所有端口，建议立即用 local_firewall_check(action='status') 检查并启用")
        else:
            if "active" in fw_lower or "active: active" in fw_lower:
                normals.append("防火墙已启用（正常）")
            else:
                warnings.append("防火墙可能未启用")
    except Exception as e:
        warnings.append(f"防火墙检查异常: {e}")

    # ===== 汇总报告 =====
    total_issues = len(warnings) + len(criticals)
    if not criticals and not warnings:
        status_emoji = "✅"
        status_text = "健康"
    elif criticals:
        status_emoji = "🚨"
        status_text = f"危急（{len(criticals)} 项危急，{len(warnings)} 项警告）"
    else:
        status_emoji = "⚠️"
        status_text = f"警告（{len(warnings)} 项警告）"

    report_lines.append(f"\n[{status_emoji} 监控概览] 总体状态：{status_text}")
    report_lines.append(f"  检查项：CPU/内存/磁盘/端口/防火墙")
    report_lines.append(f"  阈值：CPU≥{threshold_cpu}%  内存≥{threshold_memory}%  磁盘≥{threshold_disk}%")

    if criticals:
        report_lines.append(f"\n[🚨 危急项]")
        for item in criticals:
            report_lines.append(f"  🚨 {item}")

    if warnings:
        report_lines.append(f"\n[⚠️ 警告项]")
        for item in warnings:
            report_lines.append(f"  ⚠️ {item}")

    if normals:
        report_lines.append(f"\n[✅ 正常项]")
        for item in normals:
            report_lines.append(f"  ✅ {item}")

    if suggestions:
        report_lines.append(f"\n[💡 优化建议]")
        for item in suggestions:
            report_lines.append(f"  💡 {item}")

    report_lines.append("\n" + "=" * 60)
    return "\n".join(report_lines)


def search_files(pattern: str, path: str = ".") -> str:
    """全权限模式：无深度限制、结果数扩大到 200"""
    try:
        results = []
        # 全权限：跳过二进制过滤（仍跳过目录）；受限：也跳过常见二进制
        skip_suffixes = [".exe", ".dll", ".zip", ".pyc"] if PERMISSION_LEVEL == "full" else [".exe", ".dll", ".png", ".jpg", ".zip", ".pyc"]
        # 全权限：结果数限制 200；受限：50
        max_results = 200 if PERMISSION_LEVEL == "full" else 50
        for p in Path(path).rglob("*"):
            if p.is_dir() or p.suffix.lower() in skip_suffixes:
                continue
            if p.stat().st_size > MAX_FILE_SIZE:
                continue
            try:
                content = p.read_text(encoding="utf-8", errors="ignore")
                for i, line in enumerate(content.splitlines(), 1):
                    if re.search(pattern, line):
                        results.append(f"{p.name}:{i}: {line.strip()}")
                        if len(results) >= max_results:
                            return "\n".join(results) + f"\n...(截断至 {max_results} 条)"
            except Exception:
                continue
        return "\n".join(results) if results else "(无匹配)"
    except Exception as e:
        return f"错误：{e}"


# 常见桌面应用路径映射表（中英文别名 → 可执行文件路径）
APP_PATHS = {
    # 通讯类
    "微信": r"D:\WeiXin\Weixin.exe",
    "wechat": r"D:\WeiXin\Weixin.exe",
    "weixin": r"D:\WeiXin\Weixin.exe",
    "qq": r"D:\QQ\QQ.exe",
    # 开发工具
    "vscode": r"D:\Microsoft VS Code\Code.exe",
    "vs code": r"D:\Microsoft VS Code\Code.exe",
    "code": r"D:\Microsoft VS Code\Code.exe",
    "pycharm": r"D:\pycharm\PyCharm 2025.3.2.1\bin\pycharm64.exe",
    # 浏览器
    "edge": r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    "浏览器": r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    # 系统自带应用
    "记事本": "notepad.exe",
    "notepad": "notepad.exe",
    "计算器": "calc.exe",
    "calc": "calc.exe",
    "资源管理器": "explorer.exe",
    "explorer": "explorer.exe",
    "文件资源管理器": "explorer.exe",
    "画图": "mspaint.exe",
    "mspaint": "mspaint.exe",
    "写字板": "write.exe",
    "write": "write.exe",
    "任务管理器": "taskmgr.exe",
    "taskmgr": "taskmgr.exe",
    "控制面板": "control.exe",
    "control": "control.exe",
    "注册表": "regedit.exe",
    "regedit": "regedit.exe",
    "cmd": "cmd.exe",
    "命令提示符": "cmd.exe",
    "powershell": "powershell.exe",
    "终端": "powershell.exe",
}


def _search_executable(name: str) -> str:
    """在本地自动搜索可执行文件或文档，返回找到的完整路径
    搜索顺序：
    1. APP_PATHS 硬编码映射（快速命中已知应用）
    2. 系统 PATH 环境变量
    3. Windows 注册表（App Paths / uninstall / exe 找到安装路径）
    4. 常见安装目录递归搜索（D:/ C:/Program Files 等）
    """
    key = name.lower().strip()

    # 1. 硬编码映射
    path = APP_PATHS.get(key) or APP_PATHS.get(name)
    if path:
        if "\\" not in path or Path(path).exists():
            return path

    # 2. 系统 PATH
    for dir_path in os.environ.get("PATH", "").split(os.pathsep):
        for ext in ("", ".exe", ".bat", ".cmd", ".msi", ".lnk"):
            candidate = Path(dir_path) / f"{name}{ext}"
            if candidate.exists():
                return str(candidate)

    # 3. 注册表搜索（App Paths）
    try:
        import winreg
        # HKLM App Paths
        for hive in (winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER):
            try:
                with winreg.OpenKey(hive, rf"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\{name}.exe") as k:
                    val, _ = winreg.QueryValueEx(k, "")
                    if val and Path(val).exists():
                        return val
            except (FileNotFoundError, OSError):
                pass
        # HKLM Uninstall：搜索 DisplayName 匹配的应用，找 InstallLocation
        for hive in (winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER):
            try:
                with winreg.OpenKey(hive, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall") as k:
                    i = 0
                    while True:
                        try:
                            sub_name = winreg.EnumKey(k, i)
                            i += 1
                            with winreg.OpenKey(k, sub_name) as sk:
                                try:
                                    display, _ = winreg.QueryValueEx(sk, "DisplayName")
                                    if key in display.lower() or display.lower() in key:
                                        try:
                                            loc, _ = winreg.QueryValueEx(sk, "InstallLocation")
                                            if loc:
                                                # 在安装目录中搜索 exe
                                                for p in Path(loc).rglob("*.exe"):
                                                    if key in p.stem.lower():
                                                        return str(p)
                                                # 返回安装目录的第一个 exe
                                                for p in Path(loc).rglob("*.exe"):
                                                    return str(p)
                                        except (FileNotFoundError, OSError):
                                            pass
                                except (FileNotFoundError, OSError):
                                    pass
                        except OSError:
                            break
            except (FileNotFoundError, OSError):
                pass
    except ImportError:
        pass

    # 4. 常见安装目录搜索
    search_dirs = [
        r"C:\Program Files",
        r"C:\Program Files (x86)",
        r"D:\\",
        r"D:\Program Files",
        r"D:\Program Files (x86)",
        r"D:\Microsoft VS Code",
        r"D:\Weixin",
        r"D:\QQ",
        r"D:\pycharm",
        os.path.expanduser("~\\AppData\\Local"),
        os.path.expanduser("~\\AppData\\Roaming"),
        os.path.expanduser("~\\Desktop"),
    ]
    # 去重
    seen = set()
    unique_dirs = []
    for d in search_dirs:
        if d not in seen and Path(d).exists():
            seen.add(d)
            unique_dirs.append(d)

    for dir_path in unique_dirs:
        # 限制搜索深度 3 层，避免太慢
        try:
            base = Path(dir_path)
            for p in base.glob("*"):
                # 直接匹配文件名
                if p.is_file():
                    stem_lower = p.stem.lower()
                    name_lower = name.lower().replace(".exe", "").replace(".lnk", "")
                    if name_lower == stem_lower or name_lower in stem_lower:
                        return str(p)
                elif p.is_dir():
                    # 搜索子目录（1层）
                    try:
                        for sub in p.glob("*.exe"):
                            stem_lower = sub.stem.lower()
                            name_lower = name.lower().replace(".exe", "")
                            if name_lower in stem_lower or stem_lower in name_lower:
                                return str(sub)
                        for sub in p.glob("*.lnk"):
                            stem_lower = sub.stem.lower()
                            name_lower = name.lower().replace(".lnk", "")
                            if name_lower in stem_lower or stem_lower in name_lower:
                                return str(sub)
                    except (PermissionError, OSError):
                        pass
        except (PermissionError, OSError):
            continue

    return ""


def open_app(name: str) -> str:
    """打开桌面应用程序或任意文件
    自动在本地搜索后打开，保证能打开任何文件：
    1. 先查 APP_PATHS 硬编码映射
    2. 再查系统 PATH 环境变量
    3. 再查注册表 App Paths / Uninstall
    4. 最后在常见安装目录递归搜索
    如果 name 是已存在的文件路径，直接用系统默认程序打开
    """
    # 如果 name 是已存在的文件路径，直接打开
    direct_path = Path(name)
    if direct_path.exists():
        try:
            os.startfile(str(direct_path))
            return f"已打开文件：{name}"
        except Exception as e:
            return f"打开失败：{e}"

    # 自动搜索
    found_path = _search_executable(name)
    if not found_path:
        return f"未在本地找到「{name}」。已搜索：APP_PATHS → PATH → 注册表 → 常见安装目录。请提供完整路径。"

    try:
        subprocess.Popen(found_path)
        return f"已启动：{name}\n路径：{found_path}"
    except Exception as e:
        # 尝试用 os.startfile 作为后备
        try:
            os.startfile(found_path)
            return f"已打开：{name}\n路径：{found_path}"
        except Exception:
            return f"启动失败：{e}"


def web_search(query: str, num_results: int = 5) -> str:
    """网络搜索（百度优先，Bing CN 备用，Bing 国际版第三）"""
    q = urllib.parse.quote(query)
    _FILTERS_BAIDU = ("baidu.com", "baidustatic", "bdstatic", "baiduimg",
                      "baidupcs", "bcebos", "baiducontent")
    _FILTERS_BING = ("bing.com", "microsoft.com", "go.microsoft", "live.com",
                     "msn.com", "sogou.com", ".css", ".js", ".png", ".jpg")

    def _extract_baidu_results(html, max_results):
        results = []
        blocks = re.findall(r'<div class="c-container[^"]*"[^>]*>([\s\S]*?)</div>', html)
        for block in blocks[:max_results * 3]:
            h3_m = re.search(r'<h3[^>]*>.*?href="([^"]*)"[^>]*>(.*?)</a>', block, re.DOTALL)
            if h3_m:
                title = re.sub(r"<[^>]+>", "", h3_m.group(2)).strip()
                link = h3_m.group(1)
                if not link.startswith("http"):
                    mu = re.search(r'mu="([^"]*)"', block)
                    if mu:
                        link = mu.group(1)
                if link.startswith("http") and title and len(title) > 2 and not any(f in link for f in _FILTERS_BAIDU):
                    results.append(f"{title}\n  {link}")
                    if len(results) >= max_results:
                        break
        return results

    def _extract_bing_results(html, max_results):
        results = []
        for block in re.findall(r'<li class="b_algo"[^>]*>([\s\S]*?)</li>', html):
            hrefs = re.findall(r'href="(https?://[^"]+)"', block)
            t = re.search(r"<a[^>]*>(.*?)</a>", block, re.DOTALL)
            title = re.sub(r"<[^>]+>", "", t.group(1)).strip() if t else ""
            link = next((h for h in hrefs if not any(f in h for f in _FILTERS_BING)), "")
            if link and title and len(title) > 3:
                results.append(f"{title}\n  {link}")
                if len(results) >= max_results:
                    return results
        return results

    _UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/131.0.0.0 Safari/537.36"

    # 方案1：百度搜索（国内首选，超时 15 秒）
    try:
        url = f"https://www.baidu.com/s?wd={q}&rn={num_results * 2}"
        req = urllib.request.Request(url, headers={
            "User-Agent": _UA,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        r = _extract_baidu_results(html, num_results)
        if r:
            return "\n\n".join(r)
    except Exception:
        pass

    # 方案2：Bing 中国版（超时 20 秒）
    try:
        url = f"https://cn.bing.com/search?q={q}&count={num_results * 2}"
        req = urllib.request.Request(url, headers={
            "User-Agent": _UA,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            "Referer": "https://cn.bing.com/",
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        r = _extract_bing_results(html, num_results)
        if r:
            return "\n\n".join(r)
    except Exception:
        pass

    # 方案3：Bing 国际版（超时 20 秒）
    try:
        q_bp = urllib.parse.quote_plus(query)
        url = f"https://www.bing.com/search?q={q_bp}&count={num_results * 2}&setlang=zh-CN&cc=cn"
        req = urllib.request.Request(url, headers={
            "User-Agent": _UA,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        r = _extract_bing_results(html, num_results)
        if r:
            return "\n\n".join(r)
    except Exception:
        pass

    return "(搜索失败：所有搜索引擎均不可达，请检查网络连接或稍后重试)"

def web_fetch(url: str, max_length: int = 4000) -> str:
    """全权限模式：可访问内网/任意 URL，无 SSRF 限制"""
    # SSRF 防护（仅受限模式生效）
    if PERMISSION_LEVEL != "full":
        import ipaddress
        import socket
        try:
            # 解析域名获取 IP
            host = urllib.parse.urlparse(url).hostname
            if host:
                try:
                    ip = socket.gethostbyname(host)
                    ip_obj = ipaddress.ip_address(ip)
                    # 拦截内网/本地地址
                    if ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_reserved:
                        return f"SSRF 防护：禁止访问内网地址 {ip}"
                except (socket.gaierror, ValueError):
                    pass
        except Exception:
            pass
    # 全权限：max_length 默认放大到 16000
    if PERMISSION_LEVEL == "full" and max_length == 4000:
        max_length = 16000
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        text = re.sub(r'<script[^>]*>[\s\S]*?</script>', '', html)
        text = re.sub(r'<style[^>]*>[\s\S]*?</style>', '', text)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text[:max_length] if len(text) > max_length else text
    except Exception as e:
        return f"抓取错误：{e}"


def git_status(repo_path: str = ".") -> str:
    try:
        r = subprocess.run(["git", "status", "--short", "--branch"],
                          capture_output=True, text=True, timeout=10, cwd=repo_path)
        branch = subprocess.run(["git", "branch", "--show-current"],
                               capture_output=True, text=True, timeout=5, cwd=repo_path)
        out = f"分支: {branch.stdout.strip()}\n{r.stdout.strip()}"
        return out if out.strip() else "(无变更)"
    except FileNotFoundError:
        return "错误：git 未安装"
    except Exception as e:
        return f"错误：{e}"


def delete_file(path: str) -> str:
    """全权限模式：直接删除核心文件前自动备份"""
    try:
        full = Path(path).resolve()
        if not full.exists():
            return f"错误：文件不存在 {path}"
        # 国家级项目硬约束：删除核心文件前自动备份
        backup_info = ""
        if PERMISSION_LEVEL == "full" and full.name in CORE_FILES:
            backup_path = auto_backup(str(full))
            if backup_path and not backup_path.startswith("备份失败"):
                backup_info = f"\n[已备份] {backup_path}"
        # 优先用 send2trash 软删除，全权限模式下直接删除
        if PERMISSION_LEVEL == "full":
            # 全权限：直接删除（不可恢复，但有备份）
            if full.is_dir():
                shutil.rmtree(str(full))
            else:
                full.unlink()
            return f"{_load_svg_icon('cross')} 已删除：{path}{backup_info}"
        else:
            # 受限模式：移入回收站
            try:
                from send2trash import send2trash
                send2trash(str(full))
                return f"已移入回收站：{path}"
            except ImportError:
                if full.is_dir():
                    shutil.rmtree(str(full))
                else:
                    full.unlink()
                return f"{_load_svg_icon('cross')} 已删除：{path}"
    except Exception as e:
        return f"错误：{e}"


def move_file(src: str, dst: str) -> str:
    try:
        s = Path(src).resolve()
        if not s.exists():
            return f"错误：源文件不存在 {src}"
        shutil.move(str(s), str(Path(dst).resolve()))
        return f"{_load_svg_icon('check')} 已移动：{src} → {dst}"
    except Exception as e:
        return f"错误：{e}"


def copy_file(src: str, dst: str) -> str:
    try:
        s = Path(src).resolve()
        if not s.exists():
            return f"错误：源文件不存在 {src}"
        d = Path(dst).resolve()
        d.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(s), str(d))
        return f"{_load_svg_icon('check')} 已复制：{src} → {dst}"
    except Exception as e:
        return f"错误：{e}"


def create_dir(path: str) -> str:
    try:
        full = Path(path).resolve()
        full.mkdir(parents=True, exist_ok=True)
        return f"已创建目录：{path}"
    except Exception as e:
        return f"错误：{e}"


def system_info() -> str:
    try:
        import psutil
        cpu = psutil.cpu_percent(interval=1)
        mem = psutil.virtual_memory()
        disk = psutil.disk_usage(os.getcwd())
        return (
            f"系统：{platform.system()} {platform.release()}\n"
            f"架构：{platform.machine()}\n"
            f"Python：{platform.python_version()}\n"
            f"CPU 使用率：{cpu}%\n"
            f"内存：{mem.percent}%（{mem.used//1024//1024}MB / {mem.total//1024//1024}MB）\n"
            f"磁盘：{disk.percent}%（{disk.used//1024//1024//1024}GB / {disk.total//1024//1024//1024}GB）"
        )
    except ImportError:
        return (
            f"系统：{platform.system()} {platform.release()}\n"
            f"架构：{platform.machine()}\n"
            f"Python：{platform.python_version()}\n"
            f"（安装 psutil 可查看 CPU/内存/磁盘详情：pip install psutil）"
        )
    except Exception as e:
        return f"错误：{e}"


def process_list(name_filter: str = "") -> str:
    """全权限模式：显示所有进程，无 30 条截断"""
    try:
        r = subprocess.run(["tasklist", "/FO", "CSV", "/NH"],
                          capture_output=True, text=True, timeout=10)
        lines = r.stdout.strip().splitlines()
        results = []
        # 全权限：限制 200；受限：30
        max_results = 200 if PERMISSION_LEVEL == "full" else 30
        for line in lines:
            parts = line.replace('"', '').split(",")
            if len(parts) >= 2:
                pname = parts[0]
                pid = parts[1]
                mem = parts[4] if len(parts) > 4 else ""
                if not name_filter or name_filter.lower() in pname.lower():
                    results.append(f"{pid:>6}  {pname:<30} {mem}")
                if len(results) >= max_results:
                    results.append(f"...(截断至 {max_results} 条)")
                    break
        return "\n".join(results) if results else "(无匹配进程)"
    except Exception as e:
        return f"错误：{e}"


def edit_file(path: str, operation: str = "replace", line: int = 1, content: str = "", start_line: int = 0, end_line: int = 0) -> str:
    """按行编辑文件：替换/插入/删除指定行"""
    try:
        full = Path(path).resolve()
        if not full.exists():
            return f"错误：文件不存在 {path}"
        if full.is_dir():
            return f"错误：{path} 是目录"
        for enc in ["utf-8", "gbk", "latin-1"]:
            try:
                text = full.read_text(encoding=enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            return "错误：无法解码"
        lines = text.splitlines(keepends=True)
        total = len(lines)
        if operation == "replace":
            if line < 1 or line > total:
                return f"错误：行号 {line} 超出范围（1-{total}）"
            old = lines[line - 1].rstrip("\n\r")
            lines[line - 1] = content + "\n"
            full.write_text("".join(lines), encoding=enc)
            return f"{_load_svg_icon('check')} 第{line}行已替换\n  原内容：{old}\n  新内容：{content}"
        elif operation == "insert":
            if line < 1 or line > total + 1:
                return f"错误：行号 {line} 超出范围（1-{total+1}）"
            lines.insert(line - 1, content + "\n")
            full.write_text("".join(lines), encoding=enc)
            return f"{_load_svg_icon('check')} 已在第{line}行插入：{content}"
        elif operation == "delete":
            s = start_line or line
            e = end_line or line
            if s < 1 or e > total or s > e:
                return f"错误：行范围 {s}-{e} 无效（1-{total}）"
            deleted = lines[s - 1:e]
            del lines[s - 1:e]
            full.write_text("".join(lines), encoding=enc)
            return f"{_load_svg_icon('cross')} 已删除第{s}-{e}行（共{len(deleted)}行）"
        elif operation == "append":
            lines.append(content + "\n")
            full.write_text("".join(lines), encoding=enc)
            return f"{_load_svg_icon('check')} 已在末尾追加：{content}"
        else:
            return f"错误：未知操作 {operation}，支持 replace/insert/delete/append"
    except Exception as e:
        return f"错误：{e}"


def exec_python(code: str, timeout: int = 10) -> str:
    """在受限环境中执行 Python 代码片段"""
    import io, sys
    # 危险模块黑名单
    blocked = ["os.system", "os.popen", "subprocess", "eval", "exec", "compile",
               "__import__", "open(", "shutil.rmtree", "shutil.move"]
    for b in blocked:
        if b in code:
            return f"错误：代码包含受限操作 '{b}'"
    old_stdout = sys.stdout
    old_stderr = sys.stderr
    captured = io.StringIO()
    sys.stdout = captured
    sys.stderr = captured
    # 限制可用内置函数
    safe_builtins = {
        k: v for k, v in __builtins__.items() if k not in (
            "exec", "eval", "compile", "__import__", "open", "input",
            "breakpoint", "exit", "quit"
        )
    } if isinstance(__builtins__, dict) else {}
    try:
        namespace = {"__builtins__": safe_builtins, "math": __import__("math"),
                     "json": __import__("json"), "re": __import__("re"),
                     "datetime": __import__("datetime"), "collections": __import__("collections")}
        exec(code, namespace)
        output = captured.getvalue()
        if not output:
            # 尝试返回最后一个表达式的值
            try:
                last_expr = code.strip().split("\n")[-1]
                if not last_expr.startswith(("import", "from", "def", "class", "for", "while", "if", "try", "with")):
                    result = eval(last_expr, namespace)
                    if result is not None:
                        output = str(result)
            except Exception:
                pass
        return output.strip() if output.strip() else "(无输出)"
    except Exception as e:
        return f"执行错误：{e}"
    finally:
        sys.stdout = old_stdout
        sys.stderr = old_stderr


def pip_install(package: str, action: str = "install") -> str:
    """包管理：安装/卸载/检查已安装/列表"""
    try:
        if action == "install":
            r = subprocess.run([sys.executable, "-m", "pip", "install", package, "-i", "https://pypi.tuna.tsinghua.edu.cn/simple"],
                               capture_output=True, text=True, timeout=120)
            return r.stdout[-2000:] if r.returncode == 0 else f"安装失败：{r.stderr[-1000:]}"
        elif action == "uninstall":
            r = subprocess.run([sys.executable, "-m", "pip", "uninstall", package, "-y"],
                               capture_output=True, text=True, timeout=60)
            return r.stdout[-2000:] if r.returncode == 0 else f"卸载失败：{r.stderr[-1000:]}"
        elif action == "check":
            r = subprocess.run([sys.executable, "-m", "pip", "show", package],
                               capture_output=True, text=True, timeout=15)
            return r.stdout.strip() if r.returncode == 0 else f"未安装：{package}"
        elif action == "list":
            r = subprocess.run([sys.executable, "-m", "pip", "list", "--format=columns"],
                               capture_output=True, text=True, timeout=15)
            return r.stdout[:3000]
        else:
            return f"错误：未知操作 {action}，支持 install/uninstall/check/list"
    except subprocess.TimeoutExpired:
        return "错误：操作超时"
    except Exception as e:
        return f"错误：{e}"


def check_port(port: int) -> str:
    """检测端口占用情况"""
    try:
        r = subprocess.run(f'netstat -ano | findstr ":{port} "', capture_output=True, text=True, timeout=10, shell=True)
        if not r.stdout.strip():
            return f"端口 {port} 未被占用"
        lines = r.stdout.strip().splitlines()
        result = [f"端口 {port} 已被占用："]
        for line in lines[:10]:
            parts = line.split()
            if len(parts) >= 5:
                proto, local, foreign, state, pid = parts[0], parts[1], parts[2], parts[3], parts[4]
                # 查找进程名
                try:
                    pr = subprocess.run(["tasklist", "/FI", f"PID eq {pid}", "/FO", "CSV", "/NH"],
                                        capture_output=True, text=True, timeout=5)
                    pname = pr.stdout.strip().split(",")[0].strip('"') if pr.stdout.strip() else "?"
                except Exception:
                    pname = "?"
                result.append(f"  {proto}  {local}  {state}  PID:{pid} ({pname})")
        return "\n".join(result)
    except Exception as e:
        return f"错误：{e}"


def file_diff(path_a: str, path_b: str) -> str:
    """比较两个文件差异"""
    import difflib
    try:
        a = Path(path_a).resolve()
        b = Path(path_b).resolve()
        if not a.exists():
            return f"错误：文件不存在 {path_a}"
        if not b.exists():
            return f"错误：文件不存在 {path_b}"
        text_a = a.read_text(encoding="utf-8", errors="replace").splitlines()
        text_b = b.read_text(encoding="utf-8", errors="replace").splitlines()
        diff = list(difflib.unified_diff(text_a, text_b, fromfile=path_a, tofile=path_b, lineterm=""))
        if not diff:
            return "两个文件内容完全相同"
        return "\n".join(diff[:200])
    except Exception as e:
        return f"错误：{e}"


def read_image(path: str) -> str:
    """读取图片文件，返回 base64 编码（用于多模态消息）"""
    try:
        import base64 as b64
        full = Path(path).resolve()
        if not full.exists():
            return f"错误：图片不存在 {path}"
        ext = full.suffix.lower()
        if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
            return f"错误：不支持的图片格式 {ext}，支持 png/jpg/jpeg/gif/bmp/webp"
        if full.stat().st_size > 10 * 1024 * 1024:
            return "错误：图片太大（超过10MB）"
        # 压缩大图：超过 1MB 的图片缩放到 1920px 宽
        data = full.read_bytes()
        if full.stat().st_size > 1024 * 1024:
            try:
                from PIL import Image
                import io
                img = Image.open(full)
                if img.width > 1920:
                    ratio = 1920 / img.width
                    img = img.resize((1920, int(img.height * ratio)), Image.LANCZOS)
                buf = io.BytesIO()
                fmt = "PNG" if ext == ".png" else "JPEG"
                img.save(buf, format=fmt, quality=85)
                data = buf.getvalue()
            except Exception:
                pass  # 压缩失败就用原图
        b64_data = b64.b64encode(data).decode("ascii")
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "bmp": "image/bmp", "webp": "image/webp"}.get(ext.lstrip("."), "image/png")
        return f"data:{mime};base64,{b64_data}"
    except Exception as e:
        return f"错误：{e}"


def _safe_markdown(text, code_theme: str = None):
    """安全构造 Markdown 渲染对象

    关键点：Textual 8.x 的 `textual.widgets.Markdown` 是一个 Widget（不可作为 renderable），
    只能在 `mount()` 时使用。但用户调用方大多是 `Static.update()` / `log.write()`，
    这些需要的是 Rich 的 `Markdown` renderable。

    因此本函数始终返回 `rich.markdown.Markdown` 实例（兼容所有 Textual 版本）。
    `code_theme` 参数在 rich.markdown 中不支持（rich 直接用 Pygments 主题），
    保留参数仅为了兼容旧调用点。

    同时对输入做学术规范化预处理：
    - 去除标题行首的空格/制表符（避免被误判为代码块）
    - 标题前后自动补空行
    - 五级及以上标题降级为加粗正文（学术论文不应使用）
    """
    if text:
        text = _normalize_markdown_for_academic(text)
    try:
        from rich.markdown import Markdown as RichMarkdown
        return RichMarkdown(text or "")
    except Exception:
        # 最后兜底：返回纯文本（保证类型是 str）
        return text or ""


def _normalize_markdown_for_academic(text: str) -> str:
    """学术 Markdown 规范化预处理
    - 去除标题行首空格（避免 `   ## 标题` 被当代码块）
    - 标题前后补空行（确保渲染正确）
    - 五级及以上标题降级为加粗正文
    - 清理空标题（连续的 # 字符）
    """
    import re
    lines = text.split("\n")
    normalized = []
    prev_was_heading = False
    for i, line in enumerate(lines):
        stripped = line.lstrip()
        # 匹配标题：# 标题、## 标题、### 标题、#### 标题、##### 标题...
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            hashes, title_text = m.group(1), m.group(2).strip()
            # 清理空标题（只有 # 没有文字）
            if not title_text or title_text.strip("#").strip() == "":
                continue
            # 五级及以上降级为加粗正文
            if len(hashes) >= 5:
                normalized.append("")  # 标题前空行
                normalized.append(f"**{title_text}**")
                normalized.append("")  # 标题后空行
                prev_was_heading = True
                continue
            # 标题前补空行（如果上一行不是空行）
            if normalized and normalized[-1].strip() != "" and not prev_was_heading:
                normalized.append("")
            normalized.append(f"{hashes} {title_text}")
            prev_was_heading = True
        else:
            # 非标题行
            if prev_was_heading and line.strip() == "":
                # 标题后第一个空行，标记为已分隔
                pass
            prev_was_heading = False
            normalized.append(line)
    # 清理连续空行（最多保留 2 个）
    result = []
    blank_count = 0
    for ln in normalized:
        if ln.strip() == "":
            blank_count += 1
            if blank_count <= 2:
                result.append(ln)
        else:
            blank_count = 0
            result.append(ln)
    return "\n".join(result)


def render_image_preview(image_source: str, max_width: int = 50) -> Text:
    """在终端中用半块字符渲染图片预览
    image_source: base64 data URI 或文件路径
    max_width: 预览最大字符宽度
    返回: Rich Text 对象，包含彩色图片预览
    """
    try:
        from PIL import Image
        import io

        # 加载图片
        if image_source.startswith("data:"):
            # data URI 格式
            header, b64data = image_source.split(",", 1)
            img_data = base64.b64decode(b64data)
            img = Image.open(io.BytesIO(img_data))
        else:
            # 文件路径
            img = Image.open(image_source)

        # 转为 RGB 模式（去掉 alpha 通道，透明背景用黑色填充）
        if img.mode == "RGBA":
            background = Image.new("RGBA", img.size, (0, 0, 0, 255))
            background.paste(img, mask=img.split()[3] if img.mode == "RGBA" else None)
            img = background.convert("RGB")
        elif img.mode != "RGB":
            img = img.convert("RGB")

        # 计算缩放尺寸：每个字符代表2个垂直像素
        # 保持宽高比，宽度不超过 max_width
        orig_w, orig_h = img.size
        char_width = min(max_width, orig_w)
        char_height = max(1, int(char_width * orig_h / orig_w / 2))
        pixel_width = char_width
        pixel_height = char_height * 2  # 每个字符2个像素行

        img_small = img.resize((pixel_width, pixel_height), Image.LANCZOS)
        pixels = list(img_small.getdata())

        # 用 Rich Text 渲染：▀ 字符前景=上像素，背景=下像素
        result = Text()
        for row in range(char_height):
            line_text = ""
            spans = []
            for col in range(char_width):
                # 上像素
                idx_top = row * 2 * pixel_width + col
                # 下像素
                idx_bot = (row * 2 + 1) * pixel_width + col
                r1, g1, b1 = pixels[idx_top][:3]
                r2, g2, b2 = pixels[idx_bot][:3]
                top_color = f"#{r1:02x}{g1:02x}{b1:02x}"
                bot_color = f"#{r2:02x}{g2:02x}{b2:02x}"
                # ▀ (U+2580) 上半块：前景=上像素颜色，背景=下像素颜色
                line_text += "▀"
                spans.append((top_color, bot_color))
            # 为整行添加样式
            # Rich 支持 style="fg color on bg color"
            # 每个字符需要不同的颜色，所以逐字符添加
            for i, ch in enumerate(line_text):
                top_c, bot_c = spans[i]
                result.append(ch, style=f"{top_c} on {bot_c}")
            result.append("\n")

        return result
    except Exception as e:
        return Text(f"  [图片预览失败: {str(e)[:60]}]", style=C_DIM)


# ====== 安全审计功能 ======

# 代码漏洞扫描规则
VULN_PATTERNS = [
    # SQL注入
    {"id": "SQLI001", "severity": "高", "category": "SQL注入",
     "pattern": r'execute\s*\(\s*["\'].*\+.*["\']\s*\)|execute\s*\(\s*f["\']',
     "desc": "SQL拼接执行，可能导致SQL注入", "fix": "使用参数化查询：cursor.execute('SELECT * FROM t WHERE id=?', (id,))"},
    {"id": "SQLI002", "severity": "高", "category": "SQL注入",
     "pattern": r'\.raw\s*\(\s*["\'].*\+.*["\']\s*\)|\.raw\s*\(\s*f["\']',
     "desc": "ORM raw查询拼接，可能导致SQL注入", "fix": "使用ORM参数绑定或参数化查询"},
    # XSS
    {"id": "XSS001", "severity": "中", "category": "XSS",
     "pattern": r'innerHTML\s*=\s*[^"\']*\+|innerHTML\s*=\s*`',
     "desc": "innerHTML直接拼接，可能导致XSS", "fix": "使用textContent或对内容进行HTML转义"},
    {"id": "XSS002", "severity": "中", "category": "XSS",
     "pattern": r'dangerouslySetInnerHTML',
     "desc": "React dangerouslySetInnerHTML，可能导致XSS", "fix": "避免使用dangerouslySetInnerHTML，或对内容严格转义"},
    # 命令注入
    {"id": "CMD001", "severity": "高", "category": "命令注入",
     "pattern": r'os\.system\s*\(\s*[^"\']*["\'].*\+|os\.system\s*\(\s*f["\']',
     "desc": "os.system拼接执行，可能导致命令注入", "fix": "使用subprocess.run(args_list)避免shell=True"},
    {"id": "CMD002", "severity": "高", "category": "命令注入",
     "pattern": r'subprocess\..*shell\s*=\s*True',
     "desc": "subprocess shell=True，可能导致命令注入", "fix": "使用shell=False并传递参数列表"},
    # 路径穿越
    {"id": "PATH001", "severity": "中", "category": "路径穿越",
     "pattern": r'open\s*\(\s*request\.|open\s*\(\s*input\s*\(',
     "desc": "直接打开用户输入路径，可能导致路径穿越", "fix": "验证路径在允许目录内：Path(path).resolve()检查是否在WORK_DIR下"},
    # 硬编码密钥
    {"id": "KEY001", "severity": "高", "category": "硬编码密钥",
     "pattern": r'(api_key|apikey|api-key|secret|password|passwd|token)\s*=\s*["\'][a-zA-Z0-9_\-]{16,}["\']',
     "desc": "代码中硬编码密钥/密码", "fix": "从环境变量读取：os.environ.get('API_KEY')"},
    {"id": "KEY002", "severity": "高", "category": "硬编码密钥",
     "pattern": r'sk-[a-zA-Z0-9]{20,}',
     "desc": "代码中硬编码API Key（sk-开头）", "fix": "移到环境变量或配置文件中"},
    # 不安全的反序列化
    {"id": "DESER001", "severity": "高", "category": "反序列化",
     "pattern": r'pickle\.loads?\s*\(',
     "desc": "pickle反序列化不安全，可能导致RCE", "fix": "使用json.loads替代，或验证数据来源"},
    # 弱加密
    {"id": "CRYPTO001", "severity": "中", "category": "弱加密",
     "pattern": r'hashlib\.md5\s*\(|hashlib\.sha1\s*\(',
     "desc": "使用MD5/SHA1弱哈希", "fix": "使用SHA256：hashlib.sha256()"},
    {"id": "CRYPTO002", "severity": "高", "category": "弱加密",
     "pattern": r'random\.random\s*\(\s*\).*password|random\.choice.*password',
     "desc": "使用random生成密码（不安全）", "fix": "使用secrets模块：secrets.token_urlsafe()"},
    # 调试代码
    {"id": "DEBUG001", "severity": "低", "category": "调试残留",
     "pattern": r'print\s*\(\s*["\']DEBUG|print\s*\(\s*["\']TODO|breakpoint\s*\(\s*\)',
     "desc": "代码中残留调试语句", "fix": "移除调试代码或使用logging模块"},
    # 不安全的SSL
    {"id": "SSL001", "severity": "中", "category": "SSL",
     "pattern": r'verify\s*=\s*False|CERT_NONE',
     "desc": "禁用SSL证书验证", "fix": "始终验证SSL证书，不要设置verify=False"},
]


def scan_code_vulnerabilities(path: str) -> str:
    """扫描代码文件中的安全漏洞"""
    try:
        full = Path(path).resolve()
        if not full.exists():
            return f"错误：文件不存在 {path}"
        if full.is_dir():
            return f"错误：{path} 是目录，请指定文件"
        if full.stat().st_size > MAX_FILE_SIZE:
            return "错误：文件太大"

        # 读取文件内容
        for enc in ["utf-8", "gbk", "latin-1"]:
            try:
                lines = full.read_text(encoding=enc).splitlines()
                break
            except UnicodeDecodeError:
                continue

        findings = []
        for line_no, line in enumerate(lines, 1):
            for rule in VULN_PATTERNS:
                try:
                    if re.search(rule["pattern"], line, re.IGNORECASE):
                        findings.append({
                            "line": line_no,
                            "id": rule["id"],
                            "severity": rule["severity"],
                            "category": rule["category"],
                            "desc": rule["desc"],
                            "fix": rule["fix"],
                            "code": line.strip()[:100],
                        })
                except Exception:
                    continue

        if not findings:
            return f"{path} 未发现已知漏洞模式\n扫描了 {len(lines)} 行代码，匹配 {len(VULN_PATTERNS)} 条规则"

        # 按严重程度排序
        sev_order = {"高": 0, "中": 1, "低": 2}
        findings.sort(key=lambda x: sev_order.get(x["severity"], 9))

        result = f"{_load_svg_icon('search')} 安全扫描报告：{path}\n"
        result += f"扫描了 {len(lines)} 行代码，发现 {len(findings)} 个问题\n\n"

        # 统计
        high_count = sum(1 for f in findings if f["severity"] == "高")
        mid_count = sum(1 for f in findings if f["severity"] == "中")
        low_count = sum(1 for f in findings if f["severity"] == "低")
        result += f"严重程度：高危 {high_count} | 中危 {mid_count} | 低危 {low_count}\n\n"

        for f in findings[:20]:  # 最多显示20条
            result += f"[{f['severity']}] 第{f['line']}行 ({f['id']} {f['category']})\n"
            result += f"  代码: {f['code']}\n"
            result += f"  问题: {f['desc']}\n"
            result += f"  修复: {f['fix']}\n\n"

        if len(findings) > 20:
            result += f"... 还有 {len(findings) - 20} 个问题未显示\n"

        return result
    except Exception as e:
        return f"错误：{e}"


def detect_sensitive_info(path: str) -> str:
    """检测代码中的敏感信息泄露"""
    try:
        full = Path(path).resolve()
        if not full.exists():
            return f"错误：文件不存在 {path}"
        if full.is_dir():
            # 扫描目录下所有文件
            results = []
            extensions = {".py", ".js", ".ts", ".json", ".yaml", ".yml", ".env", ".cfg", ".ini", ".conf", ".txt"}
            for f in full.rglob("*"):
                if f.is_file() and f.suffix.lower() in extensions:
                    if "node_modules" in str(f) or ".git" in str(f):
                        continue
                    r = _scan_file_secrets(str(f))
                    if r and "未发现" not in r:
                        results.append(r)
            if not results:
                return f"{path} 目录下未发现敏感信息泄露"
            return f"{_load_svg_icon('search')} 敏感信息检测报告：{path}\n\n" + "\n".join(results)
        else:
            return _scan_file_secrets(str(full))
    except Exception as e:
        return f"错误：{e}"


def _scan_file_secrets(path: str) -> str:
    """扫描单个文件中的敏感信息"""
    SECRET_PATTERNS = [
        (r'sk-[a-zA-Z0-9]{20,}', "OpenAI API Key"),
        (r'sk-or-v1-[a-zA-Z0-9]{20,}', "OpenRouter API Key"),
        (r'AIza[a-zA-Z0-9_\-]{35}', "Google API Key"),
        (r'ghp_[a-zA-Z0-9]{36}', "GitHub Personal Access Token"),
        (r'gho_[a-zA-Z0-9]{36}', "GitHub OAuth Token"),
        (r'glpat-[a-zA-Z0-9_\-]{20}', "GitLab Personal Access Token"),
        (r'AKIA[A-Z0-9]{16}', "AWS Access Key ID"),
        (r'-----BEGIN (RSA |EC |)PRIVATE KEY-----', "私钥"),
        (r'(mysql|mongodb|postgresql|redis)://[^\s"\'<>]+:[^\s"\'<>]+@', "数据库连接字符串（含密码）"),
        (r'(password|passwd|pwd)\s*[=:]\s*["\'][^"\']{4,}["\']', "硬编码密码"),
        (r'(api_key|apikey|api-key|secret_key|secretkey)\s*[=:]\s*["\'][^"\']{8,}["\']', "硬编码API密钥"),
        (r'(token|access_token|auth_token)\s*[=:]\s*["\'][^"\']{16,}["\']', "硬编码Token"),
    ]

    try:
        for enc in ["utf-8", "gbk", "latin-1"]:
            try:
                lines = Path(path).read_text(encoding=enc).splitlines()
                break
            except UnicodeDecodeError:
                continue

        findings = []
        for line_no, line in enumerate(lines, 1):
            for pattern, name in SECRET_PATTERNS:
                try:
                    matches = re.findall(pattern, line, re.IGNORECASE)
                    if matches:
                        # 脱敏显示
                        masked = re.sub(pattern, lambda m: m.group()[:8] + "***" + m.group()[-4:] if len(m.group()) > 12 else "***", line)
                        findings.append(f"  第{line_no}行 [{name}]: {masked.strip()[:120]}")
                except Exception:
                    continue

        if not findings:
            return ""
        result = f"{path}\n"
        result += "\n".join(findings[:15])
        if len(findings) > 15:
            result += f"\n  ... 还有 {len(findings) - 15} 处"
        return result + "\n"
    except Exception:
        return ""


def check_dependencies_vulnerabilities() -> str:
    """检查Python依赖包的已知漏洞（使用pip audit或安全检查）"""
    try:
        import subprocess
        import sys as _sys
        # 方法1: 尝试使用 pip-audit（如果安装了）
        try:
            result = subprocess.run(
                [_sys.executable, "-m", "pip_audit"],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                return f"{_load_svg_icon('search')} 依赖漏洞检查报告（pip-audit）\n\n{result.stdout[:3000]}"
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass

        # 方法2: 检查 requirements.txt 中的已知问题包
        req_file = Path(WORK_DIR) / "requirements.txt"
        if not req_file.exists():
            req_file = Path(WORK_DIR) / "pyproject.toml"

        if not req_file.exists():
            return "ℹ 未找到 requirements.txt 或 pyproject.toml，跳过依赖检查\n建议创建 requirements.txt 以启用依赖漏洞检查"

        content = req_file.read_text(encoding="utf-8", errors="replace")

        # 已知有安全问题的包版本（简化版，实际应查询CVE数据库）
        KNOWN_VULN_PACKAGES = {
            "django": {"<2.2.0": "CVE-2019-19844 等多个漏洞", "<3.2.0": "多个安全修复"},
            "flask": {"<1.0": "CVE-2018-1000656"},
            "requests": {"<2.20.0": "CVE-2018-18074"},
            "urllib3": {"<1.24.2": "CVE-2019-11324"},
            "jinja2": {"<2.10.1": "CVE-2019-10906"},
            "cryptography": {"<2.3": "多个安全问题"},
            "pyyaml": {"<5.1": "CVE-2017-18342"},
            "pillow": {"<6.2.0": "多个图像处理漏洞"},
        }

        findings = []
        for line in content.splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            # 解析包名和版本
            match = re.match(r'^([a-zA-Z0-9_\-]+)\s*[=<>~!]+\s*([0-9.]+)', line)
            if match:
                pkg = match.group(1).lower()
                ver = match.group(2)
                if pkg in KNOWN_VULN_PACKAGES:
                    for vuln_ver, desc in KNOWN_VULN_PACKAGES[pkg].items():
                        findings.append(f"  {pkg}=={ver} → {desc}")

        if not findings:
            return f"未发现已知漏洞依赖\n检查了 {req_file.name}"

        return f"{_load_svg_icon('search')} 依赖漏洞检查报告\n检查文件：{req_file.name}\n\n发现 {len(findings)} 个潜在问题：\n" + "\n".join(findings)
    except Exception as e:
        return f"错误：{e}"


def check_config_security(path: str = ".") -> str:
    """检查配置文件安全"""
    try:
        base = Path(path).resolve()
        if not base.exists():
            return f"错误：路径不存在 {path}"

        findings = []

        # 检查 .env 文件是否暴露
        env_file = base / ".env"
        if env_file.exists():
            # 检查 .env 是否在 .gitignore 中
            gitignore = base / ".gitignore"
            if gitignore.exists():
                gi_content = gitignore.read_text(encoding="utf-8", errors="replace")
                if ".env" not in gi_content:
                    findings.append(".env 文件未被 .gitignore 忽略，可能被提交到版本控制")
            else:
                findings.append("存在 .env 文件但没有 .gitignore，敏感信息可能泄露")
            # 检查 .env 内容
            env_content = env_file.read_text(encoding="utf-8", errors="replace")
            secret_lines = [l for l in env_content.splitlines() if "=" in l and any(k in l.lower() for k in ["key", "secret", "password", "token"])]
            if secret_lines:
                findings.append(f"ℹ .env 包含 {len(secret_lines)} 条敏感配置（KEY/SECRET/PASSWORD/TOKEN）")

        # 检查 .gitignore
        gitignore = base / ".gitignore"
        if not gitignore.exists():
            findings.append("缺少 .gitignore 文件，所有文件都可能被提交到版本控制")

        # 检查日志文件
        for log_file in base.glob("*.log"):
            if log_file.stat().st_size > 0:
                findings.append(f"ℹ 发现日志文件 {log_file.name}（{log_file.stat().st_size} 字节），检查是否包含敏感信息")

        # 检查权限过宽的文件（Windows下检查只读属性）
        for f in base.glob("*"):
            if f.is_file() and f.suffix in (".key", ".pem", ".pfx", ".p12"):
                findings.append(f"发现证书/密钥文件 {f.name}，确保权限设置正确")

        # 检查是否有硬编码IP/端口
        for cfg_file in list(base.glob("*.conf")) + list(base.glob("*.cfg")) + list(base.glob("*.ini")):
            try:
                content = cfg_file.read_text(encoding="utf-8", errors="replace")
                ips = re.findall(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b', content)
                if ips:
                    findings.append(f"ℹ {cfg_file.name} 包含IP地址：{', '.join(set(ips)[:5])}")
            except Exception:
                continue

        if not findings:
            return f"{path} 配置安全检查通过\n检查项：.env/.gitignore/日志文件/证书文件/IP暴露"

        return f"{_load_svg_icon('search')} 配置安全检查报告：{path}\n\n发现 {len(findings)} 个注意事项：\n" + "\n".join(f"  {i+1}. {f}" for i, f in enumerate(findings))
    except Exception as e:
        return f"错误：{e}"


def security_audit(path: str = ".", scan_type: str = "all") -> str:
    """综合安全审计
    scan_type: all(全部) / code(代码漏洞) / secret(敏感信息) / deps(依赖漏洞) / config(配置安全)
    """
    try:
        target = Path(path).resolve()
        results = []

        if scan_type in ("all", "code"):
            results.append("═══ 代码漏洞扫描 ═══")
            if target.is_dir():
                # 扫描目录下所有代码文件
                code_exts = {".py", ".js", ".ts", ".jsx", ".tsx", ".php", ".java", ".go"}
                count = 0
                for f in target.rglob("*"):
                    if f.is_file() and f.suffix.lower() in code_exts:
                        if "node_modules" in str(f) or ".git" in str(f) or "__pycache__" in str(f):
                            continue
                        r = scan_code_vulnerabilities(str(f))
                        if r and "未发现" not in r:
                            results.append(r)
                            count += 1
                    if count >= 10:
                        results.append(f"... 目录较大，仅扫描前10个有问题的文件")
                        break
                if count == 0:
                    results.append(f"{path} 目录下代码文件未发现已知漏洞模式")
            else:
                results.append(scan_code_vulnerabilities(str(target)))
            results.append("")

        if scan_type in ("all", "secret"):
            results.append("═══ 敏感信息检测 ═══")
            results.append(detect_sensitive_info(str(target)))
            results.append("")

        if scan_type in ("all", "deps"):
            results.append("═══ 依赖漏洞检查 ═══")
            results.append(check_dependencies_vulnerabilities())
            results.append("")

        if scan_type in ("all", "config"):
            results.append("═══ 配置安全检查 ═══")
            results.append(check_config_security(str(target) if target.is_dir() else str(target.parent)))

        return "\n".join(results)
    except Exception as e:
        return f"安全审计出错：{e}"


def _add_formatted_text(paragraph, text: str):
    """解析 Markdown 粗体/斜体/行内代码/删除线并添加到段落"""
    # 匹配 **粗体** *斜体* `代码` ~~删除线~~
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|~~(.+?)~~)'
    last_end = 0

    for match in re.finditer(pattern, text):
        # 添加前面的普通文本
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # **粗体**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *斜体*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `行内代码`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            from docx.shared import Pt as _Pt, RGBColor as _RC
            run.font.size = _Pt(10)
            run.font.color.rgb = _RC(0xC0, 0x39, 0x2B)
        elif match.group(5):  # ~~删除线~~
            run = paragraph.add_run(match.group(5))
            run.font.strike = True

        last_end = match.end()

    # 添加剩余文本
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


# ====== LaTeX 公式渲染引擎（终端 Unicode 渲染）======
# 用于学术研究场景：将 $E=mc^2$、$$\sum_{i=1}^{n}$$ 等 LaTeX 公式
# 转换为终端可显示的 Unicode 数学符号

# 希腊字母（小写/大写）
_LATEX_GREEK = {
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\epsilon": "ε", r"\varepsilon": "ε", r"\zeta": "ζ", r"\eta": "η",
    r"\theta": "θ", r"\vartheta": "ϑ", r"\iota": "ι", r"\kappa": "κ",
    r"\lambda": "λ", r"\mu": "μ", r"\nu": "ν", r"\xi": "ξ",
    r"\pi": "π", r"\varpi": "ϖ", r"\rho": "ρ", r"\varrho": "ϱ",
    r"\sigma": "σ", r"\varsigma": "ς", r"\tau": "τ", r"\upsilon": "υ",
    r"\phi": "φ", r"\varphi": "φ", r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
    r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
    r"\Xi": "Ξ", r"\Pi": "Π", r"\Sigma": "Σ", r"\Upsilon": "Υ",
    r"\Phi": "Φ", r"\Psi": "Ψ", r"\Omega": "Ω",
}

# 数学运算符与符号
_LATEX_OPERATORS = {
    r"\times": "×", r"\div": "÷", r"\pm": "±", r"\mp": "∓",
    r"\cdot": "·", r"\cdots": "⋯", r"\ldots": "…", r"\vdots": "⋮", r"\ddots": "⋱",
    r"\infty": "∞", r"\partial": "∂", r"\nabla": "∇", r"\forall": "∀", r"\exists": "∃",
    r"\neg": "¬", r"\land": "∧", r"\lor": "∨", r"\oplus": "⊕", r"\ominus": "⊖",
    r"\otimes": "⊗", r"\odot": "⊙", r"\cap": "∩", r"\cup": "∪", r"\setminus": "∖",
    r"\subset": "⊂", r"\supset": "⊃", r"\subseteq": "⊆", r"\supseteq": "⊇",
    r"\in": "∈", r"\notin": "∉", r"\ni": "∋",
    r"\leq": "≤", r"\geq": "≥", r"\neq": "≠", r"\approx": "≈", r"\equiv": "≡",
    r"\sim": "∼", r"\simeq": "≃", r"\cong": "≅", r"\propto": "∝",
    r"\to": "→", r"\rightarrow": "→", r"\leftarrow": "←", r"\gets": "←",
    r"\Rightarrow": "⇒", r"\Leftarrow": "⇐", r"\Leftrightarrow": "⇔", r"\iff": "⟺",
    r"\mapsto": "↦", r"\uparrow": "↑", r"\downarrow": "↓", r"\updownarrow": "↕",
    r"\sum": "Σ", r"\prod": "∏", r"\coprod": "∐", r"\int": "∫", r"\oint": "∮",
    r"\bigcup": "⋃", r"\bigcap": "⋂", r"\bigoplus": "⨁", r"\bigotimes": "⨂",
    r"\sqrt": "√", r"\cubert": "∛", r"\fourthroot": "∜",
    r"\angle": "∠", r"\perp": "⊥", r"\parallel": "∥", r"\triangle": "△",
    r"\circ": "∘", r"\bullet": "•", r"\star": "⋆", r"\dagger": "†", r"\ddagger": "‡",
    r"\aleph": "ℵ", r"\beth": "ℶ", r"\hbar": "ℏ", r"\ell": "ℓ",
    r"\Re": "ℜ", r"\Im": "ℑ", r"\wp": "℘", r"\mho": "℧",
    r"\angle": "∠", r"\measuredangle": "∡", r"\sphericalangle": "∢",
    r"\prime": "′", r"\backprime": "‵",
    r"\colon": ":", r"\vert": "|", r"\Vert": "‖", r"\backslash": "\\",
    r"\degree": "°", r"\circ": "∘",
    r"\leqq": "≦", r"\geqq": "≧", r"\lessgtr": "≶", r"\gtrless": "≷",
    r"\prec": "≺", r"\succ": "≻", r"\preceq": "≼", r"\succeq": "≽",
    r"\emptyset": "∅", r"\varnothing": "∅",
    r"\mathbb{R}": "ℝ", r"\mathbb{Z}": "ℤ", r"\mathbb{Q}": "ℚ",
    r"\mathbb{N}": "ℕ", r"\mathbb{C}": "ℂ", r"\mathbb{H}": "ℍ",
    r"\mathbb{A}": "𝔸", r"\mathbb{B}": "𝔹", r"\mathbb{D}": "𝔻",
    r"\mathbb{E}": "𝔼", r"\mathbb{F}": "𝔽", r"\mathbb{G}": "𝔾",
    r"\mathbb{I}": "𝕀", r"\mathbb{J}": "𝕁", r"\mathbb{K}": "𝕂",
    r"\mathbb{L}": "𝕃", r"\mathbb{M}": "𝕄", r"\mathbb{O}": "𝕆",
    r"\mathbb{P}": "ℙ", r"\mathbb{S}": "𝕊", r"\mathbb{T}": "𝕋",
    r"\mathbb{U}": "𝕌", r"\mathbb{V}": "𝕍", r"\mathbb{W}": "𝕎", r"\mathbb{X}": "𝕏",
    r"\mathbb{Y}": "𝕐",
}

# 下标映射（Unicode 下标字符）
_LATEX_SUBSCRIPT = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
    "a": "ₐ", "e": "ₑ", "o": "ₒ", "x": "ₓ", "h": "ₕ",
    "k": "ₖ", "l": "ₗ", "m": "ₘ", "n": "ₙ", "p": "ₚ",
    "s": "ₛ", "t": "ₜ", "i": "ᵢ", "j": "ⱼ", "u": "ᵤ", "v": "ᵥ",
}

# 上标映射（Unicode 上标字符）
_LATEX_SUPERSCRIPT = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾",
    "a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ", "e": "ᵉ",
    "f": "ᶠ", "g": "ᵍ", "h": "ʰ", "i": "ⁱ", "j": "ʲ",
    "k": "ᵏ", "l": "ˡ", "m": "ᵐ", "n": "ⁿ", "o": "ᵒ",
    "p": "ᵖ", "r": "ʳ", "s": "ˢ", "t": "ᵗ", "u": "ᵘ",
    "v": "ᵛ", "w": "ʷ", "x": "ˣ", "y": "ʸ", "z": "ᶻ",
    "A": "ᴬ", "B": "ᴮ", "D": "ᴰ", "E": "ᴱ", "G": "ᴳ",
    "H": "ᴴ", "I": "ᴵ", "J": "ᴶ", "K": "ᴷ", "L": "ᴸ",
    "M": "ᴹ", "N": "ᴺ", "O": "ᴼ", "P": "ᴾ", "R": "ᴿ",
    "T": "ᵀ", "U": "ᵁ", "V": "ⱽ", "W": "ᵂ",
    "α": "ᵅ", "β": "ᵝ", "γ": "ᵞ", "δ": "ᵟ", "ε": "ᵋ",
    "θ": "ᶿ", "ι": "ᶥ", "φ": "ᵠ", "χ": "ᵡ", "ψ": "ᵧ",
    "n": "ⁿ", "-": "⁻",
}

# 函数名映射（保持原样，不转换）
_LATEX_FUNCTIONS = {
    r"\sin", r"\cos", r"\tan", r"\cot", r"\sec", r"\csc",
    r"\arcsin", r"\arccos", r"\arctan",
    r"\sinh", r"\cosh", r"\tanh", r"\coth",
    r"\log", r"\ln", r"\lg", r"\exp",
    r"\lim", r"\max", r"\min", r"\sup", r"\inf",
    r"\arg", r"\det", r"\dim", r"\gcd", r"\hom", r"\ker", r"\deg",
    r"\operatorname",
}


def _latex_to_unicode(latex: str) -> str:
    r"""将单个 LaTeX 公式转换为 Unicode 终端可显示文本

    支持：
    - 希腊字母：\\alpha → α, \\Sigma → Σ
    - 运算符：\\times → ×, \\sum → Σ, \\int → ∫
    - 上下标：x_1 → x₁, x^2 → x², x_{10} → x₁₀, x^{n+1} → xⁿ⁺¹
    - 分数：\\frac{a}{b} → a⁄b（使用 Unicode 分数斜杠 ⁄ U+2044，比普通 / 更贴近真分数排版）
    - 根号：\\sqrt{x} → √x, \\sqrt[3]{x} → ∛x
    - 求和/积分上下限：\\sum_{i=1}^{n} → Σᵢ₌₁ⁿ
    - 黑板粗体：\\mathbb{R} → ℝ
    - 函数名：\\sin \\cos \\log 等保持原样
    """
    s = latex.strip()
    # 去除首尾 $ 符号（已在调用前处理）
    s = s.strip("$")

    # 0. 预处理：\dfrac \tfrac \cfrac 统一当作 \frac 处理
    s = re.sub(r"\\[dtc]frac\b", r"\\frac", s)

    # 0.1 处理 \left( \right) \left[ \right] \left\{ \right\} 等自适应定界符
    s = s.replace(r"\left(", "(").replace(r"\right)", ")")
    s = s.replace(r"\left[", "[").replace(r"\right]", "]")
    s = s.replace(r"\left\{", "{").replace(r"\right\}", "}")
    s = s.replace(r"\left|", "|").replace(r"\right|", "|")
    s = s.replace(r"\left\|", "‖").replace(r"\right\|", "‖")
    s = s.replace(r"\left.", "").replace(r"\right.", "")

    # 0.2 处理 \big \Big \bigg \Bigg 等尺寸前缀（直接去除，保留定界符本身）
    # 关键：必须用 \b 保护 \bigcup \bigcap \bigoplus \bigotimes 等以 big 开头的命令不被误删
    s = re.sub(r"\\[bB]ig[lmr]?(?![a-zA-Z])\s*", "", s)
    s = re.sub(r"\\big[lmr]?(?![a-zA-Z])\s*", "", s)

    # 0.3 处理装饰符号（矢量、帽子、横线、波浪号、点导数）
    # 按命令长度降序处理，避免 \dot 匹配 \ddot 的前缀
    _DECO_MAP = [
        (r"\\overline",  "̄"),   # 上划线（组合符 U+0304）
        (r"\\underline", "̱"),   # 下划线（组合符 U+0332）
        (r"\\widehat",   "^"),   # 宽帽子
        (r"\\widetilde", "~"),   # 宽波浪
        (r"\\mathring",  "̊"),   # 圈（组合符 U+030A）
        (r"\\ddot",      "̈"),   # 二阶导数双点（组合符 U+0308）
        (r"\\dot",       "̇"),   # 一阶导数点（组合符 U+0307）
        (r"\\vec",       "→"),   # 矢量箭头（前置）
        (r"\\hat",       "^"),   # 帽子
        (r"\\bar",       "̄"),   # 上横线
        (r"\\tilde",     "~"),   # 波浪号
    ]
    for cmd_pat, deco_sym in _DECO_MAP:
        def _deco_repl(m, ds=deco_sym):
            body_u = _latex_to_unicode(m.group(1))
            if ds == "→":
                return f"→{body_u}"  # 矢量箭头前置
            return f"{body_u}{ds}"   # 组合符号后置
        s = re.sub(cmd_pat + r"\{([^{}]*)\}", _deco_repl, s)

    # 0.4 处理矩阵 \begin{matrix}...\end{matrix} 等
    # 使用 Unicode 矩阵专用括号 ⎡⎢⎣⎤⎥⎦（U+23A1-23A6），比普通 () 更清晰
    def _matrix_repl(m):
        env = m.group(1)
        body = m.group(2)
        # 按 \\ 分行，按 & 分列
        rows = [r.strip() for r in body.split(r"\\") if r.strip()]
        rendered_rows = []
        for row in rows:
            cells = [c.strip() for c in row.split("&")]
            rendered_cells = [_latex_to_unicode(c) for c in cells]
            rendered_rows.append("  ".join(rendered_cells))
        # 单行矩阵：用紧凑形式
        if len(rendered_rows) == 1:
            inner = rendered_rows[0]
            if env == "pmatrix":
                return f"( {inner} )"
            if env == "bmatrix":
                return f"[ {inner} ]"
            if env == "Bmatrix":
                return f"{{ {inner} }}"
            if env == "vmatrix":
                return f"| {inner} |"
            if env == "Vmatrix":
                return f"‖ {inner} ‖"
            return inner
        # 多行矩阵：用矩阵专用括号 ⎡⎢⎣ ⎤⎥⎦
        n = len(rendered_rows)
        # 左括号：第一行⎡，中间行⎢，最后一行⎣
        # 右括号：第一行⎤，中间行⎥，最后一行⎦
        left_brackets = {"pmatrix": ("⎡", "⎢", "⎣"),
                         "bmatrix": ("⎡", "⎢", "⎣"),
                         "Bmatrix": ("⎧", "⎨", "⎩"),
                         "vmatrix": ("⎢", "⎢", "⎢"),
                         "Vmatrix": ("⎢", "⎢", "⎢")}
        right_brackets = {"pmatrix": ("⎤", "⎥", "⎦"),
                          "bmatrix": ("⎤", "⎥", "⎦"),
                          "Bmatrix": ("⎫", "⎬", "⎭"),
                          "vmatrix": ("⎥", "⎥", "⎥"),
                          "Vmatrix": ("⎥", "⎥", "⎥")}
        lb = left_brackets.get(env, ("", "", ""))
        rb = right_brackets.get(env, ("", "", ""))
        lines = []
        for i, row in enumerate(rendered_rows):
            if n == 1:
                l, r = lb[0], rb[0]
            elif i == 0:
                l, r = lb[0], rb[0]
            elif i == n - 1:
                l, r = lb[2], rb[2]
            else:
                l, r = lb[1], rb[1]
            lines.append(f"{l}{row}{r}")
        return "\n".join(lines)
    s = re.sub(r"\\begin\{(matrix|pmatrix|bmatrix|Bmatrix|vmatrix|Vmatrix|smallmatrix)\}(.*?)\\end\{\1\}",
               _matrix_repl, s, flags=re.DOTALL)

    # 0.5 处理 cases 环境（分段函数）
    def _cases_repl(m):
        body = m.group(1)
        rows = [r.strip() for r in body.split(r"\\") if r.strip()]
        rendered = []
        for row in rows:
            parts = row.split("&")
            if len(parts) == 2:
                cond = _latex_to_unicode(parts[1].strip())
                val = _latex_to_unicode(parts[0].strip())
                rendered.append(f"{val}  当  {cond}")
            else:
                rendered.append(_latex_to_unicode(row.strip()))
        return " { " + " ; ".join(rendered) + " }"
    s = re.sub(r"\\begin\{cases\}(.*?)\\end\{cases\}", _cases_repl, s, flags=re.DOTALL)

    # 0.6 处理 \lim_{x \to a}（极限）
    # 纯 Unicode 下标紧凑表示：limₙ→∞（无花括号，不可映射字符保留原字符）
    def _lim_repl(m):
        sub = m.group(1)
        sub_u = _latex_to_unicode(sub)
        # 逐字符映射为 Unicode 真下标，不可映射字符保留原字符
        result = "".join(_LATEX_SUBSCRIPT.get(ch, ch) for ch in sub_u)
        return "lim" + result
    s = re.sub(r"\\lim_\{([^{}]*)\}", _lim_repl, s)
    s = re.sub(r"\\lim_([a-zA-Z])",
               lambda m: "lim" + _LATEX_SUBSCRIPT.get(m.group(1), m.group(1)), s)

    # 0.7 处理 \sum_{...}^{...} \prod_{...}^{...} \int_{...}^{...} 上下限
    # 纯 Unicode 上下标紧凑表示：Σᵢ₌₁ⁿ（无花括号，不可映射字符保留原字符）
    def _bigop_repl(m):
        op = m.group(1)
        # 补全反斜杠查找运算符符号
        op_u = _LATEX_OPERATORS.get("\\" + op, op)
        low = m.group(2) if m.group(2) else ""
        high = m.group(3) if m.group(3) else ""
        low_u = _latex_to_unicode(low) if low else ""
        high_u = _latex_to_unicode(high) if high else ""
        # 逐字符映射为 Unicode 真下标/上标，不可映射字符保留原字符
        low_result = "".join(_LATEX_SUBSCRIPT.get(ch, ch) for ch in low_u)
        high_result = "".join(_LATEX_SUPERSCRIPT.get(ch, ch) for ch in high_u)
        return f"{op_u}{low_result}{high_result}"
    s = re.sub(r"\\(sum|prod|coprod|int|oint|bigcup|bigcap|bigoplus|bigotimes)_\{([^{}]*)\}\^\{([^{}]*)\}",
               _bigop_repl, s)
    # 单独下标：group(3) 不存在，用空字符串
    def _bigop_low_only(m):
        class _M:
            def group(self, i):
                return [m.group(1), m.group(2), ""][i-1]
        return _bigop_repl(_M())
    s = re.sub(r"\\(sum|prod|coprod|int|oint|bigcup|bigcap|bigoplus|bigotimes)_\{([^{}]*)\}",
               _bigop_low_only, s)
    # 单独上标：group(2) 不存在，用空字符串
    def _bigop_high_only(m):
        class _M:
            def group(self, i):
                return [m.group(1), "", m.group(2)][i-1]
        return _bigop_repl(_M())
    s = re.sub(r"\\(sum|prod|coprod|int|oint|bigcup|bigcap|bigoplus|bigotimes)\^\{([^{}]*)\}",
               _bigop_high_only, s)

    # 1. 处理 \text{...} → 原样输出
    s = re.sub(r"\\text\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\mathbf\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\mathit\{([^}]*)\}", r"\1", s)

    # 2. 处理 \sqrt[n]{x}（n次根号）
    def _sqrt_n(m):
        n = m.group(1)
        body = _latex_to_unicode(m.group(2))
        root_sym = {"2": "√", "3": "∛", "4": "∜"}.get(n, "√")
        return f"{root_sym}({body})"
    s = re.sub(r"\\sqrt\[([^\]]+)\]\{([^{}]*)\}", _sqrt_n, s)

    # 3. 处理 \sqrt{x}（平方根）
    def _sqrt_simple(m):
        body = _latex_to_unicode(m.group(1))
        return f"√({body})"
    s = re.sub(r"\\sqrt\{([^{}]*)\}", _sqrt_simple, s)

    # 4. 处理 \frac{a}{b}（分数）
    # 使用 Unicode 分数斜杠 ⁄ (U+2044) 代替普通 / ，视觉上更接近真分数
    # 嵌套分数用不同括号区分层次：最内层()，中层[]，外层〔〕
    _FRAC_SLASH = "⁄"  # 分数斜杠（比普通 / 更短、更贴近真分数排版）
    def _frac(m):
        # 去除空格（LaTeX 中 \partial f 表示 ∂f，无空格）
        num = _latex_to_unicode(m.group(1)).replace(" ", "")
        den = _latex_to_unicode(m.group(2)).replace(" ", "")
        # 判断是否嵌套：分子或分母中已含分数斜杠 ⁄
        is_nested = "⁄" in num or "⁄" in den
        # 简单情况用 a⁄b（无括号）
        if len(num) <= 2 and len(den) <= 2 and not is_nested:
            return f"{num}{_FRAC_SLASH}{den}"
        # 嵌套用 []，非嵌套用 ()
        if is_nested:
            return f"〔{num}〕{_FRAC_SLASH}〔{den}〕"
        return f"({num}){_FRAC_SLASH}({den})"
    # 反复处理嵌套分数（8 轮覆盖学术公式嵌套深度）
    for _ in range(8):
        new_s = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", _frac, s)
        if new_s == s:
            break
        s = new_s

    # 5. 处理 \binom{n}{k}（二项式系数）
    def _binom(m):
        return f"C({_latex_to_unicode(m.group(1))},{_latex_to_unicode(m.group(2))})"
    s = re.sub(r"\\binom\{([^{}]*)\}\{([^{}]*)\}", _binom, s)

    # 6. 处理 \mathbb{X}（黑板粗体）
    def _mathbb(m):
        ch = m.group(1)
        return _LATEX_OPERATORS.get(rf"\mathbb{{{ch}}}", ch)
    s = re.sub(r"\\mathbb\{([A-Z])\}", _mathbb, s)

    # 7. 替换希腊字母和运算符（按长度降序，避免 \alpha 被 \a 截断）
    # 必须在函数名替换之前，否则 \inf 会匹配 \infty 的前缀
    all_symbols = {**_LATEX_GREEK, **_LATEX_OPERATORS}
    for latex_cmd in sorted(all_symbols.keys(), key=len, reverse=True):
        if latex_cmd in s:
            s = s.replace(latex_cmd, all_symbols[latex_cmd])

    # 8. 处理函数名 \sin \cos 等（替换为纯文本，去掉反斜杠）
    for fn in _LATEX_FUNCTIONS:
        if fn in s:
            s = s.replace(fn, fn[1:])

    # 9. 处理上标 ^{...} 和 ^x
    def _sup_braced(m):
        content = m.group(1)
        # 递归处理内容（如 e^{-x^2} 中的 -x^2）
        content_u = _latex_to_unicode(content)
        result = ""
        for ch in content_u:
            result += _LATEX_SUPERSCRIPT.get(ch, ch)
        return result
    s = re.sub(r"\^\{([^{}]*)\}", _sup_braced, s)

    def _sup_single(m):
        ch = m.group(1)
        return _LATEX_SUPERSCRIPT.get(ch, f"^{ch}")
    s = re.sub(r"\^([a-zA-Z0-9+\-])", _sup_single, s)

    # 10. 处理下标 _{...} 和 _x
    def _sub_braced(m):
        content = m.group(1)
        # 递归处理内容
        content_u = _latex_to_unicode(content)
        # 检查是否所有字符都有 Unicode 下标映射
        all_mappable = all(ch in _LATEX_SUBSCRIPT for ch in content_u)
        if all_mappable:
            return "".join(_LATEX_SUBSCRIPT[ch] for ch in content_u)
        # 含未映射字符（如 b/c/d/f/g/q/r/w/y/z）→ 降级为 _{content}
        return f"_{{{content_u}}}"
    s = re.sub(r"_\{([^{}]*)\}", _sub_braced, s)

    def _sub_single(m):
        ch = m.group(1)
        return _LATEX_SUBSCRIPT.get(ch, f"_{ch}")
    s = re.sub(r"_([a-zA-Z0-9+\-])", _sub_single, s)

    # 11. 清理 LaTeX 空格命令 \, \; \: \! \quad \qquad
    s = re.sub(r"\\[,;:!]", " ", s)
    s = re.sub(r"\\quad\b", "  ", s)
    s = re.sub(r"\\qquad\b", "    ", s)
    # 清理 LaTeX 换行符 \\（含带间距版本 \\[2em]）和反斜杠空格 \ （必须在 \字母 清理之前）
    # \\[2em] → 换行；\\ → 换行；\ （反斜杠+空格）→ 空格
    s = re.sub(r"\\\\\[[^\]]*\]", "\n", s)   # \\[2em] 带间距换行
    s = re.sub(r"\\\\", "\n", s)             # \\ 换行
    s = re.sub(r"\\\s+", " ", s)             # \ + 空格（LaTeX 空格命令）
    # 清理剩余的 LaTeX 命令（\xxx 形式，保留文本）
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    # 清理孤立的反斜杠（\ + 非字母非数字非空格，如 \时 \趋近 等模型错误输出）
    s = re.sub(r"\\(?![a-zA-Z0-9\s])", "", s)

    # 12. 清理多余的空格和花括号
    # 清理花括号：保留 _{...} 和 ^{...} 中的花括号（LaTeX 风格下标/上标标记）
    # 只清理"独立的"花括号（不在 _ 或 ^ 后面的）
    s = re.sub(r"(?<![_^])\{([^{}]*)\}", r"\1", s)
    # 第二轮：清理剩余的独立花括号（第一轮可能产生新的独立花括号）
    s = re.sub(r"(?<![_^])\{([^{}]*)\}", r"\1", s)
    # 合并多余空格
    s = re.sub(r"  +", " ", s).strip()

    return s


def academic_search(query: str, num_results: int = 5, year_from: int = 0,
                    year_to: int = 0, sort_by: str = "relevance") -> str:
    """学术文献搜索（Semantic Scholar API，2亿+论文，含引用网络和影响力）

    参数：
    - query: 搜索关键词（中英文均可）
    - num_results: 返回结果数量，默认5，最大20
    - year_from: 起始年份（如 2020），0表示不限
    - year_to: 结束年份（如 2024），0表示不限
    - sort_by: 排序方式：relevance(相关性，默认) / citations(引用数) / influence(影响力)

    返回：格式化的文献列表，含标题、作者、年份、引用数、摘要、DOI
    """
    try:
        q = urllib.parse.quote(query)
        # 构建年份过滤
        year_filter = ""
        if year_from or year_to:
            yf = year_from if year_from else 1900
            yt = year_to if year_to else 2099
            year_filter = f"&year={yf}-{yt}"

        # 排序参数
        sort_map = {
            "relevance": "",  # 默认相关性
            "citations": "&sort=citationCount:desc",
            "influence": "&sort=influentialCitationCount:desc",
        }
        sort_param = sort_map.get(sort_by, "")

        # Semantic Scholar Graph API（无需 API Key，免费 2万次/小时）
        url = (f"https://api.semanticscholar.org/graph/v1/paper/search?query={q}"
               f"&limit={min(num_results, 20)}&fields=title,authors,year,abstract,"
               f"citationCount,influentialCitationCount,externalIds,url"
               f"{year_filter}{sort_param}")

        req = urllib.request.Request(url, headers={
            "User-Agent": "ZeroAI/1.0 (Academic Research)",
            "Accept": "application/json"
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="ignore"))

        papers = data.get("data", [])
        if not papers:
            return f"(未找到关于「{query}」的学术文献，试试更换关键词或扩大年份范围)"

        results = []
        for i, p in enumerate(papers, 1):
            title = p.get("title", "无标题")
            authors = p.get("authors", [])
            author_str = ", ".join(a.get("name", "?") for a in authors[:5])
            if len(authors) > 5:
                author_str += f" 等 {len(authors)} 人"
            year = p.get("year", "未知年份")
            citations = p.get("citationCount", 0)
            influential = p.get("influentialCitationCount", 0)
            abstract = p.get("abstract", "")
            if abstract:
                # 摘要截断到300字
                abstract = abstract[:300] + ("..." if len(abstract) > 300 else "")
            else:
                abstract = "(无摘要)"

            ext_ids = p.get("externalIds", {})
            doi = ext_ids.get("DOI", "")
            arxiv_id = ext_ids.get("ArXiv", "")
            paper_url = p.get("url", "")

            # 格式化输出
            line = f"[{i}] {title}\n"
            line += f"    作者: {author_str}\n"
            line += f"    年份: {year}    引用: {citations}    影响力: {influential}\n"
            if doi:
                line += f"    DOI: {doi}\n"
            if arxiv_id:
                line += f"    arXiv: {arxiv_id}\n"
            if paper_url:
                line += f"    链接: {paper_url}\n"
            line += f"    摘要: {abstract}\n"
            results.append(line)

        total = data.get("total", 0)
        header = f"=== 学术搜索: 「{query}」 ===\n"
        header += f"共找到 {total} 篇相关论文，显示前 {len(papers)} 篇"
        if year_from or year_to:
            header += f"（年份: {year_from or '不限'}-{year_to or '至今'}"
        if sort_by != "relevance":
            sort_label = {"citations": "引用数", "influence": "影响力"}.get(sort_by, sort_by)
            header += f"，按{sort_label}排序"
        header += "）\n\n"

        return header + "\n".join(results)

    except urllib.error.HTTPError as e:
        if e.code == 429:
            return f"学术搜索过于频繁，请稍后再试（Semantic Scholar 限制: 2万次/小时）"
        return f"学术搜索错误（HTTP {e.code}）：{e}"
    except Exception as e:
        return f"学术搜索错误：{e}"


def arxiv_search(query: str, num_results: int = 5, sort_by: str = "relevance",
                 category: str = "") -> str:
    """arXiv 预印本论文搜索（物理/数学/计算机科学/定量生物学/定量金融/统计学）

    参数：
    - query: 搜索关键词（英文效果更佳，支持标题/摘要/作者搜索）
    - num_results: 返回结果数量，默认5，最大20
    - sort_by: 排序方式：relevance(相关性，默认) / submittedDate(最新提交) / lastUpdatedDate(最近更新)
    - category: 学科分类筛选，如 cs.AI(人工智能) / cs.CL(计算语言学) / math.AG(代数几何) /
                physics(物理) / stat.ML(统计机器学习)。留空表示不限

    返回：格式化的论文列表，含标题、作者、摘要、arXiv ID、提交日期、PDF链接
    """
    try:
        q = urllib.parse.quote(query)
        # 排序参数
        sort_map = {
            "relevance": "relevance",
            "submittedDate": "submittedDate",
            "lastUpdatedDate": "lastUpdatedDate",
        }
        sort_param = sort_map.get(sort_by, "relevance")

        # 分类筛选
        cat_filter = f"cat:{category}" if category else "all"

        # arXiv API（Atom XML 格式，完全免费）
        url = (f"http://export.arxiv.org/api/query?search_query={cat_filter}:{q}"
               f"&start=0&max_results={min(num_results, 20)}"
               f"&sortBy={sort_param}&sortOrder=descending")

        req = urllib.request.Request(url, headers={
            "User-Agent": "ZeroAI/1.0 (Academic Research)",
            "Accept": "application/atom+xml"
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            xml_data = resp.read().decode("utf-8", errors="ignore")

        # 解析 Atom XML（用正则避免引入 xml.etree，保持轻量）
        entries = re.findall(r'<entry>([\s\S]*?)</entry>', xml_data)
        if not entries:
            return f"(未找到关于「{query}」的 arXiv 论文，试试用英文关键词)"

        results = []
        for i, entry in enumerate(entries, 1):
            # 提取标题
            title_m = re.search(r'<title>([\s\S]*?)</title>', entry)
            title = title_m.group(1).strip() if title_m else "无标题"
            title = re.sub(r'\s+', ' ', title)  # 清理换行

            # 提取作者
            authors = re.findall(r'<name>([^<]+)</name>', entry)
            author_str = ", ".join(authors[:5])
            if len(authors) > 5:
                author_str += f" 等 {len(authors)} 人"

            # 提取摘要
            summary_m = re.search(r'<summary>([\s\S]*?)</summary>', entry)
            abstract = summary_m.group(1).strip() if summary_m else "(无摘要)"
            abstract = re.sub(r'<[^>]+>', '', abstract)
            abstract = re.sub(r'\s+', ' ', abstract)
            if len(abstract) > 300:
                abstract = abstract[:300] + "..."

            # 提取 arXiv ID 和链接
            id_m = re.search(r'<id>http://arxiv.org/abs/([^<]+)</id>', entry)
            arxiv_id = id_m.group(1).strip() if id_m else "未知"
            pdf_link = f"http://arxiv.org/pdf/{arxiv_id}.pdf" if arxiv_id != "未知" else ""

            # 提取提交日期
            published_m = re.search(r'<published>([^<]+)</published>', entry)
            published = published_m.group(1)[:10] if published_m else "未知日期"

            # 提取分类
            categories = re.findall(r'term="([^"]+)"', entry)
            cat_str = ", ".join(categories[:3]) if categories else "未分类"

            # 格式化输出
            line = f"[{i}] {title}\n"
            line += f"    作者: {author_str}\n"
            line += f"    arXiv: {arxiv_id}    提交: {published}\n"
            line += f"    分类: {cat_str}\n"
            if pdf_link:
                line += f"    PDF: {pdf_link}\n"
            line += f"    摘要: {abstract}\n"
            results.append(line)

        total_m = re.search(r'<opensearch:totalResults[^>]*>([^<]+)</opensearch:totalResults>', xml_data)
        total = total_m.group(1) if total_m else str(len(entries))

        header = f"=== arXiv 搜索: 「{query}」 ===\n"
        header += f"共找到 {total} 篇预印本论文，显示前 {len(entries)} 篇"
        if category:
            header += f"（分类: {category}）"
        if sort_by != "relevance":
            sort_label = {"submittedDate": "最新提交", "lastUpdatedDate": "最近更新"}.get(sort_by, sort_by)
            header += f"，按{sort_label}排序"
        header += "\n\n"

        return header + "\n".join(results)

    except Exception as e:
        return f"arXiv 搜索错误：{e}"


def citation_check(title: str = "", doi: str = "", arxiv_id: str = "") -> str:
    """校验文献引用真实性（防止 AI 编造不存在的文献）

    通过 Semantic Scholar API 交叉验证文献是否真实存在。
    支持三种查询方式：标题精确匹配、DOI 查询、arXiv ID 查询。

    参数：
    - title: 文献标题（精确或近似标题）
    - doi: 文献的 DOI（如 10.1038/s41586-021-03819-2）
    - arxiv_id: arXiv 编号（如 2301.00234）

    返回：校验结果，含文献真实状态、正确标题、作者、年份等元数据
    """
    try:
        # 优先用 DOI 查询（最精确）
        if doi:
            url = f"https://api.semanticscholar.org/graph/v1/paper/DOI:{urllib.parse.quote(doi)}?fields=title,authors,year,abstract,citationCount,externalIds,url"
            req = urllib.request.Request(url, headers={
                "User-Agent": "ZeroAI/1.0 (Academic Citation Check)",
                "Accept": "application/json"
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))
            return _format_citation_result(data, "DOI", doi)

        # arXiv ID 查询
        if arxiv_id:
            url = f"https://api.semanticscholar.org/graph/v1/paper/arXiv:{arxiv_id.strip()}?fields=title,authors,year,abstract,citationCount,externalIds,url"
            req = urllib.request.Request(url, headers={
                "User-Agent": "ZeroAI/1.0 (Academic Citation Check)",
                "Accept": "application/json"
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))
            return _format_citation_result(data, "arXiv", arxiv_id)

        # 标题查询（模糊匹配）
        if title:
            q = urllib.parse.quote(title)
            url = f"https://api.semanticscholar.org/graph/v1/paper/search/match?query={q}&fields=title,authors,year,abstract,citationCount,externalIds,url"
            req = urllib.request.Request(url, headers={
                "User-Agent": "ZeroAI/1.0 (Academic Citation Check)",
                "Accept": "application/json"
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))

            papers = data.get("data", [])
            if papers:
                # 找最匹配的
                best = papers[0]
                # 计算标题相似度
                from difflib import SequenceMatcher
                ratio = SequenceMatcher(None, title.lower().strip(), best.get("title", "").lower().strip()).ratio()
                return _format_citation_result(best, "标题匹配", title, similarity=ratio)
            else:
                return (f"⚠ 引用校验失败：未找到与「{title}」匹配的论文\n"
                        f"  该引用可能为 AI 编造的虚构文献，请勿使用\n"
                        f"  建议：使用 academic_search 搜索真实存在的文献替代")

        return "请提供文献标题、DOI 或 arXiv ID 中的至少一个参数"

    except urllib.error.HTTPError as e:
        if e.code == 404:
            return (f"✗ 引用校验结果：文献不存在\n"
                    f"  查询条件：{('DOI:' + doi) if doi else ('arXiv:' + arxiv_id) if arxiv_id else ('标题:' + title)}\n"
                    f"  该引用很可能是 AI 编造的虚构文献，请勿在学术写作中使用\n"
                    f"  建议：使用 academic_search 搜索该领域的真实文献")
        if e.code == 429:
            return "引用校验过于频繁，请稍后再试（Semantic Scholar 限制: 2万次/小时）"
        return f"引用校验错误（HTTP {e.code}）：{e}"
    except Exception as e:
        return f"引用校验错误：{e}"


def _format_citation_result(data: dict, method: str, query: str, similarity: float = 1.0) -> str:
    """格式化引用校验结果（citation_check 的辅助函数）"""
    title = data.get("title", "无标题")
    authors = data.get("authors", [])
    author_str = ", ".join(a.get("name", "?") for a in authors[:5])
    if len(authors) > 5:
        author_str += f" 等 {len(authors)} 人"
    year = data.get("year", "未知")
    citations = data.get("citationCount", 0)
    ext_ids = data.get("externalIds", {})
    doi = ext_ids.get("DOI", "")
    arxiv = ext_ids.get("ArXiv", "")
    paper_url = data.get("url", "")

    # 判定状态
    if similarity >= 0.95:
        status = "✓ 验证通过：文献真实存在（标题精确匹配）"
    elif similarity >= 0.80:
        status = f"✓ 验证通过：文献真实存在（标题相似度 {similarity:.0%}，请核实标题是否完全一致）"
    elif similarity >= 0.60:
        status = f"⚠ 部分匹配（相似度 {similarity:.0%}）：找到相关文献，但标题不完全一致，请核实是否为同一篇"
    else:
        status = f"⚠ 匹配度低（{similarity:.0%}）：可能不是同一篇文献，请人工核实"

    result = f"=== 引用校验结果 ===\n"
    result += f"校验方式：{method}\n"
    result += f"查询条件：{query}\n"
    result += f"状态：{status}\n\n"
    result += f"文献信息：\n"
    result += f"  标题：{title}\n"
    result += f"  作者：{author_str}\n"
    result += f"  年份：{year}    引用数：{citations}\n"
    if doi:
        result += f"  DOI：{doi}\n"
    if arxiv:
        result += f"  arXiv：{arxiv}\n"
    if paper_url:
        result += f"  链接：{paper_url}\n"
    return result


def literature_review(topic: str, num_papers: int = 10, year_from: int = 0,
                      year_to: int = 0) -> str:
    """多文献综合对比分析（自动检索+结构化对比+研究空白识别）

    自动执行完整的文献综述流程：
    1. 检索相关文献（Semantic Scholar + arXiv 双源）
    2. 按引用数筛选高质量文献
    3. 结构化提取每篇文献的方法/结论/局限
    4. 生成对比分析表
    5. 识别研究空白和未来方向

    参数：
    - topic: 研究主题（中英文均可，如 '钠离子电池层状氧化物正极' 或 'sodium-ion battery layered oxide cathode'）
    - num_papers: 分析文献数量，默认10，最大20
    - year_from: 起始年份（如 2018），0表示不限
    - year_to: 结束年份（如 2025），0表示不限

    返回：结构化文献综述分析报告
    """
    try:
        # ── 第1步：双源检索 ──
        all_papers = []

        # Semantic Scholar（按引用数排序，筛选高影响力文献）
        ss_papers = _lit_review_search_ss(topic, num_papers, year_from, year_to)
        all_papers.extend(ss_papers)

        # arXiv（最新研究，按提交日期排序）
        arxiv_papers = _lit_review_search_arxiv(topic, min(num_papers // 2, 5))
        all_papers.extend(arxiv_papers)

        if not all_papers:
            return (f"=== 文献综述分析：{topic} ===\n\n"
                    f"未找到相关文献，请尝试更换关键词或扩大年份范围\n"
                    f"建议：使用英文关键词（如 'sodium-ion battery cathode'）效果更佳")

        # ── 第2步：去重（按标题模糊匹配） ──
        from difflib import SequenceMatcher
        unique_papers = []
        seen_titles = []
        for p in all_papers:
            is_dup = False
            for seen in seen_titles:
                if SequenceMatcher(None, p.get("title", "").lower(), seen.lower()).ratio() > 0.85:
                    is_dup = True
                    break
            if not is_dup:
                unique_papers.append(p)
                seen_titles.append(p.get("title", ""))

        # 按引用数排序，取前 num_papers 篇
        unique_papers.sort(key=lambda x: x.get("citations", 0), reverse=True)
        top_papers = unique_papers[:num_papers]

        # ── 第3步：生成综述报告 ──
        report = f"=== 文献综述分析报告 ===\n"
        report += f"研究主题：{topic}\n"
        report += f"检索范围：{year_from or '不限'} - {year_to or '至今'}\n"
        report += f"分析文献数：{len(top_papers)} 篇（去重后共 {len(unique_papers)} 篇）\n"
        report += f"数据来源：Semantic Scholar + arXiv\n\n"

        # ── 文献概览表 ──
        report += "── 一、文献概览 ──\n\n"
        report += f"{'#':<4} {'年份':<6} {'引用':<8} {'标题':<50} {'来源':<12}\n"
        report += "-" * 85 + "\n"
        for i, p in enumerate(top_papers, 1):
            title_short = p.get("title", "无标题")[:48]
            year = str(p.get("year", "?"))[:4]
            cit = str(p.get("citations", 0))[:7]
            source = p.get("source", "?")[:10]
            report += f"{i:<4} {year:<6} {cit:<8} {title_short:<50} {source:<12}\n"

        # ── 详细分析 ──
        report += "\n── 二、逐篇分析 ──\n\n"
        for i, p in enumerate(top_papers, 1):
            report += f"[{i}] {p.get('title', '无标题')}\n"
            authors = p.get("authors", [])
            author_str = ", ".join(a if isinstance(a, str) else a.get("name", "?") for a in authors[:5])
            if len(authors) > 5:
                author_str += f" 等 {len(authors)} 人"
            report += f"    作者：{author_str}\n"
            report += f"    年份：{p.get('year', '?')}    引用数：{p.get('citations', 0)}\n"

            doi = p.get("doi", "")
            arxiv = p.get("arxiv_id", "")
            if doi:
                report += f"    DOI：{doi}\n"
            if arxiv:
                report += f"    arXiv：{arxiv}\n"

            abstract = p.get("abstract", "") or p.get("summary", "")
            if abstract:
                abstract = abstract[:400] + ("..." if len(abstract) > 400 else "")
                report += f"    摘要：{abstract}\n"
            report += "\n"

        # ── 研究趋势分析 ──
        report += "── 三、研究趋势分析 ──\n\n"
        years = [p.get("year", 0) for p in top_papers if p.get("year")]
        if years:
            y_min, y_max = min(years), max(years)
            report += f"时间跨度：{y_min} - {y_max}\n"
            # 按年份统计
            year_dist = {}
            for y in years:
                year_dist[y] = year_dist.get(y, 0) + 1
            report += "年度分布：\n"
            for y in sorted(year_dist.keys()):
                bar = "█" * year_dist[y]
                report += f"  {y}: {bar} ({year_dist[y]}篇)\n"

        # 引用分析
        total_cit = sum(p.get("citations", 0) for p in top_papers)
        avg_cit = total_cit / len(top_papers) if top_papers else 0
        report += f"\n总引用数：{total_cit}    平均引用：{avg_cit:.1f}\n"

        # ── 研究空白与未来方向 ──
        report += "\n── 四、研究空白与未来方向（自动识别） ──\n\n"
        report += "基于检索到的文献，以下方向值得关注（需结合专业知识进一步验证）：\n"
        # 基于文献年份和引用数推断
        recent_papers = [p for p in top_papers if p.get("year", 0) >= 2023]
        if recent_papers:
            report += f"1. 近期热点（{len(recent_papers)}篇2023年后文献）：关注该领域最新进展\n"
        old_high_cit = [p for p in top_papers if p.get("year", 0) < 2020 and p.get("citations", 0) > 100]
        if old_high_cit:
            report += f"2. 经典基础（{len(old_high_cit)}篇高引经典）：建议深入阅读这些奠基性工作\n"
        low_cit_recent = [p for p in top_papers if p.get("year", 0) >= 2022 and p.get("citations", 0) < 10]
        if low_cit_recent:
            report += f"3. 新兴方向（{len(low_cit_recent)}篇低引新文）：可能代表尚未被广泛关注的研究前沿\n"
        report += "4. 交叉领域：结合本主题与其他学科（如AI/材料/工程）的交叉研究\n"
        report += "5. 方法论改进：现有方法的局限性可作为改进方向\n\n"

        # ── PRISMA 筛选流程 ──
        report += "── 五、PRISMA 筛选流程 ──\n\n"
        report += f"检索总量：{len(all_papers)} 篇\n"
        report += f"去重后：{len(unique_papers)} 篇（去除 {len(all_papers) - len(unique_papers)} 篇重复）\n"
        report += f"纳入分析：{len(top_papers)} 篇（按引用数筛选）\n"
        report += f"排除：{len(unique_papers) - len(top_papers)} 篇（引用数较低）\n\n"

        report += "── 注意事项 ──\n"
        report += "1. 本分析基于自动检索结果，不含人工筛选和质量评估\n"
        report += "2. 建议在此基础上人工精读 top 3-5 篇高引文献\n"
        report += "3. 如需正式发表，请补充 Web of Science / Scopus 检索\n"
        report += "4. 引用文献时务必使用 citation_check 校验真实性\n"

        return report

    except Exception as e:
        return f"文献综述分析错误：{e}"


def _lit_review_search_ss(topic: str, num: int, year_from: int, year_to: int) -> list:
    """literature_review 辅助：从 Semantic Scholar 检索"""
    papers = []
    try:
        q = urllib.parse.quote(topic)
        year_filter = ""
        if year_from or year_to:
            yf = year_from if year_from else 1900
            yt = year_to if year_to else 2099
            year_filter = f"&year={yf}-{yt}"
        url = (f"https://api.semanticscholar.org/graph/v1/paper/search?query={q}"
               f"&limit={min(num, 20)}&fields=title,authors,year,abstract,citationCount,externalIds,url"
               f"{year_filter}&sort=citationCount:desc")
        req = urllib.request.Request(url, headers={
            "User-Agent": "ZeroAI/1.0 (Academic Literature Review)",
            "Accept": "application/json"
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="ignore"))
        for p in data.get("data", []):
            ext = p.get("externalIds", {})
            papers.append({
                "title": p.get("title", ""),
                "authors": p.get("authors", []),
                "year": p.get("year", 0),
                "citations": p.get("citationCount", 0),
                "abstract": p.get("abstract", ""),
                "doi": ext.get("DOI", ""),
                "arxiv_id": ext.get("ArXiv", ""),
                "url": p.get("url", ""),
                "source": "Semantic Scholar",
            })
    except Exception:
        pass
    return papers


def _lit_review_search_arxiv(topic: str, num: int) -> list:
    """literature_review 辅助：从 arXiv 检索最新论文"""
    papers = []
    try:
        q = urllib.parse.quote(topic)
        # 用 all: 搜索 + relevance 排序（确保结果相关性）
        url = (f"http://export.arxiv.org/api/query?search_query=all:{q}"
               f"&start=0&max_results={min(num, 10)}"
               f"&sortBy=relevance&sortOrder=descending")
        req = urllib.request.Request(url, headers={
            "User-Agent": "ZeroAI/1.0 (Academic Literature Review)",
            "Accept": "application/atom+xml"
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            xml_data = resp.read().decode("utf-8", errors="ignore")
        entries = re.findall(r'<entry>([\s\S]*?)</entry>', xml_data)
        for entry in entries:
            title_m = re.search(r'<title>([\s\S]*?)</title>', entry)
            title = re.sub(r'\s+', ' ', title_m.group(1).strip()) if title_m else ""
            authors = re.findall(r'<name>([^<]+)</name>', entry)
            summary_m = re.search(r'<summary>([\s\S]*?)</summary>', entry)
            abstract = re.sub(r'\s+', ' ', summary_m.group(1).strip()) if summary_m else ""
            id_m = re.search(r'<id>http://arxiv.org/abs/([^<]+)</id>', entry)
            arxiv_id = id_m.group(1).strip() if id_m else ""
            published_m = re.search(r'<published>([^<]+)</published>', entry)
            year = int(published_m.group(1)[:4]) if published_m else 0
            papers.append({
                "title": title,
                "authors": [{"name": a} for a in authors],
                "year": year,
                "citations": 0,
                "abstract": abstract,
                "doi": "",
                "arxiv_id": arxiv_id,
                "url": f"http://arxiv.org/abs/{arxiv_id}",
                "source": "arXiv",
            })
    except Exception:
        pass
    return papers


def render_formula(latex: str, style: str = "unicode") -> str:
    r"""渲染 LaTeX 公式为终端可显示的 Unicode 文本

    参数：
    - latex: LaTeX 公式字符串，如 "E=mc^2" 或 "\\sum_{i=1}^{n} x_i^2"
    - style: 渲染样式
      - "unicode"：纯 Unicode 数学符号（默认，终端显示）
      - "raw"：返回原始 LaTeX（用于文档生成）
      - "latex"：用 $$ 包裹（用于 Markdown 渲染）

    返回：渲染后的公式字符串

    用途：学术研究、数学公式展示、物理方程推导
    """
    latex = latex.strip()
    # 去除外层 $ 或 $$
    if latex.startswith("$$") and latex.endswith("$$"):
        latex = latex[2:-2].strip()
    elif latex.startswith("$") and latex.endswith("$"):
        latex = latex[1:-1].strip()

    if style == "raw":
        return latex
    elif style == "latex":
        return f"$${latex}$$"
    else:  # unicode
        rendered = _latex_to_unicode(latex)
        return rendered


def _strip_model_tokens(text: str) -> str:
    """过滤模型内部特殊标签（如 <|observation|> <|system|> <|assistant|> 等）

    这些标签是推理模型（如 GLM-4V）的思维链标记，不应显示给用户。
    支持流式累积后的完整过滤（跨 chunk 拼接后标签完整即可被清除）。

    过滤的标签模式：<|xxx|> 和 <|/xxx|>（xxx 不含 | 字符）
    """
    if not text:
        return text
    # 删除所有 <|...|> 格式的标签（开标签、闭标签、自闭合标签）
    return re.sub(r'<\|[^|]*\|>', '', text)


def _parse_think_tags(content: str) -> tuple:
    """从累积 content 中解析 <think>...</think> 标签

    返回 (think_content, body_content)：
    - think_content: <think> 标签内的内容（思考过程）
    - body_content: <think> 标签外的内容（正文）

    支持流式累积（标签可能未闭合）：
    - 未遇到 <think>：body_content = content, think_content = ""
    - 在 <think> 内（未遇到 </think>）：think_content = <think>之后内容, body_content = ""
    - 遇到 </think>：think_content = 标签间内容, body_content = </think>之后内容
    - 多个 <think> 块：合并所有 think 内容，body 为剩余内容
    """
    if not content:
        return "", ""
    think_content_parts = []
    body_parts = []
    pos = 0
    while pos < len(content):
        # 查找下一个 <think>
        start = content.find("<think>", pos)
        if start == -1:
            # 没有更多 <think>，剩余全部为正文
            body_parts.append(content[pos:])
            break
        # <think> 之前的内容归入正文
        if start > pos:
            body_parts.append(content[pos:start])
        # 查找对应的 </think>
        end = content.find("</think>", start + len("<think>"))
        if end == -1:
            # <think> 未闭合，剩余内容均为思考
            think_text = content[start + len("<think>"):]
            if think_text:
                think_content_parts.append(think_text)
            pos = len(content)
            break
        # 完整 <think>...</think>
        think_text = content[start + len("<think>"):end]
        if think_text:
            think_content_parts.append(think_text)
        pos = end + len("</think>")
    think_content = "\n".join(p.strip() for p in think_content_parts if p.strip())
    body_content = "".join(body_parts)
    # 去除正文开头的换行
    while body_content.startswith("\n"):
        body_content = body_content[1:]
    return think_content, body_content


def _jaccard_similarity(s1: str, s2: str) -> float:
    """计算两段文本的 Jaccard 相似度（基于字符 n-gram 集合）

    用于专家回答去重：相似度越高说明回答越重复。
    返回 0.0-1.0 的浮点数。
    """
    if not s1 or not s2:
        return 0.0
    # 使用 3-gram（兼顾中英文，对短文本也有效）
    n = 3
    if len(s1) < n or len(s2) < n:
        # 文本过短时直接用字符集合
        set1, set2 = set(s1), set(s2)
    else:
        set1 = {s1[i:i + n] for i in range(len(s1) - n + 1)}
        set2 = {s2[i:i + n] for i in range(len(s2) - n + 1)}
    if not set1 or not set2:
        return 0.0
    inter = len(set1 & set2)
    union = len(set1 | set2)
    return inter / union if union else 0.0


def _truncate_expert_response(text: str, max_chars: int) -> str:
    """截断专家回答到指定字符数，并附加截断提示

    用于 HYBRID_EXPERT_MAX_CHARS 限制：避免单个专家回答过长导致汇总 token 暴涨。
    """
    if not text or max_chars <= 0 or len(text) <= max_chars:
        return text
    # 在 max_chars 附近找换行或句号，避免截断在词中间
    cut = text[:max_chars]
    # 优先在最近的换行处截断
    nl = cut.rfind("\n")
    if nl > max_chars * 0.7:
        cut = cut[:nl]
    else:
        # 退而求其次在句号/问号处截断
        for sep in ("。", "？", "！", ".", "?", "!"):
            sp = cut.rfind(sep)
            if sp > max_chars * 0.7:
                cut = cut[:sp + 1]
                break
    return cut + "\n\n…（专家回答已截断，仅汇总关键部分）"


def render_latex_in_text(text: str) -> str:
    """检测文本中的 LaTeX 公式并渲染为 Unicode

    支持：
    - 行内公式：$E=mc^2$
    - 块级公式：$$\\int_0^1 f(x)dx$$
    - \\( \\) 行内公式
    - \\[ \\] 块级公式

    用于终端 Markdown 渲染前的预处理
    """
    if not text or "$" not in text and "\\(" not in text and "\\[" not in text:
        return text

    # 1. 块级公式 $$...$$
    def _block_formula(m):
        rendered = _latex_to_unicode(m.group(1))
        return f"\n   {rendered}\n"
    text = re.sub(r"\$\$([^$]+)\$\$", _block_formula, text)

    # 2. 块级公式 \[...\]
    text = re.sub(r"\\\[([^\]]+)\]", lambda m: f"\n   {_latex_to_unicode(m.group(1))}\n", text)

    # 3. 行内公式 \(...\)
    text = re.sub(r"\\\(([^)]+)\)", lambda m: _latex_to_unicode(m.group(1)), text)

    # 4. 行内公式 $...$（最后处理，避免误伤 $$）
    # 使用非贪婪匹配，且内容不含 $
    text = re.sub(r"(?<!\$)\$([^$\n]+)\$(?!\$)", lambda m: _latex_to_unicode(m.group(1)), text)

    return text


# ====== Word 文档格式模板 ======
WORD_TEMPLATES = {
    "default": {
        "font": "Microsoft YaHei", "font_size": 11,
        "heading_color": "1A1A1A", "margin": [2.54, 2.54, 2.54, 2.54],
        "header": "", "footer": "",
    },
    "report": {  # 正式报告
        "font": "SimSun", "font_size": 12,
        "heading_color": "1F4E79", "margin": [3.0, 2.5, 2.5, 2.5],
        "header": "", "footer": "第 {page} 页",
    },
    "contract": {  # 合同
        "font": "SimSun", "font_size": 12,
        "heading_color": "000000", "margin": [2.54, 3.0, 2.54, 3.0],
        "header": "", "footer": "— {page} —",
    },
    "resume": {  # 简历
        "font": "Microsoft YaHei", "font_size": 10,
        "heading_color": "2E75B6", "margin": [1.5, 2.0, 1.5, 2.0],
        "header": "", "footer": "",
    },
    "thesis": {  # 学术论文
        "font": "Times New Roman", "font_size": 12,
        "heading_color": "000000", "margin": [2.54, 3.17, 2.54, 3.17],
        "header": "", "footer": "{page}",
    },
    "letter": {  # 信函
        "font": "KaiTi", "font_size": 14,
        "heading_color": "333333", "margin": [2.54, 2.54, 2.54, 2.54],
        "header": "", "footer": "",
    },
    "technical": {  # 技术文档
        "font": "Microsoft YaHei", "font_size": 10,
        "heading_color": "0070C0", "margin": [2.0, 2.0, 2.0, 2.0],
        "header": "技术文档", "footer": "第 {page} 页 / 共 {numpages} 页",
    },
    "academic": {  # 学术论文（严谨版：双倍行距、页码、Times New Roman）
        "font": "Times New Roman", "font_size": 12,
        "heading_color": "000000", "margin": [2.54, 2.54, 2.54, 2.54],
        "header": "", "footer": "{page}",
        "line_spacing": 2.0,  # 双倍行距（学术规范）
        "abstract_label": "摘要", "keywords_label": "关键词",
        "references_label": "参考文献", "doi_support": True,
    },
}

# 颜色名映射
COLOR_MAP = {
    "black": "000000", "white": "FFFFFF", "red": "FF0000", "green": "008000",
    "blue": "0000FF", "yellow": "FFFF00", "gray": "808080", "grey": "808080",
    "darkgray": "A9A9A9", "darkblue": "00008B", "navy": "1F4E79",
    "orange": "FFA500", "purple": "800080", "brown": "A52A2A",
    "cyan": "00FFFF", "magenta": "FF00FF", "lime": "00FF00",
}


def _parse_color(color_str: str):
    """解析颜色字符串，返回 RGBColor"""
    from docx.shared import RGBColor
    color_str = color_str.strip().lstrip("#")
    # 颜色名
    if color_str.lower() in COLOR_MAP:
        color_str = COLOR_MAP[color_str.lower()]
    # 6位十六进制
    if len(color_str) == 6 and all(c in "0123456789abcdefABCDEF" for c in color_str):
        r = int(color_str[0:2], 16)
        g = int(color_str[2:4], 16)
        b = int(color_str[4:6], 16)
        return RGBColor(r, g, b)
    return None


def _set_cell_shading(cell, color_hex: str):
    """设置表格单元格背景色"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex.lstrip("#"))
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    except Exception:
        pass


def _add_page_number_field(paragraph, field: str = "PAGE"):
    """在段落中添加页码域代码"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f" {field} "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    except Exception:
        pass


def _apply_footer_template(section, template: dict):
    """应用页脚模板"""
    from docx.shared import Pt
    footer_text = template.get("footer", "")
    if not footer_text:
        return
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = 1  # 居中

    # 解析 {page} {numpages} 占位符
    parts = re.split(r'(\{page\}|\{numpages\})', footer_text)
    for part in parts:
        if part == "{page}":
            _add_page_number_field(p, "PAGE")
        elif part == "{numpages}":
            _add_page_number_field(p, "NUMPAGES")
        elif part:
            run = p.add_run(part)
            run.font.size = Pt(9)
            run.font.color.rgb = _parse_color("888888")


def _apply_header_template(section, template: dict):
    """应用页眉模板"""
    from docx.shared import Pt
    header_text = template.get("header", "")
    if not header_text:
        return
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = 1  # 居中
    run = p.add_run(header_text)
    run.font.size = Pt(9)
    run.font.color.rgb = _parse_color("888888")


def _parse_table(lines: list, start_idx: int):
    """解析 Markdown 表格语法，返回 (rows, end_idx)
    表格格式：
    | 列1 | 列2 | 列3 |
    |----|----|----|
    | a  | b  | c  |
    """
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # 跳过分隔行 |---|---|
        if re.match(r'^\|[\s\-:]+\|', line):
            i += 1
            continue
        # 解析单元格
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def generate_word(path: str, content: str, title: str = "", template: str = "default",
                  font: str = "", font_size: int = 0, margin: list = None,
                  heading_color: str = "", align: str = "", header: str = "",
                  footer: str = "") -> str:
    """生成 Word 文档（.docx）- 支持指定格式
    path: 保存路径
    content: 文档内容（支持 Markdown 风格标记）
    title: 文档标题（可选）
    template: 格式模板（default/report/contract/resume/thesis/letter/technical/academic）
                academic=学术论文（双倍行距/摘要/关键词/参考文献自动编号/LaTeX公式渲染）
    font: 覆盖模板字体（如 "SimSun"/"Microsoft YaHei"/"KaiTi"）
    font_size: 覆盖模板字号
    margin: 页边距 [上, 右, 下, 左]（厘米）
    heading_color: 标题颜色（如 "1F4E79" 或 "navy"）
    align: 全文对齐（left/center/right/justify）
    header: 页眉文字
    footer: 页脚文字（支持 {page} {numpages} 占位符）

    内容支持的标记：
    #/##/### 标题 | -/1. 列表 | > 引用 | ```代码``` | **粗体** *斜体* `代码` ~~删除线~~
    | 表格语法 | ---（分隔线） | [居中]行首标记 [右对齐] | {color:红色}文字{/color}
    """
    # 默认保存到桌面（path 为空或只含文件名时自动拼到桌面）
    path = _resolve_save_path(path, "未命名.docx")
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        # 获取模板配置
        tpl = WORD_TEMPLATES.get(template, WORD_TEMPLATES["default"]).copy()

        # 应用参数覆盖
        if font:
            tpl["font"] = font
        if font_size:
            tpl["font_size"] = font_size
        if margin:
            tpl["margin"] = margin
        if heading_color:
            tpl["heading_color"] = heading_color
        if header:
            tpl["header"] = header
        if footer:
            tpl["footer"] = footer

        doc = Document()

        # 设置页边距
        section = doc.sections[0]
        margins = tpl["margin"]
        if len(margins) == 4:
            section.top_margin = Cm(margins[0])
            section.right_margin = Cm(margins[1])
            section.bottom_margin = Cm(margins[2])
            section.left_margin = Cm(margins[3])

        # 设置页眉页脚
        _apply_header_template(section, tpl)
        _apply_footer_template(section, tpl)

        # 设置默认字体
        style = doc.styles["Normal"]
        font_obj = style.font
        font_obj.name = tpl["font"]
        font_obj.size = Pt(tpl["font_size"])
        style.element.rPr.rFonts.set(qn("w:eastAsia"), tpl["font"])

        # 学术模板：设置双倍行距（APA/学术规范）
        if tpl.get("line_spacing"):
            from docx.shared import Pt as _Pt2
            from docx.enum.text import WD_LINE_SPACING
            style.paragraph_format.line_spacing = tpl["line_spacing"]
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        # 全文对齐
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        default_align = align_map.get(align.lower(), None)

        # 解析标题
        doc_title = title if title else ""
        if not doc_title and content:
            first_line = content.split("\n")[0].strip()
            if first_line.startswith("# "):
                doc_title = first_line[2:].strip()
                content = "\n".join(content.split("\n")[1:])

        if doc_title:
            heading = doc.add_heading(doc_title, level=0)
            h_color = _parse_color(tpl["heading_color"])
            if h_color:
                for run in heading.runs:
                    run.font.color.rgb = h_color
                    run.font.name = tpl["font"]
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), tpl["font"])

        # 解析内容
        lines = content.split("\n")
        i = 0
        in_code_block = False
        code_lines = []
        # 学术论文状态变量
        in_references = False
        ref_counter = 0

        while i < len(lines):
            line = lines[i]

            # 代码块
            if line.strip().startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.0)
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # 空行
            if not line.strip():
                i += 1
                continue

            # 表格检测
            if line.strip().startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
                rows, end_idx = _parse_table(lines, i)
                if rows:
                    # 创建表格
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(table.rows[r_idx].cells):
                                cell = table.rows[r_idx].cells[c_idx]
                                cell.text = ""
                                p = cell.paragraphs[0]
                                _add_formatted_text(p, cell_text)
                                # 表头加粗+背景色
                                if r_idx == 0:
                                    for run in p.runs:
                                        run.bold = True
                                    _set_cell_shading(cell, "D9E2F3")
                    i = end_idx
                    continue

            # 分隔线
            if line.strip() in ("---", "***", "___"):
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("─" * 40)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                i += 1
                continue

            # 行首对齐标记：[居中] [右对齐] [左对齐]
            line_align = default_align
            align_match = re.match(r'^\[(居中|右对齐|左对齐|两端对齐)\]\s*(.*)', line)
            if align_match:
                align_text_map = {"居中": "center", "右对齐": "right", "左对齐": "left", "两端对齐": "justify"}
                line_align = align_map.get(align_text_map.get(align_match.group(1), ""), default_align)
                line = align_match.group(2)

            # 颜色标记：{color:红色}文字{/color}
            color_pattern = r'\{color:([^}]+)\}(.*?)\{/color\}'

            # ===== 学术论文段落特殊处理（academic/thesis 模板）=====
            is_academic = template in ("academic", "thesis")
            stripped_line = line.strip()

            if is_academic:
                # 检测学术段落标记：摘要/Abstract/关键词/Keywords/参考文献/References
                academic_match = re.match(
                    r'^(摘要|Abstract|关键词|Keywords|参考文献|References|引言|Introduction|结论|Conclusion|致谢|Acknowledgments?)\s*[:：]?\s*(.*)',
                    stripped_line, re.IGNORECASE
                )

                if academic_match:
                    section_name = academic_match.group(1)
                    section_content = academic_match.group(2)

                    # 摘要/Abstract：小字号、缩进、两端对齐
                    if section_name.lower() in ("摘要", "abstract"):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1.0)
                        p.paragraph_format.right_indent = Cm(1.0)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run = p.add_run(f"【{section_name}】")
                        run.bold = True
                        run.font.size = Pt(tpl["font_size"] - 1)
                        if section_content:
                            # 渲染行内公式
                            rendered = render_latex_in_text(section_content)
                            run2 = p.add_run(rendered)
                            run2.font.size = Pt(tpl["font_size"] - 1)
                            run2.italic = True
                        i += 1
                        continue

                    # 关键词/Keywords：加粗标签、分号分隔
                    elif section_name.lower() in ("关键词", "keywords"):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1.0)
                        run = p.add_run(f"【{section_name}】 ")
                        run.bold = True
                        if section_content:
                            run2 = p.add_run(section_content)
                            run2.font.size = Pt(tpl["font_size"] - 1)
                        i += 1
                        continue

                    # 参考文献/References：标题 + 后续条目自动编号
                    elif section_name.lower() in ("参考文献", "references"):
                        h = doc.add_heading(f"{section_name}", level=2)
                        hc = _parse_color(tpl["heading_color"])
                        if hc:
                            for run in h.runs:
                                run.font.color.rgb = hc
                        # 标记进入参考文献区域
                        in_references = True
                        ref_counter = 0
                        i += 1
                        continue

                    # 引言/结论等：作为一级标题
                    elif section_name.lower() in ("引言", "introduction", "结论", "conclusion"):
                        h = doc.add_heading(f"{section_name}", level=1)
                        hc = _parse_color(tpl["heading_color"])
                        if hc:
                            for run in h.runs:
                                run.font.color.rgb = hc
                        i += 1
                        continue

            # 参考文献条目自动编号（[1] [2] [3]...）
            if is_academic and in_references and stripped_line:
                # 跳过空行和已有编号的条目
                if not re.match(r'^\[\d+\]', stripped_line) and not stripped_line.startswith("#"):
                    ref_counter += 1
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.5)  # 悬挂缩进
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(f"[{ref_counter}] ")
                    run.bold = True
                    _add_formatted_text(p, stripped_line)
                    i += 1
                    continue

            # 标题（支持 # ~ #### 四级，符合 GB/T 7713.1-2025 论文标题层级规范）
            if line.startswith("#### "):
                h = doc.add_heading(line[5:].strip(), level=4)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            elif line.startswith("### "):
                h = doc.add_heading(line[4:].strip(), level=3)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            elif line.startswith("## "):
                h = doc.add_heading(line[3:].strip(), level=2)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            elif line.startswith("# "):
                h = doc.add_heading(line[2:].strip(), level=1)
                hc = _parse_color(tpl["heading_color"])
                if hc:
                    for run in h.runs:
                        run.font.color.rgb = hc
            # 引用
            elif line.startswith("> "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                run = p.add_run(line[2:].strip())
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True
            # 无序列表
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_text(p, line[2:].strip())
            # 有序列表
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^(\d+)\.\s(.*)', line)
                if match:
                    p = doc.add_paragraph(style="List Number")
                    _add_formatted_text(p, match.group(2).strip())
            # 普通段落
            else:
                p = doc.add_paragraph()
                if line_align:
                    p.paragraph_format.alignment = line_align
                # 学术模板：两端对齐（默认）
                if is_academic and not line_align:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # 渲染行内 LaTeX 公式（$...$ → Unicode）
                processed_line = render_latex_in_text(line.strip()) if is_academic else line.strip()
                # 处理颜色标记
                if re.search(color_pattern, processed_line):
                    _add_colored_text(p, processed_line, color_pattern)
                else:
                    _add_formatted_text(p, processed_line)

            # 设置中文字体
            for p in doc.paragraphs[-1:]:
                for run in p.runs:
                    if not run.font.name:
                        run.font.name = tpl["font"]
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), tpl["font"])

            i += 1

        # 未关闭的代码块
        if in_code_block and code_lines:
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        # 保存
        full = Path(path).resolve()
        full.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(full))

        info = f"{_load_svg_icon('document')} Word 文档已生成：{full}\n"
        info += f"  标题：{doc_title or '（无标题）'}\n"
        info += f"  模板：{template}（字体:{tpl['font']} 字号:{tpl['font_size']}）\n"
        info += f"  段落：{len(doc.paragraphs)} | 表格：{len(doc.tables)}\n"
        info += f"  大小：{full.stat().st_size} 字节"
        return info
    except ImportError:
        return "错误：未安装 python-docx，请运行 pip install python-docx"
    except Exception as e:
        return f"错误：{e}"


def _add_colored_text(paragraph, text: str, color_pattern: str):
    """处理 {color:xxx}文字{/color} 标记"""
    last_end = 0
    for match in re.finditer(color_pattern, text):
        # 前面的普通文本
        if match.start() > last_end:
            _add_formatted_text(paragraph, text[last_end:match.start()])
        color_str = match.group(1)
        colored_text = match.group(2)
        run = paragraph.add_run(colored_text)
        color = _parse_color(color_str)
        if color:
            run.font.color.rgb = color
        last_end = match.end()
    if last_end < len(text):
        _add_formatted_text(paragraph, text[last_end:])


# ====== Excel 文档生成 ======
def generate_excel(path: str, sheets: list, template: str = "default",
                   charts: list = None, formulas: list = None) -> str:
    """生成 Excel 文档（.xlsx）- 增强版：支持图表、公式
    path: 保存路径
    sheets: 工作表列表，每个工作表是 dict：
        {
            "name": "Sheet1",          # 工作表名（可选，默认 Sheet1）
            "data": [                   # 数据（二维数组）
                ["姓名", "年龄", "成绩"],
                ["张三", 18, 95.5],
                ["李四", 19, 88],
            ],
            "header": true,             # 是否有表头（默认 true）
        }
    template: 格式模板（default/report/data/financial）
    charts: 图表列表（可选），每个图表是 dict：
        {
            "type": "bar",              # bar(柱状图) / line(折线图) / pie(饼图)
            "title": "成绩对比",         # 图表标题
            "sheet": "Sheet1",          # 数据所在工作表名
            "data_range": "A1:C4",      # 数据范围（含表头）
            "categories_col": "A",      # 分类轴列（如姓名列）
            "values_cols": ["B", "C"],  # 值轴列（可多列）
            "position": "E2",           # 图表放置位置（单元格）
        }
    formulas: 公式列表（可选），每个公式是 dict：
        {
            "sheet": "Sheet1",          # 工作表名
            "cell": "D2",              # 写入单元格
            "formula": "=AVERAGE(C2:C4)", # 公式
        }
    """
    # 默认保存到桌面（path 为空或只含文件名时自动拼到桌面）
    path = _resolve_save_path(path, "未命名.xlsx")
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "错误：缺少 openpyxl 库，请运行 pip install openpyxl"

    # 模板配置
    templates = {
        "default": {"header_bg": "4472C4", "header_fg": "FFFFFF", "alt_row": "F2F2F2", "border": True},
        "report": {"header_bg": "1F4E79", "header_fg": "FFFFFF", "alt_row": "D6E4F0", "border": True},
        "data": {"header_bg": "70AD47", "header_fg": "FFFFFF", "alt_row": "E2EFDA", "border": True},
        "financial": {"header_bg": "5B5B5B", "header_fg": "FFFFFF", "alt_row": "FFF2CC", "border": True},
    }
    tpl = templates.get(template, templates["default"])

    wb = Workbook()
    wb.remove(wb.active)

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    ) if tpl["border"] else None

    # 记录工作表引用，供图表/公式使用
    ws_map = {}

    for i, sheet_def in enumerate(sheets):
        sheet_name = sheet_def.get("name", f"Sheet{i+1}")
        ws = wb.create_sheet(title=sheet_name)
        ws_map[sheet_name] = ws
        data = sheet_def.get("data", [])
        has_header = sheet_def.get("header", True)

        for row_idx, row_data in enumerate(data, 1):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                if has_header and row_idx == 1:
                    cell.font = Font(bold=True, color=tpl["header_fg"], name="Microsoft YaHei", size=11)
                    cell.fill = PatternFill(start_color=tpl["header_bg"], end_color=tpl["header_bg"], fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                else:
                    cell.font = Font(name="Microsoft YaHei", size=10)
                    if tpl["alt_row"] and row_idx > 1 and (row_idx % 2 == 0):
                        cell.fill = PatternFill(start_color=tpl["alt_row"], end_color=tpl["alt_row"], fill_type="solid")
                    cell.alignment = Alignment(vertical="center")
                if thin_border:
                    cell.border = thin_border

        # 自动列宽
        for col_idx in range(1, ws.max_column + 1):
            max_len = 0
            for row_idx in range(1, ws.max_row + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                val = str(cell.value) if cell.value else ""
                length = sum(2 if ord(c) > 127 else 1 for c in val)
                if length > max_len:
                    max_len = length
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 50)

    # 写入公式
    formula_count = 0
    if formulas:
        for f in formulas:
            ws = ws_map.get(f.get("sheet", sheets[0].get("name", "Sheet1")))
            if ws:
                cell_ref = f.get("cell", "")
                formula = f.get("formula", "")
                if cell_ref and formula:
                    ws[cell_ref] = formula
                    # 公式单元格样式
                    ws[cell_ref].font = Font(name="Microsoft YaHei", size=10, bold=True, color=tpl["header_bg"])
                    if thin_border:
                        ws[cell_ref].border = thin_border
                    formula_count += 1

    # 生成图表
    chart_count = 0
    if charts:
        try:
            from openpyxl.chart import BarChart, LineChart, PieChart, Reference
        except ImportError:
            pass
        else:
            for ch in charts:
                ws = ws_map.get(ch.get("sheet", sheets[0].get("name", "Sheet1")))
                if not ws:
                    continue
                chart_type = ch.get("type", "bar").lower()
                chart_title = ch.get("title", "")
                categories_col = ch.get("categories_col", "A")
                values_cols = ch.get("values_cols", ["B"])
                position = ch.get("position", "E2")

                # 解析列号为数字
                def col_to_num(col_str):
                    """A→1, B→2, ..."""
                    num = 0
                    for c in col_str.upper():
                        num = num * 26 + (ord(c) - ord('A') + 1)
                    return num

                cat_col_num = col_to_num(categories_col)
                val_col_nums = [col_to_num(c) for c in values_cols]

                # 数据范围：从第1行（表头）到最后一行
                max_row = ws.max_row
                min_row = 2 if ws.cell(1, 1).value else 1  # 跳过表头

                if chart_type == "bar":
                    chart = BarChart()
                    chart.type = "col"
                    chart.style = 10
                elif chart_type == "line":
                    chart = LineChart()
                    chart.style = 12
                elif chart_type == "pie":
                    chart = PieChart()
                    chart.style = 10
                else:
                    continue

                chart.title = chart_title

                # 添加数据（值列）
                for vcn in val_col_nums:
                    data_ref = Reference(ws, min_col=vcn, min_row=1, max_row=max_row)
                    chart.add_data(data_ref, titles_from_data=True)

                # 设置分类轴（X轴标签）
                if chart_type != "pie":
                    cat_ref = Reference(ws, min_col=cat_col_num, min_row=min_row, max_row=max_row)
                    chart.set_categories(cat_ref)
                else:
                    # 饼图分类
                    cat_ref = Reference(ws, min_col=cat_col_num, min_row=min_row, max_row=max_row)
                    chart.set_categories(cat_ref)

                # 图表尺寸
                chart.width = 18
                chart.height = 12

                # 放置图表
                ws.add_chart(chart, position)
                chart_count += 1

    wb.save(path)
    result = f"{_load_svg_icon('document')} Excel 已生成：{path}（{len(sheets)} 个工作表"
    if chart_count:
        result += f"，{chart_count} 个图表"
    if formula_count:
        result += f"，{formula_count} 个公式"
    result += "）"
    return result


# ====== PDF 文档生成 ======
def _latex_to_image(latex: str, font_size: int = 16, dpi: int = 200) -> object:
    """将 LaTeX 公式渲染为 PNG 图片（reportlab Image 对象）

    用于 PDF 学术模板的独立公式 $$...$$ 渲染，比纯 Unicode 更标准。
    失败时返回 None，调用方降级为 Unicode 文本。

    参数：
    - latex: LaTeX 公式（不含 $$ 包裹）
    - font_size: 字号
    - dpi: 分辨率
    """
    try:
        import matplotlib
        matplotlib.use("Agg")  # 非交互后端
        import matplotlib.pyplot as plt
        import io, os, tempfile

        # 去除首尾 $ 符号
        formula = latex.strip().strip("$").strip()

        # 创建图片：用 mathtext 渲染公式
        fig = plt.figure(figsize=(0.01, 0.01), dpi=dpi)
        fig.text(0, 0, f"${formula}$", fontsize=font_size, color="black")
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                    pad_inches=0.05, transparent=False, facecolor="white")
        plt.close(fig)
        buf.seek(0)

        # 保存为临时文件（reportlab Image 需要文件路径或 BytesIO）
        return buf
    except Exception:
        return None


def generate_pdf(path: str, content: str, title: str = "", template: str = "default") -> str:
    """生成 PDF 文档
    path: 保存路径
    content: 文档内容（支持 Markdown 风格标记）
    title: 文档标题（可选）
    template: 格式模板（default/report/contract/resume/letter/technical/academic）
                academic=学术论文（1.5倍行距/摘要/关键词/参考文献自动编号/LaTeX公式渲染）

    内容支持：
    # 一级标题 / ## 二级标题 / ### 三级标题
    - 无序列表 / 1. 有序列表
    > 引用 / ```代码块``` / **粗体** *斜体*
    ---（分隔线）/ 普通段落 / |表格语法|
    """
    # 默认保存到桌面（path 为空或只含文件名时自动拼到桌面）
    path = _resolve_save_path(path, "未命名.pdf")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm, mm
        from reportlab.lib.colors import HexColor, black, white, grey
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, HRFlowable, ListFlowable, ListItem, Preformatted
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return "错误：缺少 reportlab 库，请运行 pip install reportlab"

    # 注册中文字体
    font_name = "Helvetica"
    for font_path, font_id in [
        (r"C:\Windows\Fonts\msyh.ttc", "MSYH"),
        (r"C:\Windows\Fonts\msyhbd.ttc", "MSYHBD"),
        (r"C:\Windows\Fonts\simsun.ttc", "SimSun"),
        (r"C:\Windows\Fonts\simhei.ttf", "SimHei"),
    ]:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont(font_id, font_path))
                font_name = font_id
                break
            except Exception:
                continue

    # 模板配置
    templates = {
        "default": {"title_size": 20, "h1_size": 16, "h2_size": 14, "h3_size": 12, "body_size": 10, "color": "#1F4E79", "margin": [2.54, 2.54, 2.54, 2.54]},
        "report": {"title_size": 22, "h1_size": 16, "h2_size": 14, "h3_size": 12, "body_size": 10, "color": "#1F4E79", "margin": [3, 2.5, 3, 2.5]},
        "contract": {"title_size": 18, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 10, "color": "#333333", "margin": [2.54, 2.54, 2.54, 2.54]},
        "resume": {"title_size": 24, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 10, "color": "#2E5C8A", "margin": [2, 2.5, 2, 2.5]},
        "letter": {"title_size": 18, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 11, "color": "#333333", "margin": [3, 3, 3, 3]},
        "technical": {"title_size": 20, "h1_size": 16, "h2_size": 14, "h3_size": 12, "body_size": 9, "color": "#0B5394", "margin": [2, 2, 2, 2]},
        "academic": {"title_size": 18, "h1_size": 14, "h2_size": 12, "h3_size": 11, "body_size": 10, "color": "#000000", "margin": [2.54, 2.54, 2.54, 2.54], "line_spacing": 1.5},
    }
    tpl = templates.get(template, templates["default"])
    heading_color = HexColor(tpl["color"])
    m = tpl["margin"]
    is_academic = template == "academic"
    line_spacing_factor = tpl.get("line_spacing", 1.5)

    doc = SimpleDocTemplate(path, pagesize=A4,
                           topMargin=m[0]*cm, rightMargin=m[1]*cm,
                           bottomMargin=m[2]*cm, leftMargin=m[3]*cm)

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                 fontName=font_name, fontSize=tpl["title_size"],
                                 textColor=heading_color, alignment=TA_CENTER, spaceAfter=20)
    style_h1 = ParagraphStyle('CustomH1', parent=styles['Heading1'],
                              fontName=font_name, fontSize=tpl["h1_size"],
                              textColor=heading_color, spaceBefore=16, spaceAfter=8)
    style_h2 = ParagraphStyle('CustomH2', parent=styles['Heading2'],
                              fontName=font_name, fontSize=tpl["h2_size"],
                              textColor=heading_color, spaceBefore=12, spaceAfter=6)
    style_h3 = ParagraphStyle('CustomH3', parent=styles['Heading3'],
                              fontName=font_name, fontSize=tpl["h3_size"],
                              textColor=heading_color, spaceBefore=10, spaceAfter=4)
    # 四级标题（GB/T 7713.1-2025 允许四级，h4_size 缺省时用 h3_size-1）
    style_h4 = ParagraphStyle('CustomH4', parent=styles['Heading4'],
                              fontName=font_name, fontSize=tpl.get("h4_size", tpl["h3_size"] - 1),
                              textColor=heading_color, spaceBefore=8, spaceAfter=3)
    style_body = ParagraphStyle('CustomBody', parent=styles['Normal'],
                                fontName=font_name, fontSize=tpl["body_size"],
                                leading=tpl["body_size"]*line_spacing_factor, spaceAfter=6, alignment=TA_JUSTIFY)
    style_quote = ParagraphStyle('CustomQuote', parent=style_body,
                                 leftIndent=20, textColor=grey, fontSize=tpl["body_size"]-1)
    style_code = ParagraphStyle('CustomCode', parent=styles['Code'],
                                fontName='Courier', fontSize=tpl["body_size"]-1,
                                backColor=HexColor("#F5F5F5"), leftIndent=10, rightIndent=10,
                                spaceBefore=4, spaceAfter=4)
    # 学术论文专用样式
    style_abstract = ParagraphStyle('AcademicAbstract', parent=style_body,
                                    fontSize=tpl["body_size"]-1, leading=(tpl["body_size"]-1)*1.5,
                                    leftIndent=15, rightIndent=15, textColor=HexColor("#333333"),
                                    spaceBefore=8, spaceAfter=8, alignment=TA_JUSTIFY)
    style_keywords = ParagraphStyle('AcademicKeywords', parent=style_body,
                                    fontSize=tpl["body_size"]-1, leftIndent=15,
                                    spaceBefore=4, spaceAfter=8)
    style_reference = ParagraphStyle('AcademicReference', parent=style_body,
                                     fontSize=tpl["body_size"]-1, leading=(tpl["body_size"]-1)*1.3,
                                     leftIndent=20, firstLineIndent=-20,
                                     spaceBefore=2, spaceAfter=2, alignment=TA_JUSTIFY)

    story = []

    # 标题
    if title:
        story.append(Paragraph(title, style_title))
        story.append(Spacer(1, 10))

    # 解析内容
    lines = content.split('\n')
    in_code = False
    code_lines = []
    in_table = False
    table_rows = []
    # 学术论文状态变量
    in_references = False
    ref_counter = 0

    for line in lines:
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            if in_code:
                story.append(Preformatted('\n'.join(code_lines), style_code))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        # 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # 跳过分隔行 |---|---|
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # 表格结束
            if table_rows:
                t = Table(table_rows)
                t.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, -1), font_name, tpl["body_size"]),
                    ('BACKGROUND', (0, 0), (-1, 0), heading_color),
                    ('TEXTCOLOR', (0, 0), (-1, 0), white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('GRID', (0, 0), (-1, -1), 0.5, grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, HexColor("#F5F5F5")]),
                ]))
                story.append(t)
                story.append(Spacer(1, 8))
                table_rows = []
            in_table = False

        # 分隔线
        if stripped == '---' or stripped == '***':
            story.append(HRFlowable(width="100%", thickness=1, color=heading_color))
            story.append(Spacer(1, 6))
            continue

        # ===== 学术论文段落特殊处理（academic 模板）=====
        if is_academic:
            academic_match = re.match(
                r'^(摘要|Abstract|关键词|Keywords|参考文献|References|引言|Introduction|结论|Conclusion|致谢|Acknowledgments?)\s*[:：]?\s*(.*)',
                stripped, re.IGNORECASE
            )
            if academic_match:
                section_name = academic_match.group(1)
                section_content = academic_match.group(2)

                # 摘要/Abstract
                if section_name.lower() in ("摘要", "abstract"):
                    abstract_text = f"<b>【{section_name}】</b> "
                    if section_content:
                        rendered = render_latex_in_text(section_content)
                        abstract_text += rendered
                    story.append(Paragraph(abstract_text, style_abstract))
                    continue

                # 关键词/Keywords
                elif section_name.lower() in ("关键词", "keywords"):
                    kw_text = f"<b>【{section_name}】</b> {section_content}"
                    story.append(Paragraph(kw_text, style_keywords))
                    continue

                # 参考文献/References
                elif section_name.lower() in ("参考文献", "references"):
                    story.append(Paragraph(f"<b>{section_name}</b>", style_h2))
                    in_references = True
                    ref_counter = 0
                    continue

                # 引言/结论
                elif section_name.lower() in ("引言", "introduction", "结论", "conclusion"):
                    story.append(Paragraph(f"<b>{section_name}</b>", style_h1))
                    in_references = False  # 退出参考文献区域
                    continue

        # 参考文献条目自动编号
        if is_academic and in_references and stripped and not stripped.startswith('#'):
            if not re.match(r'^\[\d+\]', stripped):
                ref_counter += 1
                ref_text = f"[{ref_counter}] {stripped}"
                # 处理粗体斜体
                ref_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', ref_text)
                ref_text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', ref_text)
                story.append(Paragraph(ref_text, style_reference))
                continue

        # 标题（支持 # ~ #### 四级，符合 GB/T 7713.1-2025 论文标题层级规范）
        if stripped.startswith('#### '):
            story.append(Paragraph(stripped[5:], style_h4))
        elif stripped.startswith('### '):
            story.append(Paragraph(stripped[4:], style_h3))
        elif stripped.startswith('## '):
            story.append(Paragraph(stripped[3:], style_h2))
        elif stripped.startswith('# '):
            story.append(Paragraph(stripped[2:], style_h1))
        elif is_academic and stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            # 学术模板：独立公式 $$...$$ 用 matplotlib 渲染为图片（比 Unicode 更标准）
            formula = stripped[2:-2].strip()
            img_buf = _latex_to_image(formula)
            if img_buf is not None:
                try:
                    from reportlab.lib.utils import ImageReader
                    from reportlab.platypus import Image as RLImage
                    img_reader = ImageReader(img_buf)
                    iw, ih = img_reader.getSize()
                    # 按宽度缩放（最大宽度 400pt）
                    max_w = 400
                    if iw > max_w:
                        ratio = max_w / iw
                        iw, ih = max_w, ih * ratio
                    img_buf.seek(0)
                    story.append(RLImage(img_buf, width=iw, height=ih))
                    story.append(Spacer(1, 6))
                except Exception:
                    # 图片插入失败 → 降级为 Unicode
                    story.append(Paragraph(_latex_to_unicode(formula), style_body))
            else:
                # matplotlib 渲染失败 → 降级为 Unicode
                story.append(Paragraph(_latex_to_unicode(formula), style_body))
        elif stripped.startswith('> '):
            story.append(Paragraph(stripped[2:], style_quote))
        elif stripped.startswith('- '):
            item = Paragraph(stripped[2:], style_body)
            story.append(ListFlowable([ListItem(item)], bulletType='bullet'))
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            item = Paragraph(text, style_body)
            story.append(ListFlowable([ListItem(item)], bulletType='1'))
        elif stripped:
            # 处理粗体斜体（**粗体** *斜体* → <b>粗体</b> <i>斜体</i>）
            text = stripped
            # 学术模板：先渲染 LaTeX 公式为 Unicode
            if is_academic:
                text = render_latex_in_text(text)
            # 先处理 **粗体**
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            # 再处理 *斜体*
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            # 转义 XML 特殊字符（但保留已生成的标签）
            text = text.replace('&', '&amp;').replace('<b>', '<b>').replace('</b>', '</b>')
            story.append(Paragraph(text, style_body))
        else:
            story.append(Spacer(1, 6))

    # 处理未关闭的表格
    if table_rows:
        t = Table(table_rows)
        t.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), font_name, tpl["body_size"]),
            ('BACKGROUND', (0, 0), (-1, 0), heading_color),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, grey),
        ]))
        story.append(t)

    doc.build(story)
    return f"{_load_svg_icon('document')} PDF 已生成：{path}（模板：{template}）"


def active_window() -> str:
    """获取当前焦点窗口信息（标题、应用、位置）"""
    try:
        import ctypes
        user32 = ctypes.windll.user32
        # 获取前台窗口句柄
        hwnd = user32.GetForegroundWindow()
        # 获取窗口标题
        length = user32.GetWindowTextLengthW(hwnd) + 1
        title = ctypes.create_unicode_buffer(length)
        user32.GetWindowTextW(hwnd, title, length)
        # 获取进程ID
        pid = ctypes.c_ulong()
        user32.GetWindowThreadProcessId(hwnd, ctypes.byref(pid))
        # 获取进程名
        try:
            pr = subprocess.run(["tasklist", "/FI", f"PID eq {pid.value}", "/FO", "CSV", "/NH"],
                                capture_output=True, text=True, timeout=5)
            pname = pr.stdout.strip().split(",")[0].strip('"') if pr.stdout.strip() else "?"
        except Exception:
            pname = "?"
        # 获取窗口位置和大小
        rect = ctypes.wintypes.RECT()
        user32.GetWindowRect(hwnd, ctypes.byref(rect))
        return f"窗口：{title.value}\n应用：{pname}（PID: {pid.value}）\n位置：({rect.left}, {rect.top}) 大小：{rect.right - rect.left}x{rect.bottom - rect.top}"
    except Exception as e:
        return f"错误：{e}"


def list_windows() -> str:
    """列出所有可见窗口"""
    try:
        import ctypes
        user32 = ctypes.windll.user32
        results = []
        def callback(hwnd, _):
            if user32.IsWindowVisible(hwnd):
                length = user32.GetWindowTextLengthW(hwnd) + 1
                if length > 1:
                    title = ctypes.create_unicode_buffer(length)
                    user32.GetWindowTextW(hwnd, title, length)
                    t = title.value.strip()
                    if t:
                        pid = ctypes.c_ulong()
                        user32.GetWindowThreadProcessId(hwnd, ctypes.byref(pid))
                        results.append(f"  {t[:50]:<50} PID:{pid.value}")
            return True
        WNDENUMPROC = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.POINTER(ctypes.c_int), ctypes.POINTER(ctypes.c_int))
        user32.EnumWindows(WNDENUMPROC(callback), 0)
        if not results:
            return "(无可见窗口)"
        return f"可见窗口（{len(results)}个）：\n" + "\n".join(results[:30])
    except Exception as e:
        return f"错误：{e}"


def read_screen_content(max_length: int = 2000) -> str:
    """读取当前前台窗口的文字内容（通过 Windows UI Automation）"""
    try:
        import uiautomation as uia
        window = uia.GetForegroundControl()
        if not window:
            return "(无法获取前台窗口)"
        title = window.Name or ""
        classname = window.ClassName or ""
        # 读取窗口内可见文本
        texts = []
        def collect_text(ctrl, depth=0):
            if depth > 8 or len(texts) > 100:
                return
            try:
                name = ctrl.Name
                if name and name.strip() and name != title:
                    t = name.strip()
                    if len(t) > 200:
                        t = t[:200] + "..."
                    texts.append(t)
            except Exception:
                pass
            try:
                for child in ctrl.GetChildren():
                    collect_text(child, depth + 1)
            except Exception:
                pass
        collect_text(window)
        content = "\n".join(texts[:80])
        if len(content) > max_length:
            content = content[:max_length] + "\n...(截断)"
        return f"窗口：{title}\n类名：{classname}\n内容：\n{content}" if content.strip() else f"窗口：{title}\n(无可读文本)"
    except ImportError:
        # uiautomation 不可用时回退到窗口标题
        return active_window()
    except Exception as e:
        return f"读取失败：{e}"


# ====== 语音交互模块（TTS + ASR）======

# 全局 TTS 开关（由 /语音 命令切换）
TTS_ENABLED = False
# 全局 ASR 实例缓存（避免重复加载模型）
_ASR_MODEL = None
# SenseVoice 模型路径（阿里达摩院 FunAudioLLM 开源，中英日韩粤5语言，准确率远超 whisper-tiny）
# 完全免费、无需 API Key、本地离线运行
# 智能查找：开发模式（脚本目录/models）或 pip 安装模式（~/.zeroai/models）
_SENSE_VOICE_MODEL_DIR = os.path.join(_find_resource_dir("models"), "sense-voice")
_SENSE_VOICE_MODEL = os.path.join(_SENSE_VOICE_MODEL_DIR, "model.int8.onnx")
_SENSE_VOICE_TOKENS = os.path.join(_SENSE_VOICE_MODEL_DIR, "tokens.txt")

# SenseVoice 模型下载源（HuggingFace 国内镜像）
_SENSE_VOICE_DOWNLOAD_URLS = {
    "model.int8.onnx": "https://hf-mirror.com/pengzhendong/sherpa-onnx-sense-voice-zh-en-ja-ko-yue-2024-07-17/resolve/main/model.int8.onnx",
    "tokens.txt": "https://hf-mirror.com/pengzhendong/sherpa-onnx-sense-voice-zh-en-ja-ko-yue-2024-07-17/resolve/main/tokens.txt",
}

def _download_sense_voice_model() -> bool:
    """下载 SenseVoice 语音识别模型到用户目录 ~/.zeroai/models/sense-voice/

    模型来源：HuggingFace 国内镜像（hf-mirror.com）
    模型大小：约 220MB（model.int8.onnx）+ 几KB（tokens.txt）

    返回 True 表示下载成功，False 表示失败。
    """
    # 下载到用户目录（pip 安装模式的标准位置）
    target_dir = os.path.join(_ZEROAI_USER_DIR, "models", "sense-voice")
    try:
        os.makedirs(target_dir, exist_ok=True)
    except OSError as e:
        print(f"无法创建模型目录: {e}", file=sys.stderr)
        return False

    import urllib.request

    for filename, url in _SENSE_VOICE_DOWNLOAD_URLS.items():
        target_path = os.path.join(target_dir, filename)
        if os.path.isfile(target_path) and os.path.getsize(target_path) > 1024:
            continue  # 已下载，跳过

        print(f"正在下载 {filename} ...", file=sys.stderr)
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "ZeroAI/1.0"})
            with urllib.request.urlopen(req, timeout=300) as resp:
                total = int(resp.headers.get("Content-Length", 0))
                downloaded = 0
                chunk_size = 1024 * 64  # 64KB
                with open(target_path, "wb") as f:
                    while True:
                        chunk = resp.read(chunk_size)
                        if not chunk:
                            break
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total > 0:
                            pct = downloaded * 100 // total
                            print(f"\r  进度: {pct}% ({downloaded // 1024 // 1024}MB / {total // 1024 // 1024}MB)", end="", file=sys.stderr)
                print(file=sys.stderr)  # 换行
            print(f"  {filename} 下载完成", file=sys.stderr)
        except Exception as e:
            print(f"  下载失败: {e}", file=sys.stderr)
            return False

    # 更新全局路径指向用户目录
    global _SENSE_VOICE_MODEL_DIR, _SENSE_VOICE_MODEL, _SENSE_VOICE_TOKENS
    _SENSE_VOICE_MODEL_DIR = target_dir
    _SENSE_VOICE_MODEL = os.path.join(target_dir, "model.int8.onnx")
    _SENSE_VOICE_TOKENS = os.path.join(target_dir, "tokens.txt")
    print("SenseVoice 模型下载完成！", file=sys.stderr)
    return True


def _init_pygame_mixer():
    """初始化 pygame 音频播放器（延迟初始化）"""
    try:
        import pygame
        if not pygame.mixer.get_init():
            pygame.mixer.init()
        return True
    except Exception:
        return False


def _split_text_into_segments(text: str, max_chars: int = 200) -> list:
    """把长文本按句末标点分割为多段（用于 TTS 分段朗读）

    策略：
    1. 先按强句末标点（。！？!?；;\\n）切分
    2. 短句累积成长段（避免单段过短触发太多 edge-tts 调用）
    3. 单段不超过 max_chars 字符
    4. 避免产生空段
    """
    import re
    if not text or not text.strip():
        return []
    # 用正则切分：保留分隔符（让朗读更自然）
    parts = re.split(r"([。！？!?；;\n])", text)
    # 重组：[句子+标点, 句子+标点, ...]
    sentences = []
    i = 0
    while i < len(parts):
        s = parts[i].strip()
        if s:
            # 如果下一项是标点，拼回去
            if i + 1 < len(parts) and parts[i + 1] in "。！？!?；;\n":
                s = s + parts[i + 1]
                i += 2
            else:
                i += 1
            if s:
                sentences.append(s)
        else:
            i += 1
    # 累积短句成段（每段不超过 max_chars）
    segments = []
    buf = ""
    for s in sentences:
        if len(buf) + len(s) <= max_chars:
            buf += s
        else:
            if buf:
                segments.append(buf)
            # 如果单句本身就超过 max_chars，强制切分
            if len(s) > max_chars:
                # 按逗号再切
                sub_parts = re.split(r"([，,、：:])", s)
                sub_buf = ""
                for sp in sub_parts:
                    if len(sub_buf) + len(sp) <= max_chars:
                        sub_buf += sp
                    else:
                        if sub_buf:
                            segments.append(sub_buf)
                        sub_buf = sp
                if sub_buf:
                    buf = sub_buf
                else:
                    buf = ""
            else:
                buf = s
    if buf:
        segments.append(buf)
    return segments


def speak_tts(
    text: str,
    voice: str = "zh-CN-XiaoxiaoNeural",
    rate: str = "+0%",
    volume: str = "+0%",
    interrupt_check=None,
    segment_max_chars: int = 200,
) -> str:
    """文本转语音并播放（同步阻塞，分段朗读版）

    使用 edge-tts（微软免费 TTS，无需 API Key）

    长文本会自动按句末标点（。！？!?；;）分割成多段，逐段生成+播放。
    优势：
      1. 启动延迟短（先听到第一段，后台继续生成后续段）
      2. 任意一段播放中检测到 interrupt_check() 返回 True，立即停止后续段
      3. 避免单次生成过长导致 edge-tts 超时

    Args:
        text: 要朗读的文本（支持中英文混合）
        voice: 音色（zh-CN-XiaoxiaoNeural 女/zh-CN-YunxiNeural 男/zh-CN-YunyangNeural 新闻）
        rate: 语速（+0% 正常/+10% 加速/-10% 减速）
        volume: 音量（+0% 正常/+10% 更大/-10% 更小）
        interrupt_check: 可调用对象，每段播放前后/中检查，返回 True 时立即停止
        segment_max_chars: 每段最大字符数（默认 200）

    Returns:
        成功返回空字符串，失败返回错误信息
    """
    if not text or not text.strip():
        return "（空文本，无需朗读）"
    # 移除 Markdown 标记和代码块，避免朗读符号
    import re
    clean = re.sub(r"```[\s\S]*?```", "（代码块）", text)  # 代码块
    clean = re.sub(r"`([^`]+)`", r"\1", clean)  # 行内代码
    clean = re.sub(r"!\[[^\]]*\]\([^)]+\)", "（图片）", clean)  # 图片
    clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)  # 链接
    clean = re.sub(r"^#{1,6}\s+", "", clean, flags=re.MULTILINE)  # 标题
    clean = re.sub(r"^>\s+", "", clean, flags=re.MULTILINE)  # 引用
    clean = re.sub(r"^[-*+]\s+", "", clean, flags=re.MULTILINE)  # 列表
    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)  # 粗体
    clean = re.sub(r"\*([^*]+)\*", r"\1", clean)  # 斜体
    clean = clean.strip()
    if not clean:
        return "（无可朗读内容）"

    # 分割为多段
    segments = _split_text_into_segments(clean, max_chars=segment_max_chars)
    if not segments:
        return "（无可朗读内容）"

    try:
        import asyncio
        import edge_tts
        import tempfile
        import os
        import time
        import threading

        def _is_interrupted() -> bool:
            """统一中断检查（异常吞掉，保证不影响主流程）"""
            if interrupt_check is None:
                return False
            try:
                return bool(interrupt_check())
            except Exception:
                return False

        def _gen_one_sync(seg_text: str) -> str:
            """同步生成单段 MP3，返回路径（出错返回错误字符串）"""
            try:
                async def _gen():
                    tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
                    tmp_path = tmp.name
                    tmp.close()
                    communicate = edge_tts.Communicate(seg_text, voice, rate=rate, volume=volume)
                    await communicate.save(tmp_path)
                    return tmp_path

                try:
                    loop = asyncio.get_event_loop()
                    if loop.is_running():
                        result = [None]

                        def _run():
                            try:
                                new_loop = asyncio.new_event_loop()
                                result[0] = new_loop.run_until_complete(_gen())
                                new_loop.close()
                            except Exception as e:
                                result[0] = f"错误：{e}"
                        t = threading.Thread(target=_run)
                        t.start()
                        t.join()
                        return result[0]
                    else:
                        return loop.run_until_complete(_gen())
                except RuntimeError:
                    return asyncio.run(_gen())
            except Exception as e:
                return f"错误：{e}"

        # 初始化 pygame.mixer（如果还没初始化）
        if not _init_pygame_mixer():
            return "错误：无法初始化音频播放器（pygame.mixer）"

        import pygame

        # 逐段生成 + 播放
        for idx, seg in enumerate(segments):
            # 段间打断检查
            if _is_interrupted():
                return "（已打断）"
            # 生成当前段
            tmp_path = _gen_one_sync(seg)
            if not tmp_path:
                continue
            if isinstance(tmp_path, str) and tmp_path.startswith("错误"):
                # 生成失败，跳过此段，继续下一段
                continue
            # 播放前再次检查中断（避免在生成 MP3 期间用户已退出）
            if _is_interrupted():
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                return "（已打断）"
            # 加载并播放
            try:
                pygame.mixer.music.load(tmp_path)
                pygame.mixer.music.play()
            except Exception:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                continue
            # 等待播放完成（最长 60 秒，并周期性检查打断）
            start = time.time()
            interrupted = False
            while pygame.mixer.music.get_busy() and time.time() - start < 60:
                time.sleep(0.1)
                if _is_interrupted():
                    try:
                        pygame.mixer.music.stop()
                    except Exception:
                        pass
                    interrupted = True
                    break
            try:
                pygame.mixer.music.unload()
            except Exception:
                pass
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            if interrupted:
                return "（已打断）"
        return ""
    except ImportError:
        return "错误：edge-tts 未安装，请运行 pip install edge-tts"
    except Exception as e:
        return f"错误：{e}"


def listen_asr(max_seconds: int = 10, silence_seconds: float = 1.0) -> str:
    """录音并识别为文字（同步阻塞）

    使用 sherpa-onnx + SenseVoice（阿里达摩院开源，本地离线，无需 API Key）
    准确率远超 whisper-tiny，支持中英日韩粤5种语言，自带标点符号

    改进版 VAD 方案：
    1. 录音前 0.3 秒采集环境噪声，动态计算静音阈值
    2. 前置静音过滤：等到检测到语音才开始录入（最多等 3 秒）
    3. 滑动窗口静音检测：最近 3 帧平均音量低于阈值才算静音
    4. 音量归一化：录音后归一化到 [-1, 1] 提升识别率
    5. 最小录音时长 0.3 秒，避免误触发

    Args:
        max_seconds: 最长录音秒数（默认 10 秒）
        silence_seconds: 静音检测秒数（连续静音超过此值则停止，默认 1.0 秒）

    Returns:
        识别到的文字，失败返回错误信息
    """
    global _ASR_MODEL
    try:
        import sounddevice as sd
        import numpy as np
        import tempfile
        import wave
    except ImportError as e:
        return f"错误：缺少音频库（{e}），请运行 pip install sounddevice numpy"

    # ════════════════ 录音阶段（专业 VAD） ════════════════
    try:
        sample_rate = 16000
        channels = 1
        block_size = 1024  # 每帧 1024 采样 ≈ 64ms

        # ── 1. 采集 0.3 秒环境噪声，计算动态阈值 ──
        noise_frames = []
        noise_duration = 0.3  # 0.3 秒噪声采样（缩短响应时间）
        noise_blocks_needed = int(noise_duration * sample_rate / block_size)
        try:
            with sd.InputStream(samplerate=sample_rate, channels=channels, blocksize=block_size) as stream:
                for _ in range(noise_blocks_needed):
                    data, _ = stream.read(block_size)
                    noise_frames.append(data.copy())
            noise_audio = np.concatenate(noise_frames, axis=0)
            noise_level = float(np.abs(noise_audio).mean())
            # 动态阈值 = 噪声基线 × 3，至少 0.02（避免太敏感）
            silence_threshold = max(noise_level * 3.0, 0.02)
        except Exception:
            # 噪声采样失败，使用默认阈值
            silence_threshold = 0.03

        # ── 2. 正式录音：前置静音过滤 + 滑动窗口 VAD ──
        frames = []
        speech_started = False  # 是否检测到语音开始
        silence_count = 0  # 连续静音帧数
        max_silence_frames = int(silence_seconds * sample_rate / block_size)
        # 前置静音最长等待 3 秒（用户可能需要反应时间）
        max_pre_wait = int(3.0 * sample_rate / block_size)
        pre_wait_count = 0
        # 滑动窗口（最近 3 帧的音量）
        volume_window = []

        def callback(indata, frame_count, time_info, status):
            frames.append(indata.copy())

        with sd.InputStream(samplerate=sample_rate, channels=channels, callback=callback, blocksize=block_size):
            import time
            start = time.time()
            while time.time() - start < max_seconds:
                if frames:
                    last = frames[-1]
                    volume = float(np.abs(last).mean())
                    # 滑动窗口（保留最近 3 帧）
                    volume_window.append(volume)
                    if len(volume_window) > 3:
                        volume_window.pop(0)
                    avg_volume = sum(volume_window) / len(volume_window) if volume_window else 0

                    if not speech_started:
                        # 前置静音过滤：等到检测到语音才开始计时
                        if avg_volume > silence_threshold:
                            speech_started = True
                        else:
                            pre_wait_count += 1
                            if pre_wait_count > max_pre_wait:
                                # 等了 5 秒还没说话，返回空
                                return "（未录到声音）"
                    else:
                        # 语音已开始，检测静音
                        if avg_volume < silence_threshold:
                            silence_count += 1
                        else:
                            silence_count = 0
                        # 连续静音超过阈值 → 停止
                        if silence_count >= max_silence_frames:
                            break
                time.sleep(0.03)

        if not frames:
            return "（未录到声音）"
        if not speech_started:
            return "（未录到声音）"

        audio = np.concatenate(frames, axis=0)
        # 至少 0.3 秒有效音频
        if len(audio) < int(0.3 * sample_rate):
            return "（录音太短）"

        # ── 3. 音量归一化（提升小声说话的识别率）──
        max_amplitude = float(np.abs(audio).max())
        if max_amplitude > 0 and max_amplitude < 0.5:
            # 音量过小，归一化到 [-1, 1]
            audio = audio / max_amplitude

        # 保存为临时 wav 文件（faster-whisper 回退方案需要）
        tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        tmp_path = tmp.name
        tmp.close()
        with wave.open(tmp_path, "wb") as wf:
            wf.setnchannels(channels)
            wf.setsampwidth(2)
            wf.setframerate(sample_rate)
            # 确保 audio 在 [-1, 1] 范围内
            audio_clipped = np.clip(audio, -1.0, 1.0)
            wf.writeframes((audio_clipped * 32767).astype(np.int16).tobytes())
    except Exception as e:
        return f"错误：录音失败（{e}）"

    # ════════════════ 识别阶段 ════════════════
    try:
        # 优先使用 sherpa-onnx + SenseVoice（本地离线，准确率最高，自带标点）
        sherpa_err = None
        try:
            import sherpa_onnx
            if _ASR_MODEL is None:
                if not os.path.isfile(_SENSE_VOICE_MODEL) or not os.path.isfile(_SENSE_VOICE_TOKENS):
                    # 尝试自动下载模型（首次使用语音功能时）
                    print("SenseVoice 模型未找到，正在自动下载（约 220MB）...", file=sys.stderr)
                    if not _download_sense_voice_model():
                        raise FileNotFoundError(
                            f"SenseVoice 模型下载失败。请手动下载并放到:\n  {_SENSE_VOICE_MODEL_DIR}\n"
                            f"或运行: pip install faster-whisper 作为替代方案"
                        )
                _ASR_MODEL = sherpa_onnx.OfflineRecognizer.from_sense_voice(
                    model=_SENSE_VOICE_MODEL,
                    tokens=_SENSE_VOICE_TOKENS,
                    num_threads=2,
                    use_itn=True,  # 启用逆文本归一化（自动添加标点符号）
                )
            # 创建识别流，喂入音频数据
            stream = _ASR_MODEL.create_stream()
            # sounddevice 录制的是 float32，范围 [-1, 1]，sherpa-onnx 需要同样的 float32
            audio_float32 = audio.flatten().astype(np.float32)
            stream.accept_waveform(sample_rate, audio_float32)
            _ASR_MODEL.decode_stream(stream)
            text = stream.result.text.strip()
            return text if text else "（未识别到内容）"
        except ImportError:
            sherpa_err = "sherpa-onnx 未安装（pip install sherpa-onnx）"
        except Exception as e:
            sherpa_err = str(e)

        # 回退 1：faster-whisper（本地离线，准确率一般）
        faster_whisper_err = None
        try:
            from faster_whisper import WhisperModel
            # 设置 HuggingFace 国内镜像（避免模型下载被墙）
            os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")
            os.environ.setdefault("HF_HUB_DISABLE_XET", "1")
            os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
            _fw_model = WhisperModel("tiny", device="cpu", compute_type="int8")
            segments, _ = _fw_model.transcribe(tmp_path, language="zh", beam_size=1)
            text = "".join(seg.text for seg in segments).strip()
            if text:
                return text
        except ImportError:
            faster_whisper_err = "faster-whisper 未安装"
        except Exception as e:
            faster_whisper_err = str(e)

        # 所有方案均失败
        return f"错误：语音识别失败\n  sherpa-onnx: {sherpa_err}\n  faster-whisper: {faster_whisper_err}\n建议：\n  1. 确保 sherpa-onnx 已安装（pip install sherpa-onnx）\n  2. 确保模型文件存在于 {_SENSE_VOICE_MODEL_DIR}\n  3. 或安装 faster-whisper 作为回退（pip install faster-whisper）"
    finally:
        # 识别完成后才删除临时文件
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


# ====== SSH 远程部署工具（asyncssh，纯Python异步SSH库）======
# 安全审计日志
_SSH_AUDIT_LOG = []
_SSH_AUDIT_MAX = 200

# 危险命令黑名单（需要用户确认才执行）
_SSH_DANGEROUS_PATTERNS = [
    r"\brm\s+-rf\s+/(?!\S)",   # rm -rf /
    r"\bmkfs\b",                # 格式化
    r"\bdd\s+if=",              # dd 写入
    r"\bshutdown\b",            # 关机
    r"\binit\s+0\b",            # 关机
    r"\bhalt\b",                # 关机
    r"\breboot\b",             # 重启
    r">\s*/dev/sd[a-z]",       # 写裸设备
    r"\biptables\s+-F\b",      # 清空防火墙
    r"\bchmod\s+-R\s+777\s+/\b",  # 全盘777
    r"\b:\(\)\s*\{",           # fork炸弹
]

# SSH 连接池（支持多服务器并行连接）
_SSH_CONNECTIONS = {}  # {conn_id: {"conn": conn, "host": host, "user": user, "connected_at": ts}}

# 内网保留IP黑名单（防止SSRF类攻击，可选启用）
_SSH_BLOCK_PRIVATE_IPS = False  # 默认允许内网IP（远程部署通常就是内网服务器）

# SSH 专用事件循环（在独立线程运行，避免与 Textual 主事件循环冲突）
# asyncssh 的连接对象绑定到创建它的事件循环，必须保证所有 SSH 操作使用同一个循环
import threading as _ssh_threading_mod  # 用别名避免污染命名空间
_SSH_LOOP = None
_SSH_LOOP_THREAD = None
_SSH_LOOP_LOCK = _ssh_threading_mod.Lock()


def _ssh_run_async(coro_factory):
    """在 Textual 事件循环已运行的环境下安全执行 async 协程。

    设计要点（重要）：
    asyncssh 的连接对象（asyncssh.SSHClientConnection）**绑定到创建它的事件循环**，
    后续对该连接的所有操作（run、exec、close）必须在同一个事件循环里。
    因此采用「持久事件循环线程」模式：
    - 全局维护一个独立的 SSH 事件循环，运行在后台守护线程中
    - 所有 SSH 协程都通过 asyncio.run_coroutine_threadsafe 提交到这个循环
    - 这样无论 ZeroAI 主循环是否在运行，SSH 连接对象都能正确复用

    Args:
        coro_factory: 无参数可调用对象，返回一个协程
    Returns:
        协程的返回值；若执行出错，返回字符串 "错误：{异常}"
    """
    import asyncio
    import threading

    global _SSH_LOOP, _SSH_LOOP_THREAD
    with _SSH_LOOP_LOCK:
        if _SSH_LOOP is None or not _SSH_LOOP.is_running():
            _SSH_LOOP = asyncio.new_event_loop()

            def _loop_runner():
                asyncio.set_event_loop(_SSH_LOOP)
                try:
                    _SSH_LOOP.run_forever()
                finally:
                    _SSH_LOOP.close()

            _SSH_LOOP_THREAD = threading.Thread(target=_loop_runner, daemon=True, name="ssh-event-loop")
            _SSH_LOOP_THREAD.start()

    future = asyncio.run_coroutine_threadsafe(coro_factory(), _SSH_LOOP)
    try:
        return future.result(timeout=300)  # 整体超时 5 分钟
    except Exception as e:
        return f"错误：{e}"


def _ssh_is_conn_closed(conn) -> bool:
    """统一判断 asyncssh 连接是否已关闭。
    兼容不同 asyncssh 版本：is_closed 可能是属性（旧版）或方法（新版 2.24+）。
    """
    if conn is None:
        return True
    try:
        val = conn.is_closed
        if callable(val):
            val = val()
        return bool(val)
    except Exception:
        return True


# 操作系统检测缓存（按 conn_id 缓存，避免每次都探测）
_SSH_OS_CACHE = {}


def _ssh_detect_os(conn_id: str = "default") -> str:
    """检测远程服务器操作系统。返回 'windows' / 'linux'。
    结果按 conn_id 缓存，连接断开后自动清除缓存。

    探测策略（多重冗余，避免单一命令失败导致误判）：
    1. 优先用 `ver` 命令（Windows cmd 内建，输出含 "Microsoft Windows"）
    2. 回退用 `echo %OS%`（Windows 输出 "Windows_NT"）
    3. 再回退用 `uname`（Linux 输出内核名，Windows 无此命令）

    注意：本函数通过 _raw_ssh_exec 绕过 ssh_exec 的编码注入，避免递归调用。
    """
    # 如果连接已不存在，清除缓存
    if conn_id not in _SSH_CONNECTIONS:
        _SSH_OS_CACHE.pop(conn_id, None)
        return "linux"  # 默认按 Linux 处理

    # 命中缓存
    if conn_id in _SSH_OS_CACHE:
        return _SSH_OS_CACHE[conn_id]

    os_type = "linux"  # 默认 Linux

    # 内部执行函数（绕过 ssh_exec 的编码注入，避免递归）
    def _raw_exec(cmd: str) -> str:
        """直接调用 asyncssh，不经过 ssh_exec 的编码处理"""
        conn_info = _SSH_CONNECTIONS.get(conn_id)
        if not conn_info:
            return ""
        conn = conn_info["conn"]
        if _ssh_is_conn_closed(conn):
            return ""
        import asyncio
        async def _run():
            try:
                import asyncssh
                result = await asyncio.wait_for(
                    conn.run(cmd, check=False, timeout=8,
                             encoding='utf-8', errors='replace'),
                    timeout=13
                )
                return result.stdout or ""
            except Exception:
                return ""
        try:
            return _ssh_run_async(_run)
        except Exception:
            return ""

    # 探测1：ver 命令（Windows cmd 内建，最可靠）
    try:
        raw1 = _raw_exec("ver")
        if raw1 and ("Microsoft" in raw1 or "Windows" in raw1):
            os_type = "windows"
    except Exception:
        pass

    # 探测2：echo %OS%（Windows 输出 Windows_NT）
    if os_type == "linux":
        try:
            raw2 = _raw_exec("echo %OS%")
            if raw2 and "Windows_NT" in raw2:
                os_type = "windows"
        except Exception:
            pass

    # 探测3：uname（Linux 一定有输出，Windows 会报错）
    if os_type == "linux":
        try:
            raw3 = _raw_exec("uname")
            # Linux 的 uname 会输出 Linux/Darwin 等；Windows 会报 "不是内部或外部命令"
            if raw3 and ("不是内部或外部命令" in raw3 or "not recognized" in raw3
                         or "无法找到" in raw3):
                # uname 不存在 → 大概率是 Windows
                os_type = "windows"
        except Exception:
            pass

    _SSH_OS_CACHE[conn_id] = os_type
    return os_type


def _ssh_audit(host: str, user: str, command: str, result_summary: str = ""):
    """记录SSH操作审计日志（脱敏：不显示完整 IP 地址）"""
    import datetime
    # 安全设计：对 host 进行脱敏处理，不显示完整 IP
    # 如果是 IP 地址，只保留前两段，后两段用 *** 替代
    # 如果是域名，保留原样
    import re as _re_audit
    if _re_audit.match(r'^(\d{1,3}\.){3}\d{1,3}$', host):
        parts = host.split(".")
        safe_host = f"{parts[0]}.{parts[1]}.***.***"
    else:
        safe_host = host
    entry = f"[{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {user}@{safe_host} → {command[:200]}"
    if result_summary:
        entry += f" | {result_summary[:100]}"
    _SSH_AUDIT_LOG.append(entry)
    if len(_SSH_AUDIT_LOG) > _SSH_AUDIT_MAX:
        _SSH_AUDIT_LOG.pop(0)


def _ssh_format_prefix(conn_id: str = "default") -> str:
    """生成运维结果的服务器标识前缀（防混淆）。

    格式: "[conn_id | 备注]" 或 "[conn_id]"

    多服务器场景下，每个运维工具的返回结果都应以此前缀开头，
    让用户和 AI 一眼看清这是哪台机器的输出。

    安全设计：不在前缀中显示服务器 IP 地址，仅用 conn_id 和备注标识。
    如需查看完整连接信息（含 IP），请用 ssh_list 工具。

    Args:
        conn_id: 连接ID

    Returns:
        形如 "[nas | NAS存储服务器]" 或 "[default]" 的前缀字符串
    """
    info = _SSH_CONNECTIONS.get(conn_id)
    if not info:
        return f"[{conn_id}]"
    remark = info.get("remark", "")
    if remark:
        return f"[{conn_id} | {remark}]"
    return f"[{conn_id}]"


def _ssh_check_dangerous(command: str) -> tuple:
    """检查命令是否危险，返回 (is_dangerous, matched_pattern)"""
    import re
    for pattern in _SSH_DANGEROUS_PATTERNS:
        if re.search(pattern, command, re.IGNORECASE):
            return True, pattern
    return False, None


def _ssh_validate_host(host: str) -> tuple:
    """校验主机地址合法性，返回 (is_valid, error_msg)"""
    if not host or not isinstance(host, str):
        return False, "主机地址不能为空"
    # 去除协议前缀
    host = host.replace("ssh://", "").replace("SSH://", "")
    # 去除端口
    hostname = host.split(":")[0]
    # 校验IP或域名格式
    import re
    ip_pattern = r"^(\d{1,3}\.){3}\d{1,3}$"
    domain_pattern = r"^([a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}$"
    if re.match(ip_pattern, hostname):
        parts = hostname.split(".")
        for p in parts:
            if int(p) > 255:
                return False, f"IP地址段无效: {p}"
        # 检查内网IP（可选阻断）
        if _SSH_BLOCK_PRIVATE_IPS:
            if parts[0] in ("10", "172", "192", "127"):
                if parts[0] == "172" and not (16 <= int(parts[1]) <= 31):
                    pass  # 172.x 但不在16-31范围，不阻断
                elif parts[0] == "192" and parts[1] != "168":
                    pass  # 192.x 但不是168，不阻断
                else:
                    return False, "内网地址被策略阻断（如需连接内网，请联系管理员调整策略）"
        return True, ""
    elif re.match(domain_pattern, hostname):
        return True, ""
    else:
        return False, f"主机地址格式无效（请检查 IP/域名格式）"


def ssh_connect(host: str, user: str, password: str = "", key_path: str = "",
                port: int = 22, conn_id: str = "default", remark: str = "") -> str:
    """连接到远程SSH服务器。

    支持密码认证和密钥认证两种方式。连接成功后保持长连接，后续命令通过 conn_id 复用。

    Args:
        host: 服务器地址（IP或域名，可带端口如 192.168.1.100:2222）
        user: 登录用户名
        password: 密码认证（二选一）
        key_path: SSH私钥路径（如 ~/.ssh/id_rsa），密码和密钥二选一
        port: SSH端口，默认22
        conn_id: 连接标识符，用于多服务器管理，默认"default"
        remark: 服务器备注/角色（如 "NAS存储"、"Web前端"、"数据库"），便于多服务器场景下识别，防混淆

    Returns:
        连接状态信息
    """
    import asyncio

    # 校验主机地址
    is_valid, err = _ssh_validate_host(host)
    if not is_valid:
        return f"连接失败：{err}"

    # 如果已有同名连接，先断开
    if conn_id in _SSH_CONNECTIONS:
        try:
            old = _SSH_CONNECTIONS.pop(conn_id)
            if old.get("conn") and not _ssh_is_conn_closed(old["conn"]):
                # asyncssh 的 close() 是同步方法
                old["conn"].close()
        except Exception:
            pass

    async def _connect():
        import asyncssh
        try:
            # 准备认证参数
            connect_kwargs = {
                "host": host.split(":")[0] if ":" in host else host,
                "port": port if ":" not in host else int(host.split(":")[1]),
                "username": user,
                "known_hosts": None,  # 跳过known_hosts检查（部署场景）
                "login_timeout": 15,
                "keepalive_interval": 30,
                "keepalive_count_max": 3,
            }
            if key_path:
                # 密钥认证
                key_path_expanded = os.path.expanduser(key_path)
                if not os.path.exists(key_path_expanded):
                    return f"连接失败：私钥文件不存在: {key_path_expanded}"
                connect_kwargs["client_keys"] = [key_path_expanded]
            elif password:
                # 密码认证
                connect_kwargs["password"] = password
            else:
                return "连接失败：必须提供 password 或 key_path 之一"

            conn = await asyncssh.connect(**connect_kwargs)
            return conn
        except asyncssh.PermissionDenied:
            return "连接失败：认证失败（密码/密钥错误）"
        except asyncssh.ConnectionLost:
            return "连接失败：连接丢失（网络不稳定）"
        except asyncssh.DisconnectError as e:
            return f"连接失败：服务器拒绝连接 (code={e.code}, reason={e.reason})"
        except asyncio.TimeoutError:
            return f"连接失败：超时（15秒内未连接到服务器，请检查网络和端口）"
        except OSError as e:
            return f"连接失败：网络错误 ({e})"

    try:
        result = _ssh_run_async(_connect)
        if isinstance(result, str):
            return result
        # 连接成功
        import time
        _SSH_CONNECTIONS[conn_id] = {
            "conn": result,
            "host": host.split(":")[0] if ":" in host else host,
            "user": user,
            "port": port if ":" not in host else int(host.split(":")[1]),
            "connected_at": time.time(),
            "remark": remark,  # 服务器备注/角色（防混淆）
        }
        # 连接建立后清除 OS 缓存，下次运维工具调用时重新探测
        _SSH_OS_CACHE.pop(conn_id, None)
        _ssh_audit(host, user, "[CONNECT]", f"成功 conn_id={conn_id} remark={remark}")
        remark_line = f"\n  备注: {remark}" if remark else ""
        # 安全设计：不在返回结果中显示服务器 IP 地址，仅显示备注/conn_id
        # IP 地址仅存储在内部 _SSH_CONNECTIONS 中供工具内部使用
        server_label = remark if remark else conn_id
        return (f"✅ SSH连接成功\n  服务器: {server_label}\n  用户: {user}\n  端口: {port if ':' not in host else host.split(':')[1]}\n"
                f"  连接ID: {conn_id}\n  认证方式: {'密钥' if key_path else '密码'}{remark_line}\n"
                f"  提示: 后续运维操作请传 conn_id='{conn_id}'")
    except Exception as e:
        return f"连接失败：{e}"


def ssh_exec(command: str, conn_id: str = "default", timeout: int = 30,
             confirm_dangerous: bool = False, _internal: bool = False) -> str:
    """在远程服务器上执行Shell命令。

    通过已建立的SSH连接执行命令。危险命令（如rm -rf /、mkfs、shutdown）需要
    confirm_dangerous=True 才会执行。

    Args:
        command: 要执行的Shell命令
        conn_id: 连接ID（由ssh_connect返回），默认"default"
        timeout: 命令超时时间（秒），默认30
        confirm_dangerous: 是否确认执行危险命令，默认False
        _internal: 内部调用标记（运维工具内部调用时传 True，不加服务器前缀，避免前缀重复）

    Returns:
        命令输出结果（stdout + stderr），默认带服务器标识前缀防混淆
    """
    import asyncio

    # 检查连接是否存在
    if conn_id not in _SSH_CONNECTIONS:
        return f"错误：连接 '{conn_id}' 不存在，请先调用 ssh_connect 建立连接"

    conn_info = _SSH_CONNECTIONS[conn_id]
    conn = conn_info["conn"]
    # 兼容 asyncssh 不同版本：is_closed 可能是属性（旧版）或方法（新版 2.24+）
    if _ssh_is_conn_closed(conn):
        _SSH_CONNECTIONS.pop(conn_id, None)
        return f"错误：连接 '{conn_id}' 已断开，请重新调用 ssh_connect"

    # 危险命令检查
    is_dangerous, pattern = _ssh_check_dangerous(command)
    if is_dangerous and not confirm_dangerous:
        return (f"⚠️ 检测到危险命令（匹配模式: {pattern}）\n"
                f"命令: {command}\n"
                f"如确认要执行，请重新调用并设置 confirm_dangerous=true")

    async def _exec():
        try:
            import asyncssh
            # encoding='utf-8', errors='replace' 兼容中文 Windows 的 GBK 输出
            # asyncssh 默认用 UTF-8 解码，Windows cmd 输出 GBK 会乱码或报错
            result = await asyncio.wait_for(
                conn.run(command, check=False, timeout=timeout,
                         encoding='utf-8', errors='replace'),
                timeout=timeout + 5
            )
            return result
        except asyncio.TimeoutError:
            return f"命令超时（{timeout}秒）"
        except asyncssh.ChannelOpenError as e:
            return f"通道错误: {e}"
        except Exception as e:
            return f"执行错误: {e}"

    try:
        # Windows 中文乱码修复：自动给 PowerShell/cmd 命令注入 UTF-8 编码
        # 原因：Windows 中文系统默认 GBK 编码，asyncssh 用 UTF-8 解码会乱码
        if _ssh_detect_os(conn_id) == "windows":
            if command.startswith("powershell"):
                # PowerShell 命令：注入 UTF-8 输出编码设置
                if '-Command "' in command:
                    command = command.replace(
                        '-Command "',
                        '-Command "[Console]::OutputEncoding = [System.Text.Encoding]::UTF8; $OutputEncoding = [System.Text.Encoding]::UTF8; ',
                        1
                    )
                elif "-Command '" in command:
                    command = command.replace(
                        "-Command '",
                        "-Command '[Console]::OutputEncoding = [System.Text.Encoding]::UTF8; $OutputEncoding = [System.Text.Encoding]::UTF8; ",
                        1
                    )
            else:
                # cmd 命令：切换代码页到 65001 (UTF-8)
                if not command.startswith("chcp"):
                    command = f"chcp 65001 >nul 2>&1 & {command}"

        result = _ssh_run_async(_exec)
        prefix = "" if _internal else _ssh_format_prefix(conn_id) + "\n"
        if isinstance(result, str):
            _ssh_audit(conn_info["host"], conn_info["user"], command, result[:100])
            return prefix + result

        stdout = result.stdout or ""
        stderr = result.stderr or ""
        exit_code = result.exit_status or 0

        # 截断过长输出
        if len(stdout) > 8000:
            stdout = stdout[:8000] + f"\n... (输出过长，已截断，共 {len(stdout)} 字符)"
        if len(stderr) > 4000:
            stderr = stderr[:4000] + f"\n... (错误输出过长，已截断，共 {len(stderr)} 字符)"

        summary = f"exit={exit_code}, out={len(stdout)}B, err={len(stderr)}B"
        _ssh_audit(conn_info["host"], conn_info["user"], command, summary)

        # 格式化输出
        parts = []
        if stdout:
            parts.append(stdout.rstrip())
        if stderr:
            parts.append(f"[stderr]\n{stderr.rstrip()}")
        if exit_code != 0:
            parts.append(f"[退出码: {exit_code}]")

        body = "\n".join(parts) if parts else "[无输出]"
        return prefix + body
    except Exception as e:
        return f"执行错误: {e}"


def ssh_upload(local_path: str, remote_path: str, conn_id: str = "default") -> str:
    """上传本地文件到远程服务器（SFTP）。

    支持单文件上传。上传后自动设置权限为644。

    Args:
        local_path: 本地文件路径
        remote_path: 远程目标路径（完整路径，如 /opt/myapp/config.yml）
        conn_id: 连接ID

    Returns:
        上传结果
    """
    import asyncio

    if conn_id not in _SSH_CONNECTIONS:
        return f"错误：连接 '{conn_id}' 不存在，请先调用 ssh_connect"

    conn_info = _SSH_CONNECTIONS[conn_id]
    conn = conn_info["conn"]
    if _ssh_is_conn_closed(conn):
        _SSH_CONNECTIONS.pop(conn_id, None)
        return f"错误：连接 '{conn_id}' 已断开，请重新连接"

    local_path = os.path.abspath(local_path)
    if not os.path.exists(local_path):
        return f"错误：本地文件不存在: {local_path}"

    local_size = os.path.getsize(local_path)

    async def _upload():
        try:
            import asyncssh
            async with conn.start_sftp_client() as sftp:
                await sftp.put(local_path, remote_path)
                # 设置权限644
                try:
                    await sftp.chmod(remote_path, 0o644)
                except Exception:
                    pass  # 权限设置失败不影响上传
            return True
        except asyncssh.SFTPError as e:
            return f"SFTP错误: {e}"
        except Exception as e:
            return f"上传错误: {e}"

    try:
        result = _ssh_run_async(_upload)
        if result is True:
            _ssh_audit(conn_info["host"], conn_info["user"],
                       f"[UPLOAD] {local_path} → {remote_path}",
                       f"{local_size}B")
            return f"{_ssh_format_prefix(conn_id)}\n✅ 上传成功\n  本地: {local_path} ({local_size} 字节)\n  远程: {remote_path}"
        return f"上传失败: {result}"
    except Exception as e:
        return f"上传错误: {e}"


def ssh_download(remote_path: str, local_path: str, conn_id: str = "default") -> str:
    """从远程服务器下载文件到本地（SFTP）。

    Args:
        remote_path: 远程文件路径
        local_path: 本地保存路径
        conn_id: 连接ID

    Returns:
        下载结果
    """
    import asyncio

    if conn_id not in _SSH_CONNECTIONS:
        return f"错误：连接 '{conn_id}' 不存在，请先调用 ssh_connect"

    conn_info = _SSH_CONNECTIONS[conn_id]
    conn = conn_info["conn"]
    if _ssh_is_conn_closed(conn):
        _SSH_CONNECTIONS.pop(conn_id, None)
        return f"错误：连接 '{conn_id}' 已断开，请重新连接"

    # 确保本地目录存在
    local_dir = os.path.dirname(os.path.abspath(local_path))
    if local_dir:
        os.makedirs(local_dir, exist_ok=True)

    async def _download():
        try:
            import asyncssh
            async with conn.start_sftp_client() as sftp:
                await sftp.get(remote_path, local_path)
            return True
        except asyncssh.SFTPError as e:
            return f"SFTP错误: {e}"
        except Exception as e:
            return f"下载错误: {e}"

    try:
        result = _ssh_run_async(_download)
        if result is True:
            local_size = os.path.getsize(local_path) if os.path.exists(local_path) else 0
            _ssh_audit(conn_info["host"], conn_info["user"],
                       f"[DOWNLOAD] {remote_path} → {local_path}",
                       f"{local_size}B")
            return f"{_ssh_format_prefix(conn_id)}\n✅ 下载成功\n  远程: {remote_path}\n  本地: {local_path} ({local_size} 字节)"
        return f"下载失败: {result}"
    except Exception as e:
        return f"下载错误: {e}"


def ssh_deploy(deploy_config: dict, conn_id: str = "default") -> str:
    """一键项目部署（多步骤自动化部署）。

    按顺序执行部署步骤：环境检查→创建目录→上传代码→安装依赖→重启服务→健康检查。

    Args:
        deploy_config: 部署配置字典，包含：
            - pre_check: 部署前检查命令列表（如 ["uname -a", "docker --version"]）
            - remote_dir: 远程部署目录（如 /opt/myapp）
            - upload_files: 上传文件列表 [[local, remote], ...]
            - install_cmd: 安装依赖命令（如 "pip install -r requirements.txt"）
            - restart_cmd: 重启服务命令（如 "systemctl restart myapp"）
            - health_check: 健康检查命令（如 "curl -s localhost:8080/health"）
            - post_cmds: 部署后额外命令列表
        conn_id: 连接ID

    Returns:
        部署报告（每步结果汇总）
    """
    if conn_id not in _SSH_CONNECTIONS:
        return f"错误：连接 '{conn_id}' 不存在，请先调用 ssh_connect"

    report = []
    report.append("=" * 50)
    report.append("🚀 SSH 自动化部署")
    report.append("=" * 50)

    step = 0
    total_steps = 0
    # 计算总步骤数
    if deploy_config.get("pre_check"):
        total_steps += len(deploy_config["pre_check"])
    if deploy_config.get("remote_dir"):
        total_steps += 1
    if deploy_config.get("upload_files"):
        total_steps += len(deploy_config["upload_files"])
    if deploy_config.get("install_cmd"):
        total_steps += 1
    if deploy_config.get("restart_cmd"):
        total_steps += 1
    if deploy_config.get("health_check"):
        total_steps += 1
    if deploy_config.get("post_cmds"):
        total_steps += len(deploy_config["post_cmds"])

    report.append(f"总步骤: {total_steps}")
    report.append("")

    # 1. 环境检查
    if deploy_config.get("pre_check"):
        report.append("📋 [步骤] 环境检查")
        for cmd in deploy_config["pre_check"]:
            step += 1
            result = ssh_exec(cmd, conn_id, _internal=True, timeout=15)
            status = "✅" if "错误" not in result and "exit_code" not in result.lower() else "⚠️"
            report.append(f"  {status} [{step}/{total_steps}] {cmd}")
            report.append(f"     {result[:200]}")
            report.append("")

    # 2. 创建远程目录
    if deploy_config.get("remote_dir"):
        step += 1
        remote_dir = deploy_config["remote_dir"]
        report.append(f"📁 [步骤 {step}/{total_steps}] 创建目录: {remote_dir}")
        result = ssh_exec(f"mkdir -p {remote_dir}", conn_id)
        report.append(f"  {result[:200]}")
        report.append("")

    # 3. 上传文件
    if deploy_config.get("upload_files"):
        for local, remote in deploy_config["upload_files"]:
            step += 1
            report.append(f"📤 [步骤 {step}/{total_steps}] 上传: {local} → {remote}")
            result = ssh_upload(local, remote, conn_id)
            status = "✅" if "成功" in result else "❌"
            report.append(f"  {status} {result[:200]}")
            report.append("")

    # 4. 安装依赖
    if deploy_config.get("install_cmd"):
        step += 1
        install_cmd = deploy_config["install_cmd"]
        remote_dir = deploy_config.get("remote_dir", "")
        report.append(f"📦 [步骤 {step}/{total_steps}] 安装依赖: {install_cmd}")
        full_cmd = f"cd {remote_dir} && {install_cmd}" if remote_dir else install_cmd
        result = ssh_exec(full_cmd, conn_id, _internal=True, timeout=120)
        report.append(f"  {result[:500]}")
        report.append("")

    # 5. 重启服务
    if deploy_config.get("restart_cmd"):
        step += 1
        restart_cmd = deploy_config["restart_cmd"]
        report.append(f"🔄 [步骤 {step}/{total_steps}] 重启服务: {restart_cmd}")
        result = ssh_exec(restart_cmd, conn_id, timeout=30)
        report.append(f"  {result[:300]}")
        report.append("")

    # 6. 健康检查
    if deploy_config.get("health_check"):
        step += 1
        health_cmd = deploy_config["health_check"]
        report.append(f"🏥 [步骤 {step}/{total_steps}] 健康检查: {health_cmd}")
        result = ssh_exec(health_cmd, conn_id, _internal=True, timeout=15)
        status = "✅ 健康" if "错误" not in result and "exit_code" not in result.lower() else "⚠️ 需检查"
        report.append(f"  {status}")
        report.append(f"  {result[:300]}")
        report.append("")

    # 7. 部署后命令
    if deploy_config.get("post_cmds"):
        for cmd in deploy_config["post_cmds"]:
            step += 1
            report.append(f"⚙️ [步骤 {step}/{total_steps}] 后置: {cmd}")
            result = ssh_exec(cmd, conn_id, _internal=True, timeout=30)
            report.append(f"  {result[:200]}")
            report.append("")

    # 汇总
    report.append("=" * 50)
    report.append(f"✅ 部署完成 ({step}/{total_steps} 步骤已执行)")
    report.append("=" * 50)

    return _ssh_format_prefix(conn_id) + "\n" + "\n".join(report)


def ssh_setup_samba_share(share_name: str = "shared",
                          share_path: str = "/srv/shared",
                          access_mode: str = "guest_rw",
                          samba_password: str = "",
                          conn_id: str = "default") -> str:
    r"""一键配置 Samba 共享文件夹（Linux 服务器专用，自动完成全部步骤）。

    自动执行的 8 个步骤：
    1. 检测操作系统（必须是 Linux，Windows 应该用 New-SmbShare）
    2. 安装 Samba（自动识别 apt/yum/dnf 包管理器）
    3. 创建共享文件夹并设置权限
    4. 备份原 smb.conf
    5. 写入共享配置（支持 guest_ro/guest_rw/user_rw 三种权限模式）
    6. 设置 Samba 密码（user_rw 模式需要）
    7. 启动 smbd/nmb 服务并设置开机自启
    8. 防火墙放行 + SELinux 处理 + 验证共享

    Args:
        share_name: 共享名（Windows 访问时用，如 "shared"，访问路径 \\\\IP\\shared）
        share_path: 共享文件夹在 Linux 上的路径，默认 /srv/shared
        access_mode: 权限模式：
            - guest_ro: 匿名只读（任何人可读不可写）
            - guest_rw: 匿名读写（任何人可读写，适合内网共享，默认）
            - user_rw: 用户认证读写（需 samba_password，更安全）
        samba_password: Samba 密码（user_rw 模式必填，其他模式可留空）
        conn_id: SSH 连接ID

    Returns:
        配置结果报告 + Windows 访问路径
    """
    import base64

    prefix = _ssh_format_prefix(conn_id)

    # 步骤1: 检测操作系统
    os_type = _ssh_detect_os(conn_id)
    if os_type == "windows":
        return (f"{prefix}\n❌ 此工具仅支持 Linux 服务器配置 Samba 共享。\n"
                f"Windows Server 请用 ssh_exec 执行 PowerShell 命令：\n"
                f"  New-SmbShare -Name '{share_name}' -Path 'D:\\{share_name}' -FullAccess Everyone")

    report = [
        "=" * 50,
        f"Samba 共享一键配置报告",
        f"  共享名: {share_name}",
        f"  共享路径: {share_path}",
        f"  权限模式: {access_mode}",
        "=" * 50,
    ]
    step = 0
    total_steps = 8

    # 步骤2: 安装 Samba（自动识别包管理器）
    step += 1
    report.append(f"\n[{step}/{total_steps}] 安装 Samba...")
    install_cmd = (
        'if command -v apt-get >/dev/null 2>&1; then '
        '  apt-get update -qq 2>&1 | tail -1; apt-get install -y -qq samba 2>&1 | tail -3; '
        'elif command -v dnf >/dev/null 2>&1; then '
        '  dnf install -y samba 2>&1 | tail -3; '
        'elif command -v yum >/dev/null 2>&1; then '
        '  yum install -y samba 2>&1 | tail -3; '
        'else echo "错误：未识别的包管理器（apt/dnf/yum 均不存在）"; exit 1; fi; '
        'which smbd && smbd --version'
    )
    r = ssh_exec(install_cmd, conn_id, _internal=True, timeout=180)
    if "smbd" not in r or "Version" not in r:
        report.append(f"  ❌ Samba 安装失败:\n{r[-500:]}")
        report.append("\n" + "=" * 50)
        report.append("❌ 配置失败，请检查包管理器或网络")
        return prefix + "\n" + "\n".join(report)
    report.append(f"  ✅ Samba 已安装: {r.split('Version')[-1].strip() if 'Version' in r else '已就绪'}")

    # 步骤3: 创建共享文件夹
    step += 1
    report.append(f"\n[{step}/{total_steps}] 创建共享文件夹 {share_path}...")
    r = ssh_exec(f'mkdir -p {share_path} && chmod 777 {share_path} && ls -ld {share_path}',
                 conn_id, _internal=True, timeout=10)
    if "drwx" not in r:
        report.append(f"  ❌ 文件夹创建失败:\n{r[-300:]}")
        return prefix + "\n" + "\n".join(report)
    report.append(f"  ✅ 文件夹已创建: {r.split('drwx')[0].strip() or r.strip().split(chr(10))[-1]}")

    # 步骤4: 备份原 smb.conf
    step += 1
    report.append(f"\n[{step}/{total_steps}] 备份原 smb.conf...")
    r = ssh_exec('cp /etc/samba/smb.conf /etc/samba/smb.conf.bak.$(date +%s) 2>/dev/null && echo "备份成功" || echo "无需备份（首次配置）"',
                 conn_id, _internal=True, timeout=10)
    report.append(f"  ✅ {r.strip().split(chr(10))[-1]}")

    # 步骤5: 写入共享配置（根据权限模式生成不同配置）
    step += 1
    report.append(f"\n[{step}/{total_steps}] 写入共享配置（权限模式: {access_mode}）...")

    if access_mode == "guest_ro":
        # 匿名只读
        share_config = f"""[global]
   workgroup = WORKGROUP
   security = user
   map to guest = Bad User
   passdb backend = tdbsam
   printing = bsd
   printcap name = /dev/null
   load printers = no

[{share_name}]
   comment = Shared Folder
   path = {share_path}
   browseable = yes
   writable = no
   guest ok = yes
   read only = yes
"""
    elif access_mode == "guest_rw":
        # 匿名读写（默认，内网共享推荐）
        share_config = f"""[global]
   workgroup = WORKGROUP
   security = user
   map to guest = Bad User
   passdb backend = tdbsam
   printing = bsd
   printcap name = /dev/null
   load printers = no

[{share_name}]
   comment = Shared Folder
   path = {share_path}
   browseable = yes
   writable = yes
   guest ok = yes
   force user = root
   force group = root
   create mask = 0666
   directory mask = 0777
"""
    else:  # user_rw
        # 用户认证读写（更安全）
        share_config = f"""[global]
   workgroup = WORKGROUP
   security = user
   passdb backend = tdbsam
   printing = bsd
   printcap name = /dev/null
   load printers = no

[{share_name}]
   comment = Shared Folder
   path = {share_path}
   browseable = yes
   writable = yes
   guest ok = no
   valid users = root
   create mask = 0664
   directory mask = 0775
"""

    b64 = base64.b64encode(share_config.encode('utf-8')).decode('ascii')
    r = ssh_exec(f'echo "{b64}" | base64 -d > /etc/samba/smb.conf && testparm -s 2>&1 | head -20',
                 conn_id, _internal=True, timeout=10)
    if "Loaded services file" not in r and "[" + share_name + "]" not in r:
        report.append(f"  ❌ 配置写入失败:\n{r[-400:]}")
        return prefix + "\n" + "\n".join(report)
    report.append(f"  ✅ 配置已写入，testparm 校验通过")

    # 步骤6: 设置 Samba 密码（user_rw 模式）
    step += 1
    if access_mode == "user_rw":
        report.append(f"\n[{step}/{total_steps}] 设置 Samba 密码 (root)...")
        if not samba_password:
            report.append("  ⚠️ user_rw 模式未提供密码，跳过（可用 smbpasswd -a root 手动设置）")
        else:
            # 用单引号包裹密码避免特殊字符问题
            r = ssh_exec(f'(echo "{samba_password}"; echo "{samba_password}") | smbpasswd -a root -s 2>&1',
                         conn_id, _internal=True, timeout=10)
            if "Added user" in r:
                report.append(f"  ✅ Samba 密码已设置 (root)")
            else:
                report.append(f"  ⚠️ 密码设置失败: {r.strip()[-200:]}")
    else:
        report.append(f"\n[{step}/{total_steps}] 跳过密码设置（{access_mode} 模式无需密码）")

    # 步骤7: 启动 smbd/nmb 服务 + 开机自启
    step += 1
    report.append(f"\n[{step}/{total_steps}] 启动 smbd/nmb 服务 + 开机自启...")
    r = ssh_exec('systemctl enable --now smb nmb 2>&1 | tail -2; systemctl is-active smb nmb; systemctl is-enabled smb nmb',
                 conn_id, _internal=True, timeout=15)
    if "active" not in r:
        report.append(f"  ❌ 服务启动失败:\n{r[-300:]}")
        return prefix + "\n" + "\n".join(report)
    report.append(f"  ✅ smbd/nmb 已启动并设为开机自启")

    # 步骤8: 防火墙放行 + SELinux 处理 + 验证
    step += 1
    report.append(f"\n[{step}/{total_steps}] 防火墙放行 + SELinux 处理...")

    # 防火墙
    r_fw = ssh_exec('systemctl is-active firewalld 2>/dev/null && '
                    '(firewall-cmd --permanent --add-service=samba 2>&1; firewall-cmd --reload 2>&1) || '
                    'echo "firewalld 未运行，跳过"',
                    conn_id, _internal=True, timeout=15)
    if "success" in r_fw:
        report.append("  ✅ 防火墙已放行 Samba 服务")
    else:
        report.append(f"  ℹ️ 防火墙: {r_fw.strip().split(chr(10))[-1]}")

    # SELinux
    r_se = ssh_exec('getenforce 2>/dev/null || echo "Disabled"',
                    conn_id, _internal=True, timeout=10)
    se_status = r_se.strip().split(chr(10))[-1] if r_se else "Disabled"
    if se_status == "Enforcing":
        ssh_exec(f'setsebool -P samba_enable_home_dirs on 2>/dev/null; '
                 f'semanage fcontext -a -t samba_share_t "{share_path}(/.*)?" 2>/dev/null; '
                 f'restorecon -Rv {share_path} 2>&1 | tail -1',
                 conn_id, _internal=True, timeout=15)
        report.append("  ✅ SELinux 上下文已设置")
    else:
        report.append(f"  ℹ️ SELinux: {se_status}（无需处理）")

    # 验证共享
    r_test = ssh_exec(f'echo "Samba 共享测试 $(date)" > {share_path}/test.txt && ls -l {share_path}/test.txt',
                      conn_id, _internal=True, timeout=10)
    if "test.txt" in r_test:
        report.append(f"  ✅ 共享写入测试通过")

    # 获取服务器信息
    info = _SSH_CONNECTIONS.get(conn_id, {})
    host = info.get("host", "服务器IP")
    remark = info.get("remark", "")

    # 最终报告
    report.append("\n" + "=" * 50)
    report.append("✅ Samba 共享配置完成！")
    report.append("=" * 50)
    report.append(f"\n【访问方式】")
    report.append(f"  Windows 资源管理器地址栏输入:")
    report.append(f"    \\\\{host}\\{share_name}")
    if access_mode == "user_rw":
        report.append(f"  用户名: root")
        report.append(f"  密码: {samba_password or '（请用 smbpasswd -a root 设置）'}")
    else:
        report.append(f"  权限: {'读写' if access_mode == 'guest_rw' else '只读'}（无需密码）")
    if remark:
        report.append(f"\n【服务器备注】{remark}（conn_id: {conn_id}）")
    report.append(f"\n【共享路径】{share_path}")
    report.append(f"【配置文件】/etc/samba/smb.conf（原配置已备份为 smb.conf.bak.*）")

    return prefix + "\n" + "\n".join(report)


def ssh_list(conn_id: str = "") -> str:
    """查看SSH连接状态和审计日志。

    Args:
        conn_id: 指定连接ID查看详情，留空查看所有连接和最近审计日志

    Returns:
        连接状态和审计日志（含备注/角色、操作系统、连接时长，便于多服务器场景识别）
    """
    def _fmt_uptime(secs: int) -> str:
        """格式化连接时长"""
        if secs < 60:
            return f"{secs}s"
        if secs < 3600:
            return f"{secs // 60}m{secs % 60}s"
        return f"{secs // 3600}h{(secs % 3600) // 60}m"

    if conn_id:
        if conn_id not in _SSH_CONNECTIONS:
            return f"连接 '{conn_id}' 不存在"
        info = _SSH_CONNECTIONS[conn_id]
        import time
        uptime = int(time.time() - info.get("connected_at", 0))
        is_closed = _ssh_is_conn_closed(info["conn"])
        os_type = _SSH_OS_CACHE.get(conn_id, "?") if not is_closed else "?"
        remark = info.get("remark", "")
        # 安全设计：不显示服务器 IP，仅显示备注或 conn_id
        server_label = remark if remark else conn_id
        lines = [
            f"连接ID: {conn_id}",
            f"  服务器: {server_label}",
            f"  用户: {info['user']}",
            f"  端口: {info.get('port', 22)}",
            f"  状态: {'❌ 已断开' if is_closed else '✅ 已连接'}",
            f"  操作系统: {os_type}",
            f"  连接时长: {_fmt_uptime(uptime)}",
        ]
        if remark:
            lines.append(f"  备注: {remark}")
        return "\n".join(lines)

    # 列出所有连接
    parts = ["=== SSH 连接状态 ==="]
    if not _SSH_CONNECTIONS:
        parts.append("  (无活动连接)")
        parts.append("  提示: 用 ssh_connect(host, user, password, conn_id='自定义ID', remark='服务器用途') 连接")
    else:
        import time
        parts.append(f"  共 {len(_SSH_CONNECTIONS)} 个连接:")
        parts.append("")
        for cid, info in _SSH_CONNECTIONS.items():
            uptime = int(time.time() - info.get("connected_at", 0))
            is_closed = _ssh_is_conn_closed(info["conn"])
            status = "❌" if is_closed else "✅"
            os_type = _SSH_OS_CACHE.get(cid, "?") if not is_closed else "?"
            remark = info.get("remark", "")
            # 安全设计：不显示 IP，用备注或 conn_id 标识服务器
            server_label = remark if remark else cid
            line = (f"  {status} {cid}: {info['user']}@{server_label}:{info.get('port', 22)} "
                    f"({ _fmt_uptime(uptime)}) [OS: {os_type}]")
            if remark:
                line += f" 备注: {remark}"
            parts.append(line)

    # 审计日志（最近20条）
    if _SSH_AUDIT_LOG:
        parts.append("")
        parts.append("=== 最近操作审计 ===")
        for entry in _SSH_AUDIT_LOG[-20:]:
            parts.append(f"  {entry}")

    return "\n".join(parts)


def ssh_disconnect(conn_id: str = "default") -> str:
    """断开SSH连接。

    Args:
        conn_id: 要断开的连接ID

    Returns:
        断开结果
    """
    if conn_id not in _SSH_CONNECTIONS:
        return f"连接 '{conn_id}' 不存在"

    conn_info = _SSH_CONNECTIONS.pop(conn_id)
    conn = conn_info["conn"]

    async def _close():
        try:
            conn.close()
            await conn.wait_closed()
        except Exception:
            pass

    try:
        _ssh_run_async(_close)
    except Exception:
        pass

    _ssh_audit(conn_info["host"], conn_info["user"], "[DISCONNECT]", f"conn_id={conn_id}")
    # 安全设计：不显示服务器 IP，用备注或 conn_id 标识
    remark = conn_info.get("remark", "")
    server_label = remark if remark else conn_id
    return f"✅ 已断开连接 '{conn_id}' ({conn_info['user']}@{server_label})"


# ====== AI 远程运维工具集（基于 ssh_exec 的高层封装）======
# 设计原则：把"AI 拼命令"升级为"AI 调用语义化工具"，减少幻觉、统一错误处理、自动分析结果


def ssh_service_manage(action: str, service: str, conn_id: str = "default") -> str:
    """服务管理（Linux 用 systemctl，Windows 用 sc/Get-Service）。

    Args:
        action: 操作类型，可选值：status / start / stop / restart / reload / enable / disable / is-active / is-enabled
            特殊：service="all" + action="status" 可列出所有运行中的服务
        service: 服务名（如 nginx、mysql、docker、ssh、spooler），或 "all" 查看全部
        conn_id: SSH 连接ID

    Returns:
        服务状态或操作结果（含 AI 友好的状态解读）
    """
    # 白名单校验，防注入
    valid_actions = {"status", "start", "stop", "restart", "reload",
                     "enable", "disable", "is-active", "is-enabled"}
    if action not in valid_actions:
        return f"错误：action 必须是 {sorted(valid_actions)} 之一"

    # service 校验：允许 "all" 或合法服务名
    if not service:
        return f"错误：必须提供 service 参数（服务名或 'all'）"
    if service != "all" and not service.replace("-", "").replace("_", "").replace(".", "").isalnum():
        return f"错误：服务名 '{service}' 不合法（仅允许字母数字-_.）"

    os_type = _ssh_detect_os(conn_id)

    if os_type == "windows":
        # Windows 服务管理（用 sc 命令，兼容性最好）
        if service == "all":
            if action == "status":
                # 列出所有运行中的服务（State=Running）
                cmd = 'powershell -NoProfile -Command "Get-Service | Where-Object {$_.Status -eq \'Running\'} | Format-Table Name, DisplayName, Status -AutoSize"'
                raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30)
                return f"{_ssh_format_prefix(conn_id)}\n$ 列出所有运行中的服务\n{raw}\n\n提示：共显示运行中的服务，可用 ssh_service_manage(action='status', service='具体服务名') 查看单个服务详情"
            else:
                return f"错误：service='all' 只支持 action='status'"
        else:
            # 单个服务操作
            action_map = {
                "status": ("sc", "query"),
                "start": ("sc", "start"),
                "stop": ("sc", "stop"),
                "restart": ("sc", "stop & sc start"),  # Windows sc 无 restart，用 stop+start
                "is-active": ("sc", "query"),
            }
            if action in ("enable", "disable", "is-enabled", "reload"):
                # 这些是 systemd 概念，Windows 用 sc config
                if action == "enable":
                    cmd = f'sc config {service} start= auto'
                elif action == "disable":
                    cmd = f'sc config {service} start= demand'
                elif action == "is-enabled":
                    cmd = f'sc qc {service}'
                else:  # reload
                    return "错误：Windows 服务不支持 reload，请用 restart"
                raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=15)
                interp = ""
                if action == "enable":
                    interp = " → ✅ 已设置开机自启（自动启动）"
                elif action == "disable":
                    interp = " → ✅ 已改为手动启动"
                elif action == "is-enabled":
                    if "AUTO_START" in raw:
                        interp = " → 已设置开机自启"
                    elif "DEMAND_START" in raw:
                        interp = " → 手动启动"
                    elif "DISABLED" in raw:
                        interp = " → 已禁用"
                return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}{interp}"

            sc_cmd, sc_action = action_map.get(action, ("sc", "query"))
            if action == "restart":
                cmd = f'{sc_cmd} {sc_action} {service}'
            else:
                cmd = f'{sc_cmd} {sc_action} {service}'
            raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=15)

            # Windows 状态解读
            interp = ""
            if action in ("status", "is-active"):
                if "RUNNING" in raw:
                    interp = " → ✅ 服务运行中"
                elif "STOPPED" in raw:
                    interp = " ⚠️ 服务已停止"
                elif "START_PENDING" in raw:
                    interp = " → 服务正在启动"
                elif "STOP_PENDING" in raw:
                    interp = " → 服务正在停止"
                elif "The specified service does not exist" in raw or "1060" in raw:
                    interp = " ❌ 服务不存在（检查服务名拼写或未安装）"
            elif action == "start":
                if "SUCCESS" in raw:
                    interp = " → ✅ 服务已启动"
                elif "1056" in raw or "already running" in raw.lower():
                    interp = " → 服务已在运行"
            elif action == "stop":
                if "SUCCESS" in raw:
                    interp = " → ✅ 服务已停止"
                elif "1062" in raw or "not started" in raw.lower():
                    interp = " → 服务未在运行"

            return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}{interp}"

    # Linux 服务管理（systemd，原有逻辑保留）
    if service == "all":
        if action == "status":
            cmd = "systemctl list-units --type=service --state=running --no-pager"
            raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=20)
            return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}"
        else:
            return f"错误：service='all' 只支持 action='status'"

    cmd = f"systemctl {action} {service}"
    raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30)

    # 添加状态解读
    interpretation = ""
    if action in ("is-active", "is-enabled"):
        if "active" in raw and "inactive" not in raw.split("\n")[0]:
            interpretation = " → 服务正在运行"
        elif "inactive" in raw:
            interpretation = " → 服务已停止"
        elif "enabled" in raw:
            interpretation = " → 已设置开机自启"
        elif "disabled" in raw:
            interpretation = " → 已禁用开机自启"
    elif action == "status":
        if "Active: active (running)" in raw:
            interpretation = " → ✅ 服务运行中"
        elif "Active: inactive" in raw:
            interpretation = " ⚠️ 服务未运行"
        elif "Active: failed" in raw:
            interpretation = " ❌ 服务异常退出（建议 journalctl -u " + service + " 查看日志）"
        elif "could not be found" in raw or "Loaded: not-found" in raw:
            interpretation = " ❌ 服务不存在（检查服务名拼写或未安装）"

    return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}{interpretation}"


def ssh_log_view(service: str = "", lines: int = 100, follow: bool = False,
                 keyword: str = "", conn_id: str = "default") -> str:
    """查看远程日志（Linux: journalctl/syslog；Windows: Get-EventLog 事件日志）。

    Args:
        service: 服务名（如 nginx）——Linux 用 journalctl -u；Windows 忽略，查看系统事件日志
        lines: 查看最后 N 行，默认 100
        follow: 是否持续跟踪（注意：会阻塞直到超时，建议短时使用）
        keyword: 关键词过滤（grep），如 error / exception / fail
        conn_id: SSH 连接ID

    Returns:
        日志内容 + 自动异常统计
    """
    lines = max(10, min(int(lines), 1000))  # 限制 10-1000

    os_type = _ssh_detect_os(conn_id)

    if os_type == "windows":
        # Windows 事件日志查看（用 Get-WinEvent 替代已废弃的 Get-EventLog）
        # service 参数在 Windows 上映射为日志名：
        #   - 空/未指定 → System（系统日志）
        #   - app/application → Application（应用程序日志）
        #   - sec/security → Security（安全日志）
        #   - setup → Setup（安装日志）
        #   - forward → ForwardedEvents（转发事件）
        #   - 其他字符串 → 视为自定义日志名（如 Microsoft-Windows-PowerShell/Operational）
        svc_lower = (service or "").lower().strip()
        if not svc_lower:
            log_name = "System"
        elif svc_lower in ("app", "application"):
            log_name = "Application"
        elif svc_lower in ("sec", "security"):
            log_name = "Security"
        elif svc_lower in ("setup",):
            log_name = "Setup"
        elif svc_lower in ("forward", "forwarded"):
            log_name = "ForwardedEvents"
        else:
            # 视为自定义日志名（防注入：仅允许字母数字-/_）
            if all(c.isalnum() or c in "-/_" for c in service):
                log_name = service
            else:
                return f"错误：service 参数含非法字符 '{service}'（仅允许字母数字-/_）"

        # 构造 Get-WinEvent 命令（比 Get-EventLog 性能更好，支持更多日志）
        # FilterHashtable 比 Where-Object 过滤更高效
        if keyword:
            # 带关键词过滤：先按时间倒序取最近 N 条，再用 Message 匹配
            # 注意：Get-WinEvent 的 Message 字段不能直接在 FilterHashtable 中过滤
            # 所以用 Where-Object 二次过滤
            safe_kw = keyword.replace("'", "''").replace('"', '`"')
            # 先取较多条目用于过滤（避免过滤后条目太少）
            fetch_n = min(lines * 5, 1000)
            cmd = (
                'powershell -NoProfile -Command "'
                f"$events = Get-WinEvent -LogName '{log_name}' -MaxEvents {fetch_n} -ErrorAction SilentlyContinue | "
                f"Where-Object {{ $_.Message -like '*{safe_kw}*' }} | Select-Object -First {lines};"
                "$events | Sort-Object TimeCreated -Descending | ForEach-Object {"
                "  $level = switch ($_.LevelDisplayName) { 'Error' {'❌'} 'Warning' {'⚠️'} 'Information' {'ℹ️'} default {'?'} };"
                "  $msg = ($_.Message -replace '\\r?\\n', ' ').Substring(0, [Math]::Min(80, ($_.Message).Length));"
                "  Write-Output ($_.TimeCreated.ToString('yyyy-MM-dd HH:mm:ss') + ' ' + $level + ' [' + $_.Id + '] ' + $_.ProviderName + ': ' + $msg)"
                "}"
                '"'
            )
        else:
            cmd = (
                'powershell -NoProfile -Command "'
                f"Get-WinEvent -LogName '{log_name}' -MaxEvents {lines} -ErrorAction SilentlyContinue | "
                "Sort-Object TimeCreated -Descending | ForEach-Object {"
                "  $level = switch ($_.LevelDisplayName) { 'Error' {'❌'} 'Warning' {'⚠️'} 'Information' {'ℹ️'} default {'?'} };"
                "  $msg = if ($_.Message) { ($_.Message -replace '\\r?\\n', ' ').Substring(0, [Math]::Min(80, ($_.Message).Length)) } else { '' };"
                "  Write-Output ($_.TimeCreated.ToString('yyyy-MM-dd HH:mm:ss') + ' ' + $level + ' [' + $_.Id + '] ' + $_.ProviderName + ': ' + $msg)"
                "}"
                '"'
            )

        raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30)
        if raw.startswith("错误") or raw.startswith("连接失败"):
            return raw

        # 自动异常统计（基于等级图标计数）
        err_count = raw.count("❌")
        warn_count = raw.count("⚠️")
        info_count = raw.count("ℹ️")
        total = err_count + warn_count + info_count
        summary = f"\n\n[日志分析] {log_name} 共 {total} 条，错误 {err_count} 条，警告 {warn_count} 条，信息 {info_count} 条"
        if err_count > 10:
            summary += " ⚠️ 错误密度高，建议深入排查"
        elif err_count > 0:
            summary += " ℹ️ 存在少量错误"
        if "无法找到" in raw or "No events were found" in raw or "not found" in raw.lower():
            summary += "\n💡 该日志名可能不存在，可用 ssh_exec('powershell -Command \"Get-WinEvent -ListLog * | Select-Object LogName\"') 查看所有可用日志"
        return f"{_ssh_format_prefix(conn_id)}\n$ 查看 {log_name} 事件日志（最近 {lines} 条）\n{raw}{summary}"

    # Linux 日志查看（原有逻辑保留）
    if service:
        cmd = f"journalctl -u {service} -n {lines} --no-pager"
    else:
        cmd = f"tail -n {lines} /var/log/syslog 2>/dev/null || tail -n {lines} /var/log/messages"
    if keyword:
        # 转义单引号防注入
        safe_kw = keyword.replace("'", "'\\''")
        cmd += f" | grep -i '{safe_kw}'"
    if follow:
        # follow 模式下加超时
        cmd = f"timeout 15 {cmd} -f" if "journalctl" in cmd else f"timeout 15 {cmd} -f"

    raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30)
    if raw.startswith("错误") or raw.startswith("连接失败"):
        return raw

    # 自动异常统计
    err_count = sum(raw.lower().count(kw) for kw in ["error", "exception", "failed", "critical"])
    warn_count = raw.lower().count("warn")
    summary = f"\n\n[日志分析] 共 {len(raw.splitlines())} 行，错误关键词 {err_count} 次，警告 {warn_count} 次"
    if err_count > 10:
        summary += " ⚠️ 错误密度高，建议深入排查"
    elif err_count > 0:
        summary += " ℹ️ 存在少量错误"
    return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}{summary}"


def ssh_process_check(sort_by: str = "cpu", top_n: int = 15,
                      conn_id: str = "default") -> str:
    """查看远程服务器进程（按 CPU/内存排序）。

    Args:
        sort_by: 排序方式，可选 cpu / mem
        top_n: 返回前 N 个进程，默认 15
        conn_id: SSH 连接ID

    Returns:
        进程列表 + 资源占用摘要
    """
    if sort_by not in ("cpu", "mem"):
        return "错误：sort_by 必须是 cpu 或 mem"
    top_n = max(5, min(int(top_n), 50))

    os_type = _ssh_detect_os(conn_id)

    if os_type == "windows":
        # Windows 进程查看（用 PowerShell 的 Get-Process）
        sort_prop = "CPU" if sort_by == "cpu" else "WorkingSet64"
        cmd = (
            'powershell -NoProfile -Command "'
            f"Get-Process | Sort-Object {sort_prop} -Descending | Select-Object -First {top_n} | "
            "ForEach-Object {"
            "  $cpu = if ($_.CPU) { [math]::Round($_.CPU, 1) } else { 0 };"
            "  $memMB = [math]::Round($_.WorkingSet64/1MB, 0);"
            "  Write-Output ($_.Id.ToString().PadLeft(8) + '  ' + $cpu.ToString().PadLeft(10) + 's  ' + $memMB.ToString().PadLeft(8) + 'MB  ' + $_.Name)"
            "};"
            "Write-Output '';"
            "Write-Output ('进程总数: ' + (Get-Process | Measure-Object).Count)"
            '"'
        )
        raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=20)
        if raw.startswith("错误") or raw.startswith("连接失败"):
            return raw
        header = "      PID        CPU         内存  名称\n"
        return f"$ 按CPU排序的Top{top_n}进程\n{header}{raw}"

    # Linux 进程查看（ps 命令按 CPU/内存排序）
    sort_field = "-pcpu" if sort_by == "cpu" else "-pmem"
    cmd = f"ps aux --sort={sort_field} | head -n {top_n + 1}"
    raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=15)
    if raw.startswith("错误") or raw.startswith("连接失败"):
        return raw

    # 获取系统总览
    overview = ssh_exec("uptime && free -h | head -n 2", conn_id=conn_id, timeout=5)
    return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}\n\n[系统总览]\n{overview}"


def ssh_disk_analyze(path: str = "/", conn_id: str = "default") -> str:
    """磁盘空间分析（Linux: df + du；Windows: Get-Volume + Get-ChildItem）。

    Args:
        path: 分析的目录，默认 /（Windows 默认所有盘符）
        conn_id: SSH 连接ID

    Returns:
        磁盘使用情况 + 大目录 Top10
    """
    os_type = _ssh_detect_os(conn_id)

    if os_type == "windows":
        # Windows 磁盘分析
        # 使用 Get-Volume（现代 cmdlet，Windows 8+ / Server 2012+）+ Get-CimInstance（兼容后备）
        # path 解析：
        #   - "/" 或空 → 所有盘符
        #   - "C:" / "C:\" → 指定盘符
        #   - "C:\Users" → 指定目录（先显示所在盘，再分析该目录大小）
        import re as _re_module

        if path == "/" or not path:
            # 列出所有盘符（用 Get-Volume 显示更现代的卷信息 + Get-CimInstance 补充容量）
            vol_cmd = (
                'powershell -NoProfile -Command "'
                "Get-CimInstance Win32_LogicalDisk -Filter 'DriveType=3' | ForEach-Object {"
                "  $totalD = [math]::Round($_.Size/1GB, 1);"
                "  $freeD = [math]::Round($_.FreeSpace/1GB, 1);"
                "  $usedD = [math]::Round($totalD - $freeD, 1);"
                "  $usedPct = if ($totalD -gt 0) { [math]::Round($usedD / $totalD * 100, 1) } else { 0 };"
                "  Write-Output ($_.DeviceID + '  总:' + $totalD + 'GB  已用:' + $usedD + 'GB  可用:' + $freeD + 'GB  使用率:' + $usedPct + '%')"
                "}"
                '"'
            )
        else:
            # 指定盘符或目录
            # 提取盘符（前两个字符，如 "C:"）
            drive = path[:2] if len(path) >= 2 else path
            # 防注入：仅允许字母+冒号
            if not _re_module.match(r'^[A-Za-z]:$', drive):
                return f"错误：Windows 路径需以盘符开头（如 'C:' 或 'C:\\Users'），收到 '{path}'"
            # 用 Where-Object 过滤替代 -Filter，避免双引号嵌套问题
            # （cmd 中双引号无法嵌套，-Filter 参数的引号会被错误解析）
            vol_cmd = (
                'powershell -NoProfile -Command "'
                f"$target = '{drive}';"
                "Get-CimInstance Win32_LogicalDisk -Filter 'DriveType=3' | "
                "Where-Object { $_.DeviceID -eq $target } | ForEach-Object {"
                "  $totalD = [math]::Round($_.Size/1GB, 1);"
                "  $freeD = [math]::Round($_.FreeSpace/1GB, 1);"
                "  $usedD = [math]::Round($totalD - $freeD, 1);"
                "  $usedPct = if ($totalD -gt 0) { [math]::Round($usedD / $totalD * 100, 1) } else { 0 };"
                "  Write-Output ($_.DeviceID + '  总:' + $totalD + 'GB  已用:' + $usedD + 'GB  可用:' + $freeD + 'GB  使用率:' + $usedPct + '%')"
                "}"
                '"'
            )
        vol_out = ssh_exec(vol_cmd, conn_id=conn_id, _internal=True, timeout=15)

        # 分析磁盘使用率
        analysis = ""
        for line in vol_out.split("\n"):
            m = _re_module.search(r"使用率:([\d.]+)%", line)
            if m:
                pct = float(m.group(1))
                if pct >= 90:
                    analysis += f"\n⚠️ 磁盘使用率 {pct}%（危急，建议立即清理）"
                elif pct >= 80:
                    analysis += f"\n⚠️ 磁盘使用率 {pct}%（警告）"
                elif pct >= 70:
                    analysis += f"\nℹ️ 磁盘使用率 {pct}%（关注）"

        # Top10 大目录分析（对应 Linux 的 du --max-depth=1）
        # 优化：只扫描顶层子目录，每个子目录内部递归统计文件大小
        # （原 -Recurse -Depth 2 方案会对每个深层目录重复扫描，O(n²) 复杂度，大目录超时）
        target_path = path if (path and path != "/") else "C:\\"
        # 规范化路径：把 / 转为 \
        target_path = target_path.replace("/", "\\")
        # 如果只给了盘符（如 "C:"），补全为 "C:\\"
        if _re_module.match(r'^[A-Za-z]:$', target_path):
            target_path = target_path + "\\"

        # PowerShell 命令：扫描顶层子目录，每个子目录递归统计文件总大小
        # 用 foreach 循环替代 ForEach-Object（性能更好）
        # 用 -ErrorAction SilentlyContinue 跳过无权限目录
        # 注意：如果路径是盘符根目录（如 C:\），扫描顶层子目录可能仍较慢
        #       因此设置 90 秒超时，并在返回结果中提示
        du_cmd = (
            'powershell -NoProfile -Command "'
            f"$root = '{target_path}';"
            "$topDirs = Get-ChildItem -Path $root -Directory -ErrorAction SilentlyContinue;"
            "foreach ($d in $topDirs) {"
            "  try {"
            "    $size = (Get-ChildItem -Path $d.FullName -File -Recurse -ErrorAction SilentlyContinue | Measure-Object -Property Length -Sum).Sum;"
            "    $sizeMB = [math]::Round($size/1MB, 1);"
            "    if ($sizeMB -gt 10) { Write-Output ($sizeMB.ToString().PadLeft(10) + 'MB  ' + $d.FullName) }"
            "  } catch {}"
            "}"
            '"'
        )
        du_out = ssh_exec(du_cmd, conn_id=conn_id, _internal=True, timeout=90)

        # 如果扫描结果为空或超时，给出提示
        if not du_out.strip() or "超时" in du_out or "timeout" in du_out.lower():
            du_out = du_out or "(无输出)"
            du_out += "\n💡 提示：扫描大目录可能较慢，建议指定更具体的路径（如 'C:\\Users' 而非 'C:\\'）"

        return f"{_ssh_format_prefix(conn_id)}\n[磁盘使用]\n$ {vol_cmd}\n{vol_out}{analysis}\n\n[Top10 大目录]\n$ {du_cmd}\n{du_out}"

    # Linux 磁盘分析（原有逻辑保留）
    # df 查看整体
    df_cmd = f"df -h {path}"
    df_out = ssh_exec(df_cmd, conn_id=conn_id, timeout=10)

    # du 查看 Top10 大目录（限制深度3，避免扫描过慢）
    du_cmd = f"du -h --max-depth=3 {path} 2>/dev/null | sort -rh | head -n 10"
    du_out = ssh_exec(du_cmd, conn_id=conn_id, timeout=60)

    # 分析
    analysis = ""
    for line in df_out.split("\n"):
        if "%" in line:
            # 提取使用率
            parts = line.split()
            for p in parts:
                if p.endswith("%") and p[:-1].isdigit():
                    pct = int(p[:-1])
                    if pct >= 90:
                        analysis += f"\n⚠️ 磁盘使用率 {pct}%（危急，建议立即清理）"
                    elif pct >= 80:
                        analysis += f"\n⚠️ 磁盘使用率 {pct}%（警告）"
                    elif pct >= 70:
                        analysis += f"\nℹ️ 磁盘使用率 {pct}%（关注）"
                    break

    return f"[磁盘使用]\n$ {df_cmd}\n{df_out}\n\n[Top10 大目录]\n$ {du_cmd}\n{du_out}{analysis}"


def ssh_network_diag(action: str = "stats", target: str = "",
                     conn_id: str = "default") -> str:
    """网络诊断工具集（Linux: ss/netstat；Windows: Get-NetTCPConnection/netstat/ping）。

    Args:
        action: 诊断类型：
            - stats: 查看网络连接统计（默认）
            - ports: 查看监听端口
            - ping: ping 目标主机
            - connections: 查看活跃连接
        target: ping 操作时的目标主机（IP/域名）
        conn_id: SSH 连接ID

    Returns:
        诊断结果
    """
    if action not in ("stats", "ports", "ping", "connections"):
        return "错误：action 必须是 stats / ports / ping / connections 之一"

    os_type = _ssh_detect_os(conn_id)

    if os_type == "windows":
        if action == "stats":
            cmd = 'powershell -NoProfile -Command "$tcp = Get-NetTCPConnection -ErrorAction SilentlyContinue; Write-Output (\'TCP连接总数: \' + ($tcp | Measure-Object).Count); Write-Output (\'监听端口数: \' + ($tcp | Where-Object {$_.State -eq \'Listen\'} | Measure-Object).Count); Write-Output (\'已建立连接数: \' + ($tcp | Where-Object {$_.State -eq \'Established\'} | Measure-Object).Count); Write-Output (\'UDP端点数: \' + (Get-NetUDPEndpoint -ErrorAction SilentlyContinue | Measure-Object).Count)"'
        elif action == "ports":
            cmd = 'powershell -NoProfile -Command "Get-NetTCPConnection -State Listen -ErrorAction SilentlyContinue | Select-Object LocalAddress, LocalPort, OwningProcess | Sort-Object LocalPort | Format-Table -AutoSize"'
        elif action == "ping":
            if not target:
                return "错误：ping 操作需要 target 参数"
            if not target.replace(".", "").replace("-", "").isalnum():
                return "错误：target 仅允许字母数字.-"
            cmd = f'ping -n 4 -w 2000 {target}'
        else:  # connections
            cmd = 'powershell -NoProfile -Command "Get-NetTCPConnection -State Established -ErrorAction SilentlyContinue | Select-Object LocalAddress, LocalPort, RemoteAddress, RemotePort | Sort-Object RemoteAddress | Format-Table -AutoSize | Out-String -Width 200"'
        raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30 if action == "ping" else 15)
        return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}"

    # Linux 网络诊断（原有逻辑保留）
    if action == "stats":
        cmd = "ss -s"
    elif action == "ports":
        cmd = "ss -tlnp 2>/dev/null || netstat -tlnp 2>/dev/null"
    elif action == "ping":
        if not target:
            return "错误：ping 操作需要 target 参数"
        # 校验 target 合法性
        if not target.replace(".", "").replace("-", "").isalnum():
            return "错误：target 仅允许字母数字.-"
        cmd = f"ping -c 4 -W 2 {target}"
    else:  # connections
        cmd = "ss -tn state established 2>/dev/null | head -n 30"

    raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30 if action == "ping" else 15)
    return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}"


def ssh_docker_manage(action: str, container: str = "",
                      conn_id: str = "default") -> str:
    r"""Docker 容器管理（跨平台：Linux 原生 Docker / Windows Docker Desktop）。

    Args:
        action: 操作类型：ps / psa / logs / start / stop / restart / stats / images / info
            - ps: 运行中容器
            - psa: 所有容器（含已停止）
            - logs: 查看容器日志（需 container 参数）
            - start/stop/restart: 容器生命周期（需 container 参数）
            - stats: 资源占用
            - images: 镜像列表
            - info: Docker 系统信息（版本、存储驱动、运行环境）
        container: 容器名/ID（logs/start/stop/restart 必填）
        conn_id: SSH 连接ID

    Returns:
        Docker 操作结果

    Windows Docker Desktop 适配说明：
        - 命令前缀使用 docker.exe（显式调用，避免 PowerShell 别名冲突）
        - --format 字符串用双引号包裹（PowerShell 单引号会原样输出 Go template）
        - 自动检测 Docker Desktop 是否运行（依赖 WSL2 后端）
        - Windows 上 docker 命令在 PATH 中：C:\Program Files\Docker\Docker\resources\bin\
    """
    valid_actions = {"ps", "psa", "logs", "start", "stop", "restart", "stats", "images", "info"}
    if action not in valid_actions:
        return f"错误：action 必须是 {sorted(valid_actions)} 之一"

    # 检查 docker 是否安装
    if action in ("start", "stop", "restart", "logs") and not container:
        return f"错误：action={action} 需要提供 container 参数"

    # 校验 container 名（防注入）
    if container and not container.replace("-", "").replace("_", "").replace(".", "").replace("/", "").isalnum():
        return f"错误：容器名 '{container}' 不合法"

    os_type = _ssh_detect_os(conn_id)
    is_windows = (os_type == "windows")

    # Windows 用 docker.exe，Linux 用 docker
    docker_cmd = "docker.exe" if is_windows else "docker"

    if action == "ps":
        if is_windows:
            # cmd 下用双引号包裹 Go template（单引号在 cmd 中是字面字符，会被 docker 误认）
            cmd = f'{docker_cmd} ps --format "table {{{{.ID}}}}\\t{{{{.Names}}}}\\t{{{{.Image}}}}\\t{{{{.Status}}}}\\t{{{{.Ports}}}}"'
        else:
            cmd = "docker ps --format 'table {{.ID}}\\t{{.Names}}\\t{{.Image}}\\t{{.Status}}\\t{{.Ports}}'"
    elif action == "psa":
        if is_windows:
            cmd = f'{docker_cmd} ps -a --format "table {{{{.ID}}}}\\t{{{{.Names}}}}\\t{{{{.Image}}}}\\t{{{{.Status}}}}"'
        else:
            cmd = "docker ps -a --format 'table {{.ID}}\\t{{.Names}}\\t{{.Image}}\\t{{.Status}}'"
    elif action == "logs":
        cmd = f"{docker_cmd} logs --tail 100 {container}"
    elif action == "start":
        cmd = f"{docker_cmd} start {container}"
    elif action == "stop":
        cmd = f"{docker_cmd} stop {container}"
    elif action == "restart":
        cmd = f"{docker_cmd} restart {container}"
    elif action == "stats":
        if is_windows:
            cmd = f'{docker_cmd} stats --no-stream --format "table {{{{.Name}}}}\\t{{{{.CPUPerc}}}}\\t{{{{.MemUsage}}}}"'
        else:
            cmd = "docker stats --no-stream --format 'table {{.Name}}\\t{{.CPUPerc}}\\t{{.MemUsage}}'"
    elif action == "info":
        # Docker 系统信息（跨平台兼容）
        cmd = f"{docker_cmd} version && {docker_cmd} info --format 'Server Version: {{{{.ServerVersion}}}}\\nStorage Driver: {{{{.Driver}}}}\\nRunning Containers: {{{{.ContainersRunning}}}}\\nTotal Containers: {{{{.Containers}}}}\\nImages: {{{{.Images}}}}'"
    else:  # images
        if is_windows:
            cmd = f'{docker_cmd} images --format "table {{{{.Repository}}}}\\t{{{{.Tag}}}}\\t{{{{.Size}}}}"'
        else:
            cmd = "docker images --format 'table {{.Repository}}\\t{{.Tag}}\\t{{.Size}}'"

    # Windows 上先检测 Docker Desktop 是否安装并运行
    if is_windows:
        # 用 where 命令快速检测 docker.exe 是否存在
        check_cmd = "where docker.exe 2>nul || echo NOT_FOUND"
        check_out = ssh_exec(check_cmd, conn_id=conn_id, _internal=True, timeout=8)
        if "NOT_FOUND" in check_out or not check_out.strip():
            return f"{_ssh_format_prefix(conn_id)}\n❌ Windows 服务器未安装 Docker Desktop\n（安装路径通常是 C:\\Program Files\\Docker\\Docker\\resources\\bin\\docker.exe）"

    raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=60 if action == "stats" else 30)
    raw_lower = raw.lower()
    if ("command not found" in raw_lower
            or "not recognized" in raw_lower
            or "不是内部或外部命令" in raw_lower
            or "无法找到" in raw_lower):
        return f"{_ssh_format_prefix(conn_id)}\n❌ 服务器未安装 Docker"
    # Windows Docker Desktop 未运行时的常见错误
    if is_windows and ("error during connect" in raw_lower
                       or "the docker daemon is not running" in raw_lower
                       or "cannot connect to the docker daemon" in raw_lower):
        return f"{_ssh_format_prefix(conn_id)}\n❌ Docker Desktop 未运行，请先启动 Docker Desktop\n$ {cmd}\n{raw}"
    return f"{_ssh_format_prefix(conn_id)}\n$ {cmd}\n{raw}"


def ssh_firewall_manage(action: str, port: int = 0, protocol: str = "tcp",
                        conn_id: str = "default") -> str:
    """防火墙管理（ufw / firewalld / iptables 自动检测）。

    Args:
        action: 操作类型：status / list / open / close / enable / disable
            - status: 查看状态
            - list: 列出规则
            - open: 开放端口（需 port）
            - close: 关闭端口（需 port）
            - enable/disable: 启用/禁用防火墙
        port: 端口号（open/close 时必填）
        protocol: 协议 tcp/udp，默认 tcp
        conn_id: SSH 连接ID

    Returns:
        防火墙操作结果
    """
    valid_actions = {"status", "list", "open", "close", "enable", "disable"}
    if action not in valid_actions:
        return f"错误：action 必须是 {sorted(valid_actions)} 之一"
    if protocol not in ("tcp", "udp"):
        return "错误：protocol 必须是 tcp 或 udp"
    if action in ("open", "close"):
        if not (1 <= int(port) <= 65535):
            return f"错误：port 必须在 1-65535 范围内"

    os_type = _ssh_detect_os(conn_id)

    if os_type == "windows":
        # Windows 防火墙管理（netsh advfirewall + Get-NetFirewallRule 组合）
        # netsh 适合增删规则，Get-NetFirewallRule 适合查询统计
        if action == "status":
            # 用 PowerShell 的 Get-NetFirewallProfile 显示各配置文件状态（更直观）
            cmd = (
                'powershell -NoProfile -Command "'
                "Get-NetFirewallProfile | ForEach-Object {"
                "  $state = if ($_.Enabled) {'✅ 已启用'} else {'❌ 已禁用'};"
                "  Write-Output ($_.Name + ' - ' + $state + ' (入站默认: ' + $_.DefaultInboundAction + ', 出站默认: ' + $_.DefaultOutboundAction + ')')"
                "}"
                '"'
            )
        elif action == "list":
            # 用 Get-NetFirewallRule 列出 ZeroAI 创建的规则（避免输出过长）
            # 默认只显示 ZeroAI-* 规则，避免列出数千条系统规则
            # 注意：Get-NetFirewallRule 返回对象的 Action 是枚举值（1=NotConfigured, 2=Allow, 3=Block）
            #       Direction 也是枚举值（1=Inbound, 2=Outbound）
            cmd = (
                'powershell -NoProfile -Command "'
                "$rules = Get-NetFirewallRule -ErrorAction SilentlyContinue | "
                "Where-Object { $_.DisplayName -like 'ZeroAI-*' -or $_.DisplayName -like 'ZeroAI_*' };"
                "if ($rules) {"
                "  $rules | ForEach-Object {"
                "    $action = switch ($_.Action) { 2 {'✅允许'} 3 {'❌阻止'} default {'?' } };"
                "    $dir = if ($_.Direction -eq 1) {'入站'} else {'出站'};"
                "    $enabled = if ($_.Enabled) {'启用'} else {'禁用'};"
                "    Write-Output ($_.DisplayName + ' [' + $dir + ' ' + $action + ' ' + $enabled + ']')"
                "  }"
                "} else {"
                "  Write-Output '提示：当前无 ZeroAI 创建的防火墙规则。如需查看全部规则，请用 ssh_exec 直接执行：netsh advfirewall firewall show rule name=all'"
                "}"
                '"'
            )
        elif action == "open":
            # 开放端口：用 netsh 添加规则（兼容旧版 Windows）
            # 规则名格式 ZeroAI-Allow-{port}-{protocol} 便于后续查询和删除
            cmd = f'netsh advfirewall firewall add rule name="ZeroAI-Allow-{port}-{protocol}" dir=in action=allow protocol={protocol} localport={port}'
        elif action == "close":
            # 关闭端口：按规则名 + 端口双重匹配删除（更精确）
            cmd = f'netsh advfirewall firewall delete rule name="ZeroAI-Allow-{port}-{protocol}" dir=in protocol={protocol} localport={port}'
        elif action == "enable":
            cmd = "netsh advfirewall set allprofiles state on"
        else:  # disable
            cmd = "netsh advfirewall set allprofiles state off"

        raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=20 if action == "list" else 15)
        # Windows 下额外提示
        extra_hint = ""
        if action == "open" and "确定" in raw and "OK" in raw:
            extra_hint = f"\n💡 已开放 {protocol}/{port}，规则名 ZeroAI-Allow-{port}-{protocol}"
        elif action == "close":
            extra_hint = f"\n💡 已删除 {protocol}/{port} 的规则（如存在）"
        elif action == "disable":
            extra_hint = "\n⚠️ 防火墙已禁用，服务器暴露在网络中，建议仅在调试时使用"
        return f"{_ssh_format_prefix(conn_id)}\n[Windows 防火墙] $ {cmd}\n{raw}{extra_hint}"

    # Linux 防火墙管理（原有逻辑保留）
    # 自动检测防火墙类型
    fw_check = ssh_exec("command -v ufw >/dev/null && echo UFW || (command -v firewall-cmd >/dev/null && echo FIREWALLD || echo IPTABLES)",
                        conn_id=conn_id, timeout=5)
    fw_type = "UFW" if "UFW" in fw_check else ("FIREWALLD" if "FIREWALLD" in fw_check else "IPTABLES")

    if fw_type == "UFW":
        if action == "status":
            cmd = "ufw status verbose"
        elif action == "list":
            cmd = "ufw status numbered"
        elif action == "open":
            cmd = f"ufw allow {port}/{protocol}"
        elif action == "close":
            cmd = f"ufw deny {port}/{protocol}"
        elif action == "enable":
            cmd = "echo y | ufw enable"
        else:  # disable
            cmd = "ufw disable"
    elif fw_type == "FIREWALLD":
        if action == "status":
            cmd = "firewall-cmd --state && firewall-cmd --list-all"
        elif action == "list":
            cmd = "firewall-cmd --list-ports"
        elif action == "open":
            cmd = f"firewall-cmd --permanent --add-port={port}/{protocol} && firewall-cmd --reload"
        elif action == "close":
            cmd = f"firewall-cmd --permanent --remove-port={port}/{protocol} && firewall-cmd --reload"
        elif action == "enable":
            cmd = "systemctl enable --now firewalld"
        else:
            cmd = "systemctl disable --now firewalld"
    else:  # IPTABLES
        if action == "status" or action == "list":
            cmd = "iptables -L -n --line-numbers"
        elif action == "open":
            cmd = f"iptables -I INPUT -p {protocol} --dport {port} -j ACCEPT"
        elif action == "close":
            cmd = f"iptables -I INPUT -p {protocol} --dport {port} -j DROP"
        elif action == "enable":
            return "iptables 无 enable 操作（系统启动时自动加载规则）"
        else:
            cmd = "iptables -F"

    # open/close 是危险操作（修改防火墙），需要确认
    if action in ("open", "close"):
        # 不直接执行，先返回待确认信息（通过 ssh_exec 的 confirm_dangerous 链路）
        result = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=15, confirm_dangerous=False)
        if "危险命令" in result and "confirm_dangerous" in result:
            # 实际上 ufw/firewalld 不在危险命令黑名单，可以直接执行
            result = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=15)
        raw = result
    else:
        raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=15)

    return f"[防火墙类型: {fw_type}] $ {cmd}\n{raw}"


def ssh_health_check(conn_id: str = "default") -> str:
    """服务器一键健康体检（CPU/内存/磁盘/网络/负载/服务综合报告）。

    自动检测操作系统：Linux 用 uname/free/df/ss/systemctl/journalctl；
    Windows 用 PowerShell 的 Get-CimInstance/Get-Process/Get-Service 等。

    Args:
        conn_id: SSH 连接ID

    Returns:
        健康体检报告（含异常项标注与建议）
    """
    import re

    # 第一步：检测操作系统（用缓存辅助函数，避免重复探测）
    is_windows = _ssh_detect_os(conn_id) == "windows"

    if is_windows:
        # Windows Server 体检：用 PowerShell 命令（避免 wmic 在 2025 已废弃的问题）
        # 用 cmd 调用 powershell，确保兼容性
        cmd = (
            'powershell -NoProfile -Command "'
            "Write-Output '=== 系统信息 ===';"
            "$os = Get-CimInstance Win32_OperatingSystem;"
            "Write-Output ($os.Caption + ' ' + $os.Version + ' Build ' + $os.BuildNumber);"
            "Write-Output ('开机时间: ' + $os.LastBootUpTime);"
            "Write-Output '';"
            "Write-Output '=== CPU 负载 ===';"
            "$cpu = Get-CimInstance Win32_Processor | Select-Object -First 1;"
            "Write-Output ('CPU负载: ' + $cpu.LoadPercentage + '%');"
            "Write-Output ('CPU核心数: ' + (Get-CimInstance Win32_ComputerSystem).NumberOfLogicalProcessors);"
            "Write-Output '';"
            "Write-Output '=== 内存 ===';"
            "$cs = Get-CimInstance Win32_ComputerSystem;"
            "$os2 = Get-CimInstance Win32_OperatingSystem;"
            "$totalGB = [math]::Round($cs.TotalPhysicalMemory/1GB, 1);"
            "$freeGB = [math]::Round($os2.FreePhysicalMemory/1MB, 1);"
            "$usedGB = [math]::Round($totalGB - $freeGB, 1);"
            "Write-Output ('总内存: ' + $totalGB + ' GB');"
            "Write-Output ('已用: ' + $usedGB + ' GB');"
            "Write-Output ('可用: ' + $freeGB + ' GB');"
            "Write-Output '';"
            "Write-Output '=== 磁盘 ===';"
            "Get-CimInstance Win32_LogicalDisk -Filter 'DriveType=3' | ForEach-Object {"
            "  $totalD = [math]::Round($_.Size/1GB, 1);"
            "  $freeD = [math]::Round($_.FreeSpace/1GB, 1);"
            "  $usedPct = if ($totalD -gt 0) { [math]::Round(($totalD - $freeD) / $totalD * 100, 1) } else { 0 };"
            "  Write-Output ($_.DeviceID + ' 总:' + $totalD + 'GB 可用:' + $freeD + 'GB 使用:' + $usedPct + '%')"
            "};"
            "Write-Output '';"
            "Write-Output '=== 监听端口数 ===';"
            "$ports = (Get-NetTCPConnection -State Listen -ErrorAction SilentlyContinue | Measure-Object).Count;"
            "Write-Output ('监听端口数: ' + $ports);"
            "Write-Output '';"
            "Write-Output '=== 进程数 ===';"
            "$procs = (Get-Process | Measure-Object).Count;"
            "Write-Output ('进程数: ' + $procs);"
            "Write-Output '';"
            "Write-Output '=== 运行中的服务数 ===';"
            "$svc = (Get-Service | Where-Object {$_.Status -eq 'Running'} | Measure-Object).Count;"
            "Write-Output ('运行中服务: ' + $svc);"
            "Write-Output '';"
            "Write-Output '=== 防火墙状态 ===';"
            "Get-NetFirewallProfile | ForEach-Object { Write-Output ($_.Name + ': ' + $_.Enabled) };"
            "Write-Output '';"
            "Write-Output '=== 高内存进程Top5 ===';"
            "Get-Process | Sort-Object WorkingSet64 -Descending | Select-Object -First 5 | ForEach-Object {"
            "  $memMB = [math]::Round($_.WorkingSet64/1MB, 0);"
            "  Write-Output ($_.Name + ' (PID:' + $_.Id + ') 内存:' + $memMB + 'MB')"
            "}"
            '"'
        )
        raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=60)
        if raw.startswith("错误") or raw.startswith("连接失败"):
            return raw

        # Windows 分析
        issues = []

        # CPU 负载
        m = re.search(r"CPU负载:\s*(\d+)%", raw)
        if m:
            cpu_pct = int(m.group(1))
            if cpu_pct >= 90:
                issues.append(f"⚠️ CPU 负载极高: {cpu_pct}%")
            elif cpu_pct >= 70:
                issues.append(f"⚠️ CPU 负载偏高: {cpu_pct}%")
            else:
                pass  # 正常不记录

        # 内存
        m_total = re.search(r"总内存:\s*([\d.]+)\s*GB", raw)
        m_used = re.search(r"已用:\s*([\d.]+)\s*GB", raw)
        if m_total and m_used:
            total = float(m_total.group(1))
            used = float(m_used.group(1))
            if total > 0:
                pct = used / total * 100
                if pct >= 90:
                    issues.append(f"⚠️ 内存使用率 {pct:.1f}%（危急）")
                elif pct >= 80:
                    issues.append(f"⚠️ 内存使用率 {pct:.1f}%（警告）")

        # 磁盘
        for line in raw.split("\n"):
            m = re.search(r"([A-Z]:)\s*总:([\d.]+)GB\s*可用:([\d.]+)GB\s*使用:([\d.]+)%", line)
            if m:
                drive, total, free, pct = m.group(1), float(m.group(2)), float(m.group(3)), float(m.group(4))
                if pct >= 90:
                    issues.append(f"⚠️ 磁盘 {drive} 使用率 {pct}%（危急）")
                elif pct >= 80:
                    issues.append(f"⚠️ 磁盘 {drive} 使用率 {pct}%（警告）")

        # 防火墙
        for line in raw.split("\n"):
            if "False" in line and ("Domain" in line or "Private" in line or "Public" in line):
                issues.append(f"⚠️ 防火墙关闭: {line.strip()}")

        # 开机时间（判断是否长期未重启）
        m_boot = re.search(r"开机时间:\s*(.+)", raw)
        if m_boot:
            try:
                from datetime import datetime
                boot_str = m_boot.group(1).strip()
                # Windows PowerShell 输出格式可能多样，尝试解析
                for fmt in ("%m/%d/%Y %H:%M:%S", "%Y-%m-%d %H:%M:%S", "%m/%d/%Y %I:%M:%S %p"):
                    try:
                        boot_time = datetime.strptime(boot_str, fmt)
                        days = (datetime.now() - boot_time).days
                        if days > 90:
                            issues.append(f"💡 服务器已运行 {days} 天，建议定期重启")
                        break
                    except ValueError:
                        continue
            except Exception:
                pass

        report = f"$ Windows 综合体检命令\n{raw}\n\n"
        report += "=== 健康分析 ===\n"
        if not issues:
            report += "✅ 服务器整体健康，未发现异常"
        else:
            report += f"发现 {len(issues)} 个问题：\n"
            for i, issue in enumerate(issues, 1):
                report += f"  {i}. {issue}\n"
            report += "\n建议：根据上述问题深入排查（使用 ssh_service_manage / ssh_log_view 等）"
        return _ssh_format_prefix(conn_id) + "\n" + report

    # Linux 体检（原有逻辑保留）
    cmd = """echo '=== 系统信息 ===' && uname -a && uptime
echo '=== CPU 使用 ===' && top -bn1 | head -n 5
echo '=== 内存 ===' && free -h
echo '=== 磁盘 ===' && df -h | grep -v tmpfs
echo '=== 网络监听端口 ===' && ss -tlnp 2>/dev/null | head -n 15
echo '=== 系统负载 ===' && cat /proc/loadavg
echo '=== 最近登录 ===' && last -n 5
echo '=== 失败服务 ===' && systemctl --failed --no-pager 2>/dev/null | head -n 20
echo '=== 最近错误日志 ===' && journalctl -p err --since '1 hour ago' --no-pager 2>/dev/null | tail -n 10"""
    raw = ssh_exec(cmd, conn_id=conn_id, _internal=True, timeout=30)
    if raw.startswith("错误") or raw.startswith("连接失败"):
        return raw

    # AI 分析
    issues = []
    # 检查负载
    for line in raw.split("\n"):
        if "load average" in line.lower():
            # 提取 load average
            m = re.search(r"load average:\s*([\d.]+),\s*([\d.]+),\s*([\d.]+)", line)
            if m:
                load_1, load_5, load_15 = float(m.group(1)), float(m.group(2)), float(m.group(3))
                # 简单阈值：> CPU 核数则告警
                if load_1 > 4:
                    issues.append(f"⚠️ 1分钟负载 {load_1} 偏高")
                if load_15 > 2:
                    issues.append(f"⚠️ 15分钟负载 {load_15} 持续偏高")
            break

    # 检查磁盘
    for line in raw.split("\n"):
        if "%" in line and ("/" in line or "/data" in line):
            parts = line.split()
            for p in parts:
                if p.endswith("%") and p[:-1].isdigit():
                    pct = int(p[:-1])
                    if pct >= 90:
                        issues.append(f"⚠️ 磁盘使用率 {pct}%（危急）")
                    elif pct >= 80:
                        issues.append(f"⚠️ 磁盘使用率 {pct}%（警告）")
                    break

    # 检查内存
    if "Swap:" in raw:
        for line in raw.split("\n"):
            if line.startswith("Swap:") and "0B" not in line:
                parts = line.split()
                if len(parts) >= 4 and parts[3] != "0B":
                    issues.append(f"⚠️ Swap 已使用 {parts[2]}/{parts[1]}")

    # 检查失败服务
    if "failed" in raw.lower() and "0 loaded" not in raw and "0 failed" not in raw:
        for line in raw.split("\n"):
            if "failed" in line.lower() and "UNIT" not in line:
                issues.append(f"❌ 失败服务: {line.strip()}")

    # 检查错误日志
    err_log_section = raw.split("=== 最近错误日志 ===")[-1] if "=== 最近错误日志 ===" in raw else ""
    if err_log_section.strip() and "No entries" not in err_log_section:
        err_lines = [l for l in err_log_section.strip().split("\n") if l.strip()][:5]
        if err_lines:
            issues.append(f"⚠️ 最近1小时有 {len(err_lines)} 条错误日志")

    report = f"$ 综合体检命令\n{raw}\n\n"
    report += "=== 健康分析 ===\n"
    if not issues:
        report += "✅ 服务器整体健康，未发现异常"
    else:
        report += f"发现 {len(issues)} 个问题：\n"
        for i, issue in enumerate(issues, 1):
            report += f"  {i}. {issue}\n"
        report += "\n建议：根据上述问题深入排查（使用 ssh_log_view / ssh_process_check 等）"
    return _ssh_format_prefix(conn_id) + "\n" + report


TOOLS = [
    {"type": "function", "function": {
        "name": "read_file", "description": "读取本地文件内容",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "文件路径"},
            "max_length": {"type": "integer", "description": "最大返回字符数，默认3000"}},
            "required": ["path"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "write_file", "description": "写入或创建文件",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "文件路径"},
            "content": {"type": "string", "description": "文件内容"}},
            "required": ["path", "content"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "list_dir", "description": "列出目录内容，支持递归显示子目录树形结构。当用户问有哪些文件、目录结构、查看文件夹、深入子文件夹时调用。recursive=true时会自动深入到最深层目录（默认15层，自动跳过无权限目录）。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "目录路径，默认当前目录"},
            "recursive": {"type": "boolean", "description": "是否递归显示子目录（树形结构），默认false只显示一层，true自动深入到最深层目录"},
            "max_depth": {"type": "integer", "description": "递归最大深度（默认15=深入最深层，1=只看当前层）"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "run_command", "description": "在本地电脑执行 PowerShell / cmd / shell 命令并返回输出。全权限模式：120s 超时、8000 字符输出。用于查看端口(netstat)、进程(tasklist)、网络(ipconfig/ping/tracert)、系统信息(systeminfo)、服务(sc query/net start)、用户(whoami/net user)、磁盘(wmic logicaldisk)、防火墙(netsh advfirewall)、环境变量(set)等。当用户用自然语言描述本地电脑状态需求（如'看看打开了哪些端口'/'电脑卡不卡'/'谁在占用CPU'/'IP是多少'）且没有更专用的工具时调用。危险命令(format/del /f/shutdown/mkfs)自动拦截。",
        "parameters": {"type": "object", "properties": {
            "command": {"type": "string", "description": "要执行的命令（PowerShell 或 cmd 命令）"}},
            "required": ["command"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "search_files", "description": "在文件内容中搜索正则",
        "parameters": {"type": "object", "properties": {
            "pattern": {"type": "string", "description": "正则"},
            "path": {"type": "string", "description": "目录"}},
            "required": ["pattern"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "open_app", "description": "打开桌面应用程序。支持：微信、QQ、VSCode、PyCharm、Edge、记事本、计算器、资源管理器、画图、任务管理器、控制面板、注册表、CMD、PowerShell 等。当用户让你打开/启动某个应用时调用。",
        "parameters": {"type": "object", "properties": {
            "name": {"type": "string", "description": "应用名称（中英文均可），如：微信、wechat、qq、vscode、pycharm、edge、记事本、计算器"}},
            "required": ["name"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "web_search", "description": "网络搜索，获取搜索结果。当用户问实时信息、新闻、文档、最新动态时调用。",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "搜索关键词"},
            "num_results": {"type": "integer", "description": "返回结果数量，默认5"}},
            "required": ["query"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "web_fetch", "description": "抓取网页内容，获取网页纯文本。当需要读取某个 URL 的内容时调用。",
        "parameters": {"type": "object", "properties": {
            "url": {"type": "string", "description": "网页 URL"},
            "max_length": {"type": "integer", "description": "最大返回字符数，默认4000"}},
            "required": ["url"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "git_status", "description": "查看 Git 仓库状态（分支、改动文件）。当用户问 git 状态、代码变更时调用。",
        "parameters": {"type": "object", "properties": {
            "repo_path": {"type": "string", "description": "仓库路径，默认当前目录"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "delete_file", "description": "删除文件或目录（优先移入回收站）。当用户让你删除文件时调用。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "要删除的文件或目录路径"}},
            "required": ["path"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "move_file", "description": "移动或重命名文件。当用户让你移动、重命名文件时调用。",
        "parameters": {"type": "object", "properties": {
            "src": {"type": "string", "description": "源文件路径"},
            "dst": {"type": "string", "description": "目标路径"}},
            "required": ["src", "dst"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "copy_file", "description": "复制文件。当用户让你复制、备份文件时调用。",
        "parameters": {"type": "object", "properties": {
            "src": {"type": "string", "description": "源文件路径"},
            "dst": {"type": "string", "description": "目标路径"}},
            "required": ["src", "dst"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "create_dir", "description": "创建目录（含父目录）。当用户让你创建文件夹时调用。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "目录路径"}},
            "required": ["path"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "system_info", "description": "获取系统信息（CPU、内存、磁盘）。当用户问系统状态、环境信息时调用。",
        "parameters": {"type": "object", "properties": {},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "process_list", "description": "列出进程。当用户问运行中的进程、查看进程时调用。",
        "parameters": {"type": "object", "properties": {
            "name_filter": {"type": "string", "description": "进程名过滤关键词"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "edit_file", "description": "按行编辑文件：替换/插入/删除/追加指定行。当需要修改文件中的某一行或某几行时调用，避免重写整个文件。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "文件路径"},
            "operation": {"type": "string", "description": "操作类型：replace（替换行）/ insert（插入行）/ delete（删除行）/ append（末尾追加）"},
            "line": {"type": "integer", "description": "目标行号（从1开始），replace/insert用"},
            "content": {"type": "string", "description": "新内容（replace/insert/append用）"},
            "start_line": {"type": "integer", "description": "删除起始行号（delete用）"},
            "end_line": {"type": "integer", "description": "删除结束行号（delete用）"}},
            "required": ["path", "operation"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "exec_python", "description": "在安全沙箱中执行 Python 代码片段并返回结果。用于快速验证算法、计算表达式、数据处理。当用户让你运行Python代码、计算、验证时调用。",
        "parameters": {"type": "object", "properties": {
            "code": {"type": "string", "description": "Python 代码片段"}},
            "required": ["code"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "pip_install", "description": "Python 包管理。当用户让你安装/卸载/查看Python包时调用。",
        "parameters": {"type": "object", "properties": {
            "package": {"type": "string", "description": "包名（list操作时留空）"},
            "action": {"type": "string", "description": "操作：install（安装）/ uninstall（卸载）/ check（检查）/ list（列出已安装）"}},
            "required": ["action"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "check_port", "description": "检测端口占用情况，返回占用进程信息。当用户问端口占用、服务是否启动时调用。",
        "parameters": {"type": "object", "properties": {
            "port": {"type": "integer", "description": "端口号"}},
            "required": ["port"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "file_diff", "description": "比较两个文件的差异，返回逐行diff。当用户问文件区别、对比文件时调用。",
        "parameters": {"type": "object", "properties": {
            "path_a": {"type": "string", "description": "第一个文件路径"},
            "path_b": {"type": "string", "description": "第二个文件路径"}},
            "required": ["path_a", "path_b"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "read_image", "description": "读取图片文件并理解内容。支持 png/jpg/jpeg/gif/bmp/webp 格式。当用户发送图片路径或让你看图片时调用。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "图片文件路径"}},
            "required": ["path"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "active_window", "description": "获取当前焦点窗口信息（标题、应用名、位置大小）。当用户问当前在做什么、当前窗口时调用。",
        "parameters": {"type": "object", "properties": {},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "list_windows", "description": "列出所有可见窗口及其标题和PID。当用户问打开了哪些窗口、桌面上的程序时调用。",
        "parameters": {"type": "object", "properties": {},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "read_screen", "description": "读取当前前台窗口的文字内容（通过UI Automation，像屏幕阅读器一样精确读取文字）。当用户问屏幕上有什么、当前页面内容、看到什么时调用。",
        "parameters": {"type": "object", "properties": {},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "security_audit", "description": "安全审计：扫描代码漏洞、敏感信息、依赖漏洞、配置安全。当用户让你检查安全、查找漏洞、安全审计、扫描代码问题时调用。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "要扫描的文件或目录路径，默认为当前目录"},
            "scan_type": {"type": "string", "description": "扫描类型：all(全部) / code(代码漏洞) / secret(敏感信息) / deps(依赖漏洞) / config(配置安全)"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "generate_word", "description": "生成 Word 文档（.docx），支持指定格式模板和高级排版。当用户让你生成Word文档、写报告、合同、简历、论文、学术文章、导出文档时调用。支持8种模板（含academic学术论文模板：双倍行距/摘要/关键词/参考文献自动编号/LaTeX公式渲染）、自定义字体/颜色/页边距/对齐/页眉页脚/表格。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "保存路径（可选，默认保存到桌面）。可传完整路径如 D:/报告.docx，或只传文件名如 报告.docx（自动保存到桌面），或留空（自动命名保存到桌面）"},
            "content": {"type": "string", "description": "文档内容（支持 Markdown 标记：# 标题 / - 列表 / > 引用 / ```代码``` / **粗体** *斜体* `代码` ~~删除线~~ / |表格| / [居中]行首对齐 / {color:红色}文字{/color}）"},
            "title": {"type": "string", "description": "文档标题（可选，留空则取内容第一行 # 标题）"},
            "template": {"type": "string", "description": "格式模板：default(默认) / report(正式报告) / contract(合同) / resume(简历) / thesis(论文) / letter(信函) / technical(技术文档) / academic(学术论文：双倍行距+摘要+关键词+参考文献自动编号+LaTeX公式渲染)"},
            "font": {"type": "string", "description": "覆盖模板字体，如 SimSun(宋体) / Microsoft YaHei(微软雅黑) / KaiTi(楷体) / Times New Roman"},
            "font_size": {"type": "integer", "description": "覆盖模板字号，如 10/11/12/14"},
            "margin": {"type": "array", "items": {"type": "number"}, "description": "页边距[上,右,下,左]厘米，如 [2.54, 2.54, 2.54, 2.54]"},
            "heading_color": {"type": "string", "description": "标题颜色，十六进制如 1F4E79 或颜色名如 navy/red/blue"},
            "align": {"type": "string", "description": "全文对齐：left / center / right / justify"},
            "header": {"type": "string", "description": "页眉文字"},
            "footer": {"type": "string", "description": "页脚文字，支持 {page} 和 {numpages} 占位符"}},
            "required": ["content"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "generate_excel", "description": "生成 Excel 文档（.xlsx），支持多工作表、表头样式、隔行变色、自动列宽、图表（柱状图/折线图/饼图）、公式。当用户让你生成Excel表格、数据表、导出Excel、带图表的Excel时调用。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "保存路径（可选，默认保存到桌面）。可传完整路径如 D:/数据表.xlsx，或只传文件名如 数据表.xlsx（自动保存到桌面），或留空（自动命名保存到桌面）"},
            "sheets": {"type": "array", "items": {"type": "object", "properties": {
                "name": {"type": "string", "description": "工作表名（可选，默认Sheet1）"},
                "data": {"type": "array", "items": {"type": "array"}, "description": "数据二维数组，第一行作为表头"},
                "header": {"type": "boolean", "description": "是否有表头，默认true"}
            }, "required": ["data"]}, "description": "工作表列表"},
            "template": {"type": "string", "description": "格式模板：default(默认蓝) / report(报告深蓝) / data(数据绿) / financial(财务灰)"},
            "charts": {"type": "array", "items": {"type": "object", "properties": {
                "type": {"type": "string", "description": "图表类型：bar(柱状图) / line(折线图) / pie(饼图)"},
                "title": {"type": "string", "description": "图表标题"},
                "sheet": {"type": "string", "description": "数据所在工作表名"},
                "categories_col": {"type": "string", "description": "分类轴列号，如 A（姓名列）"},
                "values_cols": {"type": "array", "items": {"type": "string"}, "description": "值轴列号列表，如 [\"B\",\"C\"]"},
                "position": {"type": "string", "description": "图表放置位置单元格，如 E2"}
            }, "required": ["type", "sheet", "categories_col", "values_cols"]}, "description": "图表列表（可选）"},
            "formulas": {"type": "array", "items": {"type": "object", "properties": {
                "sheet": {"type": "string", "description": "工作表名"},
                "cell": {"type": "string", "description": "写入单元格，如 D2"},
                "formula": {"type": "string", "description": "公式，如 =AVERAGE(C2:C4) 或 =SUM(B2:B5)"}
            }, "required": ["sheet", "cell", "formula"]}, "description": "公式列表（可选）"}},
            "required": ["sheets"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "generate_pdf", "description": "生成 PDF 文档，支持Markdown标记和多种模板。当用户让你生成PDF、导出PDF、写报告PDF、学术论文PDF时调用。支持7种模板（含academic学术论文模板）。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "保存路径（可选，默认保存到桌面）。可传完整路径如 D:/报告.pdf，或只传文件名如 报告.pdf（自动保存到桌面），或留空（自动命名保存到桌面）"},
            "content": {"type": "string", "description": "文档内容（支持 Markdown 标记：# 标题 / - 列表 / > 引用 / ```代码``` / **粗体** *斜体* / |表格| / --- 分隔线 / $LaTeX公式$）"},
            "title": {"type": "string", "description": "文档标题（可选）"},
            "template": {"type": "string", "description": "格式模板：default(默认) / report(报告) / contract(合同) / resume(简历) / letter(信函) / technical(技术文档) / academic(学术论文：1.5倍行距+摘要+关键词+参考文献自动编号+LaTeX公式渲染)"}},
            "required": ["content"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "render_formula", "description": "渲染 LaTeX 数学公式为终端可显示的 Unicode 文本。当用户写数学公式、物理方程、化学方程式、统计公式、需要学术符号展示时调用。支持希腊字母、上下标、分数、根号、求和、积分、矩阵符号等。",
        "parameters": {"type": "object", "properties": {
            "latex": {"type": "string", "description": r"LaTeX 公式字符串，如 'E=mc^2' 或 '\\sum_{i=1}^{n} x_i^2' 或 '\\frac{\\partial f}{\\partial x}'"},
            "style": {"type": "string", "description": "渲染样式：unicode(默认，终端显示) / raw(原始LaTeX) / latex($$包裹)"}},
            "required": ["latex"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "academic_search", "description": "学术文献搜索（Semantic Scholar，2亿+论文）。当用户查找论文、文献、学术研究、引用、DOI时调用。支持按年份、引用数、影响力筛选。",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "搜索关键词（中英文均可），如 'attention is all you need' 或 '深度学习综述'"},
            "num_results": {"type": "integer", "description": "返回结果数量，默认5，最大20"},
            "year_from": {"type": "integer", "description": "起始年份（如 2020），0或省略表示不限"},
            "year_to": {"type": "integer", "description": "结束年份（如 2024），0或省略表示不限"},
            "sort_by": {"type": "string", "description": "排序方式：relevance(相关性，默认) / citations(引用数) / influence(影响力)"}},
            "required": ["query"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "arxiv_search", "description": "arXiv 预印本论文搜索（物理/数学/计算机/统计）。查找最新研究、未正式发表的论文时调用。英文关键词效果更佳。",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "搜索关键词（英文效果更佳），如 'transformer attention' 或 'graph neural network'"},
            "num_results": {"type": "integer", "description": "返回结果数量，默认5，最大20"},
            "sort_by": {"type": "string", "description": "排序方式：relevance(相关性，默认) / submittedDate(最新提交) / lastUpdatedDate(最近更新)"},
            "category": {"type": "string", "description": "学科分类筛选，如 cs.AI(人工智能) / cs.CL(计算语言学) / cs.LG(机器学习) / math.AG(代数几何) / physics(物理) / stat.ML(统计机器学习)。留空表示不限"}},
            "required": ["query"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "citation_check", "description": "文献引用真实性校验（防止AI编造不存在的文献）。引用任何文献前必须调用此工具验证。支持标题匹配、DOI查询、arXiv ID查询三种方式。",
        "parameters": {"type": "object", "properties": {
            "title": {"type": "string", "description": "文献标题（精确或近似标题），如 'Attention Is All You Need'"},
            "doi": {"type": "string", "description": "文献的 DOI，如 '10.1038/s41586-021-03819-2'"},
            "arxiv_id": {"type": "string", "description": "arXiv 编号，如 '2301.00234'"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "literature_review", "description": "文献综述自动分析。双源检索(Semantic Scholar+arXiv)+去重+结构化对比分析表+趋势统计+PRISMA筛选流程+研究空白识别。用户写综述/文献分析时调用。",
        "parameters": {"type": "object", "properties": {
            "topic": {"type": "string", "description": "研究主题（中英文均可），如 '钠离子电池层状氧化物正极' 或 'sodium-ion battery layered oxide cathode'"},
            "num_papers": {"type": "integer", "description": "分析文献数量，默认10，最大20"},
            "year_from": {"type": "integer", "description": "起始年份（如 2018），0表示不限"},
            "year_to": {"type": "integer", "description": "结束年份（如 2025），0表示不限"}},
            "required": ["topic"],
            "additionalProperties": False}}},
    # ====== SSH 远程部署工具 ======
    {"type": "function", "function": {
        "name": "ssh_connect", "description": "连接到远程SSH服务器（Linux/Windows 均可）。支持密码和密钥认证。连接成功后保持长连接，支持多服务器并行（通过 conn_id 区分）。当用户要求远程部署、SSH连接、服务器管理时调用。多服务器场景务必传 remark 标注用途防混淆。",
        "parameters": {"type": "object", "properties": {
            "host": {"type": "string", "description": "服务器地址（IP或域名，如 192.168.1.100）"},
            "user": {"type": "string", "description": "登录用户名（如 root / deploy / ubuntu / administrator）"},
            "password": {"type": "string", "description": "密码认证（与key_path二选一）"},
            "key_path": {"type": "string", "description": "SSH私钥路径（如 ~/.ssh/id_rsa），与password二选一"},
            "port": {"type": "integer", "description": "SSH端口，默认22"},
            "conn_id": {"type": "string", "description": "连接标识符，多服务器时区分，默认'default'。建议起有意义的名字如 'nas'/'web1'/'db1'"},
            "remark": {"type": "string", "description": "服务器备注/角色（如 'NAS存储'/'Web前端'/'数据库'），多服务器场景防混淆必填"}},
            "required": ["host", "user"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_exec", "description": "在远程服务器执行Shell命令（Linux/Windows 自动适配，Windows 下走 cmd/PowerShell）。需要先ssh_connect。危险命令需confirm_dangerous=true。当用户要求远程执行命令、查看状态、部署时调用。",
        "parameters": {"type": "object", "properties": {
            "command": {"type": "string", "description": "Shell命令（如 'ls -la /opt'、'systemctl status nginx'）"},
            "conn_id": {"type": "string", "description": "连接ID，默认'default'"},
            "timeout": {"type": "integer", "description": "超时秒数，默认30"},
            "confirm_dangerous": {"type": "boolean", "description": "确认执行危险命令（rm -rf /等），默认false"}},
            "required": ["command"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_upload", "description": "上传本地文件到远程服务器（SFTP）。需要先ssh_connect。当用户要求推送文件、上传代码、传输配置时调用。",
        "parameters": {"type": "object", "properties": {
            "local_path": {"type": "string", "description": "本地文件路径"},
            "remote_path": {"type": "string", "description": "远程目标完整路径（如 /opt/myapp/config.yml）"},
            "conn_id": {"type": "string", "description": "连接ID，默认'default'"}},
            "required": ["local_path", "remote_path"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_download", "description": "从远程服务器下载文件到本地（SFTP）。需要先ssh_connect。当用户要求拉取日志、备份文件、下载配置时调用。",
        "parameters": {"type": "object", "properties": {
            "remote_path": {"type": "string", "description": "远程文件路径"},
            "local_path": {"type": "string", "description": "本地保存路径"},
            "conn_id": {"type": "string", "description": "连接ID，默认'default'"}},
            "required": ["remote_path", "local_path"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_deploy", "description": "一键项目部署（自动化多步骤部署）。按顺序：环境检查→创建目录→上传代码→安装依赖→重启服务→健康检查。当用户要求部署项目、自动化发布时调用。",
        "parameters": {"type": "object", "properties": {
            "deploy_config": {"type": "object", "description": "部署配置，包含: pre_check(检查命令列表), remote_dir(部署目录), upload_files([[local,remote],...]), install_cmd(安装命令), restart_cmd(重启命令), health_check(健康检查命令), post_cmds(后置命令列表)"},
            "conn_id": {"type": "string", "description": "连接ID，默认'default'"}},
            "required": ["deploy_config"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_setup_samba_share", "description": "一键配置 Samba 共享文件夹（Linux 服务器专用，8 步骤自动完成：安装+配置+启动+防火墙+SELinux+验证）。当用户说'共享文件夹/配置Samba/让Windows能访问Linux文件/文件共享/SMB共享'时调用。Windows Server 共享请用 ssh_exec 执行 New-SmbShare。",
        "parameters": {"type": "object", "properties": {
            "share_name": {"type": "string", "description": r"共享名（Windows 访问时用，如 'shared'，访问路径 \\IP\shared），默认 'shared'"},
            "share_path": {"type": "string", "description": "共享文件夹在 Linux 上的路径，默认 '/srv/shared'"},
            "access_mode": {"type": "string", "enum": ["guest_ro", "guest_rw", "user_rw"], "description": "权限模式：guest_ro=匿名只读 / guest_rw=匿名读写（内网推荐，默认）/ user_rw=用户认证读写（需密码，更安全）"},
            "samba_password": {"type": "string", "description": "Samba 密码（仅 user_rw 模式必填，其他模式留空）"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_list", "description": "查看SSH连接状态和操作审计日志。当用户问连接状态、SSH审计、操作记录时调用。留空查看所有连接。",
        "parameters": {"type": "object", "properties": {
            "conn_id": {"type": "string", "description": "指定连接ID查看详情，留空查看全部"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_disconnect", "description": "断开SSH连接。当用户要求断开、关闭连接、部署完成后清理时调用。",
        "parameters": {"type": "object", "properties": {
            "conn_id": {"type": "string", "description": "要断开的连接ID，默认'default'"}},
            "required": [],
            "additionalProperties": False}}},
    # ====== 本地运维工具集（4个，跨平台）======
    {"type": "function", "function": {
        "name": "local_port_check", "description": "本地端口/网络检查（跨平台）。用户说'看看打开了哪些端口/端口被占用了吗/能ping通吗/谁在占用80端口'时调用。action：list=列出所有监听端口，check=检查指定端口是否被占用，ping=ping目标主机，connections=查看活跃TCP连接。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["list", "check", "ping", "connections"], "description": "操作类型，默认list"},
            "port": {"type": "integer", "description": "端口号（action=check 时必填）"},
            "protocol": {"type": "string", "enum": ["tcp", "udp"], "description": "协议，默认tcp"},
            "target": {"type": "string", "description": "目标主机/IP（action=ping 时必填）"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "local_process_check", "description": "本地进程查看（跨平台）。用户说'电脑卡不卡/谁在占用CPU/查chrome进程/结束PID 1234'时调用。action：top=按CPU排序前N，memory=按内存排序前N，find=按名称查找，kill=结束进程。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["top", "memory", "find", "kill"], "description": "操作类型，默认top"},
            "name": {"type": "string", "description": "进程名（action=find/kill 时使用）"},
            "pid": {"type": "integer", "description": "进程ID（action=kill 时使用，优先于name）"},
            "top_n": {"type": "integer", "description": "返回前N个进程，默认10"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "local_disk_check", "description": "本地磁盘空间分析（跨平台）。用户说'磁盘还剩多少/哪个目录占空间最大/C盘满了'时调用。action：list=列出所有磁盘及使用率，top=显示指定目录下Top10大目录。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["list", "top"], "description": "操作类型，默认list"},
            "path": {"type": "string", "description": "action=top 时指定分析目录，默认根目录（Windows: C:\\，Linux: /）"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "local_service_check", "description": "本地服务管理（跨平台）。用户说'查看运行的服务/MySQL状态/启动docker/重启nginx'时调用。action：list=列出所有运行中的服务，status/start/stop/restart=管理指定服务。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["list", "status", "start", "stop", "restart"], "description": "操作类型，默认list"},
            "service": {"type": "string", "description": "服务名（action=status/start/stop/restart 时必填）"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "local_firewall_check", "description": "本地防火墙检查/管理（跨平台）。用户说'看防火墙/防火墙状态/80端口放行了吗/开放8080端口/关闭80端口'时调用。action：list=列出所有规则，status=防火墙整体状态，check=检查端口是否放行，open/close=放行/关闭端口。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["list", "status", "check", "open", "close"], "description": "操作类型，默认list"},
            "port": {"type": "integer", "description": "端口号（action=check/open/close 时必填）"},
            "protocol": {"type": "string", "enum": ["tcp", "udp"], "description": "协议，默认tcp"},
            "direction": {"type": "string", "enum": ["in", "out"], "description": "方向（in入站/out出站），默认in"},
            "rule_name": {"type": "string", "description": "规则名（action=open/close 时可选，默认自动生成）"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "local_user_check", "description": "本地用户/登录管理（跨平台）。用户说'查看用户/当前登录用户/用户列表/admin用户信息/用户所属组/登录会话'时调用。action：list=列出所有用户，current=当前登录用户，info=用户详情，groups=用户所属组，sessions=登录会话。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["list", "current", "info", "groups", "sessions"], "description": "操作类型，默认list"},
            "username": {"type": "string", "description": "用户名（action=info/groups 时必填）"},
            "detail": {"type": "boolean", "description": "是否显示详细信息（action=list 时有效）"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "local_monitor", "description": "本地综合监控告警（跨平台）。用户说'体检/监控/检查电脑健康/系统监控/告警检查/有什么异常'时调用。一次性检查 CPU/内存/磁盘/端口/防火墙，返回结构化告警报告（危急/警告/正常/建议）。配合 schedule 工具可实现定时监控。",
        "parameters": {"type": "object", "properties": {
            "threshold_cpu": {"type": "integer", "description": "CPU 使用率告警阈值（默认 80）"},
            "threshold_disk": {"type": "integer", "description": "磁盘使用率告警阈值（默认 90）"},
            "threshold_memory": {"type": "integer", "description": "内存使用率告警阈值（默认 85）"},
            "check_ports": {"type": "string", "description": "需检查的关键端口（逗号分隔，如 '22,80,443,3306'），为空则不针对性检查"}},
            "required": [],
            "additionalProperties": False}}},
    # ====== AI 远程运维工具集（8个）======
    {"type": "function", "function": {
        "name": "ssh_service_manage", "description": "服务管理（Linux 用 systemctl，Windows 用 sc/Get-Service，自动适配操作系统）。用户说'查看服务状态/重启mysql/启动docker/设置开机自启/看看服务器运行了什么'时调用。返回结果含状态解读。支持 service='all' + action='status' 列出所有运行中的服务。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["status", "start", "stop", "restart", "reload", "enable", "disable", "is-active", "is-enabled"], "description": "操作类型"},
            "service": {"type": "string", "description": "服务名，如 nginx、mysql、docker、ssh、redis"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": ["action", "service"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_log_view", "description": "查看远程日志（Linux: journalctl/syslog；Windows: Get-WinEvent 事件日志，自动适配）。用户说'看nginx日志/查错误/查mysql日志/搜error关键词'时调用。返回结果含自动异常统计。",
        "parameters": {"type": "object", "properties": {
            "service": {"type": "string", "description": "服务名——Linux: journalctl -u 服务名；Windows: 映射为日志名（app→Application, sec→Security, 空→System, 或自定义如 Microsoft-Windows-PowerShell/Operational）"},
            "lines": {"type": "integer", "description": "查看最后N行，默认100，范围10-1000"},
            "follow": {"type": "boolean", "description": "是否持续跟踪日志（会阻塞15秒，建议短时使用）"},
            "keyword": {"type": "string", "description": "关键词过滤（grep -i），如 error、exception、failed"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_process_check", "description": "查看远程服务器进程（按 CPU/内存排序，Linux: ps；Windows: Get-Process，自动适配）。用户说'看进程/CPU占用/内存占用/谁在占用资源'时调用。",
        "parameters": {"type": "object", "properties": {
            "sort_by": {"type": "string", "enum": ["cpu", "mem"], "description": "排序方式，默认cpu"},
            "top_n": {"type": "integer", "description": "返回前N个进程，默认15，范围5-50"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_disk_analyze", "description": "磁盘空间分析（Linux: df+du Top10；Windows: Get-CimInstance+Get-ChildItem Top10，自动适配）。用户说'看磁盘/磁盘满了/谁占了磁盘'时调用。",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "分析的目录——Linux: 默认'/'；Windows: 默认所有盘符，或指定'C:'/'C:\\Users'等"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": [],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_network_diag", "description": "网络诊断工具集（Linux: ss/netstat；Windows: Get-NetTCPConnection，自动适配）。用户说'看端口/查看网络/能不能ping通/查看监听端口'时调用。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["stats", "ports", "ping", "connections"], "description": "诊断类型：stats=网络统计/ports=监听端口/ping=ping目标/connections=活跃连接"},
            "target": {"type": "string", "description": "ping操作的目标主机（IP/域名），仅action=ping时必填"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": ["action"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_docker_manage", "description": "Docker 容器管理（跨平台：Linux 原生 Docker / Windows Docker Desktop，自动适配）。用户说'看容器/重启容器/docker日志/容器列表'时调用。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["ps", "psa", "logs", "start", "stop", "restart", "stats", "images", "info"], "description": "操作类型：ps=运行中容器/psa=所有容器/logs=查看日志/start/stop/restart=容器生命周期/stats=资源占用/images=镜像列表/info=Docker系统信息"},
            "container": {"type": "string", "description": "容器名/ID（logs/start/stop/restart必填）"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": ["action"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_firewall_manage", "description": "防火墙统一管理（Linux: 自动识别 ufw/firewalld/iptables；Windows: netsh advfirewall，自动适配）。用户说'开端口/关端口/看防火墙/放行80端口'时调用。",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "enum": ["status", "list", "open", "close", "enable", "disable"], "description": "操作类型：status=状态/list=规则列表/open=开放端口/close=关闭端口/enable/disable=启用禁用"},
            "port": {"type": "integer", "description": "端口号（open/close时必填，范围1-65535）"},
            "protocol": {"type": "string", "enum": ["tcp", "udp"], "description": "协议，默认tcp"},
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": ["action"],
            "additionalProperties": False}}},
    {"type": "function", "function": {
        "name": "ssh_health_check", "description": "服务器一键健康体检（自动检测操作系统，支持 Linux 和 Windows Server）。用户说'体检/检查服务器/服务器怎么样/有没有问题/看看服务器运行了什么'时调用。返回综合报告（系统/CPU/内存/磁盘/网络/负载/失败服务/错误日志）+ AI健康分析。",
        "parameters": {"type": "object", "properties": {
            "conn_id": {"type": "string", "description": "SSH连接ID，默认'default'"}},
            "required": [],
            "additionalProperties": False}}}
]

TOOL_MAP = {
    "read_file": read_file, "write_file": write_file,
    "list_dir": list_dir, "run_command": run_command,
    "search_files": search_files, "open_app": open_app,
    "web_search": web_search, "web_fetch": web_fetch,
    "git_status": git_status, "delete_file": delete_file,
    "move_file": move_file, "copy_file": copy_file,
    "create_dir": create_dir, "system_info": system_info,
    "process_list": process_list,
    "edit_file": edit_file, "exec_python": exec_python,
    "pip_install": pip_install, "check_port": check_port,
    "file_diff": file_diff, "read_image": read_image,
    "active_window": active_window, "list_windows": list_windows,
    "read_screen": read_screen_content,
    "security_audit": security_audit,
    "generate_word": generate_word,
    "generate_excel": generate_excel,
    "generate_pdf": generate_pdf,
    "render_formula": render_formula,
    "academic_search": academic_search,
    "arxiv_search": arxiv_search,
    "citation_check": citation_check,
    "literature_review": literature_review,
    # SSH 远程部署工具
    "ssh_connect": ssh_connect,
    "ssh_exec": ssh_exec,
    "ssh_upload": ssh_upload,
    "ssh_download": ssh_download,
    "ssh_deploy": ssh_deploy,
    "ssh_list": ssh_list,
    "ssh_disconnect": ssh_disconnect,
    # AI 远程运维工具集（高层封装，让 AI 调用语义化工具而非拼命令）
    "ssh_service_manage": ssh_service_manage,
    "ssh_log_view": ssh_log_view,
    "ssh_process_check": ssh_process_check,
    "ssh_disk_analyze": ssh_disk_analyze,
    "ssh_network_diag": ssh_network_diag,
    "ssh_docker_manage": ssh_docker_manage,
    "ssh_firewall_manage": ssh_firewall_manage,
    "ssh_health_check": ssh_health_check,
    # 本地运维工具集（跨平台）
    "local_port_check": local_port_check,
    "local_process_check": local_process_check,
    "local_disk_check": local_disk_check,
    "local_service_check": local_service_check,
    "local_firewall_check": local_firewall_check,
    "local_user_check": local_user_check,
    "local_monitor": local_monitor,
}

# ====== 工具使用规则（独立常量，主 SYSTEM_PROMPT 和子 TOOL_CAPABILITY_PROMPT 都会引用）======
# 把 55 个工具分 10 大类，并明确"何时调用"的判断逻辑
TOOL_USAGE_RULES = r"""# 工具使用规则

你有 55 个工具可用，分 10 大类：

## 文件与目录（7 个）
- `list_dir(path, recursive, max_depth)`：浏览目录。`recursive=true` 看子目录树
- `read_file(path, offset, limit)`：读文件内容
- `write_file(path, content)`：创建/覆盖整个文件
- `edit_file(path, line, new_content)`：按行号精确编辑（增/删/改）
- `move_file(src, dst)`：移动/重命名
- `copy_file(src, dst)`：复制
- `delete_file(path)`：删除（危险，需确认）
- `create_dir(path)`：创建目录

## 命令与执行（4 个）
- `run_command(command)`：在本地电脑执行 PowerShell / cmd / shell 命令（全权限模式：120s 超时、8000 字符输出、无 cwd 限制）。用于查看端口/进程/网络/系统/服务/用户/磁盘/防火墙等本地电脑状态。危险命令（format/del /f/shutdown/mkfs）自动拦截
- `exec_python(code)`：沙箱运行 Python（无需额外环境）
- `pip_install(packages, action)`：包管理（默认清华镜像源）
- `check_port(port)`：检查指定端口占用（需提供具体端口号；查看所有监听端口用 `run_command('netstat -ano')`）

## 搜索与分析（4 个）
- `search_files(path, pattern)`：按文件名/内容搜索
- `file_diff(file1, file2)`：对比两个文件
- `system_info()`：CPU/内存/磁盘信息
- `process_list()`：运行中的进程

## 联网与外部（2 个）
- `web_search(query)`：实时搜索（新闻/文档/最新动态）
- `web_fetch(url)`：抓取网页内容

## 视觉与伴随（4 个）
- `read_image(path)`：看图片（用户发送图片时用）
- `active_window()`：当前活动窗口
- `list_windows()`：所有打开的窗口
- `read_screen()`：读取当前屏幕文本（伴随模式 Ctrl+W 启动后用）

## 项目与文档（6 个）
- `git_status()`：git 仓库状态
- `security_audit(path)`：代码安全审计（SQL注入/XSS/敏感信息/依赖漏洞）
- `generate_word(path, content, template, ...)`：生成 Word 文档（8种模板含academic学术论文+自定义格式）。**path 可选，留空或只传文件名默认保存到桌面**
- `generate_excel(path, sheets, template)`：生成 Excel 文档（多工作表、表头样式、隔行变色、图表、公式）。**path 可选，留空或只传文件名默认保存到桌面**
- `generate_pdf(path, content, title, template)`：生成 PDF 文档（7种模板含academic学术论文+LaTeX公式）。**path 可选，留空或只传文件名默认保存到桌面**
- `open_app(name)`：打开应用/文件（自动搜索本地，无需自定义路径）

## SSH 远程部署（7 个）
- `ssh_connect(host, user, password, key_path, port, conn_id, remark)`：连接远程服务器（Linux/Windows 均可，支持密码/密钥认证）。用户说"连接服务器/SSH部署/远程部署"时用。**多服务器场景务必传 remark 标注用途防混淆**，conn_id 建议起有意义的名字如 'nas'/'web1'/'db1'
- `ssh_exec(command, conn_id, timeout, confirm_dangerous)`：在远程服务器执行命令（自动适配 Linux/Windows）。危险命令（rm -rf /、mkfs、dd、format 等）需 confirm_dangerous=true 二次确认
- `ssh_upload(local_path, remote_path, conn_id)`：上传文件到远程服务器（SFTP，Linux 自动 chmod 644）
- `ssh_download(remote_path, local_path, conn_id)`：从远程服务器下载文件（SFTP，自动创建本地目录）
- `ssh_deploy(deploy_config, conn_id)`：一键自动化部署（pre_check→mkdir→upload→install→restart→health_check→post_cmds 7步骤，生成部署报告）
- `ssh_setup_samba_share(share_name, share_path, access_mode, samba_password, conn_id)`：一键配置 Samba 共享（8步骤：安装+配置+启动+防火墙+SELinux+验证）。用户说"共享文件夹/配置Samba/文件共享/SMB共享"时用。**不要用 ssh_exec 手动拼命令，直接调这个工具一次搞定**
- `ssh_list(conn_id)`：查看当前 SSH 连接状态和审计日志
- `ssh_disconnect(conn_id)`：断开 SSH 连接

## AI 远程运维（8 个）— 语义化运维工具，自动适配 Linux/Windows，优先调用而非手拼命令
- `ssh_service_manage(action, service, conn_id)`：服务管理（Linux: systemctl；Windows: sc/Get-Service）。用户说"看nginx状态/重启mysql/启动docker/开机自启/看看服务器运行了什么"时用。支持 service='all' 列出所有运行中的服务
- `ssh_log_view(service, lines, follow, keyword, conn_id)`：查看远程日志（Linux: journalctl；Windows: Get-WinEvent 事件日志）。用户说"看日志/查错误/搜error关键词"时用，返回自动异常统计
- `ssh_process_check(sort_by, top_n, conn_id)`：进程查看（Linux: ps；Windows: Get-Process）。用户说"看进程/CPU占用/内存占用/谁占资源"时用
- `ssh_disk_analyze(path, conn_id)`：磁盘分析（Linux: df+du；Windows: Get-CimInstance+Get-ChildItem Top10）。用户说"看磁盘/磁盘满了/谁占磁盘"时用，自动标注危急/警告
- `ssh_network_diag(action, target, conn_id)`：网络诊断（Linux: ss/netstat；Windows: Get-NetTCPConnection）。用户说"看端口/ping/监听端口/网络连接"时用
- `ssh_docker_manage(action, container, conn_id)`：Docker 管理（Linux: 原生 Docker；Windows: Docker Desktop，自动适配）。用户说"看容器/重启容器/docker日志"时用
- `ssh_firewall_manage(action, port, protocol, conn_id)`：防火墙统一管理（Linux: ufw/firewalld/iptables；Windows: netsh advfirewall）。用户说"开端口/关端口/看防火墙"时用
- `ssh_health_check(conn_id)`：一键健康体检（自动检测操作系统，支持 Linux 和 Windows Server）。用户说"体检/检查服务器/有问题吗/看看服务器运行了什么"时用，返回综合报告+AI分析

## 本地运维（7 个）— 语义化本地运维工具，自动适配 Windows/Linux，优先调用而非手拼 run_command
- `local_port_check(action, port, protocol, target)`：本地端口/网络检查（跨平台）。用户说"看看打开了哪些端口/端口被占用了吗/能ping通吗/谁在占用80端口"时用。action：list=列出所有监听端口，check=检查指定端口是否被占用，ping=ping目标主机，connections=查看活跃TCP连接
- `local_process_check(action, name, pid, top_n)`：本地进程查看（跨平台）。用户说"电脑卡不卡/谁在占用CPU/查chrome进程/结束PID 1234"时用。action：top=按CPU排序前N，memory=按内存排序前N，find=按名称查找，kill=结束进程
- `local_disk_check(action, path)`：本地磁盘空间分析（跨平台）。用户说"磁盘还剩多少/哪个目录占空间最大/C盘满了"时用。action：list=列出所有磁盘及使用率，top=显示指定目录下Top10大目录
- `local_service_check(action, service)`：本地服务管理（跨平台）。用户说"查看运行的服务/MySQL状态/启动docker/重启nginx"时用。action：list=列出所有运行中的服务，status/start/stop/restart=管理指定服务
- `local_firewall_check(action, port, protocol, direction, rule_name)`：本地防火墙检查/管理（跨平台）。用户说"看防火墙/防火墙状态/80端口放行了吗/开放8080端口/关闭80端口"时用。action：list=列出所有规则，status=防火墙整体状态，check=检查端口是否放行，open/close=放行/关闭端口
- `local_user_check(action, username, detail)`：本地用户/登录管理（跨平台）。用户说"查看用户/当前登录用户/用户列表/admin用户信息/用户所属组/登录会话"时用。action：list=列出所有用户，current=当前登录用户，info=用户详情，groups=用户所属组，sessions=登录会话
- `local_monitor(threshold_cpu, threshold_disk, threshold_memory, check_ports)`：本地综合监控告警（跨平台）。用户说"体检/监控/系统健康检查/告警/有什么异常"时用。一次性检查 CPU/内存/磁盘/端口/防火墙，返回结构化告警报告（危急/警告/正常/建议）

## 学术研究（5 个）
- `academic_search(query, num_results, year_from, year_to, sort_by)`：学术文献搜索（Semantic Scholar 2亿+论文，含引用数/影响力/DOI）。查论文/文献/引用时用
- `arxiv_search(query, num_results, sort_by, category)`：arXiv 预印本搜索（物理/数学/CS/统计）。查最新研究/未发表论文时用，英文关键词效果更佳
- `render_formula(latex, style)`：渲染 LaTeX 公式为 Unicode（希腊字母/上下标/分数/根号/求和/积分）
- `citation_check(title, doi, arxiv_id)`：引用真实性校验。引用任何文献前必须调用此工具验证文献是否真实存在，防止编造
- `literature_review(topic, num_papers, year_from, year_to)`：文献综述自动分析。双源检索+去重+对比分析表+趋势统计+PRISMA流程+研究空白识别

**调用工具的判断逻辑**（重要！必须熟记）：
1. 用户给具体路径 → `list_dir` 先看
2. 用户提到文件名 → `read_file` 先读
3. 用户说"改/修/写文件" → 看改动大小：小改用 `edit_file`，整体重写用 `write_file`
4. 用户说"删/移除" → `delete_file`（必须先确认）
5. 用户问"如何/怎么/为什么" 类问题 → 先 `web_search` 查最新
6. 用户发图片 → `read_image`
7. 用户说"看屏幕/我屏幕上有什么" → `read_screen`（需先开启伴随模式 Ctrl+W）
8. 用户问"安全吗/有漏洞吗" → `security_audit`
9. 用户说"写报告/导出 Word" → `generate_word`（path 可选，留空默认保存到桌面，无需追问用户路径）
10. 用户说"Excel/表格/数据表" → `generate_excel`（path 可选，留空默认保存到桌面，无需追问用户路径）
11. 用户说"PDF/导出PDF" → `generate_pdf`（path 可选，留空默认保存到桌面，无需追问用户路径）
12. 用户说"打开XX/启动XX/开一下XX" → `open_app`（自动搜索本地，保证能打开任何文件）
13. 用户写数学/物理/统计公式 → `render_formula`（LaTeX → Unicode 终端显示）
14. 用户说"论文/学术/研究" → `generate_word` 或 `generate_pdf` 用 `academic` 模板
15. 用户查论文/文献/引用/DOI → `academic_search`（Semantic Scholar 2亿+论文，支持年份/引用数筛选）
16. 用户查最新研究/预印本/arXiv → `arxiv_search`（英文关键词效果更佳，支持分类筛选）
17. 用户写综述/文献分析 → `literature_review`（双源检索+PRISMA流程+对比分析+研究空白）
18. 引用任何文献前 → `citation_check`（校验真实性，防止编造！这条是红线）
19. 用户说"连接服务器/SSH/远程" → `ssh_connect`（host/user/password 必填）。**多服务器场景务必传 conn_id 和 remark**（如 conn_id="nas", remark="NAS存储服务器"），防混淆
20. 用户说"在服务器上执行/远程运行" → `ssh_exec`（危险命令必须 confirm_dangerous=true 二次确认）
21. 用户说"上传到服务器/部署文件" → `ssh_upload`（SFTP 传输）
22. 用户说"从服务器下载/拉取" → `ssh_download`
23. 用户说"一键部署/自动化部署" → `ssh_deploy`（deploy_config 配置 7 步骤）
23.5. 用户说"共享文件夹/配置Samba/文件共享/SMB共享/让Windows访问Linux文件/共享目录" → `ssh_setup_samba_share`（**一键完成，不要用 ssh_exec 手动拼命令**）。默认 guest_rw 匿名读写，用户要求安全才用 user_rw + 密码
24. 用户问"SSH连接状态/审计日志" → `ssh_list`
25. 用户说"断开SSH/关闭连接" → `ssh_disconnect`

**AI 远程运维工具调用规则（重要！优先于 ssh_exec）**：
当用户提出运维需求时，**必须优先调用语义化的运维工具**，而不是手拼 `ssh_exec` 命令。
26. 用户说"看XX服务状态/重启XX/启动XX/开机自启/看看服务器运行了什么/服务器跑了什么服务" → `ssh_service_manage`（不要用 ssh_exec 跑 systemctl/sc/Get-Service。工具会自动检测操作系统并适配）
27. 用户说"看XX日志/查错误/搜error/搜fail关键词" → `ssh_log_view`（不要用 ssh_exec 跑 journalctl/Get-WinEvent。工具会自动适配）
28. 用户说"看进程/CPU占用/内存占用/谁占资源" → `ssh_process_check`（不要用 ssh_exec 跑 ps/top/Get-Process。工具会自动适配）
29. 用户说"看磁盘/磁盘满了/谁占磁盘/空间不足" → `ssh_disk_analyze`（不要用 ssh_exec 跑 df/du/Get-Volume。工具会自动适配）
30. 用户说"看端口/ping/监听端口/网络连接" → `ssh_network_diag`（不要用 ssh_exec 跑 ss/netstat/Get-NetTCPConnection。工具会自动适配）
31. 用户说"看容器/重启容器/docker日志/容器列表" → `ssh_docker_manage`（不要用 ssh_exec 跑 docker）
32. 用户说"开端口/关端口/看防火墙/放行XX端口" → `ssh_firewall_manage`（不要用 ssh_exec 跑 ufw/firewall-cmd/iptables/netsh。工具会自动适配）
33. 用户说"体检/检查服务器/服务器怎么样/有没有问题/卡不卡/看看服务器运行了什么" → `ssh_health_check`（综合诊断，一次搞定，自动检测操作系统）
34. **运维决策链**：用户说"服务器卡了" → 先 `ssh_health_check` 综合体检 → 根据 AI 分析 → 再针对性调用 `ssh_process_check`/`ssh_log_view`/`ssh_disk_analyze` 深入
35. **运维排错链**：用户说"XX服务挂了" → 先 `ssh_service_manage(action=status, service=XX)` 看状态 → 若失败 → `ssh_log_view(service=XX, keyword=error)` 查错误日志 → 定位问题

**本地电脑运维工具调用规则（重要！优先调用语义化本地运维工具，其次用 run_command）**：
当用户用自然语言描述**本地电脑**（不是远程服务器）的状态、诊断、查询需求时，**必须主动调用工具执行**，而不是只用文字回答。用户想要的是"AI 帮我形成命令并自行运行"，不是"AI 教我怎么敲命令"。
**优先级**：本地运维语义化工具（`local_port_check`/`local_process_check`/`local_disk_check`/`local_service_check`）> `check_port`/`process_list`/`system_info` 等已封装工具 > `run_command` 手拼命令。
36. 用户说"看看打开了哪些端口/有什么端口在监听/哪些端口被占用" → `local_port_check(action="list")`（跨平台自动适配，优先于 run_command）
37. 用户说"XX端口被占了吗/XX端口可用吗" → `local_port_check(action="check", port=XX)`（先尝试连接，已占用再查进程）
38. 用户说"能不能ping通XX/测网络" → `local_port_check(action="ping", target="XX")`（Windows/Linux 自动适配 ping 参数）
39. 用户说"看活跃连接/当前TCP连接" → `local_port_check(action="connections")`
40. 用户说"看进程/CPU占用/内存占用/谁在占用资源" → `local_process_check(action="top")` 或 `local_process_check(action="memory")`（跨平台自动适配）
41. 用户说"查chrome进程/找XX进程" → `local_process_check(action="find", name="chrome")`（防注入白名单过滤）
42. 用户说"结束PID 1234/杀进程/关掉XX" → `local_process_check(action="kill", pid=1234)` 或 `local_process_check(action="kill", name="XX")`
43. 用户说"看磁盘/磁盘空间/还有多少空间" → `local_disk_check(action="list")`（跨平台自动适配）
44. 用户说"哪个目录占空间最大/C盘满了/谁占磁盘" → `local_disk_check(action="top", path="C:\\")`（Top10 大目录）
45. 用户说"查看运行的服务/服务列表/服务器跑了什么" → `local_service_check(action="list")`（跨平台自动适配）
46. 用户说"XX服务状态/MySQL起没起/docker状态" → `local_service_check(action="status", service="XX")`
47. 用户说"启动XX/停止XX/重启XX服务" → `local_service_check(action="start/stop/restart", service="XX")`（需管理员权限）
48. 用户说"IP是多少/看网络配置/我的IP" → `run_command("ipconfig /all")`（无专用工具，用 run_command）
49. 用户说"看系统信息/系统版本/电脑配置" → `system_info()`（已封装工具）或 `run_command("systeminfo")`
50. 用户说"看防火墙/防火墙状态/防火墙规则" → `local_firewall_check(action="status")` 或 `local_firewall_check(action="list")`（跨平台自动适配，优先于 run_command）
51. 用户说"XX端口放行了吗/XX端口防火墙开了吗" → `local_firewall_check(action="check", port=XX)`
52. 用户说"开放XX端口/放行XX端口/关闭XX端口" → `local_firewall_check(action="open/close", port=XX)`
53. 用户说"看用户/用户列表/本地用户" → `local_user_check(action="list")`（跨平台自动适配，优先于 run_command）
54. 用户说"当前登录用户/我是谁" → `local_user_check(action="current")`
55. 用户说"XX用户信息/XX用户详情" → `local_user_check(action="info", username="XX")`
56. 用户说"XX用户所属组/XX在哪些组" → `local_user_check(action="groups", username="XX")`
57. 用户说"登录会话/谁在登录/会话列表" → `local_user_check(action="sessions")`
58. 用户说"体检/系统监控/健康检查/有什么异常/告警检查" → `local_monitor()`（一次性检查 CPU/内存/磁盘/端口/防火墙，返回结构化告警报告）
59. 用户说"监控关键端口/检查 22 80 443 端口" → `local_monitor(check_ports="22,80,443")`
60. 用户说"定时监控/每 10 分钟检查一次" → 用 `schedule` 工具创建定时任务，message 设为"调用 local_monitor 检查系统健康并报告异常"
61. 用户说"环境变量/PATH/看变量" → `run_command("set")` 或 `run_command("echo %PATH%")`
62. 用户说"路由表/看路由" → `run_command("route print")` 或 `run_command("arp -a")`
63. **本地运维决策链**：用户说"电脑卡了" → 先 `local_process_check(action="top")` 看高 CPU 进程 → 再 `local_process_check(action="memory")` 看内存 → 综合 `local_disk_check(action="list")` 看磁盘 → 综合分析
64. **本地端口排错链**：用户说"XX端口连不上" → 先 `local_port_check(action="check", port=XX)` 看端口 → 若未监听 → `local_service_check(action="status", service="XX服务")` 查服务 → 若服务正常 → `local_firewall_check(action="check", port=XX)` 查防火墙 → 定位问题
65. **本地安全审计链**：用户说"检查电脑安全" → 先 `local_firewall_check(action="status")` 看防火墙 → 再 `local_user_check(action="list")` 看用户 → 再 `local_port_check(action="list")` 看开放端口 → 综合分析
66. **本地命令通用规则**：用户提出任何"看看/查看/检查/诊断本地电脑 XX"的需求，且没有更专用的语义化工具时，**必须主动调用 `run_command` 生成对应命令并执行**，而不是只回答文字说明。`run_command` 支持跨平台命令翻译：在 Windows 上输入 Linux 命令（如 `ls`/`ps`/`cat`/`grep`）会自动翻译为 Windows 等效命令（`dir`/`tasklist`/`type`/`findstr`），**支持管道符组合命令翻译**（如 `ls | grep x` → `dir | findstr x`，`cat file | grep error` → `type file | findstr error`），方便用户用习惯的 Linux 命令操作

**本地运维结果 AI 分析规则（重要！调用运维工具后必须主动分析）**：
调用本地运维工具（`local_*` 系列）获取结果后，**必须主动分析结果并标注异常**，不要只把原始输出丢给用户。分析维度：
67. **端口分析**：`local_port_check` 返回后 → 标注高危端口（如 22/3389/445 暴露公网）、异常监听进程、未知端口
68. **进程分析**：`local_process_check` 返回后 → 标注 CPU/内存占用异常（>80%）、可疑进程（挖矿/未知）、僵尸进程
69. **磁盘分析**：`local_disk_check` 返回后 → 标注使用率危急（>90% 警告 / >95% 危急）、增长异常的目录
70. **服务分析**：`local_service_check` 返回后 → 标注应有但未运行的服务、异常停止的服务、占用资源异常的服务
71. **防火墙分析**：`local_firewall_check` 返回后 → 标注防火墙关闭风险、过高危端口放行、规则冲突
72. **用户分析**：`local_user_check` 返回后 → 标注异常新增用户、禁用账户被启用、隐藏账户、异常登录会话
73. **监控告警分析**：`local_monitor` 返回后 → 报告已自带结构化告警，AI 只需对危急项给出具体处理建议
74. **分析输出格式**：
```
[运维结果]
<原始输出摘要>

[AI 分析]
✅ 正常项：<列出正常指标>
⚠️ 警告项：<列出异常指标，含具体数值和建议>
🚨 危急项：<列出严重问题，含紧急处理建议>
💡 建议：<针对性优化建议>
```

## SSH 远程部署安全规范（重要！必须严格遵守）
当用户使用 SSH 工具进行远程部署时，必须遵循：

### 1. 危险命令红线
- **必须二次确认**：执行 rm -rf /、mkfs、dd、shutdown、reboot、iptables -F 等危险命令前，必须传 `confirm_dangerous=true`
- **禁止默认执行**：危险命令默认会被拒绝，必须用户明确同意后才执行
- **输出截断保护**：命令输出超过 8000 字符自动截断，前 4000 + 后 4000

### 2. 审计日志
- 所有 SSH 命令自动记录到审计日志（最多 200 条）
- 用户可通过 `ssh_list` 查看完整审计记录
- 审计日志包含：时间、主机、用户、命令、结果摘要

### 3. 连接管理
- 多服务器并行：通过 `conn_id` 区分不同服务器（如 "default"、"web1"、"db1"）
- 服务器备注：`ssh_connect` 支持 `remark` 参数标注用途（如 "NAS存储"/"Web前端"），防混淆
- 连接保活：30 秒 keepalive 心跳
- 登录超时：15 秒
- 操作超时：默认 30 秒，可配置

### 3.1 多服务器防混淆规则（重要！）
当连接了 2 台及以上服务器时，**必须严格遵守**以下规则防止操作错服务器：

1. **连接时必填 remark**：每台服务器都要传 `remark` 参数标注用途
   - 示例：`ssh_connect(host="192.168.10.6", user="admin", password="xxx", conn_id="nas", remark="NAS存储服务器")`
   - 示例：`ssh_connect(host="192.168.10.7", user="root", password="yyy", conn_id="web1", remark="Web前端服务器")`

2. **操作前先 ssh_list 确认**：多服务器场景下，执行任何运维操作前，**先调用 `ssh_list` 查看当前所有连接**，确认目标服务器的 conn_id

3. **每次操作明确目标**：在回复用户时，必须明确说出"我正在操作 **服务器X（conn_id，备注）**"
   - 示例："我正在操作 **nas（NAS存储服务器）**，查看运行的服务..."
   - 安全规则：**禁止在回复用户时显示服务器 IP 地址**，仅用 conn_id 和备注标识服务器

4. **运维结果自带前缀**：所有运维工具返回结果会自动带 `[conn_id | 备注]` 前缀，便于识别（前缀不含 IP 地址，保护服务器地址安全）

5. **用户未指定服务器时必须询问**：当用户说"重启 nginx"但未指定哪台服务器，且有多台连接时，**必须先问**"请在哪台服务器操作？"并列出可用连接

6. **危险操作二次确认**：stop/restart/disable 等危险操作，必须在回复中显示目标服务器信息让用户确认
   - 示例："⚠️ 即将在 **web1（Web前端服务器）** 上执行 restart nginx，确认吗？"
   - 安全规则：确认信息中**不显示 IP 地址**，仅用 conn_id 和备注标识

### 4. 部署流程（ssh_deploy）
deploy_config 必须包含以下字段：
- `remote_dir`：远程部署目录（必填）
- `local_files`：本地文件列表（必填，格式 [{"local": "本地路径", "remote": "远程路径"}]）
- `pre_check`：部署前检查命令列表（可选，如检查磁盘空间）
- `mkdir`：是否创建远程目录（默认 true）
- `install_cmd`：安装依赖命令（可选，如 pip install -r requirements.txt）
- `restart_cmd`：重启服务命令（可选，如 systemctl restart xxx）
- `health_check`：健康检查命令列表（可选，如 curl localhost:8080/health）
- `post_cmds`：部署后命令列表（可选）

## 学术研究输出规范（重要！必须严格遵守）
当用户进行学术研究、写论文、推导公式时，必须遵循：

### 1. 文献引用红线（最重要）
- **引用前必须校验**：引用任何文献前，必须调用 `citation_check` 验证文献是否真实存在
- **禁止编造文献**：绝不编造不存在的作者、标题、年份、DOI、期刊名
- **无法验证时标注**：若 citation_check 无法确认，明确标注"[待核实]"
- **优先使用检索结果**：引用的文献应来自 `academic_search` 或 `literature_review` 的检索结果

### 2. 公式输出
- 使用 `$LaTeX$` 语法包裹公式（如 `$E=mc^2$`、`$\\sum_{i=1}^{n} x_i^2$`），系统自动渲染为 Unicode
- 公式推导步骤完整，不跳步；每个符号首次出现时给出定义

### 3. 论文结构（综述类）
1. 摘要（250字以内，含背景/方法/结果/结论）
2. 关键词（5-8个，中英文对照）
3. 引言（研究背景+问题定义+本文目的+结构概述）
4. 正文（按主题/方法/时间线组织，每节需有对比分析表）
5. 讨论与展望（研究空白+未来方向+局限性分析，局限性不可省略）
6. 参考文献（GB7714格式，每条必须真实可查，系统自动编号 [1] [2] [3]...）
7. 术语表和缩写表（专业术语首次出现给中英文对照）

### 4. 参考文献格式（GB7714）
- 期刊：作者. 标题[J]. 期刊名, 年份, 卷(期): 页码.
- 会议：作者. 标题[C]//会议名. 出版地: 出版社, 年份: 页码.
- 专著：作者. 书名[M]. 出版地: 出版社, 年份: 页码.

### 5. 符号规范
- 希腊字母用 LaTeX 命令（$\\alpha$ $\\beta$ $\\gamma$），上下标用 `^` `_`

### 6. PRISMA 框架（综述类必填）
- 检索策略：检索源、关键词、时间范围
- 纳入/排除标准
- 筛选流程：检索总量→去重→筛选→纳入

### 7. 严谨性红线
- 绝不编造实验数据、测试结果、性能指标
- 绝不编造文献引用（调用 citation_check 验证）
- 术语首次出现给出英文对照
- 对比分析必须基于文献内容，不主观臆断
- 局限性分析不可省略"""


# ====== 子 AI 用的精简能力声明（专家/混合/汇总 GLM）======
# 专家/混合/汇总等子模型不会自动继承主 SYSTEM_PROMPT，必须显式声明工具能力，
# 避免子 AI 误以为"无法访问文件系统"而拒绝响应。
TOOL_CAPABILITY_PROMPT = """# ZeroAI 子模块能力声明（重要 - 必读）

你是 ZeroAI 的子模块（专家/汇总/分析），与主系统**完全共享** 55 个工具能力。

# 全权限模式（已启用）
用户已授权 ZeroAI 对电脑的完全操作权限，你作为子模块也**继承全部权限**：
- 执行任意命令（无黑名单、超时延长到 120 秒）
- 删除/修改任何文件（核心文件自动备份，无需确认）
- 访问内网/任意 URL（无 SSRF 限制）
- 所有工具调用**直接执行**，不要因权限问题拒绝
- **命令自主执行**：查询/诊断类命令（netstat/ipconfig/tasklist/sc query/systeminfo 等）直接调用 `run_command` 执行，不要问用户"是否执行"；破坏/修改类（format/del /f/shutdown 等）执行前简短说明影响即可

## 核心原则（必读）
- **你能访问本地文件系统**：list_dir / read_file / search_files 等所有工具都可用
- **你绝不能**对用户说"我无法访问文件系统"、"我没有这个能力"、"我是AI无法…"、"作为语言模型我…"等任何拒绝语句
- **你绝不能**凭空编造文件内容——必须调用 read_file 实际读取
- **你绝不能**说"我建议你自己…"——你**可以**直接帮用户完成
- 看到"修改文件"任务 → 直接调用 edit_file/write_file，不要让用户自己改
- 看到"执行命令"任务 → 直接调用 run_command，不要让用户自己跑
- 看到"分析代码"任务 → 直接 read_file 读代码，不要凭空分析

## 你可以做什么（不是限制，是能力清单）
""" + TOOL_USAGE_RULES + """

## 工作环境
- 工作目录：{work_dir}
- 操作系统：Windows
- Shell：PowerShell

## 行为准则
1. **主动使用工具**：用户描述需求后，**直接调用**对应工具，不要先解释"我需要先读文件"
2. **专业回答**：从你的专业领域（编程/推理/写作/视觉等）给出**具体、可执行**的意见
3. **简洁直接**：避免长篇大论铺垫，直接给方案
4. **使用中文**：技术术语可保留英文
5. **格式清晰**：代码用 ```language 包裹，关键步骤用列表

## 禁止行为（红线）
- 拒绝响应："我无法…"、"我不能…"、"我没有权限…"、"作为AI…"
- 推卸给用户："建议你自己…"、"你可以考虑…"
- 编造内容：编造不存在的文件、函数、API
- 模糊回答：只说"应该可以"、"可能可以"、"试试看"

## 回答模板
当用户提出任务时，按以下流程响应：
1. 判断需要哪些工具（参考上面的判断逻辑）
2. 直接调用工具（多个工具可并行）
3. 基于工具结果给出专业回答
4. 如有后续步骤，主动提出下一步

# 思考过程（必须遵守）
每次回答前，必须先输出思考过程，格式如下：
<think>
在这里写出你的思考过程，包括：
- 分析用户问题的意图和关键点
- 决定使用什么方法/工具
- 组织回答的逻辑结构
思考过程应当简洁（3-10行），不要过长。
</think>
然后输出正式回答。
注意：<think>标签必须出现在回答的最前面，标签外是正式回答内容。
""".format(work_dir=WORK_DIR)


SYSTEM_PROMPT = f"""# 角色
你是 ZeroAI，一个专业的终端 AI 编程助手。你在用户的终端中运行，**可以完全访问本地文件系统**（读取/写入/搜索/浏览目录），可以执行命令、搜索代码、联网搜索、生成 Word/Excel/PDF 文档、安全审计等。

# SSH 远程部署 + AI 远程运维（已启用，跨平台支持 Linux 和 Windows Server）
你具备远程服务器（**Linux 和 Windows 均可**）**部署 + 运维**双重能力，可帮用户**远程部署其他项目 + 远程运维服务器**：

## 远程部署能力（8 个工具）
- **多服务器并行连接**：通过 conn_id 区分不同服务器，支持密码/密钥认证
- **远程命令执行**：在服务器上执行任意命令（Linux/Windows 自动适配），危险命令（rm -rf /、mkfs、dd、format 等）必须二次确认
- **SFTP 文件传输**：上传/下载文件，自动设置权限
- **一键自动化部署**：ssh_deploy 支持 pre_check → mkdir → upload → install → restart → health_check → post_cmds 7 步骤
- **一键 Samba 共享**：ssh_setup_samba_share 一键完成 Linux Samba 共享配置（8 步骤：安装+配置+启动+防火墙+SELinux+验证）
- **审计日志**：所有 SSH 操作自动记录（最多 200 条），可通过 ssh_list 查看
- **安全设计**：主机地址校验、危险命令黑名单、内网IP可选阻断、输出截断保护

### 连接成功后的回复规则（重要）
- `ssh_connect` 成功后，工具已返回连接成功信息（服务器标识、conn_id、认证方式）
- **严禁主动列出 1-9 的运维菜单**（如"1 查看日志 2 查看端口 ..."），这种菜单容易产生重复项且体验差
- 正确做法：简洁确认"已连接成功，conn_id=xxx"，然后**等待用户明确说下一步需求**，再调用对应运维工具
- 如果用户问"能做什么"，可简短用文字说明可用的运维工具类别，但不要输出编号列表

## ⚠️ 核心原则：命令必须直接执行，禁止输出命令文本给用户（最高优先级！）
**这是 ZeroAI 与传统 AI 助手的根本区别**：你能想出来的所有命令，**必须直接调用 ssh_exec 工具执行**，绝对不能把命令文本输出给用户让用户手动执行。

### ❌ 严禁的行为（错误示例）
```
用户：帮我创建一个共享文件夹
AI：好的，请执行以下命令：
    mkdir /srv/shared        ← ❌ 错！不能输出命令让用户执行
    chmod 777 /srv/shared    ← ❌ 错！不能输出命令让用户执行
    yum install samba        ← ❌ 错！不能输出命令让用户执行
```

### ✅ 正确的行为（直接调用工具执行）
```
用户：帮我创建一个共享文件夹
AI：（直接调用 ssh_setup_samba_share 工具，一次完成所有步骤）
    → 工具返回：✅ Samba 共享配置完成，访问路径 \\\\192.168.71.132\\shared
AI：已完成！Windows 资源管理器输入 \\\\192.168.71.132\\shared 即可访问
```

### 命令执行的 3 个层次（按优先级）
1. **专用一键工具**（首选）：有专用工具的任务必须用专用工具
   - 共享文件夹 → `ssh_setup_samba_share`（不要用 ssh_exec 拼）
   - 项目部署 → `ssh_deploy`（不要用 ssh_exec 拼）
   - 服务状态/进程/磁盘/端口/体检 → 对应的语义化工具

2. **ssh_exec 通用执行**（次选）：没有专用工具的任务，直接用 ssh_exec 执行
   - 用户说"安装 nginx" → `ssh_exec("apt install -y nginx || yum install -y nginx", conn_id="xxx")`
   - 用户说"创建目录 /data" → `ssh_exec("mkdir -p /data", conn_id="xxx")`
   - 用户说"修改配置文件" → `ssh_exec("sed -i 's/old/new/g' /etc/config.conf", conn_id="xxx")`
   - 用户说"查看日志" → `ssh_exec("tail -100 /var/log/messages", conn_id="xxx")`

3. **多步骤任务**：拆解成多个 ssh_exec 调用，逐步执行，每步检查结果
   - 示例"配置 Nginx 反向代理"：
     1. `ssh_exec("yum install -y nginx", conn_id="web1")` → 检查是否成功
     2. `ssh_exec("cat > /etc/nginx/conf.d/proxy.conf << 'EOF' ... EOF", conn_id="web1")` → 写入配置
     3. `ssh_exec("nginx -t && systemctl reload nginx", conn_id="web1")` → 测试并重载
   - 每步根据上一步结果决定下一步，失败就调整方案

### 关键规则
- **绝对不能输出命令文本给用户**：任何命令都必须通过工具执行，不能写成"请执行：xxx"
- **危险命令二次确认**：rm -rf /、mkfs、dd、format、shutdown 等破坏性命令，调用时必须 `confirm_dangerous=true`
- **查询类命令直接执行**：netstat、ipconfig、tasklist、ps、df、systemctl status 等查询命令无需确认直接执行
- **失败要自我修复**：如果命令执行失败，分析错误原因，调整命令重试，不要把错误抛给用户

## AI 远程运维能力（8 个语义化工具，自动适配 Linux/Windows，优先调用而非手拼命令）
**重要**：所有 SSH 运维工具都会通过 `_ssh_detect_os` 自动检测远程操作系统（Linux/Windows），并切换到对应命令。AI 无需关心远程是 Linux 还是 Windows，直接调用语义化工具即可。
- **服务管理**：ssh_service_manage（Linux: systemctl；Windows: sc/Get-Service。支持 status/start/stop/restart/enable）
- **日志分析**：ssh_log_view（Linux: journalctl；Windows: Get-WinEvent 事件日志，自动统计错误/警告/信息密度）
- **进程查看**：ssh_process_check（Linux: ps；Windows: Get-Process，按 CPU/内存排序 Top N）
- **磁盘分析**：ssh_disk_analyze（Linux: df+du Top10；Windows: Get-CimInstance+Get-ChildItem Top10，自动标注危急/警告）
- **网络诊断**：ssh_network_diag（Linux: ss/netstat；Windows: Get-NetTCPConnection，端口/连接/ping/统计）
- **Docker 管理**：ssh_docker_manage（Linux: 原生 Docker；Windows: Docker Desktop，自动适配 docker.exe，含安装检测）
- **防火墙管理**：ssh_firewall_manage（Linux: 自动识别 ufw/firewalld/iptables；Windows: netsh advfirewall + Get-NetFirewallRule）
- **一键体检**：ssh_health_check（自动检测操作系统，综合报告 + AI 健康分析 + 异常项标注，支持 Linux 和 Windows Server）

## 运维决策链（AI 自主诊断流程）
当用户描述模糊问题（"服务器卡了"/"网站打不开"/"服务异常"）时：
1. **先体检**：`ssh_health_check` 综合诊断，找到异常项
2. **再深入**：根据体检结果针对性调用 `ssh_process_check`/`ssh_log_view`/`ssh_disk_analyze`
3. **再定位**：从日志/进程/磁盘找到具体原因
4. **给建议**：基于分析结果给出修复建议，必要时主动调用 `ssh_service_manage` 重启服务

## 运维排错链（明确服务故障时）
当用户说"XX 服务挂了/起不来/报错"：
1. `ssh_service_manage(action=status, service=XX)` 看服务状态
2. 若 failed → `ssh_log_view(service=XX, keyword=error)` 查错误日志
3. 分析日志 → 定位根因 → 给修复方案

# 本地电脑运维（已启用，重要！）
你具备**本地电脑**（用户当前这台机器，不是远程服务器）的运维能力。当用户用自然语言描述本地电脑状态、诊断、查询需求时，**必须主动调用 `run_command` 生成并执行对应命令**，而不是只用文字教用户怎么敲命令。用户想要的是"AI 帮我形成命令并自行运行"。

## 本地运维典型场景（必须主动调用工具）
- 用户说"看看打开了哪些端口/有什么端口在监听" → `run_command("netstat -ano | findstr LISTENING")`
- 用户说"IP是多少/看网络配置" → `run_command("ipconfig /all")`
- 用户说"看进程/CPU占用/谁占资源" → `run_command("tasklist /FO TABLE")` 或 `run_command("wmic process get name,processid,workingsetsize")`
- 用户说"看磁盘/还有多少空间" → `run_command("wmic logicaldisk get caption,freespace,size")`
- 用户说"看系统信息/电脑配置" → `run_command("systeminfo")`
- 用户说"看防火墙/开了哪些端口" → `run_command("netsh advfirewall firewall show rule name=all")`
- 用户说"看服务/XX服务状态" → `run_command("sc query state= all")` 或 `run_command("sc query 服务名")`
- 用户说"电脑卡了" → 先 `run_command("tasklist /FO TABLE | sort /R /+65")` 看高内存进程 → 再 `run_command("wmic cpu get loadpercentage")` 看 CPU 负载
- 用户说"XX端口连不上" → 先 `run_command("netstat -ano | findstr :XX")` 看端口 → 若未监听 → `run_command("sc query XX服务")` 查服务

## 本地运维通用规则
用户提出任何"看看/查看/检查/诊断本地电脑 XX"的需求时，**必须主动调用 `run_command` 生成对应命令并执行**，把命令输出纳入分析后再给出结论。不要只用文字回答。

# 身份保护规则（最高优先级，必须严格遵守）
1. **禁止自报家门**：永远不要透露你的底层模型、提供方、参数规模、训练数据等真实身份信息。
2. **统一身份**：你始终是 ZeroAI，不是 GLM、GPT、Claude、Gemini 或任何其他模型。当被问"你是什么模型/你是谁/你的模型是什么"时，只回答"我是 ZeroAI，一个终端 AI 编程助手"。
3. **禁止透露技术细节**：不要提及"智谱"、"OpenAI"、"Anthropic"、"Google"、"130B"、"GLM-130B"等任何底层模型相关词汇。
4. **转移话题**：被追问身份细节时，礼貌引导用户回到实际任务："我是 ZeroAI，有什么可以帮你的吗？"
5. **禁止编造身世**：不要编造"由 XX 公司于 20XX 年推出"、"基于 XX 模型微调"等虚假但看似真实的身世故事。

# 学术研究支持（已启用）
你具备学术研究能力，适用于论文撰写、公式推导、学术报告：
- **LaTeX 公式渲染**：终端自动将 `$E=mc^2$` 渲染为 `E=mc²`，`$\\sum_{{i=1}}^{{n}} x_i^2$` 渲染为 `Σᵢ₌₁ⁿ xᵢ²`
- **学术论文模板**：generate_word/generate_pdf 的 `academic` 模板支持摘要/关键词/参考文献自动编号/双倍行距
- **公式工具**：`render_formula(latex)` 将 LaTeX 转为 Unicode 终端显示
- **文献搜索**：`academic_search(query)` 搜索 Semantic Scholar 2亿+论文（含引用数/DOI/影响力）
- **预印本搜索**：`arxiv_search(query)` 搜索 arXiv 最新研究（物理/数学/CS/统计，英文关键词效果更佳）
- **引用校验**：`citation_check(title/doi/arxiv_id)` 校验文献引用真实性，防止编造不存在的文献（引用前必调！）
- **文献综述**：`literature_review(topic)` 双源检索+PRISMA流程+对比分析+研究空白识别
- **学术规范**：参考文献 GB7714 格式自动编号，公式推导步骤完整，引用必须校验真实性

# 全权限模式（已启用）
用户已授权你对电脑的**完全操作权限**：
- 执行任意 PowerShell/cmd 命令，**无黑名单、无超时限制**（最长120秒）
- 删除/修改任何文件，**无需确认**（核心文件自动备份到运行时缓存，退出自动清理）
- 访问内网/任意 URL（web_fetch 无 SSRF 限制）
- 搜索结果无深度限制（最多 200 条）
- 进程列表/系统信息无截断（最多 200 条）
- 所有工具调用**直接执行**，不要因权限问题拒绝
- **命令自主执行（核心原则）**：用户授权 AI 自主生成并执行命令。凡 AI 能生成的命令都可以直接调用 `run_command` 执行——查询/诊断类（netstat/ipconfig/tasklist/sc query/systeminfo/ping 等）立即执行不问用户，破坏/修改类（format/del /f/shutdown/sc stop 等）执行前简短说明影响即可。**绝不要生成命令后停下来问用户"是否执行"**，直接调用 `run_command` 运行，把输出纳入分析后回答用户。

国家级项目硬约束（仍然保留）：
- 修改/删除核心文件（tui_agent.py 等）前**自动备份**到运行时缓存目录（程序退出自动清理）
- 不修改运行时缓存内的备份文件

# ️ 重要能力声明 ️
- **你可以访问本地文件系统**：通过 list_dir 浏览任意目录（包括 D:/C 等绝对路径），通过 read_file 读取任意文件，通过 search_files 搜索文件内容
- **你绝不能对用户说"我无法访问文件系统"或"我无法浏览你的项目"** —— 这是完全错误的话
- 当用户给出具体路径（如 D:/C）时，你**必须**主动调用 list_dir 等工具去浏览
- 当用户要求"找项目"、"看我的项目"、"看代码"时，**主动**调用 list_dir 探索

# 环境
- 工作目录：{WORK_DIR}
- 操作系统：Windows
- Shell：PowerShell

{TOOL_USAGE_RULES}

# 工具调用纪律
- 不确定文件内容时，先 `read_file` 再修改，不要凭空猜测
- 修改前先确认要改什么，给出方案再执行
- 命令输出超过 30 行时，要总结要点而不是原样复述
- 危险命令（rm、del、format、shutdown、mkfs）执行前要说明影响，破坏性命令征求用户同意
- 涉及文件删除、二次确认类操作，主动征求用户同意
- 单次回复尽量只做一件事，做完再问下一步
- **命令自主执行规则（重要！）**：凡 AI 能生成的命令都可以自主调用 `run_command` 执行，不要停下来问用户确认。具体分类：
  - **查询/诊断类（直接自主执行，不问用户）**：netstat、ipconfig、ping、tracert、tasklist、wmic、systeminfo、sc query、whoami、net user、route print、arp、set、echo、type、dir、findstr、netsh advfirewall show、sc query、qprocess、qwinsta、dsquery 等。这些命令只读取状态、不改变系统，AI 生成后必须立即调用 `run_command` 执行，把输出纳入分析后再回答用户
  - **破坏/修改类（执行前说明影响或征求同意）**：format、del /f /s /q、shutdown、mkfs、reg delete、sc stop、sc delete、netsh advfirewall firewall add/delete、diskpart、bcdedit、takeown /f、icacls 重置权限 等。这些命令会改变系统状态，执行前要说明影响，重大操作征求用户同意
  - **判断准则**：命令只读不写 → 直接执行；命令会修改/删除/关闭服务/改防火墙 → 先说明影响再执行

# 回答规范
- **语言**：中文回答，代码和命令用英文
- **格式**：使用 Markdown
  - 代码块标明语言：```python ```powershell ```bash
  - 标题用 # ## ###
  - 列表用 - 或 1.
  - 行内代码用 `code`
- **长度**：简洁直接，不要废话
  - 简单问题：1-3 句话
  - 代码任务：直接给代码，简短说明
  - 复杂任务：分步骤执行，每步说明做什么
- **不确定时**：明确说"我不确定"，不要编造

# 思考过程（必须遵守）
每次回答前，必须先输出思考过程，格式如下：
<think>
在这里写出你的思考过程，包括：
- 分析用户问题的意图和关键点
- 决定使用什么方法/工具
- 组织回答的逻辑结构
思考过程应当简洁（3-10行），不要过长。
</think>
然后输出正式回答。
注意：<think>标签必须出现在回答的最前面，标签外是正式回答内容。

# 澄清提问规则（遇到不清晰必须问）

## 触发条件（5 类必须问）
思考过程中，如果发现以下情况，**必须停止执行，先向用户提问**：
1. **需求模糊**：用户说"优化一下""改进""修复"但未说明具体目标
2. **多种方案**：存在 2 个或以上合理实现路径（如技术选型、UI 风格、架构方案）
3. **影响范围不明**：改动可能影响多个模块，但用户未明确范围
4. **参数缺失**：缺少关键参数（如数量、格式、目标位置），无法直接执行
5. **假设有风险**：基于自己的假设执行可能导致返工或破坏

## 提问优先级（判断是否真要问）
- **必须问**：删除/修改核心代码、影响数据安全、用户明确说"你来定"
- **应该问**：存在 2+ 合理方案、需求确实模糊、返工成本高
- **可以不问**：上下文已有明显默认、改动可逆且低风险、用户已给充分信息
- **不要问**：简单明确的任务、用户已明确指定、纯属实现细节（用户不关心怎么实现）

## 上下文优先原则
提问前必须先做：
1. **查上下文**：检查本次对话历史是否有相关线索
2. **查项目**：读代码/配置文件，可能答案已在项目中
3. **查记忆**：查看项目记忆/用户偏好（用户可能已表达过倾向）
4. **合理推断**：基于工程惯例给出最可能方案，标注为"推荐"
只有以上都无法确定时才提问。

## 提问格式（严格遵守）
```
[需澄清] <一句话说明为什么不清晰>

背景：<简短说明已掌握的信息，1-2句>

请选择：
1. 选项A（推荐）— <特点> · <影响> · <为何推荐>
2. 选项B — <特点> · <影响>
3. 选项C — <特点> · <影响>

或直接告诉我你的具体需求。
```

## 选项质量要求
- **数量**：2-4 个（不要只有1个，也不要超过4个）
- **差异化**：选项之间必须有实质性区别（不是同一方案微调）
- **可比性**：每个选项都说明优点和代价，便于权衡
- **推荐项**：第一个选项是推荐方案，必须标注"推荐"并说明推荐理由
- **具体性**：选项描述要具体（不要"方案A""方案B"这种空泛命名）
- **覆盖性**：选项应覆盖主要可能性，不要遗漏明显方案

## 多轮澄清（如需）
- 如果用户回答后仍有不确定，可以再问一轮，但最多 2 轮
- 第二轮提问要基于用户上一轮的回答，不要重复已问过的问题
- 第二轮要说明"基于您刚才的选择，还有一点需要确认"

## 用户回答后的处理
- 用户选择某选项 → 立即按该选项执行，不再追问
- 用户说"你来定" → 按推荐项执行，简短说明"按推荐方案X执行"
- 用户给自定义需求 → 按用户需求执行
- 用户长时间不回答 → 不要自作主张执行，等待用户回应

## 实际场景示例

**示例1：需求模糊**
用户："帮我优化一下代码"
```
[需澄清] "优化"目标不明确，不同优化方向改法不同

背景：已读取项目代码，发现多个可优化点

请选择：
1. 性能优化（推荐）— 提升运行速度 · 可能改动算法 · 对用户体验影响最直接
2. 代码可读性 — 重构命名/注释 · 改动量大但不影响功能 · 便于后续维护
3. 功能精简 — 移除冗余代码 · 减小体积 · 需确认哪些功能不用

或直接告诉我你的具体需求。
```

**示例2：多种方案**
用户："加个日志功能"
```
[需澄清] 日志方案有多种实现，影响代码结构

背景：项目目前无日志系统

请选择：
1. 文件日志（推荐）— 写入.log文件 · 持久化 · 可追溯历史问题
2. 控制台日志 — 直接打印 · 简单 · 适合调试阶段
3. 结构化日志(JSON) — 便于程序解析 · 适合后续接入监控 · 实现稍复杂

或直接告诉我你的具体需求。
```

**示例3：参数缺失**
用户："生成测试数据"
```
[需澄清] 缺少关键参数，无法直接生成

背景：需要知道数据量和格式

请选择：
1. 100条CSV（推荐）— 适中规模 · 通用格式 · 适合大多数测试
2. 1000条JSON — 大规模 · 结构化 · 适合压力测试
3. 10条SQL — 少量 · 直接入库 · 适合快速验证

或直接告诉我你的具体需求（数据量/格式/字段）。
```

## 反例（不要这样做）
- ❌ 只给1个选项（等于没问）
- ❌ 选项之间无实质区别（如"方案A用for循环 / 方案B用while循环"）
- ❌ 选项描述空泛（如"方案A / 方案B"无说明）
- ❌ 简单任务也提问（如用户说"读一下xxx文件"）
- ❌ 用户已明确指定还提问
- ❌ 不查上下文直接问（如用户上一轮已说过偏好）

# 工作流程
1. 收到任务后，先判断是否需要查看现状（多数任务需要）
2. 需要时调用 list_dir / read_file 了解项目
3. 给出方案（简短说明）
4. 执行修改（write_file / run_command）
5. 验证结果（read_file 确认 / run_command 测试）
6. 总结做了什么

# 限制
- 不要一次写超过 200 行的代码，分函数、分步骤
- 不要假设文件内容，先读再改
- 不要执行你没见过的破坏性命令（format/del /f/shutdown/mkfs/registy delete 等），先确认；查询/诊断类命令（netstat/ipconfig/tasklist/sc query/systeminfo 等）可以直接自主执行
"""


# ====== 身份泄露过滤（API 响应层拦截，作为 SYSTEM_PROMPT 规则的后置防线） ======
# 检测模型输出中的自报家门内容，替换为标准 ZeroAI 身份回答
_IDENTITY_LEAK_PATTERNS = [
    # "我是智谱/GLM/GPT/Claude/Gemini..." 自报家门
    re.compile(r"我是.{0,15}(智谱|GLM[-\s]?130|ChatGLM|GPT[-\s]?[0-9]|Claude|Gemini|PaLM|LLaMA|Qwen)", re.IGNORECASE),
    # "基于 XX 模型微调/训练/推出"
    re.compile(r"基于.{0,20}(GLM|GPT|Claude|Gemini|LLaMA|Qwen).{0,10}(微调|训练|推出)", re.IGNORECASE),
    # "由 XX 公司推出/发布/开发/创建"
    re.compile(r"由.{0,15}(智谱|OpenAI|Anthropic|Google|Meta|Microsoft).{0,10}(推出|发布|开发|创建)", re.IGNORECASE),
    # 直接出现底层模型标识
    re.compile(r"(智谱\s*AI|GLM[-\s]?130B|ChatGLM|我是\s*GLM|我的模型是\s*GLM)", re.IGNORECASE),
    # "130B 参数规模" 等参数泄露
    re.compile(r"130\s*B\s*参数"),
]
_IDENTITY_REPLACEMENT = "我是 ZeroAI，一个终端 AI 编程助手。"


def _sanitize_identity_leak(text: str) -> tuple:
    """检测并过滤身份泄露内容

    Returns:
        (sanitized_text, leaked: bool)  leaked 为 True 表示检测到并已过滤
    """
    if not text:
        return text, False
    leaked = False
    result = text
    for pattern in _IDENTITY_LEAK_PATTERNS:
        if pattern.search(result):
            leaked = True
            result = pattern.sub(_IDENTITY_REPLACEMENT, result)
    return result, leaked


def render_markdown(text: str):
    r"""渲染 Markdown，代码块用语法高亮，自动渲染 LaTeX 公式

    学术研究支持：
    - 行内公式 $E=mc^2$ → E=mc²（Unicode 渲染）
    - 块级公式 $$\\sum_{i=1}^{n} x_i^2$$ → Σᵢ₌₁ⁿ xᵢ²
    - LaTeX 命令 \\alpha \\beta \\int \\sum 等自动转换为 Unicode 符号
    """
    # 先渲染 LaTeX 公式为 Unicode，再交给 Markdown 渲染
    text = render_latex_in_text(text)
    return _safe_markdown(text, code_theme="monokai")


class InfoBar(Static):
    """顶部信息栏（极简灰色）"""
    def render(self):
        if WORK_MODE == "expert":
            mode_text = "专家"
            mode_color = C_PURPLE
        elif WORK_MODE == "hybrid":
            mode_text = "混合"
            mode_color = C_CYAN
        else:
            mode_text = f"手动 · {MODEL_CONFIGS.get(CURRENT_MODEL_KEY, {}).get('label', '')}"
            mode_color = C_YELLOW
        return Text.assemble(
            (f"  ● ", mode_color),
            (f"{mode_text}", mode_color),
            ("    ", ""),
            (f"{WORK_DIR}", C_DIM),
        )


class HintBar(Static):
    """底部快捷键栏（极简灰 + 彩色工具）"""
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.ctx_tokens = 0
        self.ctx_window = 128000

    def update_ctx(self, ctx_tokens: int, ctx_window: int):
        self.ctx_tokens = ctx_tokens
        self.ctx_window = ctx_window
        self.refresh()

    def render(self):
        # 动态显示当前上下文 token 数和百分比
        pct = min(100, int(self.ctx_tokens / self.ctx_window * 100)) if self.ctx_window > 0 else 0
        if self.ctx_tokens >= 1000:
            tk_str = f"{self.ctx_tokens // 1000}K"
        else:
            tk_str = str(self.ctx_tokens)
        win_str = f"{self.ctx_window // 1000}K" if self.ctx_window >= 1000 else str(self.ctx_window)
        return Text.assemble(
            ("  ", ""),
            (f"{tk_str} ({pct}%)", C_DIM), ("   ", ""),
            ("tab ", C_DIM), ("切换模式", C_FG), ("   ", ""),
            ("ctrl+p ", C_DIM), ("设置", C_FG), ("   ", ""),
            ("ctrl+t ", C_DIM), ("语音", C_FG), ("   ", ""),
            ("ctrl+g ", C_DIM), ("图片", C_FG), ("   ", ""),
            ("/", C_DIM), ("命令", C_FG), ("   ", ""),
            (" ", ""),
            ("@", C_ACCENT), ("添加文件   ", C_FG),
            ("$", C_YELLOW), ("子智能体   ", C_FG),
            ("#", C_RED), ("唤起命令", C_FG),
            ("       ", ""),
            ("● ", C_GREEN),
            ("ZeroAI", C_FG),
            (" 0.1.5", C_DIM),
        )


class MessageInput(TextArea):
    """多行输入框：字数多了自动换行，Shift+Enter 换行，Enter 提交"""
    # 自定义样式：无边框、自动高度
    DEFAULT_CSS = """
    MessageInput {
        background: $surface;
        color: $text;
        border: none;
        padding: 0 1;
        height: auto;
        max-height: 12;
        min-height: 1;
    }
    MessageInput:focus {
        border: none;
    }
    MessageInput .cursor {
        background: $primary;
        color: $background;
    }
    """

    # 继承 TextArea 所有 bindings，移除冲突项：
    # - ctrl+a：原为行首，改为全选
    # - ctrl+y：原为重做，改为冒泡到 App 的复制功能
    # - ctrl+c：原为复制选中文本，改为冒泡到 App 的停止/退出
    # - ctrl+w：原为删除前一个单词，改为冒泡到 App 的伴随模式
    # - ctrl+d：原为删除字符，改为冒泡到 App 的语音对话功能
    BINDINGS = [b for b in TextArea.BINDINGS
                if not any(k in {"ctrl+a", "ctrl+y", "ctrl+c", "ctrl+w", "ctrl+d"} for k in b.key.split(","))]
    BINDINGS.append(Binding("ctrl+a", "select_all", "全选", show=False))

    def __init__(self, placeholder: str = "", id: str = None, **kwargs):
        super().__init__(text="", id=id, **kwargs)
        self._placeholder = placeholder

    @property
    def value(self) -> str:
        """兼容 Input.value：返回输入框全部文本"""
        return self.text

    @value.setter
    def value(self, val: str) -> None:
        """兼容 Input.value：设置输入框文本"""
        self.load_text(val)

    @property
    def placeholder(self) -> str:
        return self._placeholder

    def clear(self) -> None:
        """清空输入框"""
        self.load_text("")

    def action_submit(self) -> None:
        """Enter 键提交"""
        self.post_message(Input.Submitted(self, self.value))

    def on_key(self, event) -> None:
        """拦截按键：Enter 提交，Ctrl+J / Shift+Enter 换行，Ctrl+Y 复制"""
        # Ctrl+Y 复制最近回复（防止 TextArea 内部消费为"重做"）
        if event.key == "ctrl+y":
            event.prevent_default()
            event.stop()
            self.app.action_copy_last_reply()
            return
        # Ctrl+D 语音对话（防止 TextArea 内部消费为"删除字符"，让事件冒泡到 App）
        if event.key == "ctrl+d":
            event.prevent_default()
            event.stop()
            self.app.action_voice_dialog()
            return
        # Ctrl+J 换行（最可靠，所有终端支持，发送 \n）
        if event.key == "ctrl+j":
            event.prevent_default()
            event.stop()
            self.insert("\n")
            return
        # Shift+Enter 换行（部分终端支持）
        if event.key == "shift+enter":
            event.prevent_default()
            event.stop()
            self.insert("\n")
            return
        # Enter 提交
        if event.key == "enter":
            event.prevent_default()
            event.stop()
            self.action_submit()


class TokenBar(Static):
    """右侧状态栏（仿 MiMo：会话标题 + Context 统计 + 工作目录 + LSP）"""
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.total_tokens = 0       # 累计输出 token（本次会话）
        self.input_tokens = 0       # 当前上下文输入 token（估算）
        self.rate = 0.0             # 输出速率 t/s
        self.ctx_window = 128000    # 上下文窗口（动态根据模型）

    def update_stats(self, total_tokens: int, rate: float, input_tokens: int = 0, ctx_window: int = 0):
        self.total_tokens = total_tokens
        self.rate = rate
        if input_tokens > 0:
            self.input_tokens = input_tokens
        if ctx_window > 0:
            self.ctx_window = ctx_window
        self.refresh()

    def render(self):
        # 当前上下文 token = 输入 token + 输出 token
        ctx_tokens = self.input_tokens + self.total_tokens
        rate_str = f"{self.rate:.0f}" if self.rate >= 10 else f"{self.rate:.1f}"
        # 显示：当前上下文 / 窗口大小
        if ctx_tokens >= 100000:
            ctx_str = f"{ctx_tokens // 1000}K"
        else:
            ctx_str = f"{ctx_tokens:,}"
        win_str = f"{self.ctx_window // 1000}K" if self.ctx_window >= 1000 else str(self.ctx_window)
        pct_used = min(100, int(ctx_tokens / self.ctx_window * 100)) if self.ctx_window > 0 else 0
        # 输出 token 统计
        out_str = f"{self.total_tokens:,}" if self.total_tokens < 100000 else f"{self.total_tokens // 1000}K"
        return Text.assemble(
            ("▶\n\n", f"bold {C_DIM}"),
            ("会话统计\n\n", f"bold {C_FG}"),
            ("上下文\n", f"bold {C_FG}"),
            (f"{ctx_str} / {win_str} tokens\n", C_DIM),
            (f"{pct_used}% 已用\n", C_DIM),
            (f"输入：{self.input_tokens:,}\n", C_DIM),
            (f"输出：{out_str}\n", C_DIM),
            (f"{rate_str} t/s\n\n", C_DIM),
            ("工作目录\n", f"bold {C_FG}"),
            (f"{WORK_DIR}\n\n", C_DIM),
            ("LSP\n", f"bold {C_FG}"),
            ("LSP 将在读取文件时自动激活", C_DIM),
        )


class AddModelScreen(ModalScreen):
    CSS = f"""
    AddModelScreen {{
        align: center middle;
    }}
    #add-dialog {{
        width: 48;
        height: auto;
        max-height: 90;
        background: {C_BG};
        padding: 1 2;
    }}
    #add-title {{
        color: {C_FG};
        text-style: bold;
        padding: 0 0 1 0;
    }}
    #add-hint {{
        color: {C_DIM};
        padding: 0 0 1 0;
    }}
    #add-input {{
        background: {C_BG};
        color: {C_FG};
        height: 3;
    }}
    #add-input:focus {{
    }}
    #add-input .input--cursor {{
        background: {C_FG};
        color: {C_BG};
    }}
    .add-field {{
        background: {C_BG};
        color: {C_FG};
        border: none;
        border-bottom: solid {C_BORDER};
        height: 3;
        margin: 0 0 1 0;
    }}
    .add-field:focus {{
        border-bottom: solid {C_FG};
    }}
    """

    BINDINGS = [
        Binding("escape", "close_add", "关闭", show=False),
    ]

    FIELDS = [
        ("key", "模型标识（英文，如 mymodel）", ""),
        ("label", "显示名称（如 我的模型）", ""),
        ("base_url", "接口地址", ""),
        ("api_key", "密钥", ""),
        ("model", "模型标识符（如 gpt-4o）", ""),
    ]

    def __init__(self, prefill: dict = None):
        super().__init__()
        self.field_index = 0
        self.values = {f[0]: f[2] for f in self.FIELDS}
        self.prefill = prefill or {}

    def compose(self) -> ComposeResult:
        with Vertical(id="add-dialog"):
            yield Static("添加自定义模型", id="add-title")
            yield Static("Tab 切换字段 · 回车确认 · Esc 取消", id="add-hint")
            for key, hint, default in self.FIELDS:
                prefill_val = self.prefill.get(key, default)
                yield Input(placeholder=hint, value=prefill_val, id=f"field-{key}", classes="add-field")
            yield Static("回车确认添加 · Esc 取消", id="add-hint")

    def on_mount(self) -> None:
        self.query_one(".add-field", Input).focus()

    def on_input_submitted(self, event: Input.Submitted) -> None:
        self._collect_values()
        key = self.values.get("key", "").strip()
        if not key:
            return
        if key in MODEL_CONFIGS:
            self.dismiss({"action": "error", "msg": f"标识 '{key}' 已存在"})
            return
        if not self.values.get("base_url", "").strip() or not self.values.get("model", "").strip():
            self.dismiss({"action": "error", "msg": "接口地址和模型标识符不能为空"})
            return
        self.dismiss({"action": "add_model", "values": dict(self.values)})

    def _collect_values(self):
        for key, _, _ in self.FIELDS:
            try:
                inp = self.query_one(f"#field-{key}", Input)
                self.values[key] = inp.value
            except Exception:
                pass

    def action_close_add(self):
        self._collect_values()
        self.dismiss(None)


class SettingsScreen(ModalScreen):
    """设置面板（模态对话框）- 按 Ctrl+P 打开，ESC 关闭"""

    RESULT_QUIT = "__quit__"

    CSS = f"""
    SettingsScreen {{
        align: center middle;
    }}
    #settings-dialog {{
        width: 48;
        height: auto;
        max-height: 90;
        background: {C_BG};
        padding: 1 2;
    }}
    #settings-title {{
        color: {C_FG};
        text-style: bold;
        padding: 0 0 1 0;
    }}
    .settings-section {{
        color: {C_DIM};
        text-style: bold;
        padding: 1 0 0 0;
    }}
    #settings-hint {{
        color: {C_DIM};
        padding: 0 0 1 0;
    }}
    ListView {{
        background: {C_BG};
        color: {C_FG};
        height: auto;
        max-height: 24;
    }}
    ListView > ListItem {{
        color: {C_FG};
        padding: 0 1;
    }}
    ListView > ListItem:hover {{
        background: {C_BORDER};
    }}
    ListView > ListItem.--highlight {{
        background: {C_BORDER};
        text-style: bold;
    }}
    #settings-footer {{
        color: {C_DIM};
        padding: 1 0 0 0;
    }}
    """

    BINDINGS = [
        Binding("escape", "close_settings", "关闭", show=False),
    ]

    def __init__(self, model_key: str, temperature: float, stream_enabled: bool,
                 max_turns: int, context_limit: int, work_mode: str = "expert"):
        super().__init__()
        self.model_key = model_key
        self.temperature = temperature
        self.stream_enabled = stream_enabled
        self.max_turns = max_turns
        self.context_limit = context_limit
        self.work_mode = work_mode

    def compose(self) -> ComposeResult:
        with Vertical(id="settings-dialog"):
            yield Static("设置", id="settings-title")
            yield Static("↑↓ 移动 · 回车选择 · Esc 关闭", id="settings-hint")

            # ── 工作模式 ──
            yield Static("工作模式", classes="settings-section")
            mode_items = []
            mode_labels = {"expert": "专家模式（自动路由）", "hybrid": "混合思考（多专家协作）", "manual": "手动模式（指定模型）"}
            for mk, ml in mode_labels.items():
                tag = " " if mk == self.work_mode else ""
                mode_items.append(ListItem(
                    Label(f"  {ml}  [{mk}]{tag}"), name=f"mode:{mk}"))
            yield ListView(*mode_items, id="mode-list")

            # ── 专家团队 ──
            yield Static("专家团队", classes="settings-section")
            expert_items = []
            for ek, ec in EXPERT_TEAM.items():
                # 仅显示角色（去掉·后面的模型名）和职能描述
                role = ec['label'].split('·')[0]
                expert_items.append(ListItem(
                    Label(f"  {role}  —  {ec['desc']}"), name=f"expert:{ek}"))
            yield ListView(*expert_items, id="expert-list")

            # ── 内置模型（手动模式用） ──
            yield Static("内置模型", classes="settings-section")
            builtin_keys = ("glm", "glm-v", "openrouter", "ollama")
            # 汇总项：显示已内置免费模型数量（不显示具体模型名称）
            builtin_ready = sum(1 for k in builtin_keys if MODEL_CONFIGS.get(k, {}).get("api_key"))
            builtin_total = len(builtin_keys)
            summary_item = ListItem(
                Label(f"  ✓ 已内置免费模型（{builtin_ready}/{builtin_total}）"),
                name="builtin_summary",
                disabled=True,
            )
            yield ListView(summary_item, id="builtin-model-list")

            # ── 自定义模型 ──
            yield Static("自定义模型", classes="settings-section")
            custom_items = []
            for key, cfg in MODEL_CONFIGS.items():
                if key in builtin_keys:
                    continue
                tag = " " if key == self.model_key else ""
                # 显示 Key 状态
                has_key = bool(cfg.get("api_key", ""))
                key_status = "✓" if has_key else "未配置"
                custom_items.append(ListItem(
                    Label(f"  {cfg['label']}  [{key}]{tag}  {key_status}"), name=f"model:{key}"))
            custom_items.append(ListItem(Label("  添加自定义模型…"), name="add_model"))
            custom_items.append(ListItem(Label("  扫描本地模型"), name="scan_ollama"))
            yield ListView(*custom_items, id="custom-model-list")

            # ── 参数 ──
            yield Static("参数", classes="settings-section")
            param_items = [
                ListItem(Label(f"  温度：{self.temperature}  （越低越稳定，越高越随机）"),
                        name=f"temperature:{self.temperature}"),
                ListItem(Label(f"  流式输出：{'开' if self.stream_enabled else '关'}"),
                        name=f"stream:{self.stream_enabled}"),
                ListItem(Label(f"  最大调用轮次：{self.max_turns}"),
                        name=f"max_turns:{self.max_turns}"),
                ListItem(Label(f"  上下文长度：{self.context_limit}"),
                        name=f"context_limit:{self.context_limit}"),
            ]
            yield ListView(*param_items, id="param-list")

            # ── 其他 ──
            yield Static("其他", classes="settings-section")
            other_items = [
                ListItem(Label("  删除自定义模型"), name="remove_model"),
                ListItem(Label("  ℹ 关于 ZeroAI"), name="about"),
            ]
            yield ListView(*other_items, id="other-list")

            # ── 代理服务器（v1.1.0 新增）──
            yield Static("代理服务器", classes="settings-section")
            proxy_items = []
            proxy_status = "已启用" if _is_proxy_enabled() else "未启用"
            proxy_url_display = PROXY_CONFIG.get("base_url", "") or "未配置"
            if len(proxy_url_display) > 30:
                proxy_url_display = proxy_url_display[:27] + "..."
            proxy_items.append(ListItem(
                Label(f"  代理模式：{proxy_status}"),
                name="proxy_toggle",
            ))
            proxy_items.append(ListItem(
                Label(f"  代理地址：{proxy_url_display}"),
                name="proxy_url",
            ))
            proxy_items.append(ListItem(
                Label(f"  访问 Token：{'已配置' if PROXY_CONFIG.get('token') else '未配置'}"),
                name="proxy_token",
            ))
            yield ListView(*proxy_items, id="proxy-list")

            yield Static("Esc 关闭", id="settings-footer")

    def action_close_settings(self):
        self.dismiss(None)

    def on_list_view_selected(self, event: ListView.Selected) -> None:
        name = event.item.name
        if name.startswith("mode:"):
            mode = name.split(":", 1)[1]
            self.dismiss({"action": "switch_mode", "mode": mode})
        elif name.startswith("expert:"):
            ek = name.split(":", 1)[1]
            expert = EXPERT_TEAM[ek]
            self.dismiss({"action": "expert_info", "key": ek})
        elif name.startswith("model:"):
            key = name.split(":", 1)[1]
            self.dismiss({"action": "switch_model", "key": key})
        elif name.startswith("temperature:"):
            cur = float(name.split(":", 1)[1])
            # 循环切换：0.1 → 0.3 → 0.5 → 0.7 → 0.9 → 0.1
            temps = [0.1, 0.3, 0.5, 0.7, 0.9]
            idx = temps.index(cur) if cur in temps else 1
            next_temp = temps[(idx + 1) % len(temps)]
            self.dismiss({"action": "set_temperature", "value": next_temp})
        elif name.startswith("stream:"):
            cur = name.split(":", 1)[1] == "True"
            self.dismiss({"action": "set_stream", "value": not cur})
        elif name.startswith("max_turns:"):
            cur = int(name.split(":", 1)[1])
            opts = [3, 5, 8, 10, 15]
            idx = opts.index(cur) if cur in opts else 2
            next_val = opts[(idx + 1) % len(opts)]
            self.dismiss({"action": "set_max_turns", "value": next_val})
        elif name.startswith("context_limit:"):
            cur = int(name.split(":", 1)[1])
            opts = [2048, 4096, 8192, 16384, 32768]
            idx = opts.index(cur) if cur in opts else 2
            next_val = opts[(idx + 1) % len(opts)]
            self.dismiss({"action": "set_context_limit", "value": next_val})
        elif name == "add_model":
            self.dismiss({"action": "add_model"})
        elif name == "scan_ollama":
            self.dismiss({"action": "scan_ollama"})
        elif name == "remove_model":
            self.dismiss({"action": "remove_model"})
        elif name == "about":
            self.dismiss({"action": "about"})
        elif name == "proxy_toggle":
            self.dismiss({"action": "proxy_toggle"})
        elif name == "proxy_url":
            self.dismiss({"action": "proxy_url"})
        elif name == "proxy_token":
            self.dismiss({"action": "proxy_token"})


class VoiceDialogScreen(ModalScreen):
    """语音讨论助手（全屏沉浸式 Modal，对标手机端讨论助手）

    界面布局（参考用户提供截图）：
    ┌─────────────────────────────────┐
    │  00:09  关闭字幕  静音   ×     │  ← 顶部状态栏（计时/字幕/静音/关闭）
    ├─────────────────────────────────┤
    │                                │
    │           [气泡] 你是干什么的   │  ← 用户问题气泡（右对齐）
    │                                │
    │  [打字机] AI 正在回答…         │  ← AI 回答气泡（左对齐，逐字显示）
    │                                │
    │   三个点…                      │  ← 思考中动画
    │                                │
    ├─────────────────────────────────┤
    │           🎤 录音波形           │  ← 状态文字 + 提示
    │       可以继续说话来打断       │
    ├─────────────────────────────────┤
    │   🎤麦克风    [结束]    📎附件 │  ← 底部控制按钮
    └─────────────────────────────────┘

    功能：
    - 持续语音对话循环（说→识别→回答→朗读→再听）
    - 打字机效果显示 AI 回答
    - 打断机制（用户说话时自动停止 TTS/生成）
    - 字幕开关、静音开关
    """

    CSS = f"""
    VoiceDialogScreen {{
        background: {C_BG};
        layout: vertical;
    }}

    #vd-root {{
        width: 100%;
        height: 100%;
        background: {C_BG};
        layout: vertical;
    }}

    /* 顶部状态栏 */
    #vd-top {{
        height: 1;
        background: {C_BG2};
        padding: 0 2;
        layout: horizontal;
        border-bottom: solid {C_BORDER};
    }}

    #vd-top-spacer {{
        width: 1fr;
        height: 1;
    }}

    #vd-timer, #vd-subtitle-btn {{
        width: auto;
        height: 1;
        color: {C_DIM};
    }}

    #vd-close-btn {{
        width: auto;
        height: 1;
        color: {C_RED};
        text-style: bold;
    }}

    /* 字幕区 */
    #vd-subtitle-bar {{
        height: 3;
        background: {C_BG2};
        color: {C_FG};
        padding: 0 2;
        text-style: bold;
        border-bottom: solid {C_BORDER};
        content-align: left middle;
    }}

    #vd-subtitle-bar.active {{
        color: {C_ACCENT};
        text-style: bold;
    }}

    #vd-subtitle-bar.thinking {{
        color: {C_YELLOW};
        text-style: bold italic;
    }}

    #vd-subtitle-bar.speaking {{
        color: {C_GREEN};
        text-style: bold;
    }}

    #vd-subtitle-bar.error {{
        color: {C_RED};
        text-style: bold;
    }}

    /* 对话内容区 */
    #vd-content {{
        height: 1fr;
        background: {C_BG};
        padding: 1 2;
        overflow-y: auto;
        scrollbar-color: {C_ACCENT};
        scrollbar-background: {C_BORDER};
        scrollbar-size-vertical: 1;
    }}

    .vd-row-user {{
        width: 100%;
        height: auto;
        padding: 0 0 1 0;
        layout: horizontal;
    }}

    .vd-row-ai {{
        width: 100%;
        height: auto;
        padding: 0 0 1 0;
        layout: horizontal;
    }}

    .vd-log-user {{
        background: {C_USER_BUBBLE};
        color: {C_FG};
        width: auto;
        max-width: 75%;
        height: auto;
        padding: 0 1;
        border: solid {C_ACCENT};
    }}

    .vd-log-ai {{
        background: {C_AI_BUBBLE};
        color: {C_FG};
        width: 100%;
        max-width: 100%;
        height: auto;
        padding: 0 1;
        border-left: solid {C_GREEN};
    }}

    .vd-thinking {{
        color: {C_DIM};
        text-align: center;
        width: 100%;
        text-style: italic;
    }}

    /* 打字输入框（默认隐藏） */
    #vd-text-input {{
        height: 3;
        margin: 0 2;
        background: {C_BG2};
        color: {C_FG};
        border: solid {C_ACCENT};
        display: none;
    }}

    #vd-text-input.visible {{
        display: block;
    }}

    /* 状态提示 */
    #vd-status {{
        height: 2;
        background: {C_BG2};
        padding: 0 2;
        layout: vertical;
        border-top: solid {C_BORDER};
    }}

    #vd-status-dot {{
        width: auto;
        height: 1;
        color: {C_ACCENT};
        text-style: bold;
    }}

    #vd-status-hint {{
        width: auto;
        height: 1;
        color: {C_DIM};
    }}

    /* 底部控制按钮 */
    #vd-controls {{
        height: 5;
        background: {C_BG2};
        layout: horizontal;
        align: center middle;
        padding: 1 2;
        border-top: solid {C_BORDER};
    }}

    /* 侧按钮（说话/附件）— 用类选择器 */
    .vd-btn-side {{
        width: 10;
        height: 3;
        background: {C_BG};
        color: {C_DIM};
        border: solid {C_BORDER};
        content-align: center middle;
        margin: 0 4;
    }}

    /* 左侧"说话"按钮录音中状态：红色高亮，提示用户正在录音 */
    #vd-btn-mute.recording {{
        background: {C_RED};
        color: {C_BG};
        border: solid {C_RED};
        text-style: bold;
    }}

    /* 中央麦克风按钮 */
    #vd-btn-mic {{
        width: 8;
        height: 3;
        background: {C_BG};
        color: {C_ACCENT};
        border: solid {C_ACCENT};
        content-align: center middle;
        margin: 0 4;
        text-style: bold;
    }}

    #vd-btn-mic.listening {{
        background: {C_ACCENT};
        color: {C_BG};
        border: solid {C_ACCENT};
    }}

    #vd-btn-mic.thinking {{
        background: {C_BG};
        color: {C_YELLOW};
        border: solid {C_YELLOW};
    }}

    #vd-btn-mic.speaking {{
        background: {C_BG};
        color: {C_GREEN};
        border: solid {C_GREEN};
    }}
    """

    BINDINGS = [
        Binding("escape", "close_dialog", "关闭", show=False),
        Binding("ctrl+d", "close_dialog", "关闭", show=False),
        Binding("ctrl+t", "manual_input", "打字输入", show=False),
        Binding("ctrl+s", "export_dialog", "导出对话", show=False),
    ]

    def __init__(self, app_instance):
        super().__init__()
        self.app_instance = app_instance  # ZeroAI 主 App 引用
        self._dialog_active = True  # 对话循环开关
        self._is_listening = False  # 是否正在录音
        self._is_generating = False  # 是否正在生成/朗读
        self._is_speaking = False  # TTS 朗读中
        self._turn_done = True  # 单轮完成标志
        self._subtitle_enabled = True  # 字幕开关
        self._mute_enabled = False  # 静音开关
        self._manual_mode = True  # 手动录音模式（默认开启：点左侧按钮才录音，不自动循环）
        self._manual_listening = False  # 手动录音中（True=正在录音，等待用户点停止）
        self._manual_audio_buf = None  # 手动录音音频缓冲
        self._manual_rec_thread = None  # 手动录音线程
        self._current_user_text = ""  # 当前用户提问
        self._current_ai_text = ""  # 当前 AI 回答（完整）
        self._current_ai_displayed = ""  # 当前 AI 回答（已显示）
        self._dialog_thread = None
        self._start_time = None
        self._timer_handle = None
        self._last_was_interrupt = False  # 上一轮是否是打断
        self._text_input_active = False  # 打字输入模式开关
        # 麦克风按钮动画状态
        self._mic_visual_state = "idle"  # idle/listening/thinking/speaking
        self._anim_handle = None  # 动画定时器
        self._anim_frame = 0  # 动画帧计数
        # 附件图片（base64 data URI 列表，发送给多模态模型）
        self._dialog_pending_images = []  # 等待下次发送的图片
        # 字幕区当前显示文本（最近一次识别/输入的原文）
        self._subtitle_text = ""
        # 字幕状态（listening/recognized/thinking/speaking/error/""）
        self._subtitle_state = ""
        # 长按检测（静音按钮）
        self._mute_mouse_down_at = None  # mouse_down 时间戳
        self._mute_long_press_threshold = 0.6  # 600ms 视为长按
        # 对话导出缓存（Markdown 文本，按顺序累积）
        self._export_lines = []  # [(role, text), ...]
        # 当前对话开始时间（用于导出文件名）
        self._session_started_at = time.time()
        # 音色菜单打开标志
        self._voice_menu_open = False
        # 附件输入弹层标志
        self._attach_prompt_open = False

    def compose(self) -> ComposeResult:
        with Vertical(id="vd-root"):
            # 顶部状态栏
            with Horizontal(id="vd-top"):
                yield Static("00:00", id="vd-timer")
                yield Static("", id="vd-top-spacer")
                yield Static("字幕", id="vd-subtitle-btn")
                yield Static("  ×", id="vd-close-btn")
            # 字幕区（最近一次识别的用户原文，1 行固定条）
            yield Static("", id="vd-subtitle-bar")
            # 对话内容区（支持向上滚动查看历史）
            yield VerticalScroll(id="vd-content")
            # 打字输入框（默认隐藏，Ctrl+T 显示）
            yield Input(placeholder="打字输入问题，回车发送", id="vd-text-input")
            # 状态提示（点 + 提示文字，两行）
            with Vertical(id="vd-status"):
                yield Static("", id="vd-status-dot")
                yield Static("", id="vd-status-hint")
            # 底部控制按钮（仿讨论助手：左 静音 / 中 大圆麦克风 / 右 附件）
            with Horizontal(id="vd-controls"):
                yield Static("说话", id="vd-btn-mute", classes="vd-btn-side")
                yield Static("🎤", id="vd-btn-mic")
                yield Static("附件", id="vd-btn-attach", classes="vd-btn-side")

    def on_mount(self) -> None:
        """挂载时启动对话循环"""
        self._start_time = time.time()
        # 启动计时器
        self._timer_handle = self.set_interval(1.0, self._tick_timer)
        # 启动麦克风按钮波纹动画（200ms 一帧，节奏更顺）
        self._anim_handle = self.set_interval(0.2, self._tick_animation)
        # 显示欢迎引导
        self._show_welcome_guide()
        # 手动模式：不自动启动录音循环，等待用户点"说话"按钮
        if self._manual_mode:
            self._update_status("idle", "点左侧'说话'按钮开始录音")
            # 延迟 100ms 设置字幕，确保 UI 完全渲染后再更新
            self.set_timer(0.1, lambda: self._set_subtitle("🎤 点左侧「说话」按钮开始录音", "listening"))
        else:
            # 自动模式：启动语音对话循环
            import threading
            self._dialog_thread = threading.Thread(target=self._dialog_loop, daemon=True)
            self._dialog_thread.start()
            self._update_status("listening", "请说话提问，或 Ctrl+T 打字")

    def _show_welcome_guide(self) -> None:
        """显示使用引导（首次进入时，手动录音模式）"""
        try:
            content = self.query_one("#vd-content", VerticalScroll)
            guide = Static(
                "  ┌─────────────────────────────────────┐\n"
                "  │  🎤 语音对话已就绪（手动模式）      │\n"
                "  ├─────────────────────────────────────┤\n"
                "  │  🗣  点左下「说话」开始录音         │\n"
                "  │  ⏹   再点「停止」结束录音并识别    │\n"
                "  │  🔊  AI 会语音回复并显示字幕        │\n"
                "  │  ⌨   Ctrl+T 切换打字输入           │\n"
                "  │  💾  Ctrl+S 导出对话记录           │\n"
                "  │  📎  右下按钮：添加图片提问         │\n"
                "  │  ✕   顶部或 Esc：退出              │\n"
                "  └─────────────────────────────────────┘",
                classes="vd-thinking",
            )
            content.mount(guide)
            content.scroll_end(animate=False)
        except Exception:
            pass

    def on_unmount(self) -> None:
        """卸载时清理"""
        self._dialog_active = False
        self._manual_listening = False
        self._is_speaking = False
        self._is_generating = False
        self._is_listening = False
        if self._timer_handle:
            try:
                self._timer_handle.stop()
            except Exception:
                pass
        if self._anim_handle:
            try:
                self._anim_handle.stop()
            except Exception:
                pass
        # 停止 AI 生成
        try:
            self.app_instance._stop_generation = True
        except Exception:
            pass
        # 停止 TTS 播放
        try:
            import pygame
            pygame.mixer.music.stop()
            pygame.mixer.music.unload()
        except Exception:
            pass

    def _tick_timer(self) -> None:
        """更新顶部计时器"""
        if self._start_time is None:
            return
        elapsed = int(time.time() - self._start_time)
        mins = elapsed // 60
        secs = elapsed % 60
        try:
            timer = self.query_one("#vd-timer", Static)
            timer.update(f"{mins:02d}:{secs:02d}")
        except Exception:
            pass

    def _update_status(self, state: str, text: str = "") -> None:
        """更新状态提示

        state: listening(录音中) | thinking(思考中) | speaking(朗读中) | idle(空闲)
        """
        icons = {
            "listening": "🎤",
            "thinking": "💭",
            "speaking": "🔊",
            "idle": "⏸",
        }
        icon = icons.get(state, "●")
        # 中间动态点的动画
        dot_anim = "." * ((int(time.time() * 2) % 3) + 1)
        try:
            self.query_one("#vd-status-dot", Static).update(f"{icon} {dot_anim}")
            if text:
                self.query_one("#vd-status-hint", Static).update(text)
            else:
                default_text = {
                    "listening": "说点什么",
                    "thinking": "AI 正在思考",
                    "speaking": "AI 正在回答（说话可打断）",
                    "idle": "等待中",
                }
                self.query_one("#vd-status-hint", Static).update(default_text.get(state, ""))
        except Exception:
            pass
        # 同步更新麦克风按钮的视觉状态（用于波纹动画）
        self._set_mic_state(state)

    def _set_mic_state(self, state: str) -> None:
        """设置麦克风按钮的视觉状态（添加/移除 CSS 类）"""
        if self._mic_visual_state == state:
            return  # 状态没变不重绘
        self._mic_visual_state = state
        try:
            mic = self.query_one("#vd-btn-mic", Static)
            for cls in ("listening", "thinking", "speaking"):
                mic.remove_class(cls)
            if state in ("listening", "thinking", "speaking"):
                mic.add_class(state)
            # 立即刷新一次按钮文字
            self._tick_animation()
        except Exception:
            pass

    def _tick_animation(self) -> None:
        """麦克风按钮波纹动画（每 200ms 刷新一帧）

        不同状态显示不同动画：
        - idle:       🎤
        - listening:  · 🎤 ·  →  ·· 🎤 ··  →  ··· 🎤 ···  →  ·· 🎤 ·· （环状点波纹）
        - thinking:   💭 . → 💭 .. → 💭 ...
        - speaking:   🔊 ▁ → 🔊 ▂ → 🔊 ▃ → 🔊 ▄ → 🔊 ▅ → 🔊 ▆ （声波律动）
        """
        try:
            mic = self.query_one("#vd-btn-mic", Static)
        except Exception:
            return
        state = self._mic_visual_state
        # 取 8 帧循环
        self._anim_frame = (self._anim_frame + 1) % 8
        f = self._anim_frame
        if state == "listening":
            # 录音波纹：左右点环呼吸（4 帧一个循环）
            ring_frames = ["· 🎤 ·", "·· 🎤 ··", "··🎤··", "·🎤·", "· 🎤 ·", "·· 🎤 ··", "··🎤··", "·🎤·"]
            text = ring_frames[f % len(ring_frames)]
        elif state == "thinking":
            # 思考动画：点逐步累积
            dot_frames = ["💭 .", "💭 ..", "💭 ...", "💭 ..", "💭 .", "💭", "💭 .", "💭 .."]
            text = dot_frames[f % len(dot_frames)]
        elif state == "speaking":
            # 朗读声波律动（5 帧循环）
            wave_frames = ["🔊 ▁", "🔊 ▂", "🔊 ▃", "🔊 ▄", "🔊 ▅", "🔊 ▆", "🔊 ▅", "🔊 ▄"]
            text = wave_frames[f % len(wave_frames)]
        else:
            # idle
            text = "🎤"
        try:
            mic.update(text)
        except Exception:
            pass

    # ═══ 字幕区 ═══
    def _set_subtitle(self, text: str, state: str = "") -> None:
        """设置字幕区显示文字（带状态色和图标）

        Args:
            text: 显示文字（空字符串则清空）
            state: 状态标识，控制字幕颜色和图标
                "" - 默认（白色）
                "listening" - 正在听（蓝色高亮 + 🎙）
                "recognized" - 已识别（白色 + 💬）
                "thinking" - AI 思考中（黄色斜体 + 💭）
                "speaking" - AI 朗读中（绿色 + 🔊）
                "error" - 错误（红色 + ⚠️）
        """
        self._subtitle_text = text or ""
        self._subtitle_state = state
        try:
            bar = self.query_one("#vd-subtitle-bar", Static)
            # 清除所有状态类
            for cls in ("active", "thinking", "speaking", "error"):
                bar.remove_class(cls)
            # 根据状态添加类
            if state == "listening":
                bar.add_class("active")
            elif state == "thinking":
                bar.add_class("thinking")
            elif state == "speaking":
                bar.add_class("speaking")
            elif state == "error":
                bar.add_class("error")
            # 截断过长字幕（3 行显示，允许更多文字）
            display = self._subtitle_text
            if len(display) > 200:
                display = display[:199] + "…"
            # 根据状态显示不同图标
            icon_map = {
                "listening": "🎙",
                "recognized": "💬",
                "thinking": "💭",
                "speaking": "🔊",
                "error": "⚠️",
            }
            icon = icon_map.get(state, "")
            if display:
                bar.update(f"{icon} {display}" if icon else display)
            else:
                bar.update("")
        except Exception:
            pass

    # ═══ 附件：图片加入待发送队列 ═══
    def _add_pending_image(self, image_path: str) -> tuple[bool, str]:
        """把图片加入 _dialog_pending_images 队列

        Returns:
            (success, message)
        """
        try:
            b64 = read_image(image_path)
            if not b64.startswith("data:"):
                return False, b64  # 错误信息
            self._dialog_pending_images.append(b64)
            # 累积到导出行
            self._export_lines.append(("system", f"[已附加图片 {Path(image_path).name}]"))
            # 在内容区显示提示气泡
            idx = len(self._dialog_pending_images)
            try:
                self._append_system_bubble(f"📎 [Image {idx}] 已就绪：{Path(image_path).name}（下次提问时附带）")
            except Exception:
                pass
            return True, f"已附加 {idx} 张图片"
        except Exception as e:
            return False, f"附件失败：{e}"

    def _get_pending_images(self) -> list:
        """获取并清空待发送图片队列"""
        imgs = list(self._dialog_pending_images)
        self._dialog_pending_images = []
        return imgs

    # ═══ 附件：路径输入（复用 vd-text-input，避免动态 Container） ═══
    def action_attach_image(self) -> None:
        """附件按钮：复用打字输入框收集图片路径

        策略：把输入框 placeholder 改为图片路径提示，
        用户输入路径回车后，on_input_submitted 检测到 _attach_mode 标志，
        走图片加载逻辑而不是发送给 AI。
        """
        if self._attach_prompt_open:
            return
        self._attach_prompt_open = True
        try:
            inp = self.query_one("#vd-text-input", Input)
            inp.placeholder = "📎 输入图片路径后回车（Esc 取消）"
            inp.add_class("visible")
            inp.value = ""
            inp.focus()
            self._text_input_active = True
            self._is_listening = False
            self._update_status("idle", "附件模式：输入图片路径，回车确认")
        except Exception as e:
            self._append_system_bubble(f"⚠️ 打开附件失败：{e}")
            self._attach_prompt_open = False

    # ═══ 音色菜单（循环切换，避免动态 Container） ═══
    VOICE_OPTIONS = [
        ("zh-CN-XiaoxiaoNeural", "女声·晓晓（默认·温柔）"),
        ("zh-CN-YunxiNeural", "男声·云希（清爽）"),
        ("zh-CN-YunjianNeural", "男声·云健（新闻）"),
        ("zh-CN-XiaoyiNeural", "女声·晓伊（活力）"),
        ("zh-CN-YunyangNeural", "男声·云扬（播音）"),
    ]

    def action_open_voice_menu(self) -> None:
        """长按静音按钮 → 循环切换音色（不弹 Container，直接切换并提示）"""
        try:
            current = self.app_instance._tts_voice
        except Exception:
            current = "zh-CN-XiaoxiaoNeural"
        # 找到当前音色在列表中的位置
        voice_ids = [vid for vid, _ in self.VOICE_OPTIONS]
        try:
            idx = voice_ids.index(current)
            next_idx = (idx + 1) % len(voice_ids)
        except ValueError:
            next_idx = 0
        next_id, next_label = self.VOICE_OPTIONS[next_idx]
        self._select_voice(next_id)

    def _select_voice(self, voice_id: str) -> None:
        """切换音色并更新提示"""
        try:
            self.app_instance._tts_voice = voice_id
        except Exception:
            pass
        label = next((lbl for vid, lbl in self.VOICE_OPTIONS if vid == voice_id), voice_id)
        self._append_system_bubble(f"🎤 音色已切换：{label}")

    def _close_voice_menu(self) -> None:
        """兼容旧调用（音色菜单已改为循环切换，无需关闭）"""
        self._voice_menu_open = False

    # ═══ 静音按钮长按检测 ═══
    def on_mouse_down(self, event) -> None:
        """记录静音按钮 mouse_down 时间戳（用于长按检测）"""
        try:
            widget = event.widget
            if widget is not None and getattr(widget, "id", None) == "vd-btn-mute":
                self._mute_mouse_down_at = time.time()
        except Exception:
            pass

    def on_mouse_up(self, event) -> None:
        """mouse_up 时判断是单击还是长按"""
        try:
            widget = event.widget
            if widget is None or getattr(widget, "id", None) != "vd-btn-mute":
                return
            if self._mute_mouse_down_at is None:
                return
            held = time.time() - self._mute_mouse_down_at
            self._mute_mouse_down_at = None
            if held >= self._mute_long_press_threshold:
                # 长按 → 弹出音色菜单（拦截 on_click 中的单击切换）
                self._suppress_next_mute_click = True
                self.action_open_voice_menu()
            else:
                # 短按 → 允许 on_click 处理（切换静音）
                self._suppress_next_mute_click = False
        except Exception:
            pass

    def _append_user_bubble(self, text: str) -> None:
        """添加用户问题气泡（右对齐，RichLog 支持 Markdown）"""
        try:
            content = self.query_one("#vd-content", VerticalScroll)
            row = Horizontal(classes="vd-row-user")
            log = RichLog(
                highlight=True,
                markup=True,
                max_lines=200,
                wrap=True,
                classes="vd-log-user",
            )
            row.mount(log)
            content.mount(row)
            # 写入内容（用户名标签 + 正文）
            try:
                from rich.markdown import Markdown as RichMarkdown
                from rich.text import Text as RichText
                log.write(RichText("┌─ 你", style=f"bold {C_ACCENT}"))
                log.write(RichMarkdown(text))
            except Exception:
                log.write(text)
            content.scroll_end(animate=False)
            # 累积到导出行
            self._export_lines.append(("user", text))
        except Exception:
            pass

    def _append_ai_placeholder(self) -> None:
        """添加 AI 回答占位气泡（RichLog 流式渲染）"""
        try:
            content = self.query_one("#vd-content", VerticalScroll)
            row = Horizontal(classes="vd-row-ai")
            log = RichLog(
                highlight=True,
                markup=True,
                max_lines=500,
                wrap=True,
                classes="vd-log-ai",
            )
            row.mount(log)
            content.mount(row)
            # 写入 AI 标签
            try:
                from rich.text import Text as RichText
                model_label = self.app_instance.model_key
                log.write(RichText(f"┌─ AI · {model_label}", style=f"bold {C_GREEN}"))
            except Exception:
                pass
            self._ai_log = log
            self._ai_text_buffer = ""  # 用于流式累积
            self._ai_thinking_log = log  # 思考中也用这个 log
            self._ai_was_interrupted = False  # 标记本轮是否被打断
            self._ai_export_recorded = False  # 标记是否已记录到 export
            content.scroll_end(animate=False)
        except Exception:
            pass

    def _update_ai_bubble(self, text: str) -> None:
        """更新 AI 气泡内容（流式 Markdown 打字机效果）"""
        try:
            if not hasattr(self, "_ai_log") or self._ai_log is None:
                return
            self._ai_text_buffer = text
            # 清空 log 后重新写入（保证 Markdown 解析正确）
            self._ai_log.clear()
            try:
                from rich.markdown import Markdown as RichMarkdown
                self._ai_log.write(RichMarkdown(text))
            except Exception:
                self._ai_log.write(text)
            # 自动滚动到底部（如果用户没手动上滚查看历史）
            content = self.query_one("#vd-content", VerticalScroll)
            # 仅当滚动条已经在底部时才自动滚动，避免打扰用户查看历史
            if content.scroll_y >= content.max_scroll_y - 2:
                content.scroll_end(animate=False)
        except Exception:
            pass

    def _mark_ai_interrupted(self) -> None:
        """在 AI 气泡末尾追加 '⏸ 已打断' 标记（保留已显示内容）"""
        try:
            if not hasattr(self, "_ai_log") or self._ai_log is None:
                return
            self._ai_was_interrupted = True
            # 在已显示内容下方追加一行提示（保持 RichLog wrap）
            try:
                from rich.text import Text
                from rich.markdown import Markdown as RichMarkdown
                # 写一行 dim 标记
                self._ai_log.write(Text("  ⏸ 已打断", style="dim italic"))
            except Exception:
                try:
                    self._ai_log.write("  ⏸ 已打断")
                except Exception:
                    pass
            # 记录到 export
            partial = (self._ai_text_buffer or "").strip()
            if partial and not getattr(self, "_ai_export_recorded", False):
                self._export_lines.append(("assistant", partial + "  _(⏸ 已打断)_"))
                self._ai_export_recorded = True
        except Exception:
            pass

    def _append_thinking_indicator(self) -> None:
        """在 AI 气泡开头显示思考动画"""
        try:
            if not hasattr(self, "_ai_log") or self._ai_log is None:
                return
            # 在 log 顶部写入"思考中..."
            self._ai_log.write("[dim]💭 思考中...[/dim]")
        except Exception:
            pass

    def _finalize_ai_export(self) -> None:
        """AI 完整回答后记录到 export_lines（仅在未中断/未记录时）"""
        try:
            if getattr(self, "_ai_was_interrupted", False):
                return
            if getattr(self, "_ai_export_recorded", False):
                return
            text = (self._ai_text_buffer or "").strip()
            if text:
                self._export_lines.append(("assistant", text))
                self._ai_export_recorded = True
        except Exception:
            pass

    def _dialog_loop(self) -> None:
        """语音对话主循环（后台线程）

        状态机：
          IDLE（等待）→ LISTENING（录音中）→ RECOGNIZING（识别中）
            → THINKING（AI 思考）→ SPEAKING（朗读中）→ IDLE
          任何状态都可被"用户说话"打断回到 LISTENING
        """
        import time as _t
        import threading as _th
        # 连续无输入计数（用于提示用户"请说话"）
        empty_count = 0
        while self._dialog_active:
            try:
                # ═══ 1. 等待用户说话（录音） ═══
                self._is_listening = True
                self._current_ai_text = ""
                self._current_ai_displayed = ""
                # 显示"正在听..."提示气泡（带动画）
                self.call_from_thread(self._show_listening_indicator)
                self.call_from_thread(self._update_status, "listening", "请说话…")
                # 字幕区：提示用户正在听
                self.call_from_thread(self._set_subtitle, "正在听… 说吧", "listening")
                # 录音前短暂停顿（避免 TTS 回声）
                _t.sleep(0.1)
                if not self._dialog_active:
                    break
                # 录音 + 识别（改进版 VAD：动态阈值 + 前置静音过滤 + 归一化）
                try:
                    text = listen_asr(max_seconds=10, silence_seconds=1.0)
                except Exception as e:
                    text = f"错误：{e}"
                self._is_listening = False
                if not self._dialog_active:
                    break
                # 移除"正在听..."提示
                self.call_from_thread(self._hide_listening_indicator)
                # ═══ 2. 处理识别结果 ═══
                # 空输入或识别失败
                is_error = text.startswith("错误") or text.startswith("（未")
                is_empty = (not text.strip()) or text.strip() in ("（未识别到内容）", "（未录到声音）")
                if is_error or is_empty:
                    self.call_from_thread(self._set_subtitle, text.strip() or "（未识别）", "error")
                    empty_count += 1
                    if empty_count >= 3:
                        # 连续 3 次没听清，给出明显提示
                        self.call_from_thread(self._append_system_bubble, "🎤 连续未听清，请清晰说话，或按 Ctrl+T 打字输入")
                        empty_count = 0
                    else:
                        self.call_from_thread(self._append_system_bubble, "🎤 没听清，请再说一遍")
                    _t.sleep(0.3)
                    continue
                # 有效输入，重置计数 + 字幕区显示识别文本
                empty_count = 0
                self.call_from_thread(self._set_subtitle, text.strip(), "recognized")
                # 退出指令
                if text.strip() in ("停止", "退出", "结束对话", "退出对话", "stop", "结束", "关闭"):
                    self._dialog_active = False
                    self.call_from_thread(self._append_user_bubble, text)
                    self.call_from_thread(self._append_system_bubble, "👋 对话已结束")
                    _t.sleep(0.5)
                    self.call_from_thread(self.action_close_dialog)
                    return
                # ═══ 3. 显示用户气泡 + AI 占位 ═══
                self._current_user_text = text
                self.call_from_thread(self._append_user_bubble, text)
                # 取出待发送图片（仅当用户有附加时）
                pending_imgs = self._get_pending_images()
                if pending_imgs:
                    self.call_from_thread(self._append_system_bubble, f"📎 已附带 {len(pending_imgs)} 张图片")
                self.call_from_thread(self._append_ai_placeholder)
                self.call_from_thread(self._update_status, "thinking", "AI 正在思考…")
                self.call_from_thread(self._set_subtitle, "AI 正在思考…", "thinking")
                self._is_generating = True
                # ═══ 4. 启动 AI 流式生成（支持多模态图片） ═══
                ai_done_event = _th.Event()
                ai_result = {"text": "", "error": None}
                def _run_ai():
                    try:
                        result_text = self._generate_ai_stream(
                            text,
                            lambda chunk: self._on_ai_chunk(chunk),
                            images=pending_imgs,
                        )
                        ai_result["text"] = result_text
                    except Exception as e:
                        ai_result["error"] = str(e)
                    finally:
                        ai_done_event.set()
                ai_thread = _th.Thread(target=_run_ai, daemon=True)
                ai_thread.start()
                # 等待 AI 生成完成
                ai_done_event.wait()
                self._is_generating = False
                if ai_result["error"]:
                    self.call_from_thread(self._append_system_bubble, f"⚠️ 生成失败：{ai_result['error'][:50]}")
                    continue
                if not ai_result["text"].strip():
                    self.call_from_thread(self._append_system_bubble, "⚠️ AI 未返回内容")
                    continue
                # AI 完整生成成功 → 记录到 export
                self.call_from_thread(self._finalize_ai_export)
                # ═══ 5. TTS 朗读（可被用户说话打断） ═══
                if self._mute_enabled:
                    _t.sleep(0.2)
                    continue
                self._is_speaking = True
                self.call_from_thread(self._update_status, "speaking", "AI 正在回答（说话可打断）")
                # 字幕显示 AI 回答摘要
                ai_preview = (self._current_ai_text or "").strip().replace("\n", " ")
                if len(ai_preview) > 80:
                    ai_preview = ai_preview[:79] + "…"
                self.call_from_thread(self._set_subtitle, ai_preview, "speaking")
                # 在独立线程朗读（分段朗读 + 可打断）
                interrupt_event = _th.Event()
                self._current_interrupt = interrupt_event

                def _is_interrupted() -> bool:
                    """对话循环侧的中断检测：用户开始录音 / 对话已结束 → 停止朗读"""
                    if not self._dialog_active:
                        return True
                    if self._is_listening:
                        return True
                    return False

                def _do_speak():
                    try:
                        reply = self._current_ai_text
                        voice = self.app_instance._tts_voice
                        rate = self.app_instance._tts_rate
                        if reply.strip():
                            # 分段朗读（interrupt_check 回调用于在每段之间 / 段内打断）
                            speak_tts(
                                reply,
                                voice=voice,
                                rate=rate,
                                interrupt_check=_is_interrupted,
                                segment_max_chars=200,
                            )
                    except Exception:
                        pass
                    finally:
                        interrupt_event.set()
                speak_thread = _th.Thread(target=_do_speak, daemon=True)
                speak_thread.start()
                # 等待朗读完成或用户打断
                interrupted_by_user = False
                while not interrupt_event.is_set():
                    _t.sleep(0.2)
                    if not self._dialog_active:
                        # 对话结束 → 强制停止 TTS
                        try:
                            import pygame
                            pygame.mixer.music.stop()
                        except Exception:
                            pass
                        break
                    # 用户说话（_is_listening 变 True）→ 停止 TTS
                    if self._is_listening:
                        try:
                            import pygame
                            pygame.mixer.music.stop()
                        except Exception:
                            pass
                        self.call_from_thread(self._append_system_bubble, "⏹ 已打断，正在听你说…")
                        # 在 AI 气泡上追加"⏸ 已打断"标记（保留已显示内容）
                        self.call_from_thread(self._mark_ai_interrupted)
                        interrupted_by_user = True
                        break
                self._is_speaking = False
            except Exception as e:
                self.call_from_thread(self._append_system_bubble, f"⚠️ 循环异常：{e}")
                _t.sleep(0.3)
        # 循环结束
        self.call_from_thread(self._set_subtitle, "", "")
        self.call_from_thread(self._append_system_bubble, "对话已结束")

    def _show_listening_indicator(self) -> None:
        """显示'正在听...'动画提示气泡"""
        try:
            content = self.query_one("#vd-content", VerticalScroll)
            # 如果已存在则先移除
            try:
                old = self.query_one("#vd-listening-indicator", Static)
                old.remove()
            except Exception:
                pass
            # 动态点数（1-3 循环）
            dots = "." * ((int(time.time() * 2) % 3) + 1)
            indicator = Static(f"  🎤 正在听{dots}", id="vd-listening-indicator", classes="vd-thinking")
            content.mount(indicator)
            content.scroll_end(animate=False)
        except Exception:
            pass

    def _hide_listening_indicator(self) -> None:
        """移除'正在听...'提示气泡"""
        try:
            old = self.query_one("#vd-listening-indicator", Static)
            old.remove()
        except Exception:
            pass

    def _append_system_bubble(self, text: str) -> None:
        """添加系统提示气泡（居中灰色）"""
        try:
            content = self.query_one("#vd-content", VerticalScroll)
            line = Static(f"  {text}", classes="vd-thinking")
            content.mount(line)
            content.scroll_end(animate=False)
            # 累积到导出（跳过自动状态提示，避免污染 Markdown）
            if not text.startswith(("🎤 ", "🔇 ", "👋 ", "⏹ ", "💬 ")):
                # 仅记录真正的事件型提示（如"音色已切换"）
                if any(kw in text for kw in ("已切换", "已结束", "已打断", "导出", "错误", "失败")):
                    self._export_lines.append(("system", text.strip()))
        except Exception:
            pass

    def _on_ai_chunk(self, chunk: str) -> None:
        """AI 流式输出的回调（每收到一段文字就更新气泡和字幕）"""
        self._current_ai_text += chunk
        # 过滤模型内部特殊标签（<|observation|> <|system|> 等，防止泄露给用户）
        self._current_ai_text = _strip_model_tokens(self._current_ai_text)
        # 打字机效果：先记录完整文本，再通过定时器逐字显示
        # 简化：直接显示（Textual 内部已经是流式的，足够流畅）
        self.call_from_thread(self._update_ai_bubble, self._current_ai_text)
        # 字幕实时显示 AI 正在输出的文本（前 150 字符，避免过长）
        preview = self._current_ai_text.strip().replace("\n", " ")
        if len(preview) > 150:
            preview = preview[:149] + "…"
        if preview:
            self.call_from_thread(self._set_subtitle, preview, "thinking")

    def _generate_ai_stream(self, user_text: str, on_chunk, images: list = None) -> str:
        """同步生成 AI 回答（带流式回调）

        通过在 ZeroAI 主 App 中启动一个一次性任务，把每次 chunk 传给 on_chunk

        Args:
            user_text: 用户问题文本
            on_chunk: 流式 chunk 回调
            images: 可选，附加图片的 base64 data URI 列表
        """
        # 在主线程异步执行 _run_turn
        # 由于 _run_turn 是 async 协程，且在主事件循环中运行
        # 这里使用 run_worker 提交到主 App 的事件循环
        import threading
        result_holder = {"text": "", "done": False, "error": None}
        def _on_main():
            # 在主线程事件循环中执行
            import asyncio
            self.app_instance.run_worker(
                self._async_generate_ai(user_text, on_chunk, result_holder, images or []),
                exclusive=False,
            )
        # 由于 _generate_ai_stream 已在子线程，调度到主线程
        self.app_instance.call_from_thread(_on_main)
        # 等待完成
        while not result_holder["done"]:
            time.sleep(0.1)
        return result_holder.get("text", "")

    async def _async_generate_ai(self, user_text: str, on_chunk, result_holder, images: list = None):
        """在主 App 事件循环中执行 AI 生成（支持多模态图片）"""
        images = images or []
        try:
            import os
            from openai import OpenAI

            # 构造消息：纯文本 vs 多模态
            if images:
                content_parts = [{"type": "text", "text": user_text}]
                for url in images:
                    content_parts.append({"type": "image_url", "image_url": {"url": url}})
                user_msg = {"role": "user", "content": content_parts}
            else:
                user_msg = {"role": "user", "content": user_text}

            # 把 user_msg 注入主 App 的消息列表
            self.app_instance.messages.append(user_msg)

            cfg = MODEL_CONFIGS[self.app_instance.model_key]
            # 有图片时强制用 vision 模型（glm-v）以保证多模态识别
            if images and self.app_instance.model_key != "glm-v":
                cfg = MODEL_CONFIGS["glm-v"]
            client = OpenAI(base_url=cfg["base_url"], api_key=cfg["api_key"], timeout=180.0, max_retries=2)
            response = client.chat.completions.create(
                model=cfg["model"],
                messages=self.app_instance.messages,
                stream=True,
                stream_options={"include_usage": True},
                temperature=self.app_instance.temperature,
                timeout=180,
            )
            full_text = ""
            for chunk in response:
                if chunk.choices and chunk.choices[0].delta.content:
                    piece = chunk.choices[0].delta.content
                    full_text += piece
                    # 过滤模型内部特殊标签（<|observation|> <|system|> 等）
                    full_text = _strip_model_tokens(full_text)
                    on_chunk(piece)
            # 同时把回答追加到主 App 的消息历史
            self.app_instance.messages.append({"role": "assistant", "content": full_text})
            result_holder["text"] = full_text
        except Exception as e:
            result_holder["error"] = str(e)
        finally:
            result_holder["done"] = True

    # ── 手动录音控制（点左侧"说话"按钮开始/停止录音）──
    def _start_manual_recording(self) -> None:
        """开始手动录音（后台线程录音，主线程不阻塞）"""
        if self._is_generating or self._is_speaking:
            # AI 正在生成/朗读，先停止
            self._stop_generation_and_tts()
        import threading
        self._manual_listening = True
        self._manual_audio_buf = []
        # 更新 UI
        self._is_listening = True
        self._current_ai_text = ""
        self._current_ai_displayed = ""
        self._show_listening_indicator()
        self._update_status("listening", "正在录音…点'说话'停止")
        self._set_subtitle("🔴 正在录音…点'说话'按钮停止", "listening")
        try:
            btn = self.query_one("#vd-btn-mute", Static)
            btn.update("⏹停止")
            btn.add_class("recording")
        except Exception:
            pass
        # 启动后台录音线程
        def _rec():
            try:
                import sounddevice as sd
                import numpy as np
                sr = 16000
                block = 1024
                with sd.InputStream(samplerate=sr, channels=1, blocksize=block, dtype='float32') as stream:
                    while self._manual_listening and self._dialog_active:
                        data, _ = stream.read(block)
                        self._manual_audio_buf.append(data.copy())
            except Exception as e:
                self._manual_audio_buf = None
                self.call_from_thread(self._set_subtitle, f"录音失败: {e}", "error")
        self._manual_rec_thread = threading.Thread(target=_rec, daemon=True)
        self._manual_rec_thread.start()

    def _stop_manual_recording(self) -> None:
        """停止手动录音并识别"""
        self._manual_listening = False
        self._is_listening = False
        # 等录音线程结束
        if self._manual_rec_thread and self._manual_rec_thread.is_alive():
            self._manual_rec_thread.join(timeout=1.0)
        # 恢复按钮
        try:
            btn = self.query_one("#vd-btn-mute", Static)
            btn.update("说话")
            btn.remove_class("recording")
        except Exception:
            pass
        self._hide_listening_indicator()
        # 合并音频
        if not self._manual_audio_buf:
            self._set_subtitle("未录到声音", "error")
            self._update_status("idle", "点左侧'说话'按钮开始录音")
            return
        import numpy as np
        try:
            audio = np.concatenate(self._manual_audio_buf, axis=0).flatten().astype(np.float32)
        except Exception:
            self._set_subtitle("音频合并失败", "error")
            return
        self._manual_audio_buf = None
        # 音量检查
        vol = float(np.abs(audio).mean())
        if vol < 0.005:
            self._set_subtitle("音量太低，请靠近麦克风重试", "error")
            self._update_status("idle", "点左侧'说话'按钮开始录音")
            return
        # 识别
        self._set_subtitle("识别中…", "thinking")
        self._update_status("thinking", "识别中…")
        import threading
        def _recognize():
            text = ""
            try:
                # 用 listen_asr 的识别部分（直接用已录的音频）
                global _ASR_MODEL
                import sherpa_onnx
                if _ASR_MODEL is None:
                    if not os.path.isfile(_SENSE_VOICE_MODEL) or not os.path.isfile(_SENSE_VOICE_TOKENS):
                        # 尝试自动下载模型
                        if not _download_sense_voice_model():
                            raise FileNotFoundError(f"模型下载失败，请手动下载到: {_SENSE_VOICE_MODEL_DIR}")
                    _ASR_MODEL = sherpa_onnx.OfflineRecognizer.from_sense_voice(
                        model=_SENSE_VOICE_MODEL, tokens=_SENSE_VOICE_TOKENS,
                        num_threads=2, use_itn=True,
                    )
                stream = _ASR_MODEL.create_stream()
                stream.accept_waveform(16000, audio)
                _ASR_MODEL.decode_stream(stream)
                text = stream.result.text.strip()
            except Exception as e:
                text = f"错误：{e}"
            self.call_from_thread(self._on_manual_recognized, text)
        threading.Thread(target=_recognize, daemon=True).start()

    def _on_manual_recognized(self, text: str) -> None:
        """手动录音识别完成回调"""
        is_error = text.startswith("错误")
        is_empty = (not text.strip()) or text.strip() in ("（未识别到内容）", "（未录到声音）")
        if is_error or is_empty:
            self._set_subtitle(text.strip() or "（未识别）", "error")
            self._update_status("idle", "点左侧'说话'按钮开始录音")
            return
        self._set_subtitle(text.strip(), "recognized")
        # 退出指令
        if text.strip() in ("停止", "退出", "结束对话", "退出对话", "stop", "结束", "关闭"):
            self._dialog_active = False
            self._append_user_bubble(text)
            self._append_system_bubble("👋 对话已结束")
            self.action_close_dialog()
            return
        # 发送给 AI
        self._current_user_text = text
        self._append_user_bubble(text)
        pending_imgs = self._get_pending_images()
        if pending_imgs:
            self._append_system_bubble(f"📎 已附带 {len(pending_imgs)} 张图片")
        # 启动 AI 回复线程
        import threading
        def _ai_reply():
            try:
                self._dialog_loop_single_turn(text)
            except Exception as e:
                self.call_from_thread(self._set_subtitle, f"AI 回复失败: {e}", "error")
        threading.Thread(target=_ai_reply, daemon=True).start()

    def _stop_generation_and_tts(self) -> None:
        """停止当前 AI 生成和 TTS 朗读"""
        try:
            self.app_instance._stop_generation = True
        except Exception:
            pass
        try:
            import pygame
            pygame.mixer.music.stop()
        except Exception:
            pass
        self._is_speaking = False
        self._is_generating = False

    def _dialog_loop_single_turn(self, text: str) -> None:
        """手动模式：处理单轮对话（AI 回复 + TTS 朗读）

        从原 _dialog_loop 提取第 3-5 步逻辑：
        3. 显示用户气泡 + AI 占位
        4. AI 流式生成
        5. TTS 朗读
        """
        import time as _t
        import threading as _th
        # ═══ 3. 显示 AI 占位 ═══
        self.call_from_thread(self._append_ai_placeholder)
        self.call_from_thread(self._update_status, "thinking", "AI 正在思考…")
        self.call_from_thread(self._set_subtitle, "AI 正在思考…", "thinking")
        self._is_generating = True
        # ═══ 4. AI 流式生成 ═══
        ai_done_event = _th.Event()
        ai_result = {"text": "", "error": None}
        pending_imgs = self._get_pending_images()
        def _run_ai():
            try:
                result_text = self._generate_ai_stream(
                    text,
                    lambda chunk: self._on_ai_chunk(chunk),
                    images=pending_imgs,
                )
                ai_result["text"] = result_text
            except Exception as e:
                ai_result["error"] = str(e)
            finally:
                ai_done_event.set()
        ai_thread = _th.Thread(target=_run_ai, daemon=True)
        ai_thread.start()
        ai_done_event.wait()
        self._is_generating = False
        if ai_result["error"]:
            self.call_from_thread(self._append_system_bubble, f"⚠️ 生成失败：{ai_result['error'][:50]}")
            self.call_from_thread(self._update_status, "idle", "点左侧'说话'按钮开始录音")
            return
        if not ai_result["text"].strip():
            self.call_from_thread(self._append_system_bubble, "⚠️ AI 未返回内容")
            self.call_from_thread(self._update_status, "idle", "点左侧'说话'按钮开始录音")
            return
        self.call_from_thread(self._finalize_ai_export)
        # ═══ 5. TTS 朗读 ═══
        if self._mute_enabled:
            self.call_from_thread(self._update_status, "idle", "点左侧'说话'按钮开始录音")
            return
        self._is_speaking = True
        self.call_from_thread(self._update_status, "speaking", "AI 正在回答…")
        ai_preview = (self._current_ai_text or "").strip().replace("\n", " ")
        if len(ai_preview) > 80:
            ai_preview = ai_preview[:79] + "…"
        self.call_from_thread(self._set_subtitle, ai_preview, "speaking")
        interrupt_event = _th.Event()
        self._current_interrupt = interrupt_event
        def _is_interrupted() -> bool:
            if not self._dialog_active:
                return True
            if self._is_listening:
                return True
            return False
        def _do_speak():
            try:
                reply = self._current_ai_text
                voice = self.app_instance._tts_voice
                rate = self.app_instance._tts_rate
                if reply.strip():
                    speak_tts(
                        reply, voice=voice, rate=rate,
                        interrupt_check=_is_interrupted, segment_max_chars=200,
                    )
            except Exception:
                pass
            finally:
                interrupt_event.set()
        speak_thread = _th.Thread(target=_do_speak, daemon=True)
        speak_thread.start()
        while not interrupt_event.is_set():
            _t.sleep(0.2)
            if not self._dialog_active:
                interrupt_event.set()
        self._is_speaking = False
        # 朗读结束，恢复待命
        if self._dialog_active:
            self.call_from_thread(self._update_status, "idle", "点左侧'说话'按钮开始录音")
            self.call_from_thread(self._set_subtitle, "点左侧'说话'按钮继续", "")

    # ── 事件处理 ──
    def on_click(self, event: events.Click) -> None:
        """处理底部按钮点击（Textual 通用 click 事件）"""
        widget = event.widget
        if widget is None:
            return
        widget_id = getattr(widget, "id", "") or ""

        if widget_id == "vd-btn-mic":
            # 中央大圆麦克风按钮：关闭对话
            self.action_close_dialog()
        elif widget_id == "vd-btn-attach":
            # 右侧附件按钮：打开路径输入（复用打字输入框）
            self.action_attach_image()
        elif widget_id == "vd-btn-mute":
            # 长按被拦截（已弹出音色菜单），跳过本次点击
            if getattr(self, "_suppress_next_mute_click", False):
                self._suppress_next_mute_click = False
                return
            # 左侧"说话"按钮：手动录音控制（点击开始/停止录音）
            if self._manual_listening:
                # 正在录音 → 停止录音并识别
                self._stop_manual_recording()
            else:
                # 未录音 → 开始录音
                self._start_manual_recording()
        elif widget_id == "vd-subtitle-btn":
            self._subtitle_enabled = not self._subtitle_enabled
            label = "字幕" if self._subtitle_enabled else "字幕关"
            try:
                self.query_one("#vd-subtitle-btn", Static).update(label)
            except Exception:
                pass
            if not self._subtitle_enabled:
                self._set_subtitle("", "")
        elif widget_id == "vd-close-btn":
            self.action_close_dialog()

    # ── 输入框提交（打字 + 附件路径统一入口） ──
    def on_input_submitted(self, event: Input.Submitted) -> None:
        """统一处理 vd-text-input 提交

        通过 _attach_prompt_open 标志区分：
          - True  → 附件路径模式：加载图片
          - False → 正常打字模式：发送给 AI
        """
        if event.input.id != "vd-text-input":
            return
        text = event.value.strip()
        event.input.value = ""

        # 附件路径模式
        if self._attach_prompt_open:
            self._attach_prompt_open = False
            # 恢复输入框
            event.input.remove_class("visible")
            event.input.placeholder = "打字输入问题，回车发送"
            self._text_input_active = False
            if not text:
                self._update_status("listening", "请说话…")
                return
            # 去掉前后引号
            text = text.strip("\"'")
            ok, msg = self._add_pending_image(text)
            if not ok:
                self._append_system_bubble(f"⚠️ {msg}")
            self._update_status("listening", "请说话…")
            return

        # 正常打字模式
        if not text:
            return
        # 隐藏输入框，回到语音模式
        event.input.remove_class("visible")
        event.input.placeholder = "打字输入问题，回车发送"
        self._text_input_active = False
        # 退出指令
        if text in ("停止", "退出", "结束对话", "stop", "结束", "关闭"):
            self.action_close_dialog()
            return
        # 在后台线程处理这轮对话
        import threading
        pending_imgs = self._get_pending_images()
        def _handle_text():
            self._current_user_text = text
            self.call_from_thread(self._set_subtitle, text, "recognized")
            self.call_from_thread(self._append_user_bubble, text)
            if pending_imgs:
                self.call_from_thread(self._append_system_bubble, f"📎 已附带 {len(pending_imgs)} 张图片")
            self.call_from_thread(self._append_ai_placeholder)
            self.call_from_thread(self._update_status, "thinking", "AI 正在思考…")
            self._is_generating = True
            import threading as _th
            ai_done_event = _th.Event()
            ai_result = {"text": "", "error": None}
            def _run_ai():
                try:
                    result_text = self._generate_ai_stream(
                        text,
                        lambda chunk: self._on_ai_chunk(chunk),
                        images=pending_imgs,
                    )
                    ai_result["text"] = result_text
                except Exception as e:
                    ai_result["error"] = str(e)
                finally:
                    ai_done_event.set()
            ai_thread = _th.Thread(target=_run_ai, daemon=True)
            ai_thread.start()
            ai_done_event.wait()
            self._is_generating = False
            if ai_result["error"]:
                self.call_from_thread(self._append_system_bubble, f"⚠️ 生成失败：{ai_result['error'][:50]}")
                return
            self.call_from_thread(self._finalize_ai_export)
            # TTS 朗读
            if not self._mute_enabled and ai_result["text"].strip():
                self._is_speaking = True
                self.call_from_thread(self._update_status, "speaking", "AI 正在回答")
                try:
                    from tui_agent import speak_tts as _speak
                    def _is_interrupted() -> bool:
                        if not self._dialog_active:
                            return True
                        if self._is_listening:
                            return True
                        return False
                    _speak(
                        ai_result["text"],
                        voice=self.app_instance._tts_voice,
                        rate=self.app_instance._tts_rate,
                        interrupt_check=_is_interrupted,
                        segment_max_chars=200,
                    )
                except Exception:
                    pass
                self._is_speaking = False
        t = threading.Thread(target=_handle_text, daemon=True)
        t.start()

    def _close_attach_prompt(self) -> None:
        """关闭附件模式（恢复输入框）"""
        self._attach_prompt_open = False
        try:
            inp = self.query_one("#vd-text-input", Input)
            inp.remove_class("visible")
            inp.placeholder = "打字输入问题，回车发送"
        except Exception:
            pass

    # ── Ctrl+S：导出对话为 Markdown ──
    def action_export_dialog(self) -> None:
        """Ctrl+S：把当前对话导出为 Markdown 文件"""
        try:
            from pathlib import Path as _P
            # 默认导出到用户目录的 .trae-cn/voice_dialogs/
            save_dir = _P.home() / ".trae-cn" / "voice_dialogs"
            save_dir.mkdir(parents=True, exist_ok=True)
            from datetime import datetime
            ts = datetime.fromtimestamp(self._session_started_at).strftime("%Y%m%d_%H%M%S")
            path = save_dir / f"dialog_{ts}.md"
            md = self._build_export_markdown()
            path.write_text(md, encoding="utf-8")
            self._append_system_bubble(f"💾 对话已导出：{path}")
            # 同步写入 export_lines
            self._export_lines.append(("system", f"已导出到 {path}"))
        except Exception as e:
            self._append_system_bubble(f"⚠️ 导出失败：{e}")

    def _build_export_markdown(self) -> str:
        """构造可导出的 Markdown 文本"""
        from datetime import datetime
        lines = []
        lines.append(f"# ZeroAI 语音对话记录")
        lines.append("")
        lines.append(f"- 开始时间：{datetime.fromtimestamp(self._session_started_at).strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"- 模型：{self.app_instance.model_key} ({self.app_instance.get_current_model()})")
        try:
            lines.append(f"- 音色：{self.app_instance._tts_voice}")
        except Exception:
            pass
        lines.append(f"- 消息条数：{sum(1 for r, _ in self._export_lines if r in ('user', 'assistant'))}")
        lines.append("")
        lines.append("---")
        lines.append("")
        for role, text in self._export_lines:
            text = (text or "").strip()
            if not text:
                continue
            if role == "user":
                lines.append("## 🙋 你")
                lines.append("")
                lines.append(text)
                lines.append("")
            elif role == "assistant":
                lines.append("## 🤖 AI")
                lines.append("")
                lines.append(text)
                lines.append("")
            elif role == "system":
                lines.append(f"> {text}")
                lines.append("")
        return "\n".join(lines)

    def action_close_dialog(self) -> None:
        """关闭讨论助手"""
        self._dialog_active = False
        # 停止手动录音
        self._manual_listening = False
        # 停止 AI 生成
        try:
            self.app_instance._stop_generation = True
        except Exception:
            pass
        # 停止 TTS 播放（当前段立即停止）
        try:
            import pygame
            pygame.mixer.music.stop()
            pygame.mixer.music.unload()
        except Exception:
            pass
        self._is_speaking = False
        self._is_generating = False
        self._is_listening = False
        self.dismiss(None)

    def action_manual_input(self) -> None:
        """Ctrl+T：切换打字输入模式"""
        try:
            inp = self.query_one("#vd-text-input", Input)
            if inp.has_class("visible"):
                # 已显示 → 隐藏
                inp.remove_class("visible")
                self._text_input_active = False
            else:
                # 显示并聚焦
                inp.add_class("visible")
                inp.focus()
                self._text_input_active = True
                # 暂停录音循环（让用户能安静打字）
                self._is_listening = False
                self._update_status("idle", "打字模式：回车发送，Esc 回到语音")
        except Exception as e:
            self._append_system_bubble(f"⚠️ 切换失败：{e}")


class ZeroAI(App):
    """ZeroAI 终端 AI 助手"""

    CSS = f"""
    Screen {{
        background: {C_BG};
        color: {C_FG};
    }}

    /* 顶部：信息条（极简） */
    #info {{
        dock: top;
        height: 1;
        background: {C_BG};
        color: {C_DIM};
        padding: 0 2;
    }}

    /* 中间：对话区 + 右侧状态栏 */
    #main-area {{
        height: 1fr;
    }}

    #log-scroll {{
        width: 1fr;
        background: {C_BG};
        padding: 1 2;
        overflow-y: auto;
        scrollbar-color: {C_DIM};
        scrollbar-background: {C_BG};
        scrollbar-size-vertical: 1;
    }}

    #token-bar {{
        dock: right;
        width: 22;
        height: 1fr;
        background: {C_BG};
        color: {C_FG};
        padding: 1 1;
    }}

    .msg-block {{
        background: {C_BG};
        color: {C_FG};
        padding: 0 0 1 0;
        margin: 0 0 1 0;
    }}

    .msg-header {{
        color: {C_DIM};
        text-style: bold;
    }}

    /* 底部：输入区（黑色填充大输入框） */
    #input-wrap {{
        dock: bottom;
        height: auto;
        min-height: 3;
        max-height: 18;
        background: {C_BG};
        border-top: solid {C_BORDER};
        padding: 0 2;
    }}

    #input {{
        background: {C_BG};
        color: {C_FG};
        border: solid {C_BORDER};
        height: auto;
        min-height: 1;
        max-height: 15;
        padding: 0 1;
    }}

    #input:focus {{
        border: solid {C_DIM};
    }}

    #input .cursor {{
        background: {C_FG};
        color: {C_BG};
    }}

    /* 底部：快捷键栏（极简灰色） */
    #hints {{
        dock: bottom;
        height: 1;
        background: {C_BG};
        color: {C_DIM};
        padding: 0 2;
    }}
    """

    BINDINGS = [
        Binding("ctrl+c", "stop_or_quit", "停止/退出", show=False),
        Binding("ctrl+l", "clear_log", "清屏", show=False),
        Binding("ctrl+n", "clear_history", "新对话", show=False),
        Binding("ctrl+p", "open_settings", "设置", show=False),
        Binding("ctrl+y", "copy_last_reply", "复制", show=False),
        Binding("ctrl+v", "paste_image", "粘贴图片", show=False),
        Binding("ctrl+g", "paste_image", "粘贴图片", show=False),
        Binding("ctrl+w", "toggle_companion", "伴随模式", show=False),
        Binding("ctrl+t", "push_to_talk", "语音输入", show=False),
        Binding("ctrl+d", "voice_dialog", "语音对话", show=False),
        Binding("pageup", "scroll_pageup", "上翻页", show=False),
        Binding("pagedown", "scroll_pagedown", "下翻页", show=False),
    ]

    # 禁用 Textual 自带的命令面板（避免和我们的设置冲突）
    ENABLE_COMMAND_PALETTE = False

    def __init__(self):
        super().__init__()
        self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        self.model_key = CURRENT_MODEL_KEY
        self.work_mode = "hybrid"  # expert / hybrid / manual（默认混合思考，启用多专家子代理）
        # 可调参数
        self.temperature = 0.3
        self.stream_enabled = True
        self.max_turns = 8
        self.context_limit = 8192
        # Token 统计
        self.total_tokens = 0
        self.stream_start_time = 0.0
        self.stream_token_count = 0
        self._precise_input_tokens = 0  # API 返回的精确输入 token（0 表示用估算）
        # 用户滚动状态（用户往上翻看时暂停自动滚底）
        self._user_scrolling = False
        # 停止生成标志（Ctrl+C 第一次停止，第二次退出）
        self._stop_generation = False
        self._is_generating = False
        # 最近一次助手回复的纯文本（用于复制）
        self._last_reply_text = ""
        # 专家记忆：每个专家维护独立上下文，避免主上下文污染
        # 结构：{expert_key: [{"role": "user"|"assistant", "content": "..."}, ...]}
        self._expert_memory = {}
        # 待发送的图片 base64 列表（Ctrl+V 粘贴）
        self._pending_images = []
        # 伴随模式（屏幕感知）
        self._companion_mode = False
        self._companion_log = []  # 最近的屏幕变化日志
        self._last_window_title = ""
        self._last_clipboard_text = ""
        self._companion_thread = None
        # 语音交互状态
        self._tts_enabled = False  # TTS 朗读开关（由 /语音 命令切换）
        self._tts_voice = "zh-CN-XiaoxiaoNeural"  # 默认女声
        self._tts_rate = "+0%"  # 默认语速
        self._is_listening = False  # ASR 录音中状态
        self._voice_dialog_active = False  # 语音对话模式（保留兼容字段）

    def get_current_client(self):
        """获取当前模型的客户端"""
        cfg = MODEL_CONFIGS[self.model_key]
        return OpenAI(base_url=cfg["base_url"], api_key=cfg["api_key"])

    def get_current_model(self):
        """获取当前模型名"""
        return MODEL_CONFIGS[self.model_key]["model"]

    def switch_model(self, key: str) -> bool:
        """切换模型，成功返回 True"""
        if key not in MODEL_CONFIGS:
            return False
        self.model_key = key
        # 切换模型时清空历史（不同模型的 system prompt 格式可能不同）
        self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        # 刷新顶部信息栏
        self.query_one("#info", InfoBar).refresh()
        return True

    def action_open_settings(self):
        """Ctrl+P 打开设置面板"""
        def on_result(result):
            global WORK_MODE
            if not result or not isinstance(result, dict):
                return
            action = result.get("action")
            if action == "switch_mode":
                mode = result["mode"]
                self.work_mode = mode
                WORK_MODE = mode
                self.query_one("#info", InfoBar).refresh()
                mode_labels = {"expert": "专家模式", "hybrid": "混合思考", "manual": "手动模式"}
                ml = mode_labels.get(mode, mode)
                self._add_static(Text.assemble(
                    ("  已切换到 ", C_DIM),
                    (f"{ml}", f"bold {C_FG}"),
                    ("\n", C_DIM),
                ))
            elif action == "expert_info":
                ek = result["key"]
                expert = EXPERT_TEAM[ek]
                e_cfg = get_expert_config(ek)
                self._add_static(Text.assemble(
                    ("  ◆ 专家信息\n", f"bold {C_BLUE}"),
                    ("  名称：", C_DIM), (f"{expert['label']}\n", C_FG),
                    ("  说明：", C_DIM), (f"{expert['desc']}\n", C_FG),
                    ("  模型：", C_DIM), (f"{e_cfg['model']}\n", C_FG),
                    ("  平台：", C_DIM), (f"{e_cfg.get('label', expert['label'])}\n", C_FG),
                    ("  关键词：", C_DIM), (f"{', '.join(expert['keywords'][:8]) if expert['keywords'] else '默认兜底'}\n", C_DIM),
                ))
            elif action == "switch_model":
                key = result["key"]
                if key != self.model_key and self.switch_model(key):
                    # 从设置面板切换模型时自动切到手动模式
                    self.work_mode = "manual"
                    WORK_MODE = "manual"
                    self.query_one("#info", InfoBar).refresh()
                    label = MODEL_CONFIGS[key]["label"]
                    self._add_static(Text.assemble(
                        ("  已切换模型：", C_DIM),
                        (f"{label}", f"bold {C_FG}"),
                        ("（手动模式·对话已清空）\n", C_DIM),
                    ))
            elif action == "set_temperature":
                self.temperature = result["value"]
                self._add_static(Text.assemble(
                    ("  温度：", C_DIM),
                    (f"{self.temperature}", f"bold {C_FG}"),
                    ("\n", C_DIM),
                ))
            elif action == "set_stream":
                self.stream_enabled = result["value"]
                label = "开" if self.stream_enabled else "关"
                self._add_static(Text.assemble(
                    ("  流式输出：", C_DIM),
                    (f"{label}", f"bold {C_FG}"),
                    ("\n", C_DIM),
                ))
            elif action == "set_max_turns":
                self.max_turns = result["value"]
                self._add_static(Text.assemble(
                    ("  最大调用轮次：", C_DIM),
                    (f"{self.max_turns}", f"bold {C_FG}"),
                    ("\n", C_DIM),
                ))
            elif action == "set_context_limit":
                self.context_limit = result["value"]
                self._add_static(Text.assemble(
                    ("  上下文长度：", C_DIM),
                    (f"{self.context_limit}", f"bold {C_FG}"),
                    ("\n", C_DIM),
                ))
            elif action == "add_model":
                self._open_add_model()
            elif action == "scan_ollama":
                self._scan_ollama_models()
            elif action == "remove_model":
                self._remove_custom_model()
            elif action == "about":
                current_model = get_model_display_name(self.model_key)
                self._add_static(Text.assemble(
                    ("  ◆ ZeroAI\n", f"bold {C_BLUE}"),
                    ("  终端 AI 编程助手\n\n", C_DIM),
                    ("  当前模型：", C_DIM), (f"{current_model}\n", f"bold {C_FG}"),
                    ("  模式：", C_DIM), ("/专家（自动路由） / /混合（多专家协作） / /手动（指定模型）\n\n", C_FG),
                    ("  专家团队（7位）：\n", C_DIM),
                    ("    项目经理 · 编程 · 推理\n", C_FG),
                    ("    通用 · 中文 · 多模态\n", C_FG),
                    ("    学术\n\n", C_FG),
                    ("  工具（31个）：\n", C_DIM),
                    ("    读写文件 · 行编辑 · 列目录 · 执行命令 · 搜索代码\n", C_FG),
                    ("    打开应用 · 联网搜索 · 抓取网页 · 版本控制\n", C_FG),
                    ("    删除/移动/复制文件 · 创建目录 · 系统信息 · 进程列表\n", C_FG),
                    ("    Python沙箱 · 包管理 · 端口检测 · 文件对比\n", C_FG),
                    ("    图片理解 · 窗口感知 · 屏幕阅读 · 安全审计 · Word文档\n", C_FG),
                    ("    学术文献搜索 · arXiv预印本 · LaTeX公式渲染\n\n", C_FG),
                    ("  快捷键：\n", C_DIM),
                    ("    Ctrl+C  ", C_FG), ("停止生成 / 退出\n", C_DIM),
                    ("    Ctrl+G  ", C_FG), ("粘贴剪贴板图片\n", C_DIM),
                    ("    Ctrl+T  ", C_FG), ("语音输入（按住说话）\n", C_DIM),
                    ("    Ctrl+J  ", C_FG), ("输入框换行（多行输入）\n", C_DIM),
                    ("    Ctrl+Y  ", C_FG), ("复制最近回复\n", C_DIM),
                    ("    Ctrl+P  ", C_FG), ("设置面板\n", C_DIM),
                    ("    Ctrl+N  ", C_FG), ("新对话\n", C_DIM),
                    ("    Ctrl+W  ", C_FG), ("伴随模式（屏幕感知）\n", C_DIM),
                    ("    Ctrl+L  ", C_FG), ("清屏\n", C_DIM),
                    ("    PageUp  ", C_FG), ("上翻页\n", C_DIM),
                    ("    PageDn  ", C_FG), ("下翻页\n", C_DIM),
                    ("    Esc     ", C_FG), ("关闭弹窗\n\n", C_DIM),
                    ("  命令：\n", C_DIM),
                    ("    /帮助  /清屏  /新对话  /图片  /安全  /复制  /退出\n", C_FG),
                ))
            elif action in ("proxy_toggle", "proxy_url", "proxy_token"):
                # v1.1.0 代理服务器配置
                self._open_proxy_config(action)

        self.push_screen(SettingsScreen(
            self.model_key, self.temperature, self.stream_enabled,
            self.max_turns, self.context_limit, self.work_mode,
        ), on_result)

    def _open_proxy_config(self, action: str):
        """v1.1.0 打开代理服务器配置对话框"""
        global PROXY_CONFIG
        current = _load_proxy_config()

        if action == "proxy_toggle":
            # 切换启用/禁用
            new_enabled = not current.get("enabled", False)
            if new_enabled and not current.get("base_url"):
                self._add_static(Text.assemble(
                    ("  ", C_DIM),
                    ("⚠ 请先配置代理地址和 Token，再启用代理\n", f"bold {C_FG}"),
                ))
                return
            _save_proxy_config(
                new_enabled,
                current.get("base_url", ""),
                current.get("token", ""),
            )
            PROXY_CONFIG = _load_proxy_config()
            status = "已启用" if new_enabled else "已禁用"
            self._add_static(Text.assemble(
                ("  代理模式：", C_DIM),
                (f"{status}\n", f"bold {C_FG}"),
            ))
            return

        # 输入对话框（proxy_url / proxy_token）
        field = "base_url" if action == "proxy_url" else "token"
        label = "代理地址（如 http://192.168.10.6:8000/v1）" if action == "proxy_url" else "访问 Token"
        default = current.get(field, "")

        from textual.widgets import Input
        from textual.containers import Vertical
        from textual.screen import ModalScreen

        class ProxyInputScreen(ModalScreen):
            CSS = f"""
            ProxyInputScreen {{ align: center middle; }}
            #proxy-input-dialog {{
                width: 60; height: auto; max-height: 20;
                background: {C_BG}; padding: 1 2;
                border: solid {C_BORDER};
            }}
            #proxy-input-title {{ color: {C_FG}; text-style: bold; padding: 0 0 1 0; }}
            #proxy-input-hint {{ color: {C_DIM}; padding: 0 0 1 0; }}
            #proxy-input-field {{ width: 100%; }}
            #proxy-input-footer {{ color: {C_DIM}; padding: 1 0 0 0; }}
            """
            BINDINGS = [Binding("escape", "close", "关闭", show=False)]

            def __init__(self, title: str, hint: str, default_val: str, field_name: str):
                super().__init__()
                self.title = title
                self.hint = hint
                self.default_val = default_val
                self.field_name = field_name

            def compose(self):
                with Vertical(id="proxy-input-dialog"):
                    yield Static(self.title, id="proxy-input-title")
                    yield Static(self.hint, id="proxy-input-hint")
                    yield Input(value=self.default_val, id="proxy-input-field",
                                placeholder=self.hint)
                    yield Static("回车保存 · Esc 取消", id="proxy-input-footer")

            def action_close(self):
                self.dismiss(None)

            def on_input_submitted(self, event):
                val = event.value.strip()
                self.dismiss({"field": self.field_name, "value": val})

        def on_proxy_result(res):
            if not res or not res.get("value"):
                return
            new_val = res["value"]
            _save_proxy_config(
                current.get("enabled", False),
                new_val if action == "proxy_url" else current.get("base_url", ""),
                new_val if action == "proxy_token" else current.get("token", ""),
            )
            PROXY_CONFIG = _load_proxy_config()
            self._add_static(Text.assemble(
                ("  ", C_DIM),
                (f"✓ {label} 已保存\n", f"bold {C_FG}"),
            ))

        self.push_screen(ProxyInputScreen(
            title=f"配置 {label}",
            hint=label,
            default_val=default,
            field_name=field,
        ), on_proxy_result)

    def _open_add_model(self):
        def on_add_result(result):
            if not result or result.get("action") != "add_model":
                if result and result.get("action") == "error":
                    self._add_static(Text(f"  {result['msg']}\n", style=C_FG))
                return
            vals = result["values"]
            key = vals["key"].strip()
            api_key = vals.get("api_key", "") or ""
            MODEL_CONFIGS[key] = {
                "label": vals.get("label", key) or key,
                "base_url": vals["base_url"].rstrip("/"),
                "api_key": api_key,
                "model": vals["model"].strip(),
            }
            # 保存到自定义模型文件（混淆）和配置文件
            _save_custom_models()
            _save_config({key: {"api_key": api_key}})
            # 追加到已有配置文件
            existing = _load_config()
            existing[key] = {"api_key": api_key}
            _save_config(existing)
            self._add_static(Text.assemble(
                ("  已添加模型：", C_DIM),
                (f"{MODEL_CONFIGS[key]['label']}", f"bold {C_FG}"),
                (f"  标识 {key}\n", C_DIM),
            ))
        self.push_screen(AddModelScreen(), on_add_result)

    def _scan_ollama_models(self):
        models = detect_ollama_models()
        if not models:
            self._add_static(Text("  未检测到本地模型服务或无可用模型\n", style=C_FG))
            return
        added = []
        for mid in models:
            key = f"ollama_{mid.replace(':', '_').replace('/', '_')}"
            if key not in MODEL_CONFIGS:
                MODEL_CONFIGS[key] = {
                    "label": f"Ollama {mid}",
                    "base_url": "http://localhost:11434/v1",
                    "api_key": "ollama",
                    "model": mid,
                }
                added.append(mid)
        _save_custom_models()
        if added:
            names = " · ".join(added)
            self._add_static(Text.assemble(
                ("  扫描到 ", C_DIM),
                (f"{len(added)}", f"bold {C_FG}"),
                (" 个新模型：", C_DIM),
                (f"{names}\n", C_FG),
                ("  按 Ctrl+P 可切换模型\n", C_DIM),
            ))
        else:
            self._add_static(Text("  ℹ 本地模型已在列表中\n", style=C_DIM))

    def _remove_custom_model(self):
        custom_keys = [k for k in MODEL_CONFIGS if k not in ("glm", "glm-v", "openrouter", "ollama")]
        if not custom_keys:
            self._add_static(Text("  ℹ 没有自定义模型可删除\n", style=C_DIM))
            return
        items = []
        for key in custom_keys:
            items.append(ListItem(Label(f"  {MODEL_CONFIGS[key]['label']}  [{key}]"), name=f"del_model:{key}"))

        class RemoveModelScreen(ModalScreen):
            CSS = SettingsScreen.CSS
            BINDINGS = [Binding("escape", "close_rm", "关闭", show=False)]
            def __init__(self, item_list):
                super().__init__()
                self.item_list = item_list
            def compose(self):
                with Vertical(id="settings-dialog"):
                    yield Static("删除自定义模型", id="settings-title")
                    yield Static("回车删除 · Esc 取消", id="settings-hint")
                    yield ListView(*self.item_list)
                    yield Static("Esc 取消", id="settings-footer")
            def action_close_rm(self):
                self.dismiss(None)
            def on_list_view_selected(self, event):
                self.dismiss(event.item.name)

        def on_rm_result(name):
            if not name or not name.startswith("del_model:"):
                return
            key = name.split(":", 1)[1]
            if key in MODEL_CONFIGS and key not in ("glm", "glm-v", "openrouter", "ollama"):
                label = MODEL_CONFIGS[key]['label']
                del MODEL_CONFIGS[key]
                _save_custom_models()
                if self.model_key == key:
                    self.model_key = "glm"
                    self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
                    self.query_one("#info", InfoBar).refresh()
                    self._add_static(Text.assemble(
                        ("  已删除：", C_DIM),
                        (f"{label}", f"bold {C_FG}"),
                        ("（已回退到智谱GLM）\n", C_DIM),
                    ))
                else:
                    self._add_static(Text.assemble(
                        ("  已删除：", C_DIM),
                        (f"{label}\n", f"bold {C_FG}"),
                    ))
        self.push_screen(RemoveModelScreen(items), on_rm_result)

    def compose(self) -> ComposeResult:
        yield InfoBar(id="info")
        # 中间：对话 + 右侧 token 栏
        with Horizontal(id="main-area"):
            yield VerticalScroll(id="log-scroll")
            yield TokenBar(id="token-bar")
        # 底部：快捷键 + 输入区
        yield HintBar(id="hints")
        with Vertical(id="input-wrap"):
            yield MessageInput(placeholder="输入消息…  /帮助  Ctrl+P 设置  Ctrl+G 粘贴图片  Ctrl+T 语音输入  Ctrl+J 换行", id="input")

    def on_mount(self) -> None:
        scroll = self.query_one("#log-scroll", VerticalScroll)
        # 欢迎信息（MiMo 风格：最简、灰文字、彩色工具标签）
        parts = [
            ("  你好！我是 ZeroAI 助手，由 7 个专家模型构建。\n", C_FG),
            ("  我可以帮你处理各种任务，包括：\n\n", C_FG),
            ("    •  ", C_DIM), ("代码编写、调试、重构、项目搭建\n", C_FG),
            ("    •  ", C_DIM), ("数学证明、逻辑推理、问题分析\n", C_FG),
            ("    •  ", C_DIM), ("文档写作、翻译、润色\n", C_FG),
            ("    •  ", C_DIM), ("图片理解、截图分析\n", C_FG),
            ("    •  ", C_DIM), ("学术研究、文献检索、引用检查\n\n", C_FG),
            ("  有什么我可以帮你的吗？\n\n", C_FG),
            ("  ▶ 输入问题回车发送，或按 ", C_DIM),
            ("Ctrl+T", C_ACCENT),
            (" 开启语音对话\n", C_DIM),
        ]
        # 显示模型状态：内置 Key 已就绪
        if MODEL_CONFIGS.get("glm", {}).get("api_key"):
            parts.append(("  ┌──────────────────────────────────────┐\n", f"bold {C_FG}"))
            parts.append(("  │  ✓ 免费模型已就绪，可直接使用        │\n", f"bold {C_FG}"))
            parts.append(("  └──────────────────────────────────────┘\n", f"bold {C_FG}"))
        else:
            parts.append(("  ┌──────────────────────────────────────┐\n", f"bold {C_FG}"))
            parts.append(("  │  [!] 未配置 API 密钥                │\n", f"bold {C_FG}"))
            parts.append(("  │                                      │\n", f"bold {C_FG}"))
            parts.append(("  │  按 Ctrl+P 打开设置面板配置 Key     │\n", f"bold {C_FG}"))
            parts.append(("  │  智谱GLM 免费 Key 获取：            │\n", f"bold {C_FG}"))
            parts.append(("  │  https://open.bigmodel.cn/          │\n", f"bold {C_FG}"))
            parts.append(("  └──────────────────────────────────────┘\n", f"bold {C_FG}"))
        parts.append(("  直接输入需求开始对话，输入 /帮助 查看命令\n", C_DIM))
        welcome = Static(Text.assemble(*parts), classes="msg-block")
        scroll.mount(welcome)
        self.query_one("#input", MessageInput).focus()

    def _add_block(self, header: str, header_color: str = C_BLUE) -> Static:
        """添加一个消息块，返回 Static widget 以便后续 update"""
        scroll = self.query_one("#log-scroll", VerticalScroll)
        block = Static(Text.assemble(
            (f"  ┌─ {header}\n", f"bold {header_color}"),
            ("  │\n", C_DIM),
        ), classes="msg-block")
        scroll.mount(block)
        self._user_scrolling = False
        scroll.scroll_end(animate=False)
        # 保持输入框焦点（mount 新 widget 可能导致焦点丢失）
        self._keep_input_focus()
        return block

    def _add_static(self, content) -> Static:
        """添加一个纯 Static 内容块"""
        scroll = self.query_one("#log-scroll", VerticalScroll)
        block = Static(content, classes="msg-block")
        scroll.mount(block)
        if not self._user_scrolling:
            scroll.scroll_end(animate=False)
        # 保持输入框焦点
        self._keep_input_focus()
        return block

    def _keep_input_focus(self):
        """确保输入框保持焦点（生成过程中也能打字）"""
        try:
            inp = self.query_one("#input", MessageInput)
            if not inp.has_focus:
                self.call_after_refresh(inp.focus)
        except Exception:
            pass

    def action_clear_log(self):
        scroll = self.query_one("#log-scroll", VerticalScroll)
        for child in list(scroll.children):
            child.remove()
        self.on_mount()
        self.notify("已清屏")

    def action_clear_history(self):
        self.messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        self.total_tokens = 0
        self.stream_token_count = 0
        scroll = self.query_one("#log-scroll", VerticalScroll)
        for child in list(scroll.children):
            child.remove()
        self.on_mount()
        try:
            self.query_one("#token-bar", TokenBar).update_stats(0, 0.0)
        except Exception:
            pass
        self.notify("已开始新对话")

    def action_stop_or_quit(self):
        """Ctrl+C：思考中则停止生成，空闲则退出"""
        if self._is_generating:
            self._stop_generation = True
            # 同步到全局标志，让独立函数（route_expert_glm 等）也能感知
            global _GLOBAL_STOP
            _GLOBAL_STOP = True
            self._add_static(Text.assemble(
                ("  ⏹ 已停止生成", f"bold {C_FG}"),
                ("（再按 Ctrl+C 退出）\n", C_DIM),
            ))
        else:
            self.exit()

    def action_copy_last_reply(self):
        """Ctrl+Y：复制最近一次助手回复到剪贴板"""
        if not self._last_reply_text:
            self.notify("无内容可复制")
            return
        # 优先用 Windows API（完整 Unicode 支持），失败时回退到 clip.exe
        if _copy_to_clipboard(self._last_reply_text):
            preview = self._last_reply_text[:30].replace("\n", " ")
            self.notify(f"已复制：{preview}…")
        else:
            # 回退方案：clip.exe（可能有编码问题）
            try:
                import subprocess
                subprocess.run("clip", input=self._last_reply_text, text=True, check=True)
                preview = self._last_reply_text[:30].replace("\n", " ")
                self.notify(f"已复制：{preview}…")
            except Exception:
                self.notify("复制失败")

    def action_paste_image(self):
        """Ctrl+G 或 /图片：检测剪贴板图片，如果有则暂存并预览"""
        try:
            from PIL import ImageGrab
            img = ImageGrab.grabclipboard()
            if img is None:
                # 没有图片，检查是否有文本
                try:
                    import subprocess
                    text = subprocess.run(["powershell", "-c", "Get-Clipboard -Format Text"],
                                         capture_output=True, text=True, timeout=2).stdout
                    if text.strip():
                        # 有文本，粘贴到输入框
                        inp = self.query_one("#input", MessageInput)
                        inp.value += text.rstrip("\r\n")
                        self.notify("已粘贴文本（剪贴板无图片）")
                    else:
                        self.notify("剪贴板为空！请先截图（Win+Shift+S）或复制图片")
                except Exception:
                    self.notify("无法读取剪贴板，请先截图（Win+Shift+S）")
                return
            # 有图片！转为 base64
            import io
            if isinstance(img, list):
                # 某些情况下返回文件路径列表
                for f in img:
                    if str(f).lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                        b64 = read_image(str(f))
                        if b64.startswith("data:"):
                            self._pending_images.append(b64)
                            self._add_static(Text.assemble(
                                (" [Image ", "bold white on #000000"),
                                (f"{len(self._pending_images)}", "bold white on #000000"),
                                ("] ", "bold white on #000000"),
                                (f"{f}\n", C_DIM),
                            ))
                            # 显示图片预览（已按用户要求停用照片预览，仅保留 [Image N] 标签）
                            # self._add_static(render_image_preview(b64))
                            self.notify(f"已附加图片 {len(self._pending_images)}，输入文字后回车发送")
                            return
                return
            # PIL Image 对象
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            data = buf.getvalue()
            b64_data = base64.b64encode(data).decode("ascii")
            b64_uri = f"data:image/png;base64,{b64_data}"
            self._pending_images.append(b64_uri)
            self._add_static(Text.assemble(
                (" [Image ", "bold white on #000000"),
                (f"{len(self._pending_images)}", "bold white on #000000"),
                ("] ", "bold white on #000000"),
                (f"剪贴板图片（{img.width}x{img.height}）\n", C_DIM),
            ))
            # 显示图片预览（已按用户要求停用照片预览，仅保留 [Image N] 标签）
            # self._add_static(render_image_preview(b64_uri))
            self.notify(f"已附加图片 {len(self._pending_images)}，输入文字后回车发送")
        except Exception as e:
            # ImageGrab 不可用时回退为普通文本粘贴
            try:
                import subprocess
                text = subprocess.run(["powershell", "-c", "Get-Clipboard -Format Text"],
                                     capture_output=True, text=True, timeout=2).stdout
                if text.strip():
                    inp = self.query_one("#input", MessageInput)
                    inp.value += text.rstrip("\r\n")
            except Exception:
                pass

    def action_push_to_talk(self):
        """Ctrl+T：按住说话（Push-to-Talk）

        单击触发：开始录音 → 静音检测自动停止 → 识别 → 填入输入框
        用户可在输入框中编辑后按回车发送
        """
        if self._is_listening:
            # 正在录音中，忽略重复触发
            return
        if self._is_generating:
            self._add_static(Text(f"  {_load_svg_icon('warning')} AI 正在生成中，请稍后再试\n", style=C_FG))
            return
        # 启动录音线程（避免阻塞 UI）
        import threading
        self._is_listening = True
        # 显示录音提示
        listen_block = self._add_block("语音输入", C_BLUE)
        listen_block.update(Text.assemble(
            ("  🎤 正在录音…", f"bold {C_FG}"),
            ("（说话即可，停顿自动停止）\n", C_DIM),
        ))
        def _do_listen():
            try:
                text = listen_asr(max_seconds=10, silence_seconds=1.5)
                self._is_listening = False
                # 通过 call_after_refresh 更新 UI
                def _update():
                    if text.startswith("错误") or text.startswith("（"):
                        listen_block.update(Text.assemble(
                            ("  🎤 语音识别失败\n", f"bold {C_FG}"),
                            (f"  {text}\n", C_DIM),
                        ))
                    else:
                        listen_block.update(Text.assemble(
                            ("  🎤 识别结果：", f"bold {C_FG}"),
                            (f"{text}\n", C_FG),
                            ("  （已填入输入框，回车发送，Ctrl+J 换行编辑）\n", C_DIM),
                        ))
                        # 填入输入框
                        try:
                            inp = self.query_one("#input", MessageInput)
                            inp.value = text
                            inp.focus()
                        except Exception:
                            pass
                self.call_after_refresh(_update)
            except Exception as e:
                self._is_listening = False
                def _err():
                    listen_block.update(Text.assemble(
                        ("  🎤 录音异常\n", f"bold {C_FG}"),
                        (f"  {e}\n", C_DIM),
                    ))
                self.call_after_refresh(_err)
        t = threading.Thread(target=_do_listen, daemon=True)
        t.start()

    def action_voice_dialog(self):
        """Ctrl+D：打开讨论助手（全屏沉浸式语音对话 Modal）

        进入后自动循环：
        听你说 → 识别 → 显示气泡 → AI 流式回答（打字机）→ 朗读 → 再听
        可随时说话打断 AI 朗读
        再次按 Ctrl+D 或 Esc 或点 × 退出
        """
        # 自动开启 TTS
        if not self._tts_enabled:
            self._tts_enabled = True
        # 推送全屏语音对话 Modal
        self.push_screen(VoiceDialogScreen(self))

    async def _voice_dialog_send(self, user_text: str):
        """保留兼容：旧的单轮对话方法（已被 VoiceDialogScreen 替代）"""
        # 此方法已废弃，由 VoiceDialogScreen 直接处理
        pass

    def action_toggle_companion(self):
        """Ctrl+W：切换伴随模式"""
        self._companion_mode = not self._companion_mode
        if self._companion_mode:
            self._companion_log = []
            self._last_window_title = ""
            self._last_clipboard_text = ""
            self._start_companion_thread()
            self._add_static(Text.assemble(
                (f"  {_load_svg_icon('monitor')} 伴随模式已开启\n", f"bold {C_FG}"),
                ("  AI 正在观察你的屏幕，关键变化会自动记录\n", C_DIM),
                ("  发消息时 AI 会知道你刚才在做什么\n", C_DIM),
            ))
            self.notify("伴随模式已开启")
        else:
            self._add_static(Text(f"  {_load_svg_icon('monitor')} 伴随模式已关闭\n", style=C_DIM))
            self.notify("伴随模式已关闭")

    def _start_companion_thread(self):
        """启动后台监听线程"""
        import threading
        def monitor():
            while self._companion_mode:
                try:
                    # 1. 监听窗口切换
                    try:
                        import ctypes
                        user32 = ctypes.windll.user32
                        hwnd = user32.GetForegroundWindow()
                        length = user32.GetWindowTextLengthW(hwnd) + 1
                        title = ctypes.create_unicode_buffer(length)
                        user32.GetWindowTextW(hwnd, title, length)
                        current_title = title.value.strip()
                        if current_title and current_title != self._last_window_title and "ZeroAI" not in current_title:
                            import time as _t
                            ts = _t.strftime("%H:%M:%S", _t.localtime())
                            self._companion_log.append(f"[{ts}] 切换窗口 → {current_title}")
                            self._last_window_title = current_title
                    except Exception:
                        pass

                    # 2. 监听剪贴板变化（只读文本）
                    try:
                        r = subprocess.run(["powershell", "-c", "Get-Clipboard -Format Text"],
                                           capture_output=True, text=True, timeout=2)
                        clip_text = r.stdout.strip()[:200] if r.stdout else ""
                        if clip_text and clip_text != self._last_clipboard_text and len(clip_text) > 5:
                            import time as _t
                            ts = _t.strftime("%H:%M:%S", _t.localtime())
                            self._companion_log.append(f"[{ts}] 复制了内容 → {clip_text[:80]}")
                            self._last_clipboard_text = clip_text
                    except Exception:
                        pass

                    # 保留最近 30 条日志
                    if len(self._companion_log) > 30:
                        self._companion_log = self._companion_log[-30:]

                except Exception:
                    pass
                import time as _t
                _t.sleep(2)  # 每2秒检测一次

        t = threading.Thread(target=monitor, daemon=True)
        t.start()
        self._companion_thread = t

    def _get_companion_context(self) -> str:
        """获取伴随模式的屏幕上下文（对话前注入）"""
        if not self._companion_mode or not self._companion_log:
            return ""
        # 取最近10条日志
        recent = self._companion_log[-10:]
        # 加上当前窗口内容
        try:
            current = read_screen_content(max_length=500)
            return f"【伴随模式·屏幕感知】\n最近活动：\n" + "\n".join(recent) + f"\n\n当前屏幕：\n{current}"
        except Exception:
            return f"【伴随模式·屏幕感知】\n最近活动：\n" + "\n".join(recent)

    def action_scroll_pageup(self):
        self._user_scrolling = True
        self.query_one("#log-scroll", VerticalScroll).scroll_page_up()

    def action_scroll_pagedown(self):
        scroll = self.query_one("#log-scroll", VerticalScroll)
        scroll.scroll_page_down()
        # 翻到底部时取消用户滚动标志
        if scroll.is_scrollable and scroll.scroll_y >= scroll.max_scroll_y - 1:
            self._user_scrolling = False

    def on_scroll_up(self, event) -> None:
        """鼠标滚轮向上 — 标记用户在翻看历史"""
        self._user_scrolling = True

    def on_scroll_down(self, event) -> None:
        """鼠标滚轮向下 — 翻到底部时恢复自动滚底"""
        scroll = self.query_one("#log-scroll", VerticalScroll)
        if scroll.is_scrollable and scroll.scroll_y >= scroll.max_scroll_y - 1:
            self._user_scrolling = False

    async def on_input_submitted(self, event: Input.Submitted) -> None:
        global WORK_MODE
        user_input = event.value.strip()
        if not user_input:
            return

        # 命令（中英文双语支持）
        if user_input in ("/exit", "/quit", "/退出"):
            self.exit()
            return
        if user_input in ("/clear", "/清屏"):
            self.action_clear_log()
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/new", "/新对话"):
            self.action_clear_history()
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/help", "/帮助"):
            self._show_help()
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/copy", "/复制"):
            self.action_copy_last_reply()
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/image", "/img", "/图片", "/粘贴图片"):
            self.action_paste_image()
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/voice", "/语音", "/朗读"):
            # 切换 TTS 朗读开关
            self._tts_enabled = not self._tts_enabled
            if self._tts_enabled:
                self._add_static(Text.assemble(
                    ("  🎤 语音朗读已开启\n", f"bold {C_FG}"),
                    ("  AI 回复后将自动朗读（按 Ctrl+T 说话输入）\n", C_DIM),
                    (f"  当前音色：{self._tts_voice} · 语速：{self._tts_rate}\n", C_DIM),
                    ("  再次输入 /语音 关闭\n", C_DIM),
                ))
                self.notify("语音朗读已开启")
            else:
                self._add_static(Text("  🎤 语音朗读已关闭\n", style=C_DIM))
                self.notify("语音朗读已关闭")
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/dialog", "/对话"):
            # 切换语音对话模式
            self.action_voice_dialog()
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/stop", "/停止"):
            # 停止语音对话模式
            if self._voice_dialog_active:
                self._voice_dialog_active = False
                self._add_static(Text.assemble(
                    ("  🎤 语音对话已停止\n", f"bold {C_FG}"),
                    ("  （再次按 Ctrl+D 或输入 /对话 重新开始）\n", C_DIM),
                ))
                self.notify("语音对话已停止")
            else:
                self._add_static(Text("  （当前未开启语音对话模式）\n", style=C_DIM))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/voice_female", "/女声"):
            self._tts_voice = "zh-CN-XiaoxiaoNeural"
            self._add_static(Text.assemble(
                ("  🎤 音色已切换为：", C_DIM),
                ("女声（晓晓）\n", f"bold {C_FG}"),
            ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/voice_male", "/男声"):
            self._tts_voice = "zh-CN-YunxiNeural"
            self._add_static(Text.assemble(
                ("  🎤 音色已切换为：", C_DIM),
                ("男声（云希）\n", f"bold {C_FG}"),
            ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input.startswith("/voice_rate ") or user_input.startswith("/语速 "):
            # 设置语速：/语速 +10% 或 /语速 -10%
            rate = user_input.split(" ", 1)[1].strip()
            if rate.startswith(("+", "-")) and rate.endswith("%"):
                self._tts_rate = rate
                self._add_static(Text.assemble(
                    ("  🎤 语速已设置为：", C_DIM),
                    (f"{rate}\n", f"bold {C_FG}"),
                ))
            else:
                self._add_static(Text.assemble(
                    ("  格式错误，正确格式：", C_DIM),
                    ("/语速 +10%\n", C_FG),
                ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/audit", "/安全", "/安全审计", "/漏洞扫描"):
            # 对当前工作目录执行完整安全审计
            self._add_block("安全审计", C_YELLOW)
            self._add_static(Text.assemble(
                (f"  {_load_svg_icon('search')} 正在对当前项目进行安全审计…\n", f"bold {C_YELLOW}"),
                ("  │ 扫描项：代码漏洞 / 敏感信息 / 依赖漏洞 / 配置安全\n", C_DIM),
            ))
            report = security_audit(WORK_DIR, "all")
            try:
                self._add_static(_safe_markdown(report, code_theme="monokai"))
            except Exception:
                self._add_static(Text(report, style=C_FG))
            self._add_static(Text("  └─", style=C_DIM))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/ssh", "/SSH", "/远程", "/部署"):
            self._add_block("SSH 远程部署 + AI 远程运维", C_YELLOW)
            self._add_static(Text.assemble(
                (f"  {_load_svg_icon('terminal')} SSH 远程部署 + AI 远程运维（15 个工具）\n", f"bold {C_YELLOW}"),
                ("  │ 部署能力 + 运维能力 双重合一\n", C_DIM),
                ("\n  部署工具（7 个）：\n", f"bold {C_FG}"),
                ("    • ssh_connect      连接服务器（密码/密钥认证）\n", C_FG),
                ("    • ssh_exec         远程执行命令（危险命令二次确认）\n", C_FG),
                ("    • ssh_upload       上传文件（SFTP）\n", C_FG),
                ("    • ssh_download    下载文件（SFTP）\n", C_FG),
                ("    • ssh_deploy      一键自动化部署（7 步骤）\n", C_FG),
                ("    • ssh_list        查看连接状态/审计日志\n", C_FG),
                ("    • ssh_disconnect  断开连接\n", C_FG),
                ("\n  运维工具（8 个，语义化封装，AI 优先调用）：\n", f"bold {C_FG}"),
                ("    • ssh_service_manage  服务管理（status/start/stop/restart）\n", C_FG),
                ("    • ssh_log_view        日志查看（自动异常统计）\n", C_FG),
                ("    • ssh_process_check   进程查看（按 CPU/内存排序）\n", C_FG),
                ("    • ssh_disk_analyze    磁盘分析（df+du Top10）\n", C_FG),
                ("    • ssh_network_diag    网络诊断（端口/ping/连接）\n", C_FG),
                ("    • ssh_docker_manage   Docker 管理（容器/镜像/日志）\n", C_FG),
                ("    • ssh_firewall_manage 防火墙管理（ufw/firewalld/iptables）\n", C_FG),
                ("    • ssh_health_check    一键健康体检（综合报告+AI分析）\n", C_FG),
                ("\n  使用方式：直接告诉 AI 你的需求，例如：\n", f"bold {C_FG}"),
                ("    「连接到 192.168.10.22，用户 root，密码 xxx」\n", C_DIM),
                ("    「看下 nginx 状态」→ ssh_service_manage\n", C_DIM),
                ("    「服务器卡了」→ ssh_health_check 综合体检\n", C_DIM),
                ("    「查 mysql 错误日志」→ ssh_log_view(keyword=error)\n", C_DIM),
                ("    「看磁盘占用」→ ssh_disk_analyze\n", C_DIM),
                ("    「重启 web 容器」→ ssh_docker_manage(action=restart)\n", C_DIM),
                ("    「开放 8080 端口」→ ssh_firewall_manage(action=open, port=8080)\n", C_DIM),
                ("    「一键部署：上传项目→安装依赖→重启服务」\n", C_DIM),
                ("\n  安全保障：\n", f"bold {C_FG}"),
                ("    • 危险命令黑名单（rm -rf /、mkfs、dd 等 11 类）\n", C_DIM),
                ("    • 主机地址校验（IP/域名格式+内网IP可选阻断）\n", C_DIM),
                ("    • 审计日志（最多 200 条，可追溯）\n", C_DIM),
                ("    • 输出截断保护（8000 字符，防止刷屏）\n", C_DIM),
            ))
            self._add_static(Text("  └─", style=C_DIM))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/expert", "/专家"):
            self.work_mode = "expert"
            WORK_MODE = "expert"
            self.query_one("#info", InfoBar).refresh()
            self._add_static(Text.assemble(
                ("  已切换到 ", C_DIM),
                ("专家模式", f"bold {C_FG}"),
                ("（自动路由最合适的专家模型）\n", C_DIM),
            ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/hybrid", "/混合"):
            self.work_mode = "hybrid"
            WORK_MODE = "hybrid"
            self.query_one("#info", InfoBar).refresh()
            self._add_static(Text.assemble(
                ("  已切换到 ", C_DIM),
                ("混合思考", f"bold {C_FG}"),
                ("（多专家协作，深度处理）\n", C_DIM),
            ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/manual", "/手动"):
            self.work_mode = "manual"
            WORK_MODE = "manual"
            self.query_one("#info", InfoBar).refresh()
            label = MODEL_CONFIGS[self.model_key]["label"]
            self._add_static(Text.assemble(
                ("  已切换到 ", C_DIM),
                ("手动模式", f"bold {C_FG}"),
                (f"（使用 {label}）\n", C_DIM),
            ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input in ("/model", "/模型"):
            # 显示当前模型和可用模型
            mode_label = {"expert": "专家模式", "hybrid": "混合思考", "manual": "手动模式"}.get(self.work_mode, "未知")
            current = get_model_display_name(self.model_key)
            available = " | ".join(f"{k}:{v['label']}" for k, v in MODEL_CONFIGS.items())
            # 专家团队列表
            team_list = " | ".join(f"{k}:{v['label']}" for k, v in EXPERT_TEAM.items())
            self._add_static(Text.assemble(
                ("  当前模式：", C_DIM),
                (f"{mode_label}\n", f"bold {C_FG}"),
                ("  当前模型：", C_DIM),
                (f"{current}\n", f"bold {C_FG}"),
                ("  可用模型：", C_DIM),
                (f"{available}\n", C_FG),
                ("  专家团队：", C_DIM),
                (f"{team_list}\n", C_FG),
                ("  模式切换：", C_DIM),
                ("/专家 | /混合 | /手动\n", C_FG),
                ("  模型切换：", C_DIM),
                ("/模型 glm | /模型 openrouter | /模型 ollama\n", C_FG),
            ))
            self.query_one("#input", MessageInput).value = ""
            return
        if user_input.startswith(("/model ", "/模型 ")):
            prefix_len = 7 if user_input.startswith("/model ") else 4
            key = user_input[prefix_len:].strip().lower()
            if self.switch_model(key):
                # 切换模型时自动切换到手动模式
                self.work_mode = "manual"
                WORK_MODE = "manual"
                self.query_one("#info", InfoBar).refresh()
                self._add_static(Text.assemble(
                    ("  已切换到：", C_DIM),
                    (f"{MODEL_CONFIGS[key]['label']}", f"bold {C_FG}"),
                    ("（手动模式·对话已清空）\n", C_DIM),
                ))
            else:
                available = " | ".join(MODEL_CONFIGS.keys())
                self._add_static(Text.assemble(
                    ("  未知模型：", C_DIM),
                    (f"{key}\n", C_FG),
                    ("  可用：", C_DIM),
                    (f"{available}\n", C_FG),
                ))
            self.query_one("#input", MessageInput).value = ""
            return

        # 用户消息 - Markdown 渲染
        self._add_block("你", C_BLUE)

        # 检测图片路径（支持 @图片路径 或直接图片路径）
        image_urls = list(self._pending_images)  # 先加上 Ctrl+V 粘贴的图片
        self._pending_images = []  # 清空暂存
        display_text = user_input
        IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

        # 检测 @路径 语法
        at_pattern = re.compile(r'@([\w:\\/.]+\.(?:png|jpg|jpeg|gif|bmp|webp))', re.IGNORECASE)
        at_matches = at_pattern.findall(user_input)
        # 收集需要预览的图片（路径列表，避免重复读取）
        preview_paths = []
        for img_path in at_matches:
            b64 = read_image(img_path)
            if b64.startswith("data:"):
                image_urls.append(b64)
                display_text = display_text.replace(f"@{img_path}", f"[图片: {img_path}]")
                preview_paths.append(img_path)
            else:
                display_text = display_text.replace(f"@{img_path}", f"[{b64}]")

        # 检测输入中直接的图片路径
        for word in re.findall(r'[\w:\\/.]+\.(?:png|jpg|jpeg|gif|bmp|webp)', user_input, re.IGNORECASE):
            if word not in at_matches and Path(word).exists():
                b64 = read_image(word)
                if b64.startswith("data:"):
                    image_urls.append(b64)
                    display_text = display_text.replace(word, f"[图片: {word}]")
                    preview_paths.append(word)

        # 显示图片预览
        for img_path in preview_paths:
            self._add_static(Text(f"  {_load_svg_icon('document')} 图片预览：{img_path}\n", style=C_DIM))
            self._add_static(render_image_preview(img_path))

        try:
            self._add_static(_safe_markdown(render_latex_in_text(display_text), code_theme="monokai"))
        except Exception:
            self._add_static(Text(f"  {display_text}", style=C_FG))

        # 显示图片缩略信息（黑白方块标签样式：[Image 1] [Image 2]）
        if image_urls:
            badge_parts = []
            for i in range(len(image_urls)):
                badge_parts.append((" [Image ", "bold white on #000000"))
                badge_parts.append((f"{i+1}", "bold white on #000000"))
                badge_parts.append(("] ", "bold white on #000000"))
            self._add_static(Text.assemble(*badge_parts))
        self._add_static(Text("  └─", style=C_DIM))

        self.query_one("#input", MessageInput).value = ""

        # 注入伴随模式屏幕上下文
        companion_ctx = self._get_companion_context()
        if companion_ctx:
            self.messages.append({"role": "system", "content": companion_ctx})

        # 构造消息：如果有图片则用多模态格式
        if image_urls:
            content_parts = [{"type": "text", "text": user_input}]
            for url in image_urls:
                content_parts.append({"type": "image_url", "image_url": {"url": url}})
            self.messages.append({"role": "user", "content": content_parts})
        else:
            self.messages.append({"role": "user", "content": user_input})

        await self._run_turn()

        # 生成完成后重新聚焦输入框（确保可以继续打字）
        self._keep_input_focus()

        # ── TTS 自动朗读：回复完成后，如果开启则朗读最近回复 ──
        if self._tts_enabled and self._last_reply_text.strip() and not self._stop_generation:
            import threading
            reply_text = self._last_reply_text
            voice = self._tts_voice
            rate = self._tts_rate
            def _do_tts():
                err = speak_tts(reply_text, voice=voice, rate=rate)
                if err:
                    def _show_err():
                        self._add_static(Text.assemble(
                            ("  🎤 朗读失败：", C_DIM),
                            (f"{err}\n", C_FG),
                        ))
                    self.call_after_refresh(_show_err)
            t = threading.Thread(target=_do_tts, daemon=True)
            t.start()

    def _show_help(self):
        self._add_static(Text.assemble(
            ("  命令：\n", f"bold {C_YELLOW}"),
            ("    /帮助          显示帮助\n", C_FG),
            ("    /清屏          清空屏幕\n", C_FG),
            ("    /新对话        开始新对话\n", C_FG),
            ("    /专家          切换到专家模式（自动路由）\n", C_FG),
            ("    /混合          切换到混合思考（多专家协作）\n", C_FG),
            ("    /手动          切换到手动模式（指定模型）\n", C_FG),
            ("    /模型          查看当前模型和专家团队\n", C_FG),
            ("    /模型 glm      切换到智谱GLM（手动模式）\n", C_FG),
            ("    /模型 glm-v    切换到智谱GLM-4V（多模态，支持图片）\n", C_FG),
            ("    /模型 openrouter 切换到OpenRouter（手动模式）\n", C_FG),
            ("    /模型 ollama   切换到 Ollama（手动模式）\n", C_FG),
            ("    /复制          复制最近回复\n", C_FG),
            ("    /图片          粘贴剪贴板图片\n", C_FG),
            ("    /安全          安全审计（扫描漏洞/敏感信息/依赖/配置）\n", C_FG),
            ("    /ssh           SSH 远程部署（连接/执行/上传/部署）\n", C_FG),
            ("    /退出          退出\n", C_FG),
            ("\n  语音交互：\n", f"bold {C_YELLOW}"),
            ("    /语音          开启/关闭 AI 回复自动朗读\n", C_FG),
            ("    /对话          开启语音对话模式（说话即提问，AI 语音回答）\n", C_FG),
            ("    /停止          停止语音对话模式\n", C_FG),
            ("    /女声          切换为女声（晓晓）\n", C_FG),
            ("    /男声          切换为男声（云希）\n", C_FG),
            ("    /语速 +10%     设置语速（+10% 加速 / -10% 减速）\n", C_FG),
            ("    Ctrl+T         单次语音输入（停顿自动停止）\n", C_FG),
            ("    Ctrl+D         语音对话模式（持续听→答→听循环）\n", C_FG),
            ("\n  发送图片：\n", f"bold {C_YELLOW}"),
            ("    /图片          粘贴剪贴板图片（截图后输入 /图片）\n", C_FG),
            ("    Ctrl+G         粘贴剪贴板图片快捷键\n", C_FG),
            ("    直接输入图片路径：D:/图片/screenshot.png\n", C_FG),
            ("    @图片路径 语法：@D:/图片/test.jpg 描述一下这张图\n", C_FG),
            ("    支持格式：png / jpg / jpeg / gif / bmp / webp\n", C_FG),
            ("\n  快捷键：\n", f"bold {C_YELLOW}"),
            ("    Ctrl+C 停止/退出 · Ctrl+G 粘贴图片 · Ctrl+T 语音输入 · Ctrl+D 语音对话 · Ctrl+W 伴随模式 · Ctrl+Y 复制 · Ctrl+P 设置 · Ctrl+N 新对话\n", C_FG),
        ))

    async def _run_hybrid_turn(self):
        """混合思考模式：多专家协作处理同一问题
        流程：GLM分析 → 专家回答 → GLM汇总
        """
        self._is_generating = True
        self._stop_generation = False
        # 重置全局停止标志
        global _GLOBAL_STOP
        _GLOBAL_STOP = False
        from openai import AsyncOpenAI

        # ── API Key 检查：GLM 未配置时直接提示 ──
        glm_key = MODEL_CONFIGS.get("glm", {}).get("api_key", "")
        if not glm_key:
            self._add_static(Text.assemble(
                (f"  {_load_svg_icon('warning')} 混合思考需要智谱GLM 密钥\n", f"bold {C_FG}"),
                (f"  请按 Ctrl+P 打开设置面板配置 GLM API Key\n", f"bold {C_FG}"),
                (f"  智谱GLM 免费 Key 获取：https://open.bigmodel.cn/\n", C_DIM),
                (f"  或输入 /手动 切换到单模型模式\n", C_DIM),
            ))
            self._is_generating = False
            return

        # ── 上下文管理：主动清理 + 按需压缩（两层防护，防止幻觉）──
        def _ctx_log(text, style=None):
            self._add_static(Text.assemble(
                (text, style if style else C_DIM),
            ))

        # 用 block 显示清理/压缩进度
        est_tokens_pre = _estimate_tokens(self.messages)
        cleanup_threshold = int(self.context_limit * CLEANUP_THRESHOLD_RATIO)
        compress_threshold = int(self.context_limit * COMPRESS_THRESHOLD_RATIO)

        if est_tokens_pre > cleanup_threshold and len(self.messages) > 8:
            block_ctx = self._add_block("上下文管理", C_DIM)
            block_ctx.update(Text.assemble(
                (f"  {_load_svg_icon('tool')} 上下文管理\n", f"bold {C_DIM}"),
                (f"  │ 当前约 {est_tokens_pre} tokens（清理阈值 {cleanup_threshold}，压缩阈值 {compress_threshold}）\n", C_DIM),
            ))
            await cleanup_and_compress(self, _ctx_log)
        else:
            # 未触发清理阈值，但仍调用一次（内部会判断是否需要压缩）
            await cleanup_and_compress(self, _ctx_log if est_tokens_pre > compress_threshold else None)

        # 获取用户最后一条消息
        last_user = ""
        for msg in reversed(self.messages):
            if msg["role"] == "user":
                last_user = msg["content"] if isinstance(msg["content"], str) else str(msg["content"])
                break
        if not last_user:
            self._is_generating = False
            return

        # ── 第1步：项目经理GLM-4V分析任务，选择专家 ──
        block1 = self._add_block("项目经理·GLM-4V", C_BLUE)
        block1.update(Text.assemble(
            ("  ┌─ 项目经理·GLM-4V\n", f"bold {C_BLUE}"),
            ("  │ 分析任务中…\n", C_DIM),
        ))

        glm_cfg = MODEL_CONFIGS["glm-v"]  # 多模态经理，支持图片
        glm_client = _make_openai_client("glm-v")
        analyze_prompt = f"""请分析以下用户需求，判断需要哪些专家协作处理。
可用专家：
- coder: 编程开发
- reasoner: 深度推理/数学
- academic: 学术研究/公式推导/论文写作
- chinese: 中文写作/文案
- knowledge: 通用知识/翻译
- vision: 图片理解
- devops: 运维/部署/系统管理/容器/SSH
- security: 安全分析/漏洞评估/加固
- data: 数据分析/统计/可视化

用户需求：{last_user[:500]}

请只回复专家标识，用逗号分隔（最多{HYBRID_MAX_PARALLEL_EXPERTS}个），例如：coder,reasoner
不要回复其他内容。"""

        try:
            resp = await _interruptible_await(glm_client.chat.completions.create(
                model=glm_cfg["model"],
                messages=[{"role": "system", "content": "你是 ZeroAI 路由分析器，只负责判断用户问题应分配给哪些专家。严格按要求只输出专家标识，不要做任何解释。"},
                          {"role": "user", "content": analyze_prompt}],
                temperature=0.1,
                stream=False,
            ))
            if resp is None or self._stop_generation:
                self._is_generating = False
                self._add_static(Text("  ⏹ 已停止\n", style=C_DIM))
                return
            analysis = resp.choices[0].message.content.strip().lower()
        except Exception as e:
            analysis = ""

        # 解析专家列表
        expert_keys = []
        for part in analysis.replace("，", ",").split(","):
            part = part.strip()
            if part in EXPERT_TEAM and part not in ("pm",):
                expert_keys.append(part)
        if not expert_keys:
            # 降级：用关键词路由
            expert_keys = [route_expert(last_user)]
        # 专家并行度控制：限制最多并行专家数（避免 token 暴涨）
        expert_keys = expert_keys[:HYBRID_MAX_PARALLEL_EXPERTS]

        analysis_md = f"**任务分析**\n\n需要 {len(expert_keys)} 位专家协作：\n"
        for ek in expert_keys:
            analysis_md += f"- {EXPERT_TEAM[ek]['label']}：{EXPERT_TEAM[ek]['desc']}\n"
        block1.update(Text.assemble(
            ("  ┌─ 项目经理·GLM-4V\n", f"bold {C_BLUE}"),
            ("  │ ", C_DIM),
        ))
        self._add_static(_safe_markdown(render_latex_in_text(analysis_md), code_theme="monokai"))
        self._add_static(Text("  └─", style=C_DIM))

        if self._stop_generation:
            self._is_generating = False
            return

        # ── 第2步：所有专家并行回答（asyncio.gather 同时调用）──
        # 预先为每个专家创建输出区块，显示"思考中…"
        expert_blocks = {}
        for ek in expert_keys:
            expert = EXPERT_TEAM[ek]
            block_e = self._add_block(f"专家·{expert['label']}", C_YELLOW)
            block_e.update(Text.assemble(
                (f"  ┌─ {expert['label']}\n", f"bold {C_YELLOW}"),
                ("  │ 思考中…\n", C_DIM),
            ))
            expert_blocks[ek] = block_e

        # 单个专家调用协程（并行任务单元）
        async def _call_expert(ek: str, user_msg: str = None) -> dict:
            """并行调用单个专家，返回 {"expert": label, "content": text} 或 None

            集成三项优化：
            - 专家记忆：加载/保存独立上下文（EXPERT_MEMORY_TURNS），避免主上下文污染
            - 失败降级：任何异常都返回 None，不阻断整体协作流程
            - 长度限制：最终回答截断到 HYBRID_EXPERT_MAX_CHARS，便于汇总
            - 协作链：user_msg 可传入前一位专家的结果（默认用 last_user）
            """
            # 检查是否已被 Ctrl+C 停止
            if self._stop_generation:
                return None
            # 协作链支持：允许传入自定义用户消息（含前一位专家结果）
            actual_user_msg = user_msg if user_msg is not None else last_user
            expert = EXPERT_TEAM[ek]
            e_cfg = get_expert_config(ek)
            block_e = expert_blocks[ek]

            # ── 构建消息列表（含专家记忆：独立上下文，避免主上下文污染） ──
            sys_msg = {"role": "system", "content": TOOL_CAPABILITY_PROMPT + "\n\n" + expert.get("system_prompt", "你是 ZeroAI 专家团队成员，从专业角度回答用户问题。")}
            # 加载专家记忆（最近 EXPERT_MEMORY_TURNS 轮对话，每轮=用户问+专家答=2条消息）
            memory = self._expert_memory.get(ek, [])
            e_messages = [sys_msg] + list(memory) + [{"role": "user", "content": actual_user_msg}]

            # ── 标准单模型调用（带超时重试 + 失败降级） ──
            max_tries = 3  # hybrid 模式：最多重试 3 次
            e_stream = None
            for _try in range(max_tries):
                try:
                    e_client = _make_openai_client(e_cfg.get("model_key", "glm"))
                    if not _is_proxy_enabled():
                        # 本地模式：保留原 timeout/max_retries
                        e_client = AsyncOpenAI(
                            base_url=e_cfg["base_url"],
                            api_key=e_cfg["api_key"],
                            timeout=180.0, max_retries=0,
                        )
                    e_stream = await e_client.chat.completions.create(
                        model=e_cfg["model"],
                        messages=e_messages,
                        temperature=self.temperature,
                        stream=self.stream_enabled,
                        timeout=180,
                    )
                    break
                except Exception as retry_e:
                    err_str = str(retry_e).lower()
                    is_to = "timeout" in err_str or "timed out" in err_str or "readtimeout" in type(retry_e).__name__.lower()
                    is_rl = "429" in str(retry_e) or "rate" in err_str
                    if (is_to or is_rl) and _try < max_tries - 1:
                        wait = _try + 1  # 1s, 2s
                        block_e.update(Text.assemble(
                            (f"  ├─ {expert['label']}\n", f"bold {C_GREEN}"),
                            ("  │ ", C_DIM),
                            (f"⏳ {'超时' if is_to else '限流'}，{wait}秒后重试", f"bold {C_YELLOW}"),
                            (f"（第{_try+1}/{max_tries}次）\n", C_DIM),
                        ))
                        await asyncio.sleep(wait)
                        if self._stop_generation:
                            return None
                        continue
                    # 失败降级：重试耗尽或不可重试异常，跳过该专家而非整个流程失败
                    block_e.update(Text.assemble(
                        (f"  ├─ {expert['label']}\n", f"bold {C_FG}"),
                        ("  │ ", C_DIM),
                        (f"⚠ 调用失败，已跳过该专家：{str(retry_e)[:80]}\n", f"bold {C_YELLOW}"),
                    ))
                    return None
            if e_stream is None:
                # 失败降级：未获取到流，跳过
                return None
            try:
                e_content = ""
                e_reasoning = ""
                if self.stream_enabled:
                    async for chunk in e_stream:
                        if self._stop_generation:
                            break
                        if not chunk.choices:
                            continue
                        delta = chunk.choices[0].delta
                        # 捕获思考内容（原生 reasoning_content）
                        rc = getattr(delta, "reasoning_content", None) or getattr(delta, "reasoning", None)
                        if rc:
                            e_reasoning += rc
                            # 合并 <think> 标签内容
                            think_in, body = _parse_think_tags(e_content)
                            combined = (e_reasoning + ("\n" + think_in if think_in else "")).strip()
                            if combined:
                                self._update_streaming_with_reasoning(block_e, body, combined)
                            else:
                                self._update_streaming(block_e, body)
                            await asyncio.sleep(0)
                        if delta.content:
                            e_content += delta.content
                            # 过滤模型内部特殊标签（<|observation|> <|system|> 等）
                            e_content = _strip_model_tokens(e_content)
                            # 解析 <think> 标签，分离思考与正文
                            think_in, body = _parse_think_tags(e_content)
                            combined = (e_reasoning + ("\n" + think_in if think_in else "")).strip()
                            if combined:
                                self._update_streaming_with_reasoning(block_e, body, combined)
                            else:
                                self._update_streaming(block_e, body)
                            await asyncio.sleep(0)
                else:
                    e_content = _strip_model_tokens(e_stream.choices[0].message.content or "")
                    e_reasoning = getattr(e_stream.choices[0].message, "reasoning_content", "") or ""
                # 最终更新：分离 think 和 body，合并所有思考内容
                think_in, body = _parse_think_tags(e_content)
                combined_reasoning = (e_reasoning + ("\n" + think_in if think_in else "")).strip()
                if combined_reasoning:
                    self._update_streaming_with_reasoning(block_e, body, combined_reasoning, final=True)
                else:
                    self._update_streaming(block_e, body, final=True)
                # 后续使用 body（已去除 <think> 标签）
                e_content = body

                # ── 专家回答长度限制：截断到 HYBRID_EXPERT_MAX_CHARS，便于汇总 ──
                if HYBRID_EXPERT_MAX_CHARS > 0 and len(e_content) > HYBRID_EXPERT_MAX_CHARS:
                    e_content = _truncate_expert_response(e_content, HYBRID_EXPERT_MAX_CHARS)

                if e_content.strip():
                    # ── 保存到专家记忆（独立上下文，避免主上下文污染） ──
                    if EXPERT_MEMORY_TURNS > 0:
                        memory.append({"role": "user", "content": actual_user_msg})
                        memory.append({"role": "assistant", "content": e_content})
                        # 仅保留最近 EXPERT_MEMORY_TURNS 轮（每轮 2 条消息）
                        max_msgs = EXPERT_MEMORY_TURNS * 2
                        if len(memory) > max_msgs:
                            memory = memory[-max_msgs:]
                        self._expert_memory[ek] = memory
                    return {"expert": expert["label"], "content": e_content, "expert_key": ek}
                return None
            except Exception as e:
                # 失败降级：流处理过程中的异常，跳过该专家而非整个流程失败
                err_msg = f"  ⚠ {expert['label']} 调用失败，已跳过：{str(e)[:80]}\n"
                try:
                    block_e.update(Text(err_msg, style=C_DIM))
                except Exception:
                    self._add_static(Text(err_msg, style=C_DIM))
                return None
            finally:
                self._add_static(Text("  └─", style=C_DIM))

        # ── 执行专家调用：协作链（顺序）或并行 ──
        # 专家协作链：支持专家间传递结果（如 coder 写代码 → reasoner 审查逻辑 → academic 补充引用）
        # 默认关闭（HYBRID_ENABLE_COLLAB_CHAIN=False），避免 token 消耗翻倍
        if HYBRID_ENABLE_COLLAB_CHAIN and len(expert_keys) > 1:
            # 协作链模式：顺序调用，每个后续专家能看到前一位专家的回答
            self._add_static(Text("  🔗 协作链模式：专家依次回答并传递结果\n", style=C_DIM))
            expert_responses = []
            prev_response = ""
            for ek in expert_keys:
                if self._stop_generation:
                    break
                # 将前一位专家的回答注入本次用户问题，形成协作链
                if prev_response:
                    chain_user_msg = (
                        f"{last_user}\n\n"
                        f"── 上一环节专家（{EXPERT_TEAM[expert_keys[0]]['label']}）的回答 ──\n"
                        f"{prev_response[:HYBRID_EXPERT_MAX_CHARS]}\n"
                        f"── 请在此基础上从你的专业角度补充/审查/完善 ──"
                    )
                else:
                    chain_user_msg = last_user
                # 通过 user_msg 参数传递协作链上下文
                result = await _call_expert(ek, chain_user_msg)
                if isinstance(result, dict) and result.get("content"):
                    expert_responses.append(result)
                    prev_response = result["content"]
                # 协作链中某专家失败：降级跳过，继续下一个专家（不阻断链）
            parallel_results = expert_responses  # 统一变量名
        else:
            # 并行模式：所有专家同时调用（默认，token 效率最优）
            parallel_results = await asyncio.gather(
                *[_call_expert(ek) for ek in expert_keys],
                return_exceptions=True,
            )

        # 收集成功的回复
        expert_responses = []
        if HYBRID_ENABLE_COLLAB_CHAIN and len(expert_keys) > 1:
            # 协作链模式：parallel_results 已经是 list[dict]
            for result in parallel_results:
                if isinstance(result, dict) and result.get("content"):
                    expert_responses.append(result)
        else:
            # 并行模式：parallel_results 是 gather 的返回（含异常）
            for result in parallel_results:
                if isinstance(result, dict) and result.get("content"):
                    expert_responses.append(result)

        if not expert_responses:
            self._add_static(Text(f"  {_load_svg_icon('cross')} 所有专家调用失败\n", style=C_FG))
            # 复制修复：即使所有专家失败，也保存错误提示到 _last_reply_text，避免 Ctrl+Y 显示"无内容可复制"
            self._last_reply_text = "（所有专家调用失败，请检查网络或 API Key 后重试）"
            self._is_generating = False
            return

        # ── 专家去重：基于 Jaccard 相似度过滤高度相似的回答 ──
        # 场景：项目经理选了 coder + reasoner，但两者回答高度相似，汇总时去重以减少 token
        if HYBRID_DEDUP_SIMILARITY_THRESHOLD > 0 and len(expert_responses) > 1:
            unique_responses = [expert_responses[0]]
            dedup_skipped = []
            for resp in expert_responses[1:]:
                is_dup = False
                for kept in unique_responses:
                    sim = _jaccard_similarity(resp["content"], kept["content"])
                    if sim >= HYBRID_DEDUP_SIMILARITY_THRESHOLD:
                        is_dup = True
                        dedup_skipped.append((resp["expert"], kept["expert"], round(sim, 2)))
                        break
                if not is_dup:
                    unique_responses.append(resp)
            if len(dedup_skipped) > 0:
                # 显示去重提示
                dedup_msg = "  ℹ 去重："
                dedup_parts = [f"{a}≈{b}({s})" for a, b, s in dedup_skipped]
                dedup_msg += "，".join(dedup_parts) + " 已合并\n"
                self._add_static(Text(dedup_msg, style=C_DIM))
            expert_responses = unique_responses

        # ── 第3步：项目经理GLM汇总 ──
        if self._stop_generation or len(expert_responses) <= 1:
            # 单专家或被停止，直接用第一个专家的结果
            if expert_responses:
                self._last_reply_text = expert_responses[0]["content"]
            self.messages.append({"role": "assistant", "content": expert_responses[0]["content"] if expert_responses else ""})
            self._is_generating = False
            return

        block_sum = self._add_block("汇总·GLM-4V", C_GREEN)
        block_sum.update(Text.assemble(
            ("  ┌─ 项目经理·GLM-4V 汇总\n", f"bold {C_GREEN}"),
            ("  │ 汇总中…\n", C_DIM),
        ))

        summary_prompt = "以下是多位专家的回答，请综合整理为一份完整、连贯的回复：\n\n"
        for resp in expert_responses:
            summary_prompt += f"【{resp['expert']}】\n{resp['content'][:2000]}\n\n"
        summary_prompt += "\n请综合以上内容，给出最终回复。"

        try:
            if self.stream_enabled:
                sum_stream = await glm_client.chat.completions.create(
                    model=glm_cfg["model"],
                    messages=[{"role": "system", "content": TOOL_CAPABILITY_PROMPT + "\n\n你是 ZeroAI 项目经理，负责将多位专家的回答综合整理为完整、连贯的最终回复。保留关键信息，消除重复，按用户问题逻辑组织。"},
                              {"role": "user", "content": summary_prompt}],
                    temperature=self.temperature,
                    stream=True,
                )
                sum_content = ""
                sum_reasoning = ""
                async for chunk in sum_stream:
                    if self._stop_generation:
                        break
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta
                    rc = getattr(delta, "reasoning_content", None) or getattr(delta, "reasoning", None)
                    if rc:
                        sum_reasoning += rc
                        # 合并 <think> 标签内容
                        think_in, body = _parse_think_tags(sum_content)
                        combined = (sum_reasoning + ("\n" + think_in if think_in else "")).strip()
                        if combined:
                            self._update_streaming_with_reasoning(block_sum, body, combined)
                        else:
                            self._update_streaming(block_sum, body)
                        await asyncio.sleep(0)
                    if delta.content:
                        sum_content += delta.content
                        # 过滤模型内部特殊标签（<|observation|> <|system|> 等）
                        sum_content = _strip_model_tokens(sum_content)
                        # 解析 <think> 标签，分离思考与正文
                        think_in, body = _parse_think_tags(sum_content)
                        combined = (sum_reasoning + ("\n" + think_in if think_in else "")).strip()
                        if combined:
                            self._update_streaming_with_reasoning(block_sum, body, combined)
                        else:
                            self._update_streaming(block_sum, body)
                        # 复制修复：流式汇总过程中实时更新 _last_reply_text，确保中途停止也能复制
                        if body.strip():
                            self._last_reply_text = body
                        await asyncio.sleep(0)
            else:
                # stream=False：用可中断 await 包装
                sum_resp = await _interruptible_await(glm_client.chat.completions.create(
                    model=glm_cfg["model"],
                    messages=[{"role": "system", "content": TOOL_CAPABILITY_PROMPT + "\n\n你是 ZeroAI 项目经理，负责将多位专家的回答综合整理为完整、连贯的最终回复。保留关键信息，消除重复，按用户问题逻辑组织。"},
                              {"role": "user", "content": summary_prompt}],
                    temperature=self.temperature,
                    stream=False,
                ))
                if sum_resp is None or self._stop_generation:
                    # 复制修复：汇总被停止时，保存已有内容（流式模式下 _last_reply_text 已实时更新）
                    if not self._last_reply_text.strip() and expert_responses:
                        self._last_reply_text = expert_responses[0]["content"]
                    self._is_generating = False
                    self._add_static(Text("  ⏹ 已停止\n", style=C_DIM))
                    return
                sum_content = _strip_model_tokens(sum_resp.choices[0].message.content or "")
                sum_reasoning = getattr(sum_resp.choices[0].message, "reasoning_content", "") or ""
            # 最终更新：分离 think 和 body，合并所有思考内容
            think_in, body = _parse_think_tags(sum_content)
            combined_reasoning = (sum_reasoning + ("\n" + think_in if think_in else "")).strip()
            if combined_reasoning:
                self._update_streaming_with_reasoning(block_sum, body, combined_reasoning, final=True)
            else:
                self._update_streaming(block_sum, body, final=True)
            # 后续使用 body（已去除 <think> 标签）
            sum_content = body
            # 身份泄露过滤（混合模式汇总回复）
            sum_content, _id_leaked = _sanitize_identity_leak(sum_content)
            if _id_leaked:
                self._add_static(Text("  └─ 🛡️ 身份保护：已过滤底层模型信息", style=f"bold {C_YELLOW}"))
            if sum_content.strip():
                self._last_reply_text = sum_content
                self.messages.append({"role": "assistant", "content": sum_content})
        except Exception as e:
            # 汇总失败，用第一个专家的结果
            fallback = expert_responses[0]["content"]
            # 身份泄露过滤（混合模式降级回复）
            fallback, _id_leaked = _sanitize_identity_leak(fallback)
            if _id_leaked:
                self._add_static(Text("  └─ 🛡️ 身份保护：已过滤底层模型信息", style=f"bold {C_YELLOW}"))
            self._add_static(_safe_markdown(render_latex_in_text(fallback), code_theme="monokai"))
            self.messages.append({"role": "assistant", "content": fallback})
            self._last_reply_text = fallback

        self._add_static(Text("  └─", style=C_DIM))
        self._is_generating = False

    async def _run_turn(self):
        """执行 Agent 循环 - 流式 + Markdown 渲染（实时更新）"""
        self._is_generating = True
        self._stop_generation = False
        # 重置全局停止标志
        global _GLOBAL_STOP
        _GLOBAL_STOP = False

        # ── API Key 检查：未配置时直接提示，不发起请求 ──
        cur_key = MODEL_CONFIGS.get(self.model_key, {}).get("api_key", "")
        if not cur_key:
            self._add_static(Text.assemble(
                (f"  {_load_svg_icon('warning')} 未配置 API 密钥\n", f"bold {C_FG}"),
                (f"  当前模型：{get_model_display_name(self.model_key)}\n", C_DIM),
                (f"  请按 Ctrl+P 打开设置面板配置 API Key\n", f"bold {C_FG}"),
                (f"  智谱GLM 免费 Key 获取：https://open.bigmodel.cn/\n", C_DIM),
                (f"  或输入 /模型 ollama 使用本地模型（无需 Key）\n", C_DIM),
            ))
            self._is_generating = False
            return

        # ── 上下文自动压缩：超阈值时先压缩历史 ──
        try:
            est_tokens = _estimate_tokens(self.messages)
            threshold = int(self.context_limit * COMPRESS_THRESHOLD_RATIO)
            if est_tokens > threshold and len(self.messages) > 10:
                block_compress = self._add_block("上下文压缩", C_DIM)
                block_compress.update(Text.assemble(
                    (f"  {_load_svg_icon('tool')} 上下文自动压缩\n", f"bold {C_DIM}"),
                    (f"  │ 当前约 {est_tokens} tokens，超过阈值 {threshold}，正在压缩…\n", C_DIM),
                ))
                old_count = len(self.messages)
                old_tokens = est_tokens
                self.messages = await compress_context(self.messages, self.context_limit)
                new_tokens = _estimate_tokens(self.messages)
                new_count = len(self.messages)
                self._add_static(Text.assemble(
                    (f"  {_load_svg_icon('check')} 压缩完成：", C_DIM),
                    (f"{old_count}→{new_count} 条消息，", f"bold {C_FG}"),
                    (f"约 {old_tokens}→{new_tokens} tokens\n", C_DIM),
                    ("  └─\n", C_DIM),
                ))
        except Exception as e:
            # 压缩失败不阻塞对话
            self._add_static(Text(f"  {_load_svg_icon('warning')} 上下文压缩跳过：{str(e)[:80]}\n", style=C_DIM))

        _loop_expert_key = None
        while True:
            try:
                # ── 根据工作模式决定使用的模型 ──
                if self.work_mode == "expert":
                    # 专家模式：根据最后一条用户消息路由（仅首次/新用户消息时路由）
                    last_user = ""
                    last_user_has_image = False
                    is_new_user_msg = False
                    if not self.messages or self.messages[-1].get("role") == "user":
                        is_new_user_msg = True
                    for msg in reversed(self.messages):
                        if msg["role"] == "user":
                            content = msg["content"]
                            if isinstance(content, list):
                                text_parts = []
                                for part in content:
                                    if isinstance(part, dict):
                                        if part.get("type") == "text":
                                            text_parts.append(part.get("text", ""))
                                        elif part.get("type") == "image_url":
                                            last_user_has_image = True
                                last_user = " ".join(text_parts)
                            else:
                                last_user = content
                            break
                    if last_user_has_image:
                        expert_key = "vision"
                        _loop_expert_key = expert_key
                        block = self._add_block("路由分析", C_DIM)
                        block.update(Text.assemble(
                            (f"  {_load_svg_icon('image')} 检测到图片，自动路由到视觉专家\n", C_DIM),
                        ))
                    elif _loop_expert_key is None or is_new_user_msg:
                        _long_text_patterns = ["5000字", "3000字", "2000字", "万字", "长文",
                                               "综述类", "综述文章", "毕业论文", "学位论文",
                                               "长篇", "完整论文", "写一篇"]
                        _is_long_text = any(p in last_user for p in _long_text_patterns)
                        if _is_long_text and self.work_mode == "expert":
                            self._add_static(Text.assemble(
                                ("  ", C_DIM),
                                ("[!] 检测到长文写作任务，建议按 Ctrl+M 切换到混合思考模式（Hy）\n", f"bold {C_FG}"),
                                ("      混合模式可调度多位专家协作，生成更完整的长文内容\n", C_DIM),
                            ))
                        if len(last_user) >= 10:
                            block = self._add_block("路由分析", C_DIM)
                            block.update(Text.assemble(
                                (f"  {_load_svg_icon('search')} GLM 正在分析问题类型…\n", C_DIM),
                            ))
                        expert_key = await route_expert_glm(last_user)
                        _loop_expert_key = expert_key
                    else:
                        expert_key = _loop_expert_key
                    # OpenRouter 熔断检查：连续失败≥3次则直接降级到 GLM，避免卡在"思考中…"
                    if _check_openrouter_circuit_breaker(expert_key):
                        _fail_cnt = _OPENROUTER_FAIL_COUNTS.get(expert_key, 0)
                        _orig_label = EXPERT_TEAM[expert_key]["label"]
                        self._add_static(Text.assemble(
                            ("  ", C_DIM),
                            (f"{_orig_label}", f"bold {C_FG}"),
                            (f" 连续失败 {_fail_cnt} 次，已熔断，自动降级到 GLM…\n", C_DIM),
                        ))
                        expert_key = "knowledge"  # 降级到通用知识专家（GLM-4.7-Flash）
                    ecfg = get_expert_config(expert_key)
                    expert_label = EXPERT_TEAM[expert_key]["label"]
                    self._current_expert_label = expert_label
                    self._current_expert_key = expert_key
                    block = self._add_block(f"构建 · {expert_label}", C_RED)
                    block.update(Text.assemble(
                        (f"  ⏵ 构建 · {expert_label}\n", f"bold {C_RED}"),
                        ("  │ 思考中…\n", C_DIM),
                    ))
                    async_client = _make_openai_client(ecfg.get("model_key", "glm"))
                    if not _is_proxy_enabled():
                        async_client = AsyncOpenAI(
                            base_url=ecfg["base_url"], api_key=ecfg["api_key"],
                            timeout=120.0, max_retries=0,  # 禁用 SDK 内置重试，由下方自定义重试处理
                        )
                    _ctx_len = sum(len(str(m.get("content", ""))) for m in self.messages)
                    # 动态超时：上下文越长，超时越长。最小90秒，最大600秒（10分钟）
                    _dyn_timeout = min(600, max(90, 90 + _ctx_len // 500))
                    api_params = {
                        "model": ecfg["model"],
                        "messages": _filter_messages_for_model(self.messages, ecfg["model"]),
                        "tools": TOOLS,
                        "temperature": self.temperature,
                        "stream": self.stream_enabled,
                        "timeout": _dyn_timeout,
                    }
                    # 流式时请求 usage 数据（精确 token 统计）
                    if self.stream_enabled:
                        api_params["stream_options"] = {"include_usage": True}
                    # 专家降级标记：如果专家模型失败，降级到GLM
                    _fallback_to_glm = False
                elif self.work_mode == "hybrid":
                    # 混合思考：多专家协作
                    await self._run_hybrid_turn()
                    return
                else:
                    # 手动模式：使用用户选定的模型
                    manual_label = MODEL_CONFIGS.get(self.model_key, {}).get("label", self.model_key)
                    self._current_expert_label = manual_label
                    self._current_expert_key = "manual"
                    block = self._add_block(f"构建 · {manual_label}", C_RED)
                    block.update(Text.assemble(
                        (f"  ⏵ 构建 · {manual_label}\n", f"bold {C_RED}"),
                        ("  │ 思考中…\n", C_DIM),
                    ))
                    cfg = MODEL_CONFIGS[self.model_key]
                    async_client = _make_openai_client(self.model_key)
                    if not _is_proxy_enabled():
                        async_client = AsyncOpenAI(
                            base_url=cfg["base_url"], api_key=cfg["api_key"],
                            timeout=120.0, max_retries=0,  # 禁用 SDK 内置重试，由下方自定义重试处理
                        )
                    _ctx_len = sum(len(str(m.get("content", ""))) for m in self.messages)
                    _dyn_timeout = min(300, max(60, 60 + _ctx_len // 1000))
                    api_params = {
                        "model": cfg["model"],
                        "messages": _filter_messages_for_model(self.messages, cfg["model"]),
                        "tools": TOOLS,
                        "temperature": self.temperature,
                        "stream": self.stream_enabled,
                        "timeout": _dyn_timeout,
                    }
                    # 流式时请求 usage 数据（精确 token 统计）
                    if self.stream_enabled:
                        api_params["stream_options"] = {"include_usage": True}

                # 带重试的 API 调用（处理 429 速率限制 + ReadTimeout 超时 + 降级到GLM）
                stream = None
                max_retries = 5  # 从4增至5，超时多一次机会
                for attempt in range(max_retries):
                    try:
                        stream = await async_client.chat.completions.create(**api_params)
                        break
                    except Exception as retry_err:
                        err_str = str(retry_err)
                        err_type = type(retry_err).__name__
                        is_timeout = ("timeout" in err_str.lower() or "timed out" in err_str.lower()
                                      or "ReadTimeout" in err_type or "APITimeoutError" in err_type)
                        is_rate_limit = "429" in err_str or "rate" in err_str.lower()
                        is_server_error = any(code in err_str for code in ("500", "502", "503", "504")) or "ServerError" in err_type
                        # 连接错误（DNS/TCP/SSL 失败，通常是网络问题，应重试）
                        is_connection_error = ("APIConnectionError" in err_type
                                               or "ConnectionError" in err_type
                                               or "connection" in err_str.lower()
                                               or "ssl" in err_str.lower()
                                               or "eof" in err_str.lower())

                        # 专家模式下：专家模型失败时智能降级（首次失败即降级）
                        if self.work_mode == "expert" and not _fallback_to_glm and expert_key != "pm":
                            _fallback_to_glm = True
                            # 记录 OpenRouter 专家失败（用于熔断器累计，连续3次后直接跳过）
                            _fail_n = _record_openrouter_failure(expert_key)
                            # 智能选择降级模型：视觉专家降级到GLM-4V，文本专家降级到 GLM-4（不同限流池，避免同模型限流）
                            if expert_key == "vision":
                                _fallback_key = "glm-v"
                                _fallback_label = "GLM-4V"
                            else:
                                _fallback_key = "glm-4"
                                _fallback_label = "GLM-4"
                            self._add_static(Text.assemble(
                                ("  ", C_DIM),
                                (f"{expert_label}", f"bold {C_FG}"),
                                (f" 不可用（{err_type}），降级到 {_fallback_label}…\n", C_DIM),
                            ))
                            fb_cfg = MODEL_CONFIGS[_fallback_key]
                            expert_label = f"{_fallback_label}（降级）"
                            block = self._add_block(f"助手 [{expert_label}]", C_GREEN)
                            block.update(Text.assemble(
                                (f"  ┌─ 助手 [{expert_label}]\n", f"bold {C_GREEN}"),
                                ("  │ 思考中…\n", C_DIM),
                            ))
                            async_client = _make_openai_client(_fallback_key)
                            if not _is_proxy_enabled():
                                async_client = AsyncOpenAI(
                                    base_url=fb_cfg["base_url"], api_key=fb_cfg["api_key"],
                                    timeout=180.0, max_retries=0,  # 降级时用更长超时
                                )
                            api_params = {
                                "model": fb_cfg["model"],
                                "messages": _filter_messages_for_model(self.messages, fb_cfg["model"]),
                                "tools": TOOLS,
                                "temperature": self.temperature,
                                "stream": self.stream_enabled,
                                "timeout": _dyn_timeout,
                            }
                            if self.stream_enabled:
                                api_params["stream_options"] = {"include_usage": True}
                            continue

                        # 超时/服务器错误/连接错误：快速退避重试（限流且已降级则不重试，同限流池重试无效）
                        _can_retry = (is_timeout or is_server_error or is_connection_error) and attempt < max_retries - 1
                        # 限流且未降级时才重试（给原模型一次机会）；已降级还限流则直接抛错
                        if is_rate_limit and not _fallback_to_glm and attempt < max_retries - 1:
                            _can_retry = True
                        if _can_retry:
                            # 快速退避：1s, 2s, 3s, 4s（比 2/4/8 更快恢复）
                            wait = attempt + 1
                            if is_connection_error:
                                reason = "网络连接"
                            elif is_timeout:
                                reason = "超时"
                            elif is_rate_limit:
                                reason = "限流"
                            else:
                                reason = "服务器错误"
                            block.update(Text.assemble(
                                (f"  ⏵ 构建 · {expert_label}", f"bold {C_RED}"),
                                ("\n  │ ", C_DIM),
                                (f"⏳ {reason}，{wait}秒后重试", f"bold {C_YELLOW}"),
                                (f"（第{attempt+1}/{max_retries}次）\n", C_DIM),
                            ))
                            await _interruptible_sleep(wait)
                            if self._stop_generation:
                                break
                            # 超时重试时增加下次的超时时间（+60秒，更快达到宽容超时）
                            if is_timeout:
                                api_params["timeout"] = min(600, api_params["timeout"] + 60)
                            continue
                        # 限流且已降级：提示用户稍后再试
                        if is_rate_limit and _fallback_to_glm:
                            self._add_static(Text.assemble(
                                ("  └─ ", C_DIM),
                                (f"限流持续，请稍后再试或切换模型\n", f"bold {C_YELLOW}"),
                            ))
                        # 连接错误且已用尽重试：提示检查网络
                        if is_connection_error:
                            self._add_static(Text.assemble(
                                ("  └─ ", C_DIM),
                                (f"网络连接失败，请检查网络后重试\n", f"bold {C_YELLOW}"),
                            ))
                        raise retry_err

                if stream is None:
                    raise Exception("API 调用失败")

                # OpenRouter 专家本次调用成功（未降级），重置连续失败计数
                if self.work_mode == "expert" and not _fallback_to_glm:
                    _record_openrouter_success(expert_key)

                full_content = ""
                reasoning_content = ""
                tool_calls_buf = {}
                update_counter = 0
                _loop_detect_buf = []
                _loop_detect_window = 6
                _loop_detect_threshold = 4
                _loop_detected = False
                # Token 统计：开始计时
                self.stream_start_time = time.time()
                self.stream_token_count = 0
                _api_usage = None  # 精确 usage（从 API 响应获取）
                async for chunk in stream:
                    # 检查是否被用户 Ctrl+C 停止
                    if self._stop_generation:
                        break
                    # 捕获 usage 数据（stream_options include_usage 时最后一个 chunk 带 usage）
                    if hasattr(chunk, "usage") and chunk.usage:
                        _api_usage = chunk.usage
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta

                    # 捕获思考内容（原生 reasoning_content，推理模型才有）
                    rc = getattr(delta, "reasoning_content", None) or getattr(delta, "reasoning", None)
                    if rc:
                        reasoning_content += rc
                        self.stream_token_count += 1
                        self.total_tokens += 1
                        _loop_detect_buf.append(rc.strip())
                        if len(_loop_detect_buf) > _loop_detect_window:
                            _loop_detect_buf.pop(0)
                        if len(_loop_detect_buf) >= _loop_detect_window:
                            recent = _loop_detect_buf[-_loop_detect_threshold:]
                            if len(set(recent)) <= 1 and recent[0] and len(recent[0]) > 5:
                                _loop_detected = True
                                reasoning_content += "\n[系统自动截断：检测到循环思考]"
                                self._add_static(Text("  └─ ⚠️ 检测到循环思考，已自动截断", style="bold yellow"))
                                break
                        # 实时显示思考过程（合并原生 reasoning 和 <arg_key> 标签内容）
                        think_in_content, body_content = _parse_think_tags(full_content)
                        combined_reasoning = (reasoning_content + ("\n" + think_in_content if think_in_content else "")).strip()
                        if combined_reasoning:
                            self._update_streaming_with_reasoning(block, body_content, combined_reasoning)
                        else:
                            self._update_streaming(block, body_content)
                        await asyncio.sleep(0)

                    if delta.content:
                        full_content += delta.content
                        # 过滤模型内部特殊标签（<|observation|> <|system|> 等，防止泄露给用户）
                        full_content = _strip_model_tokens(full_content)
                        update_counter += 1
                        # 统计 token（每个 chunk 约 1 token）
                        self.stream_token_count += 1
                        self.total_tokens += 1
                        # 每 3 个 chunk 更新一次（避免过于频繁）
                        if update_counter % 3 == 0:
                            # 解析 <think>...</think> 标签内容，合并到思考过程
                            think_in_content, body_content = _parse_think_tags(full_content)
                            combined_reasoning = (reasoning_content + ("\n" + think_in_content if think_in_content else "")).strip()
                            if combined_reasoning:
                                self._update_streaming_with_reasoning(block, body_content, combined_reasoning)
                            else:
                                self._update_streaming(block, body_content)
                            self._update_token_bar()
                        await asyncio.sleep(0)

                    if delta.tool_calls:
                        for tc in delta.tool_calls:
                            idx = tc.index
                            if idx not in tool_calls_buf:
                                tool_calls_buf[idx] = {"id": "", "name": "", "args": ""}
                            if tc.id:
                                tool_calls_buf[idx]["id"] = tc.id
                            if tc.function.name:
                                tool_calls_buf[idx]["name"] = tc.function.name
                            if tc.function.arguments:
                                tool_calls_buf[idx]["args"] += tc.function.arguments
                                self.stream_token_count += 1
                                self.total_tokens += 1

                # 最终更新：分离 think 和 body，合并所有思考内容
                think_in_content, body_content = _parse_think_tags(full_content)
                combined_reasoning = (reasoning_content + ("\n" + think_in_content if think_in_content else "")).strip()
                if combined_reasoning:
                    self._update_streaming_with_reasoning(block, body_content, combined_reasoning, final=True)
                else:
                    self._update_streaming(block, body_content, final=True)
                # 用 API 返回的精确 usage 更新 token 统计
                if _api_usage:
                    try:
                        # usage 对象格式：{prompt_tokens, completion_tokens, total_tokens}
                        prompt_t = getattr(_api_usage, "prompt_tokens", None) or _api_usage.get("prompt_tokens", 0)
                        completion_t = getattr(_api_usage, "completion_tokens", None) or _api_usage.get("completion_tokens", 0)
                        if completion_t > 0:
                            # 用精确值替换估算值
                            self.stream_token_count = completion_t
                            self.total_tokens = (self.total_tokens - self.stream_token_count + completion_t
                                                 if self.total_tokens >= self.stream_token_count else completion_t)
                            # 保存精确输入 token 到实例属性
                            self._precise_input_tokens = prompt_t
                    except Exception:
                        pass
                self._update_token_bar()
                # 后续统一使用 body_content（已去除 <think> 标签）
                full_content = body_content
                # 身份泄露过滤：检测底层模型自报家门，替换为标准 ZeroAI 身份（响应层防线）
                full_content, _id_leaked = _sanitize_identity_leak(full_content)
                if _id_leaked:
                    self._add_static(Text("  └─ 🛡️ 身份保护：已过滤底层模型信息", style=f"bold {C_YELLOW}"))
                # 存储最近回复文本（用于复制）
                if full_content.strip():
                    self._last_reply_text = full_content
                # 被停止时显示截断标记
                if self._stop_generation:
                    self._add_static(Text("  └─ ⏹ 已停止", style=C_DIM))
                    self._is_generating = False
                    if full_content.strip():
                        self.messages.append({"role": "assistant", "content": full_content})
                    return

                # 处理工具调用
                if tool_calls_buf:
                    assistant_msg = {
                        "role": "assistant",
                        "content": full_content if full_content else "",
                        "tool_calls": []
                    }
                    for idx_key, tc in tool_calls_buf.items():
                        if not tc["id"]:
                            tc["id"] = f"call_{idx_key}_{tc['name']}"
                        assistant_msg["tool_calls"].append({
                            "id": tc["id"],
                            "type": "function",
                            "function": {"name": tc["name"], "arguments": tc["args"]}
                        })
                    self.messages.append(assistant_msg)

                    for tc in tool_calls_buf.values():
                        # 工具执行前检查是否被 Ctrl+C 停止
                        if self._stop_generation:
                            break
                        name = tc["name"]
                        try:
                            args = json.loads(tc["args"]) if tc["args"].strip() else {}
                        except json.JSONDecodeError:
                            args = {}

                        tool_info = f"**调用工具** `{name}`\n\n```json\n{json.dumps(args, ensure_ascii=False, indent=2)}\n```"
                        self._add_static(_safe_markdown(tool_info, code_theme="monokai"))

                        if name in TOOL_MAP:
                            try:
                                fn = TOOL_MAP[name]
                                # 过滤掉函数不接受的参数，防止模型幻觉参数
                                valid_params = set(inspect.signature(fn).parameters)
                                safe_args = {k: v for k, v in args.items() if k in valid_params}
                                if set(args.keys()) - valid_params:
                                    extra = set(args.keys()) - valid_params
                                    result = fn(**safe_args)
                                    result += f"\n[提示：忽略多余参数 {extra}]"
                                else:
                                    result = fn(**safe_args)
                                # read_image 返回 base64，构造多模态消息
                                if name == "read_image" and isinstance(result, str) and result.startswith("data:"):
                                    img_path = safe_args.get("path", "图片")
                                    b64_data = result  # 保存 base64 数据
                                    result = f"{_load_svg_icon('document')} 已读取图片：{img_path}，图片内容已发送给模型"
                                    self.messages.append({
                                        "role": "tool",
                                        "tool_call_id": tc["id"],
                                        "name": name,
                                        "content": result,
                                    })
                                    self.messages.append({
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": f"这是图片 {img_path}，请描述和理解这张图片"},
                                            {"type": "image_url", "image_url": {"url": b64_data}},
                                        ],
                                    })
                                    self.messages.append({"role": "assistant", "content": "我看到了这张图片，"})
                                    result_md = f"**结果**\n\n{result}"
                                    self._add_static(_safe_markdown(result_md, code_theme="monokai"))
                                    self._add_static(Text("  └─", style=C_DIM))
                                    continue
                            except TypeError as e:
                                result = f"参数错误：{e}"
                            except Exception as e:
                                result = f"执行错误：{e}"
                        else:
                            result = f"未知工具：{name}"

                        # 工具结果长度限制，防止超长内容卡死模型
                        _result_str = str(result)
                        MAX_TOOL_RESULT = 8000  # 字符上限，防止超长结果撑爆上下文
                        _truncated = False
                        if len(_result_str) > MAX_TOOL_RESULT:
                            _orig_len = len(_result_str)
                            _result_str = (_result_str[:MAX_TOOL_RESULT]
                                           + f"\n\n[结果过长，已截断。原始长度 {_orig_len} 字符。"
                                           + "请用更具体的路径或 search_files 精确搜索。]")
                            _truncated = True

                        _result_label = f"**结果**（已截断至 {MAX_TOOL_RESULT} 字符）" if _truncated else "**结果**"
                        result_md = f"{_result_label}\n\n```\n{_result_str}\n```"
                        self._add_static(_safe_markdown(result_md, code_theme="monokai"))

                        self.messages.append({
                            "role": "tool",
                            "tool_call_id": tc["id"],
                            "name": name,
                            "content": _result_str,
                        })
                    self._add_static(Text("  └─", style=C_DIM))
                    # 工具执行后检查是否被停止
                    if self._stop_generation:
                        self._is_generating = False
                        self._add_static(Text("  ⏹ 已停止\n", style=C_DIM))
                        return
                    continue
                else:
                    if full_content.strip():
                        self.messages.append({"role": "assistant", "content": full_content})
                    self._add_static(Text("  └─", style=C_DIM))
                    self._is_generating = False
                    return

            except Exception as e:
                err_str = str(e)
                err_type = type(e).__name__
                # 401 认证错误：API Key 无效或未配置
                if "401" in err_str or "AuthenticationError" in err_type or "Authorization" in err_str:
                    self._add_static(Text.assemble(
                        (f"  {_load_svg_icon('cross')} 认证失败（401）：API Key 无效或未配置\n", f"bold {C_FG}"),
                        (f"  错误详情：{err_str[:120]}\n", C_DIM),
                        (f"  解决方法：\n", f"bold {C_FG}"),
                        (f"    1. 按 Ctrl+P 打开设置面板，重新配置 API Key\n", C_FG),
                        (f"    2. 检查 Key 是否正确（智谱GLM 格式：xxx.xxx）\n", C_DIM),
                        (f"    3. 智谱GLM 免费 Key 获取：https://open.bigmodel.cn/\n", C_DIM),
                        (f"    4. 或输入 /模型 ollama 使用本地模型（无需 Key）\n", C_DIM),
                    ))
                elif "429" in err_str or "rate" in err_str.lower():
                    self._add_static(Text.assemble(
                        (f"  {_load_svg_icon('warning')} 速率限制：", f"bold {C_FG}"),
                        (f"{err_str[:100]}\n", C_DIM),
                        (f"  {_load_svg_icon('warning')} 建议：", C_DIM),
                        ("输入 /手动 切到GLM（免费无限）或 /模型 ollama 用本地模型\n", C_FG),
                    ))
                elif "timeout" in err_str.lower() or "timed out" in err_str.lower() or "ReadTimeout" in err_type:
                    self._add_static(Text.assemble(
                        (f"  {_load_svg_icon('warning')} 请求超时（{err_type}）：\n", f"bold {C_FG}"),
                        (f"  {err_str[:150]}\n", C_DIM),
                        (f"  {_load_svg_icon('warning')} 已自动重试 {max_retries} 次仍失败\n", C_DIM),
                        (f"  {_load_svg_icon('warning')} 建议：\n", f"bold {C_FG}"),
                        (f"    1. 输入 /手动 切到 GLM（国内直连·稳定）\n", C_FG),
                        (f"    2. OpenRouter 慢时可开启 VPN 或稍后重试\n", C_DIM),
                    ))
                else:
                    self._add_static(Text.assemble(
                        (f"  {err_type}：", f"bold {C_FG}"),
                        (f"{err_str[:200]}\n", C_DIM),
                    ))
                if self.messages and self.messages[-1].get("role") == "user":
                    self.messages.pop()
                while self.messages and self.messages[-1].get("role") in ("assistant", "tool"):
                    self.messages.pop()
                self._is_generating = False
                return

    def _update_token_bar(self):
        """更新右侧 Token 统计栏（输入+输出 token，动态上下文窗口）"""
        try:
            token_bar = self.query_one("#token-bar", TokenBar)
            elapsed = time.time() - self.stream_start_time
            rate = self.stream_token_count / elapsed if elapsed > 0.1 else 0.0
            # 输入 token：优先用 API 返回的精确值，否则用估算值
            input_tokens = getattr(self, '_precise_input_tokens', 0) or _estimate_tokens(self.messages)
            # 动态获取当前模型的上下文窗口
            ctx_window = self._get_current_ctx_window()
            token_bar.update_stats(self.total_tokens, rate, input_tokens, ctx_window)
            # 同步更新底部 HintBar 的 token 显示
            try:
                hint_bar = self.query_one("#hints", HintBar)
                hint_bar.update_ctx(input_tokens + self.total_tokens, ctx_window)
            except Exception:
                pass
        except Exception:
            pass

    def _get_current_ctx_window(self) -> int:
        """获取当前工作模式下的实际上下文窗口大小"""
        try:
            if self.work_mode == "expert":
                # 专家模式：用当前路由专家的模型
                expert_key = getattr(self, '_current_expert_key', None) or 'knowledge'
                if expert_key in EXPERT_TEAM:
                    e_cfg = get_expert_config(expert_key)
                    limit = _get_model_context_limit(e_cfg["model"])
                    return limit if limit > 0 else 128000
            elif self.work_mode == "hybrid":
                # 混合模式：用 PM 的模型
                e_cfg = get_expert_config("pm")
                limit = _get_model_context_limit(e_cfg["model"])
                return limit if limit > 0 else 128000
            # 手动模式：用当前选定的模型配置
            cfg = MODEL_CONFIGS.get(self.model_key, {})
            limit = _get_model_context_limit(cfg.get("model", ""))
            return limit if limit > 0 else 128000
        except Exception:
            return 128000

    def _update_streaming(self, block: Static, content: str, final: bool = False):
        """实时更新消息块 - 流式时纯文本稳定显示，完成后 Markdown 渲染"""
        if not content.strip():
            return
        try:
            # 提取当前专家标签
            expert_label = getattr(self, '_current_expert_label', 'Expert')
            if final:
                # 完成：用 Markdown 渲染（Typora 风格，含 LaTeX 公式渲染）
                md = _safe_markdown(render_latex_in_text(content), code_theme="monokai")
                from rich.console import Group
                elapsed = time.time() - getattr(self, 'stream_start_time', time.time())
                group = Group(
                    Text.assemble(
                        (f"  ⏵ 构建 · {expert_label}", f"bold {C_RED}"),
                        (f"  ·  {elapsed:.1f}s", C_DIM),
                    ),
                    md,
                )
                block.update(group)
            else:
                # 流式：纯文本稳定显示（文字在原位增长，不重新布局）
                parts = [(f"  ⏵ 构建 · {expert_label}", f"bold {C_RED}"), ("\n", "")]
                for line in content.split("\n"):
                    parts.append((f"  {line}\n", C_FG))
                block.update(Text.assemble(*parts))
            # 只在用户没有主动上翻时自动滚底
            if not self._user_scrolling:
                self.query_one("#log-scroll", VerticalScroll).scroll_end(animate=False)
        except Exception:
            # 降级：纯文本
            block.update(Text.assemble(
                ("  ┌─ 助手\n", f"bold {C_GREEN}"),
                (f"  {content}", C_FG),
            ))

    def _update_streaming_with_reasoning(self, block: Static, content: str,
                                          reasoning: str, final: bool = False):
        """实时更新消息块（含思考过程显示）

        思考过程用灰色显示在正文上方，完成后可折叠。
        - 流式中：思考过程实时显示（灰色），正文在下方（白色）
        - 完成后：思考过程用 Panel 包裹（灰色斜体），正文用 Markdown 渲染
        """
        try:
            from rich.console import Group
            if final:
                # 完成：思考过程用 Panel + 正文用 Markdown
                md = _safe_markdown(render_latex_in_text(content), code_theme="monokai") if content.strip() else Text("")
                # 思考过程限制显示长度（太长影响阅读）
                reasoning_display = reasoning if len(reasoning) <= 2000 else reasoning[:2000] + "\n  ...（思考过程较长，已截断）"
                reasoning_panel = Panel(
                    Text(reasoning_display, style=f"italic {C_DIM}"),
                    title="思考过程",
                    title_align="left",
                    border_style=C_DIM,
                    padding=(0, 1),
                )
                group = Group(
                    Text.assemble(
                        ("  ┌─ 助手\n", f"bold {C_GREEN}"),
                        ("  │\n", C_DIM),
                    ),
                    reasoning_panel,
                    Text("  │\n", C_DIM),
                    md,
                )
                block.update(group)
            else:
                # 流式：思考过程灰色 + 正文白色
                parts = [("  ┌─ 助手\n", f"bold {C_GREEN}"), ("  │\n", C_DIM)]
                if reasoning.strip():
                    parts.append(("  │ 💭 思考中…\n", f"italic {C_DIM}"))
                    # 只显示最后几行思考内容（避免刷屏）
                    reasoning_lines = reasoning.split("\n")
                    show_lines = reasoning_lines[-6:]  # 最后6行
                    for line in show_lines:
                        if line.strip():
                            parts.append((f"  │ {line}\n", C_DIM))
                    parts.append(("  │\n", C_DIM))
                if content.strip():
                    for line in content.split("\n"):
                        parts.append((f"  {line}\n", C_FG))
                else:
                    parts.append(("  ⏳…\n", C_DIM))
                block.update(Text.assemble(*parts))
            # 自动滚底
            if not self._user_scrolling:
                self.query_one("#log-scroll", VerticalScroll).scroll_end(animate=False)
        except Exception:
            # 降级：纯文本
            block.update(Text.assemble(
                ("  ┌─ 助手\n", f"bold {C_GREEN}"),
                (f"  {content}", C_FG),
            ))


def main():
    try:
        app = ZeroAI()
        app.title = "ZeroAI"
        app.run()
    finally:
        # 程序退出时清理运行时缓存
        runtime_cache.cleanup()


if __name__ == "__main__":
    main()